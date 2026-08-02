# ================================================================
# AVANCE DEL TEMARIO — ACADEMIA PREUNIVERSITARIA
# Módulo para SISTEMA YACHAY PRO
# Basado en el Temario de Admisión UNSAAC (Res. CU-575-2024-UNSAAC)
# ================================================================
"""
Cómo se integra en sistema_web.py (3 líneas, ver GUIA-AVANCE-TEMARIO.md):

    from avance_temario import tab_avance_docente, tab_avance_coordinacion

Este módulo no depende del resto del sistema: si google_sync existe lo usa,
si no, guarda todo en un archivo local avance_temario.json.
"""

import json
import io
from datetime import datetime, date, timedelta
from pathlib import Path

import streamlit as st
import pandas as pd

try:
    from google_sync import get_google_sync
except Exception:
    get_google_sync = None


# ================================================================
# 1. TEMARIO OFICIAL UNSAAC
# ================================================================
# Cada curso es una lista de temas. El nombre corto es lo que ve el
# docente; el número es el del temario oficial, para que cualquier
# reclamo se pueda contrastar con la resolución.

TEMARIO = {
    "Aritmética": [
        "Teoría de conjuntos",
        "Sistema de números naturales y enteros",
        "Sistema de números racionales",
        "Sucesiones y sumatorias notables",
        "Sistemas de numeración",
        "Divisibilidad",
        "Números primos",
        "Máximo común divisor y mínimo común múltiplo",
        "Razones y proporciones",
        "Magnitudes proporcionales",
        "Regla de tres y regla de tanto por ciento",
        "Regla de interés simple y compuesto",
        "Introducción a la estadística",
        "Introducción a las probabilidades",
    ],
    "Álgebra": [
        "Potenciación",
        "Polinomios",
        "Productos notables",
        "División de polinomios, teorema del resto y cocientes notables",
        "Factorización de polinomios",
        "Racionalización",
        "Ecuaciones de primer y segundo grado con una variable real",
        "Inecuaciones de primer y segundo grado",
        "Ecuaciones e inecuaciones con valor absoluto",
        "Matrices y determinantes",
        "Relaciones",
        "Función",
        "Funciones especiales",
        "Clases de funciones",
        "Operaciones con funciones",
        "Función exponencial",
        "Función logarítmica",
    ],
    "Geometría y Trigonometría": [
        "Nociones básicas de la geometría",
        "Segmento de recta",
        "Ángulos",
        "Triángulos",
        "Congruencia y semejanza de triángulos",
        "Relaciones métricas de triángulos rectángulos y oblicuángulos",
        "Cuadriláteros",
        "Circunferencia",
        "Polígonos",
        "Áreas de regiones",
        "Fundamentos de la trigonometría",
        "Razones trigonométricas de ángulos agudos",
        "Ángulo en posición normal",
        "Identidades trigonométricas",
        "Resolución de triángulos y ángulos verticales y horizontales",
        "Funciones trigonométricas",
    ],
    "Física": [
        "La física y magnitudes",
        "Vectores en el plano y el espacio",
        "Cinemática",
        "Dinámica lineal",
        "Estática",
        "Trabajo y energía",
        "Dinámica de rotación",
        "Movimiento oscilatorio",
        "Mecánica de fluidos",
        "Temperatura, dilatación y calor",
        "Termodinámica",
        "Electrostática",
        "Electrodinámica",
        "Electromagnetismo",
        "Movimiento ondulatorio",
        "Óptica",
        "Física moderna",
    ],
    "Química": [
        "Química y materia",
        "Estructura atómica",
        "Números cuánticos y configuración electrónica",
        "Tabla periódica moderna",
        "Enlace químico",
        "Nomenclatura de compuestos inorgánicos",
        "Masa atómica",
        "Reacciones químicas",
        "Reacciones de oxidación y reducción",
        "Estequiometría de las reacciones químicas",
        "Soluciones",
        "Equilibrio químico",
        "Hidrocarburos",
        "Alcoholes, fenoles y éteres",
        "Aldehídos y cetonas",
        "Ácidos carboxílicos y ésteres",
    ],
    "Biología": [
        "Concepto de biología y niveles de organización",
        "Composición química de la materia viviente",
        "Biomoléculas inorgánicas",
        "Biomoléculas orgánicas",
        "Célula",
        "Célula eucariota",
        "Nutrición",
        "Nivel sistémico",
        "Coordinación",
        "Reproducción",
        "Genética",
        "Evolución y origen de la vida",
        "Ecología, factores ecológicos y ecosistemas",
        "Flujo de energía y ciclos biogeoquímicos",
        "Diversidad biológica y deterioro de la flora y la fauna",
        "Contaminación, problemas ambientales y conservación",
    ],
    "Competencia Comunicativa": [
        "La comunicación",
        "El lenguaje",
        "Fonología y fonética",
        "Sílaba",
        "Acentuación gráfica o tildación",
        "Uso de las letras mayúsculas y minúsculas",
        "Signos de puntuación",
        "Sustantivo",
        "El pronombre",
        "El adjetivo",
        "El artículo y el adverbio",
        "El verbo",
        "Conectores lógico semánticos",
        "La sintaxis y oración gramatical",
        "El texto y la lectura",
        "Relaciones semánticas",
    ],
    "Historia": [
        "Ciencia histórica",
        "Hombre de la prehistoria",
        "Grandes culturas de la antigüedad",
        "Mundo greco romano",
        "Primeras culturas andinas",
        "Culturas preincas",
        "Civilización inca",
        "Mundo medieval y el tránsito al mundo moderno",
        "Expansión europea",
        "Conquista del Perú",
        "El periodo colonial peruano",
        "El mundo durante el siglo XVIII",
        "Movimientos sociales en el mundo colonial americano",
        "Tiempo de las revoluciones",
        "Crisis del orden colonial e independencia",
        "Construcción de la república peruana",
        "Estado peruano en transformación",
        "El mundo entre guerras",
        "Entre dictaduras y democracias: gobernantes del Perú siglos XX-XXI",
    ],
    "Geografía": [
        "Geografía y espacio geográfico",
        "Geosistema y espacio exterior",
        "Cartografía y sistemas de información geográfica",
        "Mapas: lectura e interpretación",
        "Relieve terrestre: origen y procesos dinámicos",
        "Espacio geográfico peruano: región andina",
        "Espacio geográfico peruano: región amazónica y costa",
        "Hidrografía del Perú: ríos y lagos",
        "Hidrografía del Perú: mar peruano",
        "Atmósfera y cambio climático",
        "Recursos naturales, conservación e impacto ambiental",
        "Riesgo de desastres en el Perú",
        "Dinámica poblacional en el Perú",
        "Actividades económicas extractivas en el Perú",
        "Actividades económicas reproductivas en el Perú",
        "Actividades del transporte en el Perú",
        "Geografía política del Perú y gestión territorial",
        "Espacio geográfico físico del Cusco",
        "Geografía de América",
        "Geografía de Europa, Asia, África, Oceanía y Antártida",
    ],
    "Economía": [
        "Conceptos generales",
        "Necesidades humanas",
        "Bienes y servicios",
        "Proceso económico",
        "Trabajo",
        "Capital",
        "Naturaleza",
        "Empresa",
        "Demanda",
        "Oferta",
        "Mercado",
        "Dinero e inflación",
        "Sistema financiero y crédito",
        "Distribución",
        "Sector público y presupuesto nacional",
        "Sector externo",
        "Crisis y ciclos",
        "Desarrollo y crecimiento económico",
    ],
    "Filosofía y Lógica": [
        "El problema del cosmos y concepciones de filosofía",
        "Historia de la filosofía: edad antigua",
        "Edad medieval y renacimiento",
        "La filosofía moderna y filosofía en el Perú",
        "Antropología filosófica: el problema del hombre",
        "Gnoseología: problema del conocimiento",
        "Corrientes del problema del conocimiento",
        "Problema de la ciencia: epistemología",
        "Problema del valor y la ética",
        "Lógica, lenguaje y pensamiento",
        "Falacias",
        "Pruebas formales en la lógica proposicional",
        "Tablas de verdad y razonamientos válidos",
        "Principios lógicos y lógica formal clásica",
        "Inferencias",
        "Lógica de clases",
        "Fórmulas booleanas y diagramas de Venn",
    ],
    "Educación Cívica": [
        "Derecho, ley y moral",
        "Valores cívicos sociales",
        "Persona y sociedad",
        "Familia",
        "Nación",
        "Estado",
        "Constitución política",
        "Derechos civiles y políticos",
        "Derechos económicos, sociales y culturales",
        "Poder legislativo",
        "Poder ejecutivo",
        "Poder judicial",
        "Organismos constitucionales autónomos",
        "Régimen económico",
        "Descentralización, gobiernos regionales y locales",
        "Derechos humanos",
        "Garantías constitucionales",
        "Sistemas de protección internacional de los DDHH",
    ],
}


# ================================================================
# 2. PESO DE CADA CURSO EN EL EXAMEN (nº de preguntas por área)
# ================================================================
# Esto es lo que convierte el módulo en algo útil: no es lo mismo
# atrasarse en Aritmética (14 preguntas) que en Cívica (8).

PESOS = {
    "A": {"Aritmética": 14, "Álgebra": 10, "Geometría y Trigonometría": 14,
          "Competencia Comunicativa": 14, "Física": 14, "Química": 14},
    "B": {"Aritmética": 14, "Álgebra": 10, "Competencia Comunicativa": 14,
          "Biología": 14, "Física": 14, "Química": 14},
    "C": {"Aritmética": 14, "Álgebra": 10, "Competencia Comunicativa": 14,
          "Historia": 12, "Geografía": 12, "Economía": 10, "Educación Cívica": 8},
    "D": {"Aritmética": 14, "Álgebra": 10, "Competencia Comunicativa": 14,
          "Historia": 12, "Geografía": 12, "Filosofía y Lógica": 10,
          "Educación Cívica": 8},
}

GRUPOS = {
    "GRUPO AB": ["A", "B"],
    "GRUPO CD": ["C", "D"],
}

ESTADOS = ["Pendiente", "En avance", "Concluido", "Reforzado"]

ESTADO_PCT = {"Pendiente": 0, "En avance": 50, "Concluido": 100, "Reforzado": 100}

ESTADO_COLOR = {
    "Pendiente": "#94a3b8",
    "En avance": "#f59e0b",
    "Concluido": "#16a34a",
    "Reforzado": "#2563eb",
}

ARCHIVO_AVANCE = "avance_temario.json"
ARCHIVO_CONFIG_CICLO = "ciclo_academia.json"

HOJA_AVANCE = "AvanceTemario"
COLS_AVANCE = ["clave", "anio", "ciclo", "grupo", "curso", "tema_num",
               "tema", "estado", "fecha", "sesion", "docente",
               "observacion", "evals", "actualizado"]

# Los exámenes se guardan con el nombre del área, que no siempre coincide
# exactamente con el nombre del curso en el temario oficial.
ALIAS_AREA = {
    "Geometría y Trigonometría": ["Geometría", "Trigonometría",
                                  "Geometria", "Trigonometria"],
    "Competencia Comunicativa": ["Competencia Lingüística", "Comunicación",
                                 "Lenguaje", "Razonamiento Verbal"],
    "Filosofía y Lógica": ["Filosofía", "Lógica"],
    "Educación Cívica": ["Cívica", "Educación Ciudadana"],
}

# Nota mínima vigesimal de aprobación usada en la academia.
# Archivo donde 'Registrar Notas → Nueva Evaluación' deja las notas.
NOTA_UMBRAL = 11.0
ARCHIVO_RESULTADOS_NOTAS = "resultados.json"


# ================================================================
# 3. CURSOS SEGÚN GRUPO
# ================================================================

def cursos_de_grupo(grupo):
    """Devuelve la lista de cursos que se dictan en un grupo."""
    areas = GRUPOS.get(grupo, [])
    cursos = []
    for a in areas:
        for c in PESOS.get(a, {}):
            if c not in cursos:
                cursos.append(c)
    return cursos


def peso_curso(grupo, curso):
    """Promedio de preguntas del curso en las áreas del grupo."""
    areas = GRUPOS.get(grupo, [])
    vals = [PESOS[a].get(curso, 0) for a in areas]
    vals = [v for v in vals if v > 0]
    return round(sum(vals) / len(vals), 1) if vals else 0


# ================================================================
# 4. PERSISTENCIA
# ================================================================

def _gs():
    if get_google_sync is None:
        return None
    try:
        return get_google_sync()
    except Exception:
        return None


def _clave(anio, ciclo, grupo, curso, tema_num):
    return f"{anio}|{ciclo}|{grupo}|{curso}|{tema_num}"


def _leer_local():
    p = Path(ARCHIVO_AVANCE)
    if not p.exists():
        return {}
    try:
        with open(p, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def _escribir_local(data):
    try:
        with open(ARCHIVO_AVANCE, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
    except Exception:
        pass


@st.cache_data(ttl=60, show_spinner=False)
def _leer_hoja_avance():
    """Lee la hoja de Google Sheets. Cacheado 60s para no saturar la API."""
    gs = _gs()
    if gs is None:
        return {}
    try:
        hojas = [w.title for w in gs.spreadsheet.worksheets()]
        if HOJA_AVANCE not in hojas:
            ws = gs.spreadsheet.add_worksheet(title=HOJA_AVANCE, rows=2000,
                                              cols=len(COLS_AVANCE))
            ws.append_row(COLS_AVANCE)
            return {}
        ws = gs.spreadsheet.worksheet(HOJA_AVANCE)
        registros = ws.get_all_records()
        return {str(r.get("clave", "")): r for r in registros if r.get("clave")}
    except Exception:
        return {}


def cargar_avance():
    """Google Sheets manda; el archivo local es respaldo."""
    remoto = _leer_hoja_avance()
    if remoto:
        local = _leer_local()
        local.update(remoto)
        _escribir_local(local)
        return local
    return _leer_local()


def guardar_avance(registros):
    """Guarda una lista de registros. Escribe local siempre y sube a Sheets."""
    data = _leer_local()
    for r in registros:
        r["actualizado"] = datetime.now().strftime("%Y-%m-%d %H:%M")
        data[r["clave"]] = r
    _escribir_local(data)

    gs = _gs()
    if gs is None:
        return False, "Guardado solo en este equipo (sin conexión a la nube)."

    try:
        hojas = [w.title for w in gs.spreadsheet.worksheets()]
        if HOJA_AVANCE not in hojas:
            ws = gs.spreadsheet.add_worksheet(title=HOJA_AVANCE, rows=2000,
                                              cols=len(COLS_AVANCE))
            ws.append_row(COLS_AVANCE)
        else:
            ws = gs.spreadsheet.worksheet(HOJA_AVANCE)

        claves_existentes = ws.col_values(1)
        nuevas, actualizaciones = [], []
        for r in registros:
            fila = [str(r.get(c, "")) for c in COLS_AVANCE]
            if r["clave"] in claves_existentes:
                idx = claves_existentes.index(r["clave"]) + 1
                actualizaciones.append((idx, fila))
            else:
                nuevas.append(fila)

        for idx, fila in actualizaciones:
            ws.update(f"A{idx}:M{idx}", [fila])
        if nuevas:
            ws.append_rows(nuevas)

        _leer_hoja_avance.clear()
        return True, f"{len(registros)} tema(s) guardado(s) y sincronizado(s)."
    except Exception as e:
        return False, f"Guardado local. La nube no respondió: {e}"


def cargar_config_ciclo():
    p = Path(ARCHIVO_CONFIG_CICLO)
    if p.exists():
        try:
            with open(p, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {"ciclo": "Ciclo Regular", "inicio": "", "examen": ""}


def guardar_config_ciclo(cfg):
    with open(ARCHIVO_CONFIG_CICLO, "w", encoding="utf-8") as f:
        json.dump(cfg, f, indent=2, ensure_ascii=False)


# ================================================================
# 5. CÁLCULO DEL AVANCE
# ================================================================

def resumen_curso(avance, anio, ciclo, grupo, curso):
    """Devuelve el estado agregado de un curso."""
    temas = TEMARIO.get(curso, [])
    total = len(temas)
    conteo = {e: 0 for e in ESTADOS}
    puntos = 0
    for i in range(total):
        r = avance.get(_clave(anio, ciclo, grupo, curso, i + 1))
        estado = (r or {}).get("estado", "Pendiente")
        if estado not in ESTADOS:
            estado = "Pendiente"
        conteo[estado] += 1
        puntos += ESTADO_PCT[estado]
    pct = round(puntos / total, 1) if total else 0.0
    return {
        "curso": curso,
        "total": total,
        "pct": pct,
        "peso": peso_curso(grupo, curso),
        **conteo,
    }


def tabla_grupo(avance, anio, ciclo, grupo):
    filas = [resumen_curso(avance, anio, ciclo, grupo, c)
             for c in cursos_de_grupo(grupo)]
    return pd.DataFrame(filas)


def dias_restantes(cfg):
    try:
        f = datetime.strptime(cfg.get("examen", ""), "%Y-%m-%d").date()
        return (f - date.today()).days
    except Exception:
        return None


def diagnostico(pct, dias):
    """Semáforo: cruza avance con tiempo restante hasta el examen."""
    if dias is None:
        return "—", "#94a3b8"
    if dias <= 0:
        return "Ciclo cerrado", "#64748b"
    # Referencia: a 30 días del examen se espera 80%; a 60 días, 60%.
    if dias <= 15:
        esperado = 95
    elif dias <= 30:
        esperado = 85
    elif dias <= 60:
        esperado = 65
    elif dias <= 90:
        esperado = 45
    else:
        esperado = 25
    brecha = pct - esperado
    if brecha >= 0:
        return "Al día", "#16a34a"
    if brecha >= -15:
        return "Ajustado", "#f59e0b"
    return "Atrasado", "#dc2626"


# ================================================================
# 6. VISTA DEL DOCENTE — marcar avance
# ================================================================

def tab_avance_docente(config=None):
    st.subheader("📚 Avance del Temario — UNSAAC")
    st.caption("Temario oficial aprobado por Res. CU-575-2024-UNSAAC")

    info = st.session_state.get("docente_info", {}) or {}
    docente = info.get("label") or st.session_state.get("usuario_actual", "—")
    cfg = cargar_config_ciclo()
    anio = date.today().year

    c1, c2 = st.columns(2)
    with c1:
        grupo = st.selectbox("Grupo", list(GRUPOS.keys()), key="av_doc_grupo")
    with c2:
        cursos = cursos_de_grupo(grupo)
        curso = st.selectbox("Curso que dictas", cursos, key="av_doc_curso")

    ciclo = cfg.get("ciclo", "Ciclo Regular")
    avance = cargar_avance()
    temas = TEMARIO.get(curso, [])

    res = resumen_curso(avance, anio, ciclo, grupo, curso)
    dias = dias_restantes(cfg)
    etiqueta, color = diagnostico(res["pct"], dias)

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Avance", f"{res['pct']}%")
    m2.metric("Temas concluidos", f"{res['Concluido'] + res['Reforzado']}/{res['total']}")
    m3.metric("Peso en el examen", f"{res['peso']} preguntas")
    m4.markdown(
        f"<div style='background:{color};color:#fff;padding:14px;border-radius:10px;"
        f"text-align:center;font-weight:bold;'>{etiqueta}"
        f"<div style='font-size:.8rem;font-weight:normal;'>"
        f"{'' if dias is None else f'{dias} días para el examen'}</div></div>",
        unsafe_allow_html=True)

    st.progress(min(res["pct"] / 100, 1.0))
    st.markdown("---")

    st.markdown("#### Marca lo que ya dictaste")
    st.caption("Solo cambia lo que avanzaste hoy y pulsa guardar al final.")

    fecha_sesion = st.date_input("Fecha de la sesión", value=date.today(),
                                 key="av_doc_fecha")

    cambios = []
    for i, tema in enumerate(temas, start=1):
        clave = _clave(anio, ciclo, grupo, curso, i)
        actual = avance.get(clave, {})
        estado_prev = actual.get("estado", "Pendiente")
        if estado_prev not in ESTADOS:
            estado_prev = "Pendiente"

        col_t, col_e, col_o = st.columns([5, 2, 3])
        with col_t:
            punto = ESTADO_COLOR[estado_prev]
            st.markdown(
                f"<div style='padding:6px 0;'>"
                f"<span style='display:inline-block;width:10px;height:10px;"
                f"border-radius:50%;background:{punto};margin-right:8px;'></span>"
                f"<b>{i}.</b> {tema}</div>", unsafe_allow_html=True)
        with col_e:
            nuevo = st.selectbox("Estado", ESTADOS,
                                 index=ESTADOS.index(estado_prev),
                                 key=f"av_est_{clave}",
                                 label_visibility="collapsed")
        with col_o:
            obs = st.text_input("Observación", value=actual.get("observacion", ""),
                                key=f"av_obs_{clave}",
                                placeholder="Sesión, dificultad, pendiente…",
                                label_visibility="collapsed")

        if nuevo != estado_prev or obs != actual.get("observacion", ""):
            cambios.append({
                "clave": clave,
                "anio": anio,
                "ciclo": ciclo,
                "grupo": grupo,
                "curso": curso,
                "tema_num": i,
                "tema": tema,
                "estado": nuevo,
                "fecha": fecha_sesion.strftime("%Y-%m-%d"),
                "sesion": actual.get("sesion", ""),
                "docente": docente,
                "observacion": obs,
            })

    st.markdown("---")
    if cambios:
        st.info(f"Tienes {len(cambios)} cambio(s) sin guardar.")
    cg1, cg2 = st.columns(2)
    with cg1:
        if st.button("💾 GUARDAR AVANCE", type="primary", use_container_width=True,
                     disabled=not cambios, key="av_doc_guardar"):
            ok, msg = guardar_avance(cambios)
            (st.success if ok else st.warning)(msg)
            st.rerun()
    with cg2:
        try:
            pdf = generar_pdf_docente(avance, anio, ciclo, grupo, curso,
                                      docente, dias)
            st.download_button(
                "📄 Descargar mi avance en PDF",
                data=pdf,
                file_name=(f"avance_{curso.replace(' ', '_')}_"
                           f"{grupo.replace(' ', '')}_{anio}.pdf"),
                mime="application/pdf", use_container_width=True,
                key="av_doc_pdf")
        except Exception as e:
            st.caption(f"No se pudo preparar el PDF: {e}")

    st.markdown("---")
    with st.expander("🎯 Cruzar con las notas de mis exámenes", expanded=False):
        nt = notas_del_curso(grupo, curso)
        if nt["promedio"] is not None:
            etiqueta, color, lectura = cuadrante(res["pct"], nt["promedio"])
            k1, k2, k3 = st.columns(3)
            k1.metric("Promedio del curso", nt["promedio"])
            k2.metric(f"Sobre {NOTA_UMBRAL}", f"{nt['pct_aprob']}%")
            k3.markdown(
                f"<div style='background:{color};color:#fff;padding:12px;"
                f"border-radius:8px;text-align:center;font-weight:bold;'>"
                f"{etiqueta}</div>", unsafe_allow_html=True)
            st.caption(lectura)
            st.markdown("---")
        seccion_vincular_evaluaciones(avance, anio, ciclo, grupo, curso, docente)


# ================================================================
# 7. VISTA DE COORDINACIÓN — panel de control
# ================================================================

def tab_avance_coordinacion(config=None):
    st.subheader("📊 Control de Avance — Academia Preuniversitaria")

    cfg = cargar_config_ciclo()
    anio = date.today().year

    with st.expander("⚙️ Configurar el ciclo", expanded=not cfg.get("examen")):
        c1, c2, c3 = st.columns(3)
        with c1:
            nciclo = st.text_input("Nombre del ciclo", value=cfg.get("ciclo", ""))
        with c2:
            ini = st.date_input(
                "Inicio de clases",
                value=datetime.strptime(cfg["inicio"], "%Y-%m-%d").date()
                if cfg.get("inicio") else date.today())
        with c3:
            exa = st.date_input(
                "Fecha del examen",
                value=datetime.strptime(cfg["examen"], "%Y-%m-%d").date()
                if cfg.get("examen") else date.today() + timedelta(days=90))
        if st.button("💾 Guardar configuración", key="av_cfg_save",
                     type="primary", use_container_width=True):
            guardar_config_ciclo({
                "ciclo": nciclo,
                "inicio": ini.strftime("%Y-%m-%d"),
                "examen": exa.strftime("%Y-%m-%d"),
            })
            st.success("Configuración guardada.")
            st.rerun()

    ciclo = cfg.get("ciclo", "Ciclo Regular")
    dias = dias_restantes(cfg)
    avance = cargar_avance()

    if dias is not None:
        if dias > 0:
            st.info(f"**{ciclo}** — faltan **{dias} días** para el examen de admisión.")
        else:
            st.warning(f"**{ciclo}** — la fecha del examen ya pasó.")

    tab1, tab4, tab2, tab5, tab6, tab3 = st.tabs(
        ["🚦 Semáforo por curso", "🎯 Dictado vs Aprendido",
         "👤 Por docente", "📋 Cronograma docentes", "📅 Calendario",
         "📥 Reportes"])

    # ---------- Semáforo ----------
    with tab1:
        for grupo in GRUPOS:
            df = tabla_grupo(avance, anio, ciclo, grupo)
            if df.empty:
                continue
            pct_grupo = round((df["pct"] * df["peso"]).sum() / df["peso"].sum(), 1) \
                if df["peso"].sum() else 0
            st.markdown(f"### {grupo} — avance ponderado: **{pct_grupo}%**")
            st.caption("Ponderado por número de preguntas del examen, "
                       "no por cantidad de temas.")

            for _, r in df.sort_values("pct").iterrows():
                etiqueta, color = diagnostico(r["pct"], dias)
                st.markdown(
                    f"""<div style='display:flex;align-items:center;gap:12px;
                    padding:10px 14px;margin-bottom:6px;border-radius:8px;
                    background:#f8fafc;border-left:6px solid {color};'>
                    <div style='flex:3;'><b>{r['curso']}</b>
                    <span style='color:#64748b;font-size:.85rem;'>
                    · {r['peso']} preguntas</span></div>
                    <div style='flex:4;background:#e2e8f0;border-radius:20px;height:14px;'>
                    <div style='width:{min(r['pct'],100)}%;background:{color};
                    height:14px;border-radius:20px;'></div></div>
                    <div style='flex:1;text-align:right;'><b>{r['pct']}%</b></div>
                    <div style='flex:1;text-align:right;color:{color};
                    font-weight:bold;font-size:.85rem;'>{etiqueta}</div>
                    </div>""", unsafe_allow_html=True)

            atrasados = [r["curso"] for _, r in df.iterrows()
                         if diagnostico(r["pct"], dias)[0] == "Atrasado"]
            if atrasados:
                st.error("Requieren intervención inmediata: " + ", ".join(atrasados))
            st.markdown("---")

    # ---------- Dictado vs Aprendido ----------
    with tab4:
        st.markdown("Cruza lo que **se dictó** con lo que los alumnos "
                    "**demostraron en los exámenes** que ya registras en el "
                    "sistema. Un curso puede ir al 90% de avance y aun así "
                    "ser el que más puntos te está costando.")
        hay_notas = False
        for grupo in GRUPOS:
            filas = tabla_dictado_vs_aprendido(avance, anio, ciclo, grupo)
            st.markdown(f"### {grupo}")
            for f in filas:
                if f["promedio"] is not None:
                    hay_notas = True
                prom_txt = "sin notas" if f["promedio"] is None else \
                    f"promedio {f['promedio']} · {f['pct_aprob']}% sobre {NOTA_UMBRAL}"
                st.markdown(
                    f"""<div style='padding:10px 14px;margin-bottom:6px;
                    border-radius:8px;background:#f8fafc;
                    border-left:6px solid {f['color']};'>
                    <div style='display:flex;justify-content:space-between;
                    align-items:center;'>
                    <div><b>{f['curso']}</b>
                    <span style='color:#64748b;font-size:.85rem;'>
                    · {int(f['peso'])} preguntas · avance {f['avance']}% ·
                    {prom_txt}</span></div>
                    <div style='color:{f['color']};font-weight:bold;'>
                    {f['estado']}</div></div>
                    <div style='color:#475569;font-size:.82rem;margin-top:4px;'>
                    {f['lectura']}</div></div>""", unsafe_allow_html=True)

            criticos = [f["curso"] for f in filas
                        if f["estado"] == "Dictado, no aprendido"]
            if criticos:
                st.error("**Se avanzó pero no se asimiló:** " +
                         ", ".join(criticos) +
                         ". Estos cursos no necesitan más velocidad, "
                         "necesitan repaso.")
            st.markdown("---")

        if not hay_notas:
            st.info("Todavía no hay notas registradas para los grupos "
                    "preuniversitarios. En cuanto apliques exámenes desde "
                    "Exámenes Semanales o YACHAY QAWAY, este panel se llena solo.")

    # ---------- Por docente ----------
    with tab2:
        filas = []
        for r in avance.values():
            if str(r.get("anio")) != str(anio):
                continue
            filas.append(r)
        if not filas:
            st.info("Todavía no hay registros de avance.")
        else:
            df = pd.DataFrame(filas)
            df["concluido"] = df["estado"].isin(["Concluido", "Reforzado"])
            resumen = (df.groupby(["docente", "grupo", "curso"])
                         .agg(temas_registrados=("tema", "count"),
                              concluidos=("concluido", "sum"),
                              ultima_actualizacion=("actualizado", "max"))
                         .reset_index())
            resumen["% del curso"] = resumen.apply(
                lambda x: round(100 * x["concluidos"] /
                                max(len(TEMARIO.get(x["curso"], [])), 1), 1), axis=1)
            st.dataframe(resumen, use_container_width=True, hide_index=True)

            st.markdown("#### Docentes sin registrar en los últimos 14 días")
            limite = (datetime.now() - timedelta(days=14)).strftime("%Y-%m-%d")
            inactivos = resumen[
                resumen["ultima_actualizacion"].astype(str) < limite]
            if inactivos.empty:
                st.success("Todos los docentes registraron avance recientemente.")
            else:
                for _, r in inactivos.iterrows():
                    st.warning(f"**{r['docente']}** — {r['curso']} ({r['grupo']}) · "
                               f"último registro: {r['ultima_actualizacion']}")

    # ---------- Cronograma para docentes ----------
    with tab5:
        st.markdown("Genera el comunicado en Word que se entrega a los "
                    "docentes: fechas del ciclo, qué temas avanza cada uno "
                    "por semana, cuántas preguntas entrega y con qué "
                    "numeración.")

        cr1, cr2 = st.columns(2)
        with cr1:
            _preset = st.selectbox("Ciclo:", list(CICLOS_PRESET.keys()),
                                   key="cr_preset")
            _area_cr = st.selectbox(
                "Área / salón:",
                ["A", "B", "C", "D", "GRUPO AB", "GRUPO CD"],
                key="cr_area",
                help="Elige un área sola si cada postulante tiene su "
                     "propio salón. Elige GRUPO AB o GRUPO CD si el "
                     "salón junta a dos áreas y comparten docente en "
                     "los cursos comunes.")
            if es_grupo_combinado(_area_cr):
                st.caption(f"Cursos comunes de {', '.join(GRUPOS[_area_cr])}: "
                           + ", ".join(cursos_de_grupo(_area_cr)))
        with cr2:
            _m1, _d1, _m2, _d2 = CICLOS_PRESET[_preset]
            _anio_cr = st.number_input("Año:", 2024, 2040, anio, key="cr_anio")
            _a2 = _anio_cr + (1 if _m2 < _m1 else 0)
            _ini_cr = st.date_input("Inicio de clases:",
                                    value=date(_anio_cr, _m1, _d1), key="cr_ini")
            _fin_cr = st.date_input("Fin del ciclo:",
                                    value=date(_a2, _m2, min(_d2, 28)),
                                    key="cr_fin")

        cr3, cr4, cr5 = st.columns(3)
        with cr3:
            _wsp = st.text_input("WhatsApp de Secretaría:", key="cr_wsp",
                                 placeholder="984 123 456")
        with cr4:
            _dia_ent = st.selectbox("Día límite de entrega:", DIAS_SEMANA[:6],
                                    index=3, key="cr_dia")
        with cr5:
            _sem_rep = st.number_input("Semanas de repaso al final:", 0, 6, 2,
                                       key="cr_repaso")

        if _fin_cr <= _ini_cr:
            st.error("El fin del ciclo debe ser posterior al inicio.")
        else:
            _fer = cargar_feriados()
            _cron = cronograma_area(_area_cr, _ini_cr, _fin_cr, _sem_rep, _fer)
            if _cron.get("omitidos"):
                st.warning("Sábados sin examen por día no lectivo: " +
                           ", ".join(f"{s.strftime('%d/%m')} ({m})"
                                     for s, m in _cron["omitidos"]))
            st.info(f"**{_cron['semanas']} semanas** — primer examen el "
                    f"{_cron['sabados'][0].strftime('%d/%m/%Y')}, último el "
                    f"{_cron['sabados'][-1].strftime('%d/%m/%Y')}. "
                    f"Los exámenes son los sábados por la tarde.")

            st.markdown("##### Docente responsable de cada curso")
            _docs = {}
            _cols_d = st.columns(2)
            for _i, _c in enumerate(_cron["cursos"]):
                with _cols_d[_i % 2]:
                    _docs[_c["curso"]] = st.text_input(
                        f"{_c['curso']} — preguntas {_c['desde']} al {_c['hasta']}",
                        key=f"cr_doc_{_area_cr}_{_c['curso']}",
                        placeholder="Apellidos y Nombres")

            st.markdown("##### Vista previa del reparto")
            _prev = []
            for _i, _sab in enumerate(_cron["sabados"]):
                _fila = {"Sem": _i + 1, "Examen": _sab.strftime("%d/%m")}
                for _c in _cron["cursos"]:
                    _t = _c["tramos"][_i] if _i < len(_c["tramos"]) else (None, None)
                    _fila[_c["curso"][:14]] = ("repaso" if _t[0] is None else
                                               (f"T{_t[0]}" if _t[0] == _t[1]
                                                else f"T{_t[0]}–{_t[1]}"))
                _prev.append(_fila)
            st.dataframe(pd.DataFrame(_prev), use_container_width=True,
                         hide_index=True, height=280)

            if st.button("📄 GENERAR COMUNICADO EN WORD", type="primary",
                         use_container_width=True, key="cr_gen"):
                try:
                    _docx = generar_comunicado_docx(
                        _area_cr, _ini_cr, _fin_cr, f"{_preset} {_anio_cr}",
                        docentes={k: v for k, v in _docs.items() if v.strip()},
                        whatsapp_secretaria=_wsp, dia_entrega=_dia_ent,
                        semanas_repaso=_sem_rep, feriados=_fer)
                    st.download_button(
                        "⬇️ Descargar comunicado (.docx)", data=_docx,
                        file_name=(f"cronograma_area_{_area_cr}_"
                                   f"{_preset.split()[1]}_{_anio_cr}.docx"),
                        mime=("application/vnd.openxmlformats-officedocument."
                              "wordprocessingml.document"),
                        use_container_width=True, key="cr_dl")
                    st.success("Listo. Ábrelo en Word, ajusta lo que necesites "
                               "y repártelo a los docentes.")
                except Exception as e:
                    st.error(f"No se pudo generar: {e}")

    # ---------- Calendario ----------
    with tab6:
        seccion_calendario_feriados()

    # ---------- Reportes ----------
    with tab3:
        filas = []
        for grupo in GRUPOS:
            for curso in cursos_de_grupo(grupo):
                for i, tema in enumerate(TEMARIO.get(curso, []), start=1):
                    r = avance.get(_clave(anio, ciclo, grupo, curso, i), {})
                    filas.append({
                        "Grupo": grupo,
                        "Curso": curso,
                        "N°": i,
                        "Tema": tema,
                        "Estado": r.get("estado", "Pendiente"),
                        "Fecha": r.get("fecha", ""),
                        "Docente": r.get("docente", ""),
                        "Observación": r.get("observacion", ""),
                    })
        df_full = pd.DataFrame(filas)
        st.dataframe(df_full, use_container_width=True, hide_index=True,
                     height=400)

        buf = io.BytesIO()
        with pd.ExcelWriter(buf, engine="openpyxl") as w:
            df_full.to_excel(w, sheet_name="Avance detallado", index=False)
            for grupo in GRUPOS:
                tabla_grupo(avance, anio, ciclo, grupo).to_excel(
                    w, sheet_name=grupo.replace(" ", "_")[:30], index=False)
        d1, d2 = st.columns(2)
        with d1:
            st.download_button(
                "📥 Descargar reporte en Excel",
                data=buf.getvalue(),
                file_name=f"avance_temario_{ciclo.replace(' ', '_')}_{anio}.xlsx",
                mime=("application/vnd.openxmlformats-officedocument."
                      "spreadsheetml.sheet"),
                use_container_width=True, key="av_coord_xlsx")
        with d2:
            try:
                pdf = generar_pdf_coordinacion(avance, anio, ciclo, dias)
                st.download_button(
                    "📄 Informe de seguimiento en PDF",
                    data=pdf,
                    file_name=f"informe_avance_{ciclo.replace(' ', '_')}_{anio}.pdf",
                    mime="application/pdf", use_container_width=True,
                    type="primary", key="av_coord_pdf")
            except Exception as e:
                st.caption(f"No se pudo preparar el PDF: {e}")

        st.markdown("---")
        st.markdown("#### Planilla para marcar a mano")
        st.caption("Hoja impresa con todos los cursos y sus temas, con "
                   "casilleros vacíos para marcar con aspa y firmar.")
        pc1, pc2 = st.columns(2)
        for col, grupo_pl in zip([pc1, pc2], list(GRUPOS.keys())):
            with col:
                try:
                    pl = generar_pdf_planilla(grupo_pl, ciclo, anio)
                    st.download_button(
                        f"🖨️ Planilla {grupo_pl}",
                        data=pl,
                        file_name=(f"planilla_temario_{grupo_pl.replace(' ', '')}"
                                   f"_{anio}.pdf"),
                        mime="application/pdf", use_container_width=True,
                        key=f"av_planilla_{grupo_pl}")
                except Exception as e:
                    st.caption(f"No se pudo generar: {e}")

        st.markdown("---")
        st.markdown("#### Temario completo por área de postulación")
        st.caption("Una hoja por área con todos sus cursos, cuánto pesa cada "
                   "uno y sus temas. Sirve para entregar al postulante.")
        _ca = st.columns(4)
        for _col, _ar in zip(_ca, ["A", "B", "C", "D"]):
            with _col:
                try:
                    _pa = generar_pdf_area(_ar, ciclo, anio)
                    st.download_button(
                        f"📘 Área {_ar}", data=_pa,
                        file_name=f"temario_area_{_ar}_{anio}.pdf",
                        mime="application/pdf", use_container_width=True,
                        key=f"av_area_{_ar}")
                except Exception as e:
                    st.caption(f"Error: {e}")

        st.caption("Para un salón que junta dos áreas (AB o CD), descarga "
                   "el temario combinado:")
        _cg1, _cg2 = st.columns(2)
        for _col, _gr in zip([_cg1, _cg2], list(GRUPOS.keys())):
            with _col:
                try:
                    _pg = generar_pdf_area(_gr, ciclo, anio)
                    st.download_button(
                        f"📘 {_gr}", data=_pg,
                        file_name=f"temario_{_gr.replace(' ', '_')}_{anio}.pdf",
                        mime="application/pdf", use_container_width=True,
                        key=f"av_grupo_{_gr}")
                except Exception as e:
                    st.caption(f"Error: {e}")

        st.markdown("---")
        st.markdown("#### Reporte individual por docente")
        cc1, cc2 = st.columns(2)
        with cc1:
            g_sel = st.selectbox("Grupo", list(GRUPOS.keys()), key="av_rep_grupo")
        with cc2:
            c_sel = st.selectbox("Curso", cursos_de_grupo(g_sel),
                                 key="av_rep_curso")
        doc_nom = ""
        for i in range(len(TEMARIO.get(c_sel, []))):
            r = avance.get(_clave(anio, ciclo, g_sel, c_sel, i + 1), {})
            if r.get("docente"):
                doc_nom = r["docente"]
                break
        try:
            pdf_ind = generar_pdf_docente(avance, anio, ciclo, g_sel, c_sel,
                                          doc_nom or "—", dias)
            st.download_button(
                f"📄 Descargar detalle de {c_sel}",
                data=pdf_ind,
                file_name=(f"avance_{c_sel.replace(' ', '_')}_"
                           f"{g_sel.replace(' ', '')}_{anio}.pdf"),
                mime="application/pdf", use_container_width=True,
                key="av_rep_pdf_ind")
        except Exception as e:
            st.caption(f"No se pudo preparar el PDF: {e}")


# ================================================================
# 8. REPORTES EN PDF
# ================================================================

PIE_LEGAL = ("Derechos reservados — I.E.P. ALTERNATIVO YACHAY · "
             "Documento generado por SISTEMA YACHAY PRO")


def _pie_pagina(canvas, doc):
    """Pie legal en todas las paginas, con numeracion."""
    from reportlab.lib.units import cm
    canvas.saveState()
    canvas.setFont("Helvetica", 6.5)
    canvas.setFillColorRGB(0.42, 0.45, 0.50)
    canvas.drawCentredString(doc.pagesize[0] / 2, 0.75 * cm, PIE_LEGAL)
    canvas.drawRightString(doc.pagesize[0] - 1.4 * cm, 0.75 * cm,
                           f"Pág. {canvas.getPageNumber()}")
    canvas.setStrokeColorRGB(0.80, 0.83, 0.87)
    canvas.setLineWidth(0.4)
    canvas.line(1.4 * cm, 1.05 * cm, doc.pagesize[0] - 1.4 * cm, 1.05 * cm)
    canvas.restoreState()


def _estilos_pdf():
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib import colors
    ss = getSampleStyleSheet()
    return {
        "titulo": ParagraphStyle("t", parent=ss["Title"], fontSize=15,
                                 textColor=colors.HexColor("#001e7c"),
                                 spaceAfter=2, alignment=TA_CENTER),
        "sub": ParagraphStyle("s", parent=ss["Normal"], fontSize=9,
                              textColor=colors.HexColor("#475569"),
                              alignment=TA_CENTER, spaceAfter=10),
        "h2": ParagraphStyle("h", parent=ss["Heading2"], fontSize=11,
                             textColor=colors.HexColor("#001e7c"),
                             spaceBefore=10, spaceAfter=4),
        "n": ParagraphStyle("n", parent=ss["Normal"], fontSize=8.5,
                            leading=11),
        "pie": ParagraphStyle("p", parent=ss["Normal"], fontSize=7.5,
                              textColor=colors.HexColor("#64748b"),
                              alignment=TA_CENTER),
    }


def _encabezado(story, titulo, subtitulo, est):
    from reportlab.platypus import Paragraph, Spacer
    story.append(Paragraph(titulo, est["titulo"]))
    story.append(Paragraph(subtitulo, est["sub"]))
    story.append(Spacer(1, 4))


def generar_pdf_coordinacion(avance, anio, ciclo, dias, institucion="ACADEMIA YACHAY"):
    """Informe de seguimiento del avance de todos los docentes."""
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle)
    from reportlab.lib import colors
    from reportlab.lib.units import cm

    est = _estilos_pdf()
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=1.6 * cm, rightMargin=1.6 * cm,
                            topMargin=1.4 * cm, bottomMargin=1.4 * cm)
    story = []

    hoy = datetime.now().strftime("%d/%m/%Y %H:%M")
    sub = f"{ciclo} {anio} &nbsp;·&nbsp; Emitido el {hoy}"
    if dias is not None and dias > 0:
        sub += f" &nbsp;·&nbsp; Faltan {dias} días para el examen"
    _encabezado(story, f"{institucion}<br/>INFORME DE AVANCE DEL TEMARIO", sub, est)
    story.append(Paragraph(
        "Temario de Admisión UNSAAC aprobado por Res. CU-575-2024-UNSAAC. "
        "El porcentaje del grupo está ponderado por el número de preguntas "
        "que cada curso tiene en el examen.", est["n"]))
    story.append(Spacer(1, 8))

    for grupo in GRUPOS:
        df = tabla_grupo(avance, anio, ciclo, grupo)
        if df.empty:
            continue
        pond = round((df["pct"] * df["peso"]).sum() / df["peso"].sum(), 1) \
            if df["peso"].sum() else 0
        story.append(Paragraph(
            f"{grupo} — avance ponderado: {pond}%", est["h2"]))

        data = [["Curso", "Preg.", "Temas", "Concl.", "En avance",
                 "Pend.", "Avance", "Estado"]]
        estilos_fila = []
        for i, (_, r) in enumerate(df.sort_values("pct").iterrows(), start=1):
            etiqueta, color = diagnostico(r["pct"], dias)
            data.append([
                Paragraph(str(r["curso"]), est["n"]),
                str(int(r["peso"])), str(int(r["total"])),
                str(int(r["Concluido"] + r["Reforzado"])),
                str(int(r["En avance"])), str(int(r["Pendiente"])),
                f"{r['pct']}%", etiqueta,
            ])
            estilos_fila.append(
                ("TEXTCOLOR", (7, i), (7, i), colors.HexColor(color)))

        t = Table(data, colWidths=[5.2 * cm, 1.3 * cm, 1.3 * cm, 1.5 * cm,
                                   1.9 * cm, 1.4 * cm, 1.7 * cm, 2.2 * cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#001e7c")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ALIGN", (1, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1),
             [colors.white, colors.HexColor("#f8fafc")]),
            ("FONTNAME", (7, 1), (7, -1), "Helvetica-Bold"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ] + estilos_fila))
        story.append(t)

        atras = [r["curso"] for _, r in df.iterrows()
                 if diagnostico(r["pct"], dias)[0] == "Atrasado"]
        if atras:
            story.append(Spacer(1, 4))
            story.append(Paragraph(
                f"<b>Requieren intervención:</b> {', '.join(atras)}", est["n"]))
        story.append(Spacer(1, 6))

    # --- Responsables por curso ---
    filas = [r for r in avance.values() if str(r.get("anio")) == str(anio)]
    if filas:
        df = pd.DataFrame(filas)
        df["concl"] = df["estado"].isin(["Concluido", "Reforzado"])
        res = (df.groupby(["docente", "grupo", "curso"])
                 .agg(concl=("concl", "sum"), ult=("actualizado", "max"))
                 .reset_index())
        story.append(Paragraph("Registro por docente", est["h2"]))
        data = [["Docente", "Grupo", "Curso", "Temas concluidos",
                 "Último registro"]]
        for _, r in res.sort_values(["grupo", "curso"]).iterrows():
            tot = len(TEMARIO.get(r["curso"], [])) or 1
            data.append([
                Paragraph(str(r["docente"]), est["n"]),
                str(r["grupo"]).replace("GRUPO ", ""),
                Paragraph(str(r["curso"]), est["n"]),
                f"{int(r['concl'])}/{tot}",
                str(r["ult"])[:16],
            ])
        t = Table(data, colWidths=[4.8 * cm, 1.6 * cm, 4.6 * cm,
                                   2.8 * cm, 3.0 * cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0891b2")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ALIGN", (1, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1),
             [colors.white, colors.HexColor("#f0f9ff")]),
        ]))
        story.append(t)

    story.append(Spacer(1, 14))
    story.append(Paragraph(
        "_______________________________<br/>"
        "Coordinación Académica", est["pie"]))
    story.append(Spacer(1, 6))
    story.append(Paragraph(
        "Documento generado automáticamente por SISTEMA YACHAY PRO", est["pie"]))

    doc.build(story, onFirstPage=_pie_pagina,
              onLaterPages=_pie_pagina)
    buf.seek(0)
    return buf.getvalue()


def generar_pdf_docente(avance, anio, ciclo, grupo, curso, docente, dias,
                        institucion="ACADEMIA YACHAY"):
    """Detalle tema por tema del curso de un docente."""
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle)
    from reportlab.lib import colors
    from reportlab.lib.units import cm

    est = _estilos_pdf()
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=1.6 * cm, rightMargin=1.6 * cm,
                            topMargin=1.4 * cm, bottomMargin=1.4 * cm)
    story = []

    res = resumen_curso(avance, anio, ciclo, grupo, curso)
    etiqueta, color = diagnostico(res["pct"], dias)
    hoy = datetime.now().strftime("%d/%m/%Y")

    _encabezado(story, f"{institucion}<br/>AVANCE DEL TEMARIO — {curso.upper()}",
                f"{grupo} &nbsp;·&nbsp; {ciclo} {anio} &nbsp;·&nbsp; "
                f"Docente: {docente} &nbsp;·&nbsp; {hoy}", est)

    resumen = Table([[
        f"Avance\n{res['pct']}%",
        f"Concluidos\n{res['Concluido'] + res['Reforzado']}/{res['total']}",
        f"Peso en el examen\n{int(res['peso'])} preguntas",
        f"Estado\n{etiqueta}",
    ]], colWidths=[4.2 * cm] * 4)
    resumen.setStyle(TableStyle([
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f1f5f9")),
        ("TEXTCOLOR", (3, 0), (3, 0), colors.HexColor(color)),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica-Bold"),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(resumen)
    story.append(Spacer(1, 10))

    data = [["N°", "Tema del temario oficial", "Estado", "Fecha", "Observación"]]
    estilos_fila = []
    for i, tema in enumerate(TEMARIO.get(curso, []), start=1):
        r = avance.get(_clave(anio, ciclo, grupo, curso, i), {})
        estado = r.get("estado", "Pendiente")
        if estado not in ESTADOS:
            estado = "Pendiente"
        data.append([
            str(i),
            Paragraph(tema, est["n"]),
            estado,
            str(r.get("fecha", "—")),
            Paragraph(str(r.get("observacion", "")), est["n"]),
        ])
        estilos_fila.append(
            ("TEXTCOLOR", (2, i), (2, i), colors.HexColor(ESTADO_COLOR[estado])))

    t = Table(data, colWidths=[1.0 * cm, 7.0 * cm, 2.2 * cm, 2.2 * cm, 5.4 * cm],
              repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#001e7c")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (2, 1), (2, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ALIGN", (0, 0), (0, -1), "CENTER"),
        ("ALIGN", (2, 0), (3, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.white, colors.HexColor("#f8fafc")]),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ] + estilos_fila))
    story.append(t)

    story.append(Spacer(1, 20))
    firmas = Table([[
        "_______________________________\nDocente del curso",
        "_______________________________\nCoordinación Académica",
    ]], colWidths=[8.4 * cm, 8.4 * cm])
    firmas.setStyle(TableStyle([
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#334155")),
    ]))
    story.append(firmas)

    doc.build(story, onFirstPage=_pie_pagina,
              onLaterPages=_pie_pagina)
    buf.seek(0)
    return buf.getvalue()


# ================================================================
# 9. CRUCE CON LAS NOTAS YA REGISTRADAS
# ================================================================
# Responde la pregunta que un checklist no puede responder:
# el tema se dictó, ¿pero el alumno lo entendió?

def _coincide_area(curso, area_examen):
    """Compara el curso del temario con el área con que se guardó el examen."""
    a = str(area_examen).strip().lower()
    if not a:
        return False
    if a == curso.lower():
        return True
    for alias in ALIAS_AREA.get(curso, []):
        if a == alias.lower():
            return True
    return curso.lower().startswith(a) or a.startswith(curso.lower()[:8])


def _coincide_grupo(grupo, grado_examen):
    """GRUPO AB coincide con 'GRUPO AB — CEPRE UNSAAC' y variantes."""
    g = str(grado_examen).upper().replace("—", "-")
    return grupo.upper().replace("GRUPO ", "").strip() in g


@st.cache_data(ttl=120, show_spinner=False)
def _leer_resultados():
    """Aplana todas las notas registradas a filas comparables.

    Fuente principal: resultados.json, que es donde "Registrar Notas →
    Nueva Evaluación" deja cada estudiante con sus áreas anidadas.
    Si no está en disco, se recupera del blob 'resultados_json' que el
    sistema guarda en la hoja Config.
    Fuente secundaria: la hoja Resultados (YACHAY QAWAY / exámenes por clave).
    """
    filas = []

    # --- Fuente 1: notas de Nueva Evaluación ---
    data = []
    p = Path(ARCHIVO_RESULTADOS_NOTAS)
    if p.exists():
        try:
            with open(p, "r", encoding="utf-8") as f:
                data = json.load(f)
        except Exception:
            data = []
    if not data:
        gs = _gs()
        if gs is not None:
            try:
                ws = gs._get_hoja("config")
                if ws:
                    for row in ws.get_all_values():
                        if row and row[0] == "resultados_json":
                            data = json.loads(row[1])
                            break
            except Exception:
                data = []

    for r in data or []:
        periodo = str(r.get("periodo", ""))
        titulo = str(r.get("titulo", ""))
        fecha = str(r.get("fecha", ""))
        eval_id = f"{fecha}|{periodo}|{titulo}"
        for a in r.get("areas", []) or []:
            try:
                nota = float(a.get("nota") or 0)
            except (TypeError, ValueError):
                continue
            if nota <= 0:
                continue
            filas.append({
                "eval_id": eval_id,
                "eval_titulo": titulo or periodo or "Evaluación",
                "periodo": periodo,
                "fecha": fecha,
                "grado": r.get("grado", ""),
                "area": a.get("nombre", ""),
                "nota": nota,
                "docente": r.get("docente_nombre") or r.get("docente", ""),
            })

    # --- Fuente 2: hoja Resultados (formato plano) ---
    gs = _gs()
    if gs is not None:
        try:
            for r in gs.leer_resultados():
                try:
                    nota = float(r.get("nota") or 0)
                except (TypeError, ValueError):
                    continue
                if nota <= 0:
                    continue
                filas.append({
                    "eval_id": str(r.get("eval_id", "")),
                    "eval_titulo": r.get("eval_titulo", ""),
                    "periodo": "",
                    "fecha": r.get("fecha", ""),
                    "grado": r.get("grado", ""),
                    "area": r.get("area", ""),
                    "nota": nota,
                    "docente": r.get("docente", ""),
                })
        except Exception:
            pass

    return filas


def notas_del_curso(grupo, curso, eval_ids=None):
    """Promedio y detalle de las evaluaciones de un curso en un grupo."""
    filas = []
    for r in _leer_resultados():
        if not _coincide_grupo(grupo, r.get("grado", "")):
            continue
        if not _coincide_area(curso, r.get("area", "")):
            continue
        if eval_ids and str(r.get("eval_id")) not in eval_ids:
            continue
        try:
            filas.append(float(r.get("nota") or 0))
        except (TypeError, ValueError):
            continue
    if not filas:
        return {"promedio": None, "alumnos": 0, "aprobados": 0, "pct_aprob": None}
    aprob = sum(1 for n in filas if n >= NOTA_UMBRAL)
    return {
        "promedio": round(sum(filas) / len(filas), 1),
        "alumnos": len(filas),
        "aprobados": aprob,
        "pct_aprob": round(100 * aprob / len(filas), 1),
    }


def evaluaciones_disponibles(grupo, curso):
    """Lista de exámenes ya aplicados que corresponden a este curso."""
    vistos, salida = set(), []
    for r in _leer_resultados():
        if not _coincide_grupo(grupo, r.get("grado", "")):
            continue
        if not _coincide_area(curso, r.get("area", "")):
            continue
        eid = str(r.get("eval_id", ""))
        if eid and eid not in vistos:
            vistos.add(eid)
            periodo = r.get("periodo", "")
            titulo = r.get("eval_titulo", "Sin título")
            salida.append({
                "eval_id": eid,
                "titulo": f"{periodo} · {titulo}" if periodo else titulo,
                "fecha": r.get("fecha", ""),
            })
    return sorted(salida, key=lambda x: str(x["fecha"]), reverse=True)


def cuadrante(pct_avance, promedio):
    """Cruza cobertura con resultados. Devuelve (etiqueta, color, lectura)."""
    if promedio is None:
        return ("Sin evaluar", "#94a3b8",
                "Se dictó pero todavía no hay notas que lo respalden.")
    alto_av = pct_avance >= 60
    alto_no = promedio >= NOTA_UMBRAL
    if alto_av and alto_no:
        return ("Consolidado", "#16a34a",
                "Buen avance y los alumnos responden. Mantener el ritmo.")
    if alto_av and not alto_no:
        return ("Dictado, no aprendido", "#dc2626",
                "Se avanzó rápido pero las notas no acompañan. "
                "Conviene frenar y reforzar antes de seguir.")
    if not alto_av and alto_no:
        return ("En ruta", "#2563eb",
                "Poco avance, pero lo dictado quedó claro. "
                "El problema es de tiempo, no de método.")
    return ("Crítico", "#ea580c",
            "Poco avance y notas bajas. Requiere intervención directa.")


def tabla_dictado_vs_aprendido(avance, anio, ciclo, grupo):
    filas = []
    for curso in cursos_de_grupo(grupo):
        res = resumen_curso(avance, anio, ciclo, grupo, curso)
        nt = notas_del_curso(grupo, curso)
        etiqueta, color, lectura = cuadrante(res["pct"], nt["promedio"])
        filas.append({
            "curso": curso,
            "peso": res["peso"],
            "avance": res["pct"],
            "promedio": nt["promedio"],
            "alumnos": nt["alumnos"],
            "pct_aprob": nt["pct_aprob"],
            "estado": etiqueta,
            "color": color,
            "lectura": lectura,
        })
    return filas


# ================================================================
# 10. VINCULAR TEMAS CON EVALUACIONES (para análisis por tema)
# ================================================================

def seccion_vincular_evaluaciones(avance, anio, ciclo, grupo, curso, docente):
    """Permite decir qué temas cubrió cada examen ya aplicado."""
    st.markdown("#### 🎯 ¿Qué temas evaluó cada examen?")
    st.caption("Al vincularlos, el sistema puede decirte si un tema se dictó "
               "pero no se entendió, en vez de solo si se dictó.")

    evals = evaluaciones_disponibles(grupo, curso)
    if not evals:
        st.info("Todavía no hay exámenes registrados para este curso y grupo. "
                "Aparecerán aquí en cuanto apliques uno desde Exámenes "
                "Semanales o YACHAY QAWAY.")
        return

    etiquetas = {f"{e['fecha']} — {e['titulo']}": e["eval_id"] for e in evals}
    sel = st.selectbox("Examen aplicado", list(etiquetas.keys()),
                       key=f"av_ev_sel_{grupo}_{curso}")
    eid = etiquetas[sel]

    temas = TEMARIO.get(curso, [])
    ya = []
    for i, tema in enumerate(temas, start=1):
        r = avance.get(_clave(anio, ciclo, grupo, curso, i), {})
        if eid in str(r.get("evals", "")).split(","):
            ya.append(f"{i}. {tema}")

    elegidos = st.multiselect(
        "Temas que entraron en este examen",
        [f"{i}. {t}" for i, t in enumerate(temas, start=1)],
        default=ya, key=f"av_ev_temas_{grupo}_{curso}_{eid}")

    if st.button("🔗 Vincular examen con estos temas",
                 key=f"av_ev_btn_{grupo}_{curso}_{eid}",
                 use_container_width=True):
        cambios = []
        indices = {int(x.split(".")[0]) for x in elegidos}
        for i, tema in enumerate(temas, start=1):
            clave = _clave(anio, ciclo, grupo, curso, i)
            r = dict(avance.get(clave, {}))
            actuales = [x for x in str(r.get("evals", "")).split(",") if x]
            nuevo = list(actuales)
            if i in indices and eid not in nuevo:
                nuevo.append(eid)
            elif i not in indices and eid in nuevo:
                nuevo.remove(eid)
            if nuevo != actuales:
                r.update({
                    "clave": clave, "anio": anio, "ciclo": ciclo,
                    "grupo": grupo, "curso": curso, "tema_num": i,
                    "tema": tema,
                    "estado": r.get("estado", "Pendiente"),
                    "fecha": r.get("fecha", ""),
                    "sesion": r.get("sesion", ""),
                    "docente": r.get("docente", docente),
                    "observacion": r.get("observacion", ""),
                    "evals": ",".join(nuevo),
                })
                cambios.append(r)
        if cambios:
            ok, msg = guardar_avance(cambios)
            (st.success if ok else st.warning)(msg)
            st.rerun()
        else:
            st.info("No hubo cambios que guardar.")

    # --- Resultado por tema vinculado ---
    st.markdown("##### Cómo les fue en los temas ya evaluados")
    hubo = False
    for i, tema in enumerate(temas, start=1):
        r = avance.get(_clave(anio, ciclo, grupo, curso, i), {})
        ids = [x for x in str(r.get("evals", "")).split(",") if x]
        if not ids:
            continue
        hubo = True
        nt = notas_del_curso(grupo, curso, eval_ids=set(ids))
        prom = nt["promedio"]
        if prom is None:
            continue
        color = "#16a34a" if prom >= NOTA_UMBRAL else "#dc2626"
        aviso = "" if prom >= NOTA_UMBRAL else " · conviene reforzar"
        st.markdown(
            f"<div style='padding:6px 10px;margin-bottom:4px;border-radius:6px;"
            f"background:#f8fafc;border-left:5px solid {color};font-size:.9rem;'>"
            f"<b>{i}. {tema}</b> — promedio "
            f"<b style='color:{color};'>{prom}</b> "
            f"({nt['alumnos']} notas, {nt['pct_aprob']}% sobre {NOTA_UMBRAL})"
            f"<span style='color:{color};'>{aviso}</span></div>",
            unsafe_allow_html=True)
    if not hubo:
        st.caption("Vincula al menos un examen arriba para ver esta sección.")


# ================================================================
# 11. PLANILLA IMPRESA PARA MARCAR A MANO
# ================================================================

def generar_pdf_planilla(grupo, ciclo, anio, institucion="ACADEMIA YACHAY"):
    """Hoja con todos los cursos y sus temas, con casilleros vacíos.

    Sirve para que dirección o el propio docente marque con aspa sobre
    papel durante la clase y luego lo pase al sistema, o para archivar
    la evidencia firmada.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, KeepTogether)
    from reportlab.lib import colors
    from reportlab.lib.units import cm

    est = _estilos_pdf()
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=1.3 * cm, rightMargin=1.3 * cm,
                            topMargin=1.2 * cm, bottomMargin=1.2 * cm)
    story = []

    _encabezado(story, f"{institucion}<br/>PLANILLA DE CONTROL DEL TEMARIO",
                f"{grupo} &nbsp;·&nbsp; {ciclo} {anio} &nbsp;·&nbsp; "
                f"Temario UNSAAC (Res. CU-575-2024)", est)
    story.append(Paragraph(
        "Marque con una <b>X</b> el estado de cada tema. "
        "<b>P</b> = pendiente &nbsp; <b>A</b> = en avance &nbsp; "
        "<b>C</b> = concluido &nbsp; <b>R</b> = reforzado.", est["n"]))
    story.append(Spacer(1, 8))

    for curso in cursos_de_grupo(grupo):
        temas = TEMARIO.get(curso, [])
        bloque = [Paragraph(
            f"{curso} &nbsp;<font size=8 color='#64748b'>"
            f"({int(peso_curso(grupo, curso))} preguntas en el examen · "
            f"{len(temas)} temas)</font>", est["h2"])]

        data = [["N°", "Tema", "P", "A", "C", "R", "Fecha", "Observación"]]
        for i, tema in enumerate(temas, start=1):
            data.append([str(i), Paragraph(tema, est["n"]),
                         "", "", "", "", "", ""])

        t = Table(data, colWidths=[0.9 * cm, 7.2 * cm, 0.75 * cm, 0.75 * cm,
                                   0.75 * cm, 0.75 * cm, 2.0 * cm, 5.0 * cm],
                  repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#001e7c")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ALIGN", (0, 0), (0, -1), "CENTER"),
            ("ALIGN", (2, 0), (6, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#94a3b8")),
            # Casilleros de marcado resaltados
            ("BACKGROUND", (2, 1), (5, -1), colors.HexColor("#f1f5f9")),
            ("ROWBACKGROUNDS", (0, 1), (1, -1),
             [colors.white, colors.HexColor("#fafafa")]),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        bloque.append(t)
        bloque.append(Spacer(1, 4))
        bloque.append(Paragraph(
            "Docente: _______________________________ &nbsp;&nbsp; "
            "Firma: _______________________ &nbsp;&nbsp; "
            "V°B° Dirección: _______________________", est["n"]))
        bloque.append(Spacer(1, 12))
        story.append(KeepTogether(bloque) if len(temas) <= 18
                     else bloque[0])
        if len(temas) > 18:
            for x in bloque[1:]:
                story.append(x)

    doc.build(story, onFirstPage=_pie_pagina,
              onLaterPages=_pie_pagina)
    buf.seek(0)
    return buf.getvalue()


# ================================================================
# 12. TEMARIO COMPLETO POR ÁREA (A, B, C, D)
# ================================================================

def tabla_numeracion_grupo(grupo):
    """Numeración de preguntas curso por curso, para un grupo combinado.

    No existe un examen único de 80 preguntas para "GRUPO AB": existen
    DOS exámenes de 80 preguntas (uno por área) que comparten casi todos
    los cursos. Por eso se muestra la numeración de cada área en su
    propia columna, en vez de sumar los pesos como si fuera un solo
    examen — sumarlos daría más de 80 y sería un dato falso.
    """
    areas = GRUPOS.get(grupo, [])
    por_area = {a: {x["curso"]: x for x in numeracion_preguntas(a)}
                for a in areas}
    cursos = cursos_de_grupo(grupo)
    # Ordenar por el peso máximo que tenga en cualquiera de las áreas
    cursos.sort(key=lambda c: -max(
        (por_area[a].get(c, {}).get("preguntas", 0) for a in areas),
        default=0))
    filas = []
    for c in cursos:
        fila = {"curso": c, "total_temas": len(TEMARIO.get(c, []))}
        for a in areas:
            info = por_area[a].get(c)
            fila[a] = (f"{info['desde']}–{info['hasta']}"
                       if info else "no entra")
        filas.append(fila)
    return filas


def _etiqueta_area(area_o_grupo):
    if es_grupo_combinado(area_o_grupo):
        return f"{area_o_grupo} (Áreas {' y '.join(GRUPOS[area_o_grupo])})"
    return f"ÁREA «{area_o_grupo}»"


def generar_pdf_area(area, ciclo, anio, institucion="ACADEMIA YACHAY"):
    """Temario completo de un área de postulación, curso por curso.

    Sirve para entregar al postulante y al docente: en una sola hoja
    saben exactamente qué cursos les tocan, cuánto pesa cada uno y
    todos los temas que entran en el examen.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle)
    from reportlab.lib import colors
    from reportlab.lib.units import cm

    est = _estilos_pdf()
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=1.5 * cm, rightMargin=1.5 * cm,
                            topMargin=1.2 * cm, bottomMargin=1.2 * cm)
    story = []

    pesos = _pesos_de(area)
    total_preg = int(round(sum(pesos.values())))

    _encabezado(story, f"{institucion}<br/>TEMARIO DE ADMISIÓN — {_etiqueta_area(area)}",
                f"{ciclo} {anio} &nbsp;·&nbsp; {total_preg} preguntas &nbsp;·&nbsp; "
                f"Res. CU-575-2024-UNSAAC", est)

    # Resumen de la distribución
    data = [["Curso", "Preguntas", "% del examen", "N° de temas"]]
    for curso, np_raw in sorted(pesos.items(), key=lambda x: -x[1]):
        np_ = int(round(np_raw))
        data.append([
            Paragraph(curso, est["n"]), str(np_),
            f"{round(100*np_/total_preg, 1)}%",
            str(len(TEMARIO.get(curso, []))),
        ])
    data.append(["TOTAL", str(total_preg), "100%",
                 str(sum(len(TEMARIO.get(c, [])) for c in pesos))])
    t = Table(data, colWidths=[8.5 * cm, 2.6 * cm, 3.0 * cm, 3.0 * cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#001e7c")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
        ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#e2e8f0")),
        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
        ("ALIGN", (1, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#94a3b8")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -2),
         [colors.white, colors.HexColor("#f8fafc")]),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(t)
    story.append(Spacer(1, 6))
    story.append(Paragraph(
        "Los cursos con más preguntas son los que más deciden el ingreso. "
        "Prioriza tu estudio en ese orden.", est["n"]))
    story.append(Spacer(1, 10))

    # Temario curso por curso
    for curso, np_raw in sorted(pesos.items(), key=lambda x: -x[1]):
        np_ = int(round(np_raw))
        temas = TEMARIO.get(curso, [])
        story.append(Paragraph(
            f"{curso} &nbsp;<font size=8 color='#64748b'>"
            f"({np_} preguntas · {len(temas)} temas)</font>", est["h2"]))
        filas, mitad = [], -(-len(temas) // 2)
        for i in range(mitad):
            izq = f"{i+1}. {temas[i]}"
            j = i + mitad
            der = f"{j+1}. {temas[j]}" if j < len(temas) else ""
            filas.append([Paragraph(izq, est["n"]), Paragraph(der, est["n"])])
        t = Table(filas, colWidths=[8.85 * cm, 8.85 * cm])
        t.setStyle(TableStyle([
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 1.5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 1.5),
            ("LINEBELOW", (0, 0), (-1, -1), 0.25, colors.HexColor("#e2e8f0")),
        ]))
        story.append(t)
        story.append(Spacer(1, 8))

    doc.build(story, onFirstPage=_pie_pagina,
              onLaterPages=_pie_pagina)
    buf.seek(0)
    return buf.getvalue()


# ================================================================
# 13. CRONOGRAMA Y COMUNICADO PARA DOCENTES
# ================================================================

CICLOS_PRESET = {
    "Ciclo Agosto – Diciembre": (8, 3, 12, 20),
    "Ciclo Marzo – Julio": (3, 1, 7, 31),
    "Verano CEPRU (Enero – Febrero)": (1, 5, 2, 28),
}

DIAS_SEMANA = ["Lunes", "Martes", "Miércoles", "Jueves", "Viernes",
               "Sábado", "Domingo"]


def sabados_del_ciclo(inicio, fin):
    """Todos los sábados entre dos fechas: son los días de examen semanal."""
    from datetime import timedelta
    d = inicio
    while d.weekday() != 5:          # 5 = sábado
        d += timedelta(days=1)
    salida = []
    while d <= fin:
        salida.append(d)
        d += timedelta(days=7)
    return salida


def repartir_temas(n_temas, n_semanas, semanas_repaso=2):
    """Reparte los temas entre las semanas, dejando las últimas para repaso.

    El reparto es proporcional, no secuencial: si un curso tiene 14 temas
    y hay 18 semanas de avance, las semanas sin tema nuevo quedan
    intercaladas y sirven de refuerzo, en vez de amontonar todo al
    principio y dejar seis semanas muertas al final.
    """
    if n_semanas <= 0 or n_temas <= 0:
        return []
    avance = max(n_semanas - semanas_repaso, 1)
    tramos, anterior = [], 0
    for i in range(n_semanas):
        if i >= avance:
            tramos.append((None, None))
            continue
        hasta = round((i + 1) * n_temas / avance)
        if hasta > anterior:
            tramos.append((anterior + 1, hasta))
            anterior = hasta
        else:
            tramos.append((None, None))
    return tramos


def _pesos_de(area_o_grupo):
    """Cursos y su peso, sirva el parametro un área sola (A, B, C, D) o
    un grupo combinado (GRUPO AB, GRUPO CD).

    En un grupo combinado el peso de cada curso es el promedio de las
    áreas que lo integran (usa peso_curso, ya usado en el resto del
    módulo), para que la numeración de un curso compartido no dependa
    de cuál de las dos áreas se mire primero.
    """
    if area_o_grupo in PESOS:
        return dict(PESOS[area_o_grupo])
    if area_o_grupo in GRUPOS:
        cursos = cursos_de_grupo(area_o_grupo)
        return {c: peso_curso(area_o_grupo, c) for c in cursos}
    return {}


def es_grupo_combinado(area_o_grupo):
    return area_o_grupo in GRUPOS


def numeracion_preguntas(area):
    """Rango fijo de numeración de cada curso dentro del examen de 80.

    Que cada docente entregue siempre el mismo tramo evita que dos
    profesores numeren igual y haya que renumerar todo al armar la prueba.
    Acepta un área sola (A, B, C, D) o un grupo combinado (GRUPO AB, CD).
    """
    pesos = _pesos_de(area)
    orden = sorted(pesos.items(), key=lambda x: -x[1])
    salida, n = [], 1
    for curso, cant in orden:
        cant_i = int(round(cant))
        salida.append({"curso": curso, "preguntas": cant_i,
                       "desde": n, "hasta": n + cant_i - 1})
        n += cant_i
    return salida


def cronograma_area(area, inicio, fin, semanas_repaso=2, feriados=None):
    """area puede ser una letra (A, B, C, D) o un grupo combinado
    (GRUPO AB, GRUPO CD)."""
    """Arma el cronograma completo: semanas, fechas y temas por curso.

    Salta los sábados que caen en día no lectivo, para no programar un
    examen un día en que el colegio está cerrado.
    """
    sabados, omitidos = sabados_habiles(inicio, fin, feriados)
    n_sem = len(sabados)
    cursos = numeracion_preguntas(area)
    for c in cursos:
        temas = TEMARIO.get(c["curso"], [])
        c["total_temas"] = len(temas)
        c["tramos"] = repartir_temas(len(temas), n_sem, semanas_repaso)
    return {"area": area, "inicio": inicio, "fin": fin,
            "sabados": sabados, "semanas": n_sem, "cursos": cursos,
            "omitidos": omitidos}


def generar_comunicado_docx(area, inicio, fin, ciclo, docentes=None,
                            institucion="ACADEMIA YACHAY",
                            whatsapp_secretaria="", dia_entrega="Jueves",
                            semanas_repaso=2, feriados=None):
    """Documento Word con el cronograma y las instrucciones para docentes."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    docentes = docentes or {}
    cron = cronograma_area(area, inicio, fin, semanas_repaso, feriados)
    combinado = es_grupo_combinado(area)
    doc = Document()

    for s in doc.sections:
        s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8)
        s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)

    est = doc.styles["Normal"]
    est.font.name = "Calibri"
    est.font.size = Pt(10.5)

    def p(txt, size=10.5, bold=False, align=None, space=4, color=None):
        par = doc.add_paragraph()
        run = par.add_run(txt)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        if align:
            par.alignment = align
        par.paragraph_format.space_after = Pt(space)
        return par

    AZUL = (0x00, 0x1E, 0x7C)

    p(institucion, 15, True, WD_ALIGN_PARAGRAPH.CENTER, 0, AZUL)
    p("CRONOGRAMA DE AVANCE Y ENTREGA DE PREGUNTAS", 13, True,
      WD_ALIGN_PARAGRAPH.CENTER, 0, AZUL)
    p(f"{_etiqueta_area(area)} — {ciclo}", 11, True, WD_ALIGN_PARAGRAPH.CENTER, 10)
    p(f"Del {inicio.strftime('%d/%m/%Y')} al {fin.strftime('%d/%m/%Y')} "
      f"· {cron['semanas']} semanas · Examen semanal los sábados por la tarde",
      9.5, False, WD_ALIGN_PARAGRAPH.CENTER, 14)

    # ── 1. Numeración fija ────────────────────────────────────────
    p("1. NUMERACIÓN DE PREGUNTAS POR CURSO", 11.5, True, None, 4, AZUL)
    p("Cada docente entrega siempre el mismo tramo de numeración. Respetarlo "
      "evita que dos cursos numeren igual y haya que rehacer la prueba.",
      9.5, False, None, 6)

    if combinado:
        _areas_g = GRUPOS[area]
        p(f"Este salón agrupa a postulantes de {' y '.join(_areas_g)}. "
          f"Cada área rinde su propio examen de 80 preguntas; por eso la "
          f"numeración se muestra por separado para cada una.", 9, False,
          None, 6)
        t = doc.add_table(rows=1, cols=2 + len(_areas_g) + 1)
        t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        _enc = ["Curso", "Docente responsable"] + \
               [f"Numeración en {a}" for a in _areas_g] + ["Temas"]
        for i, h in enumerate(_enc):
            c = t.rows[0].cells[i]
            c.text = h
            c.paragraphs[0].runs[0].bold = True
            c.paragraphs[0].runs[0].font.size = Pt(9)
        for fila in tabla_numeracion_grupo(area):
            f = t.add_row().cells
            f[0].text = fila["curso"]
            f[1].text = docentes.get(fila["curso"], "______________________")
            for i, a in enumerate(_areas_g):
                f[2 + i].text = fila[a]
            f[-1].text = str(fila["total_temas"])
            for cc in f:
                for pp in cc.paragraphs:
                    for r in pp.runs:
                        r.font.size = Pt(8.5)
    else:
        t = doc.add_table(rows=1, cols=5)
        t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(["Curso", "Docente responsable", "N° preguntas",
                               "Numeración", "Temas"]):
            c = t.rows[0].cells[i]
            c.text = h
            c.paragraphs[0].runs[0].bold = True
            c.paragraphs[0].runs[0].font.size = Pt(9.5)
        for c in cron["cursos"]:
            f = t.add_row().cells
            f[0].text = c["curso"]
            f[1].text = docentes.get(c["curso"], "______________________")
            f[2].text = str(c["preguntas"])
            f[3].text = f"{c['desde']} al {c['hasta']}"
            f[4].text = str(c["total_temas"])
            for cc in f:
                for pp in cc.paragraphs:
                    for r in pp.runs:
                        r.font.size = Pt(9)
    p("", 6, space=8)

    if cron.get("omitidos"):
        p("Sábados sin examen por día no lectivo:", 10, True, None, 2)
        for _s, _mot in cron["omitidos"]:
            par = doc.add_paragraph(
                f"{_s.strftime('%d/%m/%Y')} — {_mot}", style="List Bullet")
            par.paragraph_format.space_after = Pt(1)
            for r in par.runs:
                r.font.size = Pt(9)
        p("", 6, space=8)

    # ── 2. Reglas de entrega ──────────────────────────────────────
    p("2. ENTREGA SEMANAL DE PREGUNTAS", 11.5, True, None, 4, AZUL)
    reglas = [
        f"Fecha límite de entrega: {dia_entrega} de cada semana hasta las 6:00 p.m.",
        (f"Enviar al WhatsApp de Secretaría: {whatsapp_secretaria}"
         if whatsapp_secretaria else
         "Enviar al WhatsApp de Secretaría."),
        "Formato: archivo Word, una pregunta por numeral, con su clave de "
        "respuesta al final.",
        "Las preguntas deben corresponder únicamente a los temas avanzados "
        "esa semana, según el cronograma de la sección 3.",
        "Respetar el tramo de numeración asignado a su curso en la tabla anterior.",
        "El examen semanal se aplica los sábados por la tarde y evalúa todo lo "
        "avanzado durante la semana.",
        "El docente que no entregue a tiempo deja su tramo vacío y perjudica "
        "a todos los postulantes del área.",
    ]
    for r in reglas:
        par = doc.add_paragraph(r, style="List Bullet")
        par.paragraph_format.space_after = Pt(2)
        for run in par.runs:
            run.font.size = Pt(9.5)
    p("", 6, space=8)

    # ── 3. Cronograma semanal ─────────────────────────────────────
    p("3. CRONOGRAMA DE AVANCE POR SEMANA", 11.5, True, None, 4, AZUL)
    p("El rango indica los números de tema del temario oficial UNSAAC "
      "(Res. CU-575-2024) que corresponden a esa semana.", 9.5, False, None, 6)

    cursos = cron["cursos"]
    t2 = doc.add_table(rows=1, cols=2 + len(cursos))
    t2.style = "Light Grid Accent 1"
    enc = ["Sem.", "Examen"] + [c["curso"][:18] for c in cursos]
    for i, h in enumerate(enc):
        c = t2.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(8)
    for i, sab in enumerate(cron["sabados"]):
        f = t2.add_row().cells
        f[0].text = str(i + 1)
        f[1].text = sab.strftime("%d/%m")
        for j, c in enumerate(cursos):
            tramo = c["tramos"][i] if i < len(c["tramos"]) else (None, None)
            if tramo[0] is None:
                f[2 + j].text = "repaso"
            elif tramo[0] == tramo[1]:
                f[2 + j].text = f"T{tramo[0]}"
            else:
                f[2 + j].text = f"T{tramo[0]}–{tramo[1]}"
        for cc in f:
            for pp in cc.paragraphs:
                for r in pp.runs:
                    r.font.size = Pt(8)
    p("", 6, space=10)

    # ── 4. Detalle de temas por curso ─────────────────────────────
    p("4. DETALLE DE TEMAS POR CURSO", 11.5, True, None, 6, AZUL)
    for c in cursos:
        temas = TEMARIO.get(c["curso"], [])
        p(f"{c['curso']} — {docentes.get(c['curso'], 'docente por asignar')} "
          f"· preguntas {c['desde']} al {c['hasta']}", 10.5, True, None, 3)
        for i, sab in enumerate(cron["sabados"]):
            tramo = c["tramos"][i] if i < len(c["tramos"]) else (None, None)
            if tramo[0] is None:
                continue
            nombres = "; ".join(
                f"{k}. {temas[k-1]}" for k in range(tramo[0], tramo[1] + 1)
                if k <= len(temas))
            par = doc.add_paragraph()
            r1 = par.add_run(f"Semana {i+1} (examen {sab.strftime('%d/%m')}): ")
            r1.bold = True
            r1.font.size = Pt(9)
            r2 = par.add_run(nombres)
            r2.font.size = Pt(9)
            par.paragraph_format.space_after = Pt(1)
            par.paragraph_format.left_indent = Cm(0.6)
        p("", 6, space=8)

    # ── Firmas ────────────────────────────────────────────────────
    p("", 6, space=20)
    tf = doc.add_table(rows=1, cols=2)
    tf.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, txt in enumerate(["_______________________________\nCoordinación Académica",
                             "_______________________________\nDirección"]):
        cel = tf.rows[0].cells[i]
        cel.text = txt
        for pp in cel.paragraphs:
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in pp.runs:
                r.font.size = Pt(9)

    # Pie de pagina legal en todas las hojas del Word
    for sec in doc.sections:
        pie = sec.footer.paragraphs[0]
        pie.text = PIE_LEGAL
        pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in pie.runs:
            r.font.size = Pt(6.5)
            r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ================================================================
# 14. CALENDARIO DE FERIADOS Y SEMANAS NO LECTIVAS
# ================================================================

ARCHIVO_FERIADOS = "feriados_academia.json"

# Feriados nacionales fijos del Perú (mes, día). Los movibles como
# Jueves y Viernes Santo cambian de fecha cada año y se agregan a mano.
FERIADOS_FIJOS = [
    (1, 1, "Año Nuevo"),
    (5, 1, "Día del Trabajo"),
    (6, 7, "Batalla de Arica"),
    (6, 29, "San Pedro y San Pablo"),
    (7, 23, "Día de la Fuerza Aérea"),
    (7, 28, "Fiestas Patrias"),
    (7, 29, "Fiestas Patrias"),
    (8, 6, "Batalla de Junín"),
    (8, 30, "Santa Rosa de Lima"),
    (10, 8, "Combate de Angamos"),
    (11, 1, "Todos los Santos"),
    (12, 8, "Inmaculada Concepción"),
    (12, 9, "Batalla de Ayacucho"),
    (12, 25, "Navidad"),
]


def cargar_feriados():
    """Devuelve {'AAAA-MM-DD': 'motivo'} con los días no lectivos."""
    p = Path(ARCHIVO_FERIADOS)
    if p.exists():
        try:
            with open(p, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {}


def guardar_feriados(data):
    try:
        with open(ARCHIVO_FERIADOS, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        return True
    except Exception:
        return False


def feriados_nacionales(anio):
    """Los feriados fijos de un año, listos para precargar."""
    return {date(anio, m, d).isoformat(): motivo
            for m, d, motivo in FERIADOS_FIJOS}


def es_no_lectivo(dia, feriados=None):
    f = feriados if feriados is not None else cargar_feriados()
    return dia.isoformat() in f


def sabados_habiles(inicio, fin, feriados=None):
    """Sábados de examen, saltando los que caen en día no lectivo.

    Devuelve (habiles, omitidos) para poder avisar en el documento qué
    semanas se perdieron y por qué.
    """
    f = feriados if feriados is not None else cargar_feriados()
    habiles, omitidos = [], []
    for s in sabados_del_ciclo(inicio, fin):
        if s.isoformat() in f:
            omitidos.append((s, f[s.isoformat()]))
        else:
            habiles.append(s)
    return habiles, omitidos


def seccion_calendario_feriados():
    """Interfaz para marcar feriados, semanas de gestión y suspensiones."""
    st.markdown("#### 📅 Calendario de días no lectivos")
    st.caption("Los sábados marcados aquí se saltan al armar el cronograma, "
               "para no programar un examen un día que el colegio está cerrado.")

    feriados = cargar_feriados()
    anio_f = st.number_input("Año:", 2024, 2040, date.today().year,
                             key="fer_anio")

    cf1, cf2 = st.columns(2)
    with cf1:
        if st.button("➕ Precargar feriados nacionales del Perú",
                     use_container_width=True, key="fer_precarga"):
            feriados.update(feriados_nacionales(int(anio_f)))
            guardar_feriados(feriados)
            st.success("Feriados nacionales agregados. "
                       "Semana Santa y feriados regionales se añaden a mano.")
            st.rerun()
    with cf2:
        st.caption("Faltan por agregar a mano: Jueves y Viernes Santo "
                   "(cambian cada año), aniversario del distrito, "
                   "semana de gestión y suspensiones de último momento.")

    st.markdown("##### Agregar un día o un rango")
    ca1, ca2, ca3 = st.columns([2, 2, 3])
    with ca1:
        _d1 = st.date_input("Desde:", value=date.today(), key="fer_d1")
    with ca2:
        _d2 = st.date_input("Hasta:", value=date.today(), key="fer_d2")
    with ca3:
        _motivo = st.text_input("Motivo:", key="fer_motivo",
                                placeholder="Semana de gestión / Feriado / Suspensión")

    if st.button("Agregar al calendario", type="primary",
                 use_container_width=True, key="fer_add"):
        if _d2 < _d1:
            st.error("La fecha final no puede ser anterior a la inicial.")
        elif not _motivo.strip():
            st.error("Escribe el motivo para que quede constancia.")
        else:
            from datetime import timedelta
            d, n = _d1, 0
            while d <= _d2:
                feriados[d.isoformat()] = _motivo.strip()
                d += timedelta(days=1)
                n += 1
            guardar_feriados(feriados)
            st.success(f"{n} día(s) marcados como no lectivos.")
            st.rerun()

    if feriados:
        st.markdown("##### Días registrados")
        filas = []
        for k in sorted(feriados):
            try:
                dd = date.fromisoformat(k)
            except ValueError:
                continue
            filas.append({
                "Fecha": dd.strftime("%d/%m/%Y"),
                "Día": DIAS_SEMANA[dd.weekday()],
                "Motivo": feriados[k],
                "Cae sábado": "⚠️ sí" if dd.weekday() == 5 else "",
            })
        st.dataframe(pd.DataFrame(filas), use_container_width=True,
                     hide_index=True, height=260)

        _quitar = st.selectbox("Quitar un día:", ["—"] + sorted(feriados),
                               key="fer_del_sel")
        if _quitar != "—" and st.button(f"🗑️ Quitar {_quitar}",
                                        key="fer_del_btn"):
            feriados.pop(_quitar, None)
            guardar_feriados(feriados)
            st.rerun()
    else:
        st.info("Todavía no hay días marcados.")


# ================================================================
# 15. SIMULACROS Y EXÁMENES DE BECAS
# ================================================================
# Registro independiente de la matrícula: los postulantes externos no
# son alumnos del colegio y no deben ensuciar la base de estudiantes.
# Cada postulante declara su área, que es el dato que la matrícula no
# tiene y sin el cual un ranking de 80 preguntas no significa nada.

ARCHIVO_SIMULACROS = "simulacros.json"


def cargar_simulacros():
    p = Path(ARCHIVO_SIMULACROS)
    if p.exists():
        try:
            with open(p, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {}


def guardar_simulacros(data):
    try:
        with open(ARCHIVO_SIMULACROS, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
    except Exception:
        pass
    gs = _gs()
    if gs is None:
        return False
    try:
        ws = gs._get_hoja("config")
        if ws is None:
            return False
        blob = json.dumps(data, ensure_ascii=False)
        for i, row in enumerate(ws.get_all_values()):
            if row and row[0] == "simulacros_json":
                ws.update_cell(i + 1, 2, blob)
                return True
        ws.append_row(["simulacros_json", blob])
        return True
    except Exception:
        return False


def restaurar_simulacros_nube():
    """Si el disco se reinició (Streamlit Cloud lo hace), recupera de la nube."""
    if Path(ARCHIVO_SIMULACROS).exists():
        return cargar_simulacros()
    gs = _gs()
    if gs is None:
        return {}
    try:
        ws = gs._get_hoja("config")
        if ws is None:
            return {}
        for row in ws.get_all_values():
            if row and row[0] == "simulacros_json":
                data = json.loads(row[1])
                with open(ARCHIVO_SIMULACROS, "w", encoding="utf-8") as f:
                    json.dump(data, f, indent=2, ensure_ascii=False)
                return data
    except Exception:
        pass
    return {}


def puntaje_postulante(post, area, descuento=0.0):
    """Puntaje sobre 80 y equivalencia vigesimal.

    descuento: puntos restados por cada respuesta incorrecta. Se deja
    configurable porque las reglas de calificación cambian entre
    procesos y una constante escondida en el código envejece mal.
    """
    pesos = PESOS.get(area, {})
    total = sum(pesos.values()) or 80
    aciertos = 0
    errores = 0
    for curso, cant in pesos.items():
        a = post.get("aciertos", {}).get(curso)
        if a is None:
            continue
        a = max(0, min(int(a), cant))
        aciertos += a
        errores += (cant - a)
    bruto = aciertos - descuento * errores
    bruto = max(bruto, 0)
    return {
        "aciertos": aciertos,
        "errores": errores,
        "total": total,
        "puntaje": round(bruto, 2),
        "pct": round(100 * bruto / total, 1) if total else 0,
        "vigesimal": round(20 * bruto / total, 1) if total else 0,
    }


def ranking_simulacro(sim, ambito="general"):
    """Devuelve la tabla ordenada por puntaje.

    ambito: 'general' (todos juntos), 'A'/'B'/'C'/'D' (un área),
    'GRUPO AB'/'GRUPO CD' (las dos áreas del salón).
    """
    desc = float(sim.get("descuento", 0) or 0)
    filas = []
    for dni, post in (sim.get("postulantes") or {}).items():
        area = post.get("area", "")
        if ambito in ("A", "B", "C", "D") and area != ambito:
            continue
        if ambito in GRUPOS and area not in GRUPOS[ambito]:
            continue
        if not area:
            continue
        r = puntaje_postulante(post, area, desc)
        # Sin ninguna respuesta cargada no entra al ranking: un cero por
        # no haber sido calificado todavía falsearía los puestos.
        if not post.get("aciertos"):
            continue
        filas.append({
            "DNI": dni,
            "Postulante": post.get("nombre", ""),
            "Área": area,
            "Colegio": post.get("colegio", ""),
            "Aciertos": r["aciertos"],
            "Puntaje": r["puntaje"],
            "% ": r["pct"],
            "Nota /20": r["vigesimal"],
        })
    filas.sort(key=lambda x: (-x["Puntaje"], x["Postulante"]))
    for i, f in enumerate(filas, start=1):
        f["Puesto"] = i
    return filas


def generar_pdf_ranking_simulacro(sim, ambito, filas,
                                  institucion="ACADEMIA YACHAY",
                                  vacantes=0):
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle)
    from reportlab.lib import colors
    from reportlab.lib.units import cm

    est = _estilos_pdf()
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=1.4 * cm, rightMargin=1.4 * cm,
                            topMargin=1.3 * cm, bottomMargin=1.6 * cm)
    story = []

    etiqueta = ("Resultado general" if ambito == "general"
                else _etiqueta_area(ambito))
    try:
        _f_txt = date.fromisoformat(sim.get("fecha", "")).strftime("%d/%m/%Y")
    except ValueError:
        _f_txt = str(sim.get("fecha", ""))
    _encabezado(story, f"{institucion}<br/>{sim.get('nombre', 'SIMULACRO').upper()}",
                f"{etiqueta} &nbsp;·&nbsp; Aplicado el {_f_txt} "
                f"&nbsp;·&nbsp; {len(filas)} postulantes", est)

    desc = float(sim.get("descuento", 0) or 0)
    story.append(Paragraph(
        "Calificación: 1 punto por respuesta correcta" +
        (f", −{desc} por incorrecta." if desc else ", sin descuento por "
         "respuesta incorrecta."), est["n"]))
    story.append(Spacer(1, 8))

    data = [["Pto.", "Postulante", "DNI", "Área", "Colegio",
             "Aciertos", "Puntaje", "Nota /20"]]
    estilos = []
    for i, f in enumerate(filas, start=1):
        data.append([
            str(f["Puesto"]),
            Paragraph(str(f["Postulante"]), est["n"]),
            str(f["DNI"]), str(f["Área"]),
            Paragraph(str(f["Colegio"])[:28], est["n"]),
            f"{f['Aciertos']}/{sum(PESOS.get(f['Área'], {}).values()) or 80}",
            str(f["Puntaje"]), str(f["Nota /20"]),
        ])
        if vacantes and i <= vacantes:
            estilos.append(("BACKGROUND", (0, i), (-1, i),
                            colors.HexColor("#dcfce7")))
        if i <= 3:
            estilos.append(("FONTNAME", (0, i), (-1, i), "Helvetica-Bold"))

    t = Table(data, colWidths=[1.0*cm, 5.4*cm, 2.0*cm, 1.2*cm, 4.2*cm,
                               1.7*cm, 1.6*cm, 1.6*cm], repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#001e7c")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ALIGN", (0, 0), (0, -1), "CENTER"),
        ("ALIGN", (2, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.white, colors.HexColor("#f8fafc")]),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ] + estilos))
    story.append(t)

    if vacantes:
        story.append(Spacer(1, 6))
        story.append(Paragraph(
            f"Las filas resaltadas corresponden a las {vacantes} primeras "
            f"posiciones, alcanzadas por el beneficio ofrecido.", est["n"]))

    story.append(Spacer(1, 22))
    firmas = Table([[
        "_______________________________\nCoordinación Académica",
        "_______________________________\nDirección",
    ]], colWidths=[8.6 * cm, 8.6 * cm])
    firmas.setStyle(TableStyle([
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#334155")),
    ]))
    story.append(firmas)

    doc.build(story, onFirstPage=_pie_pagina, onLaterPages=_pie_pagina)
    buf.seek(0)
    return buf.getvalue()


def _etiqueta_simulacro(sim):
    """Etiqueta legible del simulacro, con fecha en formato peruano."""
    try:
        f = date.fromisoformat(sim.get("fecha", "")).strftime("%d/%m/%Y")
    except ValueError:
        f = str(sim.get("fecha", ""))
    return f"{f} · {sim.get('nombre', '(sin nombre)')}"


def tab_simulacro_becas(config=None):
    """Módulo completo: crear simulacro, registrar postulantes, calificar
    y publicar el ranking. Sirve tanto para exámenes de beca con público
    externo como para simulacros internos del ciclo."""
    st.subheader("🏆 Simulacros y Examen de Becas")
    st.caption("Registro independiente de la matrícula: los postulantes "
               "externos no se mezclan con los alumnos del colegio.")

    sims = restaurar_simulacros_nube()

    # ── Selección / creación ──────────────────────────────────────
    col_s1, col_s2 = st.columns([3, 2])
    with col_s1:
        opciones = ["➕ Crear nuevo"] + [
            _etiqueta_simulacro(v)
            for k, v in sorted(sims.items(),
                               key=lambda x: x[1].get("fecha", ""),
                               reverse=True)]
        sel = st.selectbox("Simulacro:", opciones, key="sim_sel")
    with col_s2:
        st.metric("Simulacros registrados", len(sims))

    if sel == "➕ Crear nuevo":
        with st.form("sim_nuevo"):
            c1, c2 = st.columns(2)
            with c1:
                nombre = st.text_input("Nombre del simulacro:",
                                       placeholder="Examen de Becas 2026 – I")
                fecha_s = st.date_input("Fecha de aplicación:", value=date.today())
            with c2:
                desc = st.number_input(
                    "Descuento por respuesta incorrecta:",
                    0.0, 1.0, 0.0, 0.25,
                    help="0 = sin penalidad. Confirma la regla del proceso "
                         "de admisión vigente antes de cambiarlo.")
                vac = st.number_input("Vacantes / becas ofrecidas:", 0, 200, 0)
            if st.form_submit_button("Crear simulacro", type="primary",
                                     use_container_width=True):
                if not nombre.strip():
                    st.error("Ponle un nombre para poder identificarlo después.")
                else:
                    sid = f"sim_{fecha_s.isoformat()}_{len(sims)+1}"
                    sims[sid] = {
                        "id": sid, "nombre": nombre.strip(),
                        "fecha": fecha_s.isoformat(),
                        "descuento": float(desc), "vacantes": int(vac),
                        "postulantes": {},
                    }
                    guardar_simulacros(sims)
                    st.success("Simulacro creado. Selecciónalo arriba para "
                               "empezar a registrar postulantes.")
                    st.rerun()
        return

    # Simulacro activo
    sid = None
    for k, v in sims.items():
        if _etiqueta_simulacro(v) == sel:
            sid = k
            break
    if sid is None:
        st.warning("No se encontró el simulacro.")
        return
    sim = sims[sid]
    sim.setdefault("postulantes", {})

    t_reg, t_cal, t_rank = st.tabs(
        ["👤 Postulantes", "✍️ Calificar", "🏅 Ranking"])

    # ── Registro de postulantes ───────────────────────────────────
    with t_reg:
        st.markdown("##### Registrar un postulante")
        with st.form("sim_add_post", clear_on_submit=True):
            r1, r2 = st.columns([3, 2])
            with r1:
                p_nom = st.text_input("Apellidos y Nombres:")
                p_col = st.text_input("Colegio de procedencia:")
            with r2:
                p_dni = st.text_input("DNI:", max_chars=12)
                p_area = st.selectbox("Área a la que postula:",
                                      ["A", "B", "C", "D"])
                p_cel = st.text_input("Celular (opcional):", max_chars=15)
            if st.form_submit_button("➕ Agregar postulante", type="primary",
                                     use_container_width=True):
                dni_l = p_dni.strip()
                if not p_nom.strip():
                    st.error("Falta el nombre.")
                elif not dni_l:
                    st.error("El DNI es obligatorio: es lo que evita "
                             "registrar dos veces a la misma persona.")
                elif dni_l in sim["postulantes"]:
                    st.error(f"El DNI {dni_l} ya está registrado en este "
                             f"simulacro ({sim['postulantes'][dni_l]['nombre']}).")
                else:
                    sim["postulantes"][dni_l] = {
                        "nombre": p_nom.strip().upper(),
                        "dni": dni_l, "area": p_area,
                        "colegio": p_col.strip(), "celular": p_cel.strip(),
                        "aciertos": {},
                    }
                    guardar_simulacros(sims)
                    st.success(f"{p_nom.strip().upper()} registrado en el Área {p_area}.")
                    st.rerun()

        if sim["postulantes"]:
            st.markdown("##### Postulantes registrados")
            df_p = pd.DataFrame([
                {"DNI": d, "Postulante": p.get("nombre", ""),
                 "Área": p.get("area", ""), "Colegio": p.get("colegio", ""),
                 "Celular": p.get("celular", ""),
                 "Calificado": "sí" if p.get("aciertos") else "no"}
                for d, p in sim["postulantes"].items()])
            st.dataframe(df_p.sort_values("Postulante"),
                         use_container_width=True, hide_index=True)

            resumen = df_p["Área"].value_counts().to_dict()
            st.caption("Por área: " + " · ".join(
                f"{a}: {n}" for a, n in sorted(resumen.items())))

            quitar = st.selectbox("Quitar postulante:", ["—"] +
                                  sorted(sim["postulantes"].keys()),
                                  key="sim_del")
            if quitar != "—" and st.button(f"🗑️ Quitar {quitar}", key="sim_del_b"):
                sim["postulantes"].pop(quitar, None)
                guardar_simulacros(sims)
                st.rerun()
        else:
            st.info("Todavía no hay postulantes registrados.")

    # ── Calificación ──────────────────────────────────────────────
    with t_cal:
        if not sim["postulantes"]:
            st.info("Registra postulantes antes de calificar.")
        else:
            area_cal = st.selectbox("Calificar el área:", ["A", "B", "C", "D"],
                                    key="sim_area_cal")
            cursos_a = list(PESOS.get(area_cal, {}).keys())
            los_de_area = {d: p for d, p in sim["postulantes"].items()
                           if p.get("area") == area_cal}
            if not los_de_area:
                st.warning(f"No hay postulantes registrados en el Área {area_cal}.")
            else:
                st.caption("Escribe cuántas respuestas acertó en cada curso. "
                           "El máximo de cada columna es el número de "
                           "preguntas que ese curso tiene en el examen.")
                filas = []
                for d, p in sorted(los_de_area.items(),
                                   key=lambda x: x[1].get("nombre", "")):
                    fila = {"DNI": d, "Postulante": p.get("nombre", "")}
                    for c in cursos_a:
                        fila[f"{c} (/{PESOS[area_cal][c]})"] = \
                            p.get("aciertos", {}).get(c, 0)
                    filas.append(fila)
                df_cal = pd.DataFrame(filas)

                edit = st.data_editor(
                    df_cal, use_container_width=True, hide_index=True,
                    disabled=["DNI", "Postulante"], key=f"sim_ed_{area_cal}",
                    column_config={
                        f"{c} (/{PESOS[area_cal][c]})": st.column_config.NumberColumn(
                            min_value=0, max_value=PESOS[area_cal][c], step=1)
                        for c in cursos_a})

                if st.button("💾 GUARDAR CALIFICACIÓN", type="primary",
                             use_container_width=True, key="sim_save_cal"):
                    for _, row in edit.iterrows():
                        d = str(row["DNI"])
                        if d not in sim["postulantes"]:
                            continue
                        ac = {}
                        for c in cursos_a:
                            v = row.get(f"{c} (/{PESOS[area_cal][c]})", 0)
                            try:
                                ac[c] = int(v)
                            except (TypeError, ValueError):
                                ac[c] = 0
                        sim["postulantes"][d]["aciertos"] = ac
                    guardar_simulacros(sims)
                    st.success("Calificación guardada.")
                    st.rerun()

    # ── Ranking ───────────────────────────────────────────────────
    with t_rank:
        ambito = st.selectbox(
            "Ver ranking de:",
            ["general", "GRUPO AB", "GRUPO CD", "A", "B", "C", "D"],
            key="sim_amb",
            help="'general' junta todas las áreas. Ten en cuenta que "
                 "compararlas entre sí solo es válido porque todas rinden "
                 "80 preguntas.")
        vacantes = int(sim.get("vacantes", 0) or 0)
        filas = ranking_simulacro(sim, ambito)

        if not filas:
            st.info("Todavía no hay postulantes calificados en este ámbito.")
        else:
            m1, m2, m3, m4 = st.columns(4)
            puntajes = [f["Puntaje"] for f in filas]
            m1.metric("Postulantes", len(filas))
            m2.metric("Puntaje más alto", max(puntajes))
            m3.metric("Promedio", round(sum(puntajes) / len(puntajes), 1))
            m4.metric("Puntaje más bajo", min(puntajes))

            df_r = pd.DataFrame(filas)[
                ["Puesto", "Postulante", "DNI", "Área", "Colegio",
                 "Aciertos", "Puntaje", "Nota /20"]]
            st.dataframe(df_r, use_container_width=True, hide_index=True,
                         height=420)

            if vacantes:
                st.success(f"Alcanzan el beneficio los primeros {vacantes} "
                           f"puestos. Corte: "
                           f"{filas[min(vacantes, len(filas))-1]['Puntaje']} puntos.")

            d1, d2 = st.columns(2)
            with d1:
                try:
                    pdf = generar_pdf_ranking_simulacro(
                        sim, ambito, filas, vacantes=vacantes)
                    st.download_button(
                        "📄 Descargar ranking en PDF", data=pdf,
                        file_name=(f"ranking_{sim.get('nombre','simulacro')[:24]}"
                                   f"_{ambito.replace(' ', '')}.pdf")
                        .replace(" ", "_"),
                        mime="application/pdf", type="primary",
                        use_container_width=True, key="sim_pdf")
                except Exception as e:
                    st.error(f"No se pudo generar el PDF: {e}")
            with d2:
                buf = io.BytesIO()
                with pd.ExcelWriter(buf, engine="openpyxl") as w:
                    df_r.to_excel(w, sheet_name="Ranking", index=False)
                st.download_button(
                    "📥 Descargar en Excel", data=buf.getvalue(),
                    file_name=f"ranking_{ambito.replace(' ', '')}.xlsx",
                    mime=("application/vnd.openxmlformats-officedocument."
                          "spreadsheetml.sheet"),
                    use_container_width=True, key="sim_xlsx")

# ================================================================
# SISTEMA YACHAY PRO v4.0 — VERSIÓN DEFINITIVA FINAL COMPLETA
# ================================================================
# Módulos: Matrícula (Alumnos + Docentes), Documentos PDF (6 tipos),
#          Carnets (individual/lote PDF 8 por hoja fotocheck),
#          Asistencia QR (Alumnos + Docentes),
#          Sistema de Calificación YACHAY (ZipGrade) — RANKING POR DOCENTE,
#          Registro Auxiliar (3 Cursos × 4 Competencias × 3 Desempeños),
#          Registro Asistencia (sin sáb/dom, sin feriados + pie feriados),
#          Gestión de Usuarios dinámicos desde Admin,
#          Protección: solo Admin puede borrar datos,
#          Links SIAGIE y Google Institucional
# ================================================================

import streamlit as st
import pandas as pd
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import Paragraph, Table, TableStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
from reportlab.lib import colors
from reportlab.lib.units import mm, cm
import qrcode
import os
import io
import textwrap
import zipfile
import time
import json
import urllib.parse
import numpy as np
import calendar
import hashlib
from datetime import datetime, timedelta, timezone, date
from PIL import Image, ImageDraw, ImageFont
from pathlib import Path

# Google Sheets sync
try:
    from google_sync import GoogleSync, get_google_sync
    GOOGLE_SYNC_DISPONIBLE = True
except ImportError:
    GOOGLE_SYNC_DISPONIBLE = False

import base64  # Para Aula Virtual

# python-docx para leer archivos Word
try:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

st.set_page_config(page_title="SISTEMA YACHAY PRO", page_icon="🎓", layout="wide")


# Estilos CSS mejorados con colores vibrantes
st.markdown("""
<style>
    /* Fondo principal con gradiente */
    .main {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    }
    
    /* Sidebar con diseño moderno */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #1e3c72 0%, #2a5298 100%);
    }
    
    /* ============================================================
       BOTONES BASE — forzado máximo
       ============================================================ */

    /* Capturar TODO con máxima especificidad */
    html body div.stApp button,
    html body div.stApp [role="button"],
    html body .stButton > button,
    html body div[data-testid="stButton"] > button,
    html body div[data-testid="column"] button,
    html body div[data-testid="stBaseButton-primary"] > button,
    html body div[data-testid="stBaseButton-secondary"] > button {
        background-color: #2563eb !important;
        background: #2563eb !important;
        color: white !important;
        border: none !important;
        border-radius: 8px !important;
        font-weight: bold !important;
        opacity: 1 !important;
        box-shadow: none !important;
        transition: background-color 0.1s !important;
        transform: none !important;
        -webkit-text-fill-color: white !important;
    }

    html body div.stApp button:hover,
    html body .stButton > button:hover,
    html body div[data-testid="stBaseButton-secondary"] > button:hover,
    html body div[data-testid="stBaseButton-primary"] > button:hover {
        background-color: #1d4ed8 !important;
        background: #1d4ed8 !important;
        color: white !important;
        -webkit-text-fill-color: white !important;
        transform: none !important;
        box-shadow: none !important;
    }

    /* Tabs — morado */
    html body div[data-testid="stTabs"] button,
    html body button[data-baseweb="tab"],
    html body [role="tab"] {
        background-color: #7c3aed !important;
        background: #7c3aed !important;
        color: white !important;
        -webkit-text-fill-color: white !important;
        border: none !important;
        font-weight: 600 !important;
        opacity: 1 !important;
    }
    html body button[data-baseweb="tab"][aria-selected="true"],
    html body [role="tab"][aria-selected="true"] {
        background-color: #4c1d95 !important;
        background: #4c1d95 !important;
        border-bottom: 3px solid #f59e0b !important;
    }

    /* Sidebar — azul oscuro */
    html body div[data-testid="stSidebar"] button,
    html body div[data-testid="stSidebar"] .stButton > button,
    html body div[data-testid="stSidebar"] div[data-testid="stBaseButton-secondary"] > button,
    html body div[data-testid="stSidebar"] div[data-testid="stBaseButton-primary"] > button {
        background-color: #dc2626 !important;
        background: #dc2626 !important;
        color: white !important;
        -webkit-text-fill-color: white !important;
        border: none !important;
        font-weight: bold !important;
    }

    /* Barra herramientas Streamlit — no tocar */
    [data-testid="stToolbar"] button,
    header button,
    header [role="button"] {
        background: transparent !important;
        background-color: transparent !important;
        border: none !important;
        box-shadow: none !important;
        color: inherit !important;
        -webkit-text-fill-color: inherit !important;
    }

    /* Barra herramientas Streamlit — no tocar */
    [data-testid="stToolbar"] button,
    [data-testid="stToolbar"] [role="button"],
    header button,
    header [role="button"] {
        background: transparent !important;
        background-color: transparent !important;
        border: none !important;
        box-shadow: none !important;
        color: inherit !important;
        font-weight: inherit !important;
        opacity: inherit !important;
        transform: none !important;
    }
    
    
    
    /* Métricas con colores */
    [data-testid="stMetricValue"] {
        font-size: 28px;
        font-weight: bold;
        color: #0072ff;
    }
    
    /* Tablas con diseño moderno */
    .dataframe {
        border-radius: 10px;
        overflow: hidden;
    }
    
    /* Headers */
    h1, h2, h3 {
        color: #1e3c72;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }
    
    /* Cards con sombra */
    .css-1r6slb0 {
        background: white;
        border-radius: 15px;
        padding: 20px;
        box-shadow: 0 8px 16px rgba(0,0,0,0.1);
        margin-bottom: 20px;
    }
    
    /* Success messages */
    .stSuccess {
        background-color: #d4edda;
        border-left: 5px solid #28a745;
        border-radius: 8px;
        padding: 15px;
    }
    
    /* Warning messages */
    .stWarning {
        background-color: #fff3cd;
        border-left: 5px solid #ffc107;
        border-radius: 8px;
        padding: 15px;
    }
    
    /* Error messages */
    .stError {
        background-color: #f8d7da;
        border-left: 5px solid #dc3545;
        border-radius: 8px;
        padding: 15px;
    }
    
    /* Input fields */
    .stTextInput>div>div>input {
        border-radius: 8px;
        border: 2px solid #e0e0e0;
        padding: 10px;
        transition: border-color 0.3s ease;
    }
    
    .stTextInput>div>div>input:focus {
        border-color: #0072ff;
        box-shadow: 0 0 0 3px rgba(0,114,255,0.1);
    }
    
    /* Selectbox */
    .stSelectbox>div>div {
        border-radius: 8px;
    }
    
    /* Tabs */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    
    .stTabs [data-baseweb="tab"] {
        border-radius: 8px 8px 0 0;
        padding: 10px 20px;
        background-color: #f0f2f6;
        font-weight: 600;
    }
    
    .stTabs [aria-selected="true"] {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        color: white;
    }
</style>
""", unsafe_allow_html=True)

# JS con MutationObserver — solo actúa cuando Streamlit cambia el DOM, sin loop
st.markdown("""
<script>
(function() {
    var css = `
        div[data-testid="stBaseButton-secondary"] button,
        div[data-testid="stBaseButton-primary"] button {
            background-color: #2563eb !important;
            background: #2563eb !important;
            color: white !important;
            -webkit-text-fill-color: white !important;
            border: none !important;
            border-radius: 8px !important;
            font-weight: bold !important;
            opacity: 1 !important;
        }
        div[data-testid="stSidebar"] div[data-testid="stBaseButton-secondary"] button,
        div[data-testid="stSidebar"] div[data-testid="stBaseButton-primary"] button {
            background-color: #dc2626 !important;
            background: #dc2626 !important;
            color: white !important;
            -webkit-text-fill-color: white !important;
        }
        div[data-testid="stTabs"] button[data-baseweb="tab"] {
            background-color: #7c3aed !important;
            background: #7c3aed !important;
            color: white !important;
            -webkit-text-fill-color: white !important;
            border: none !important;
        }
        div[data-testid="stTabs"] button[data-baseweb="tab"][aria-selected="true"] {
            background-color: #4c1d95 !important;
            background: #4c1d95 !important;
        }
    `;

    function inyectar(doc) {
        if (!doc) return;
        var id = 'yachay-btn-override';
        var existing = doc.getElementById(id);
        if (existing) return; // Ya inyectado, no hacer nada
        var style = doc.createElement('style');
        style.id = id;
        style.textContent = css;
        (doc.head || doc.body).appendChild(style);
    }

    // Inyectar una vez al cargar
    inyectar(document);
    try { inyectar(window.parent.document); } catch(e) {}

    // MutationObserver: solo actúa si Streamlit agrega/quita nodos
    var observer = new MutationObserver(function(mutations) {
        for (var m of mutations) {
            if (m.addedNodes.length > 0) {
                // Streamlit re-renderizó algo — re-inyectar si fue removido
                inyectar(document);
                try { inyectar(window.parent.document); } catch(e) {}
                break;
            }
        }
    });

    observer.observe(document.body, { childList: true, subtree: true });
})();
</script>
""", unsafe_allow_html=True)


# ================================================================
# INICIALIZAR GOOGLE SHEETS
# ================================================================
def _gs():
    """Obtener instancia de Google Sync (o None si no está disponible)"""
    if not GOOGLE_SYNC_DISPONIBLE:
        return None
    try:
        gs = get_google_sync()
        return gs if gs.conectado else None
    except Exception:
        return None

# ================================================================
# ZONA HORARIA PERÚ (UTC-5)
# ================================================================

PERU_TZ = timezone(timedelta(hours=-5))


def hora_peru():
    return datetime.now(PERU_TZ)


def hora_peru_str():
    return hora_peru().strftime('%H:%M:%S')


def fecha_peru_str():
    return hora_peru().strftime('%Y-%m-%d')


# ================================================================
# FUNCIÓN PARA REDUCIR PESO DE PDFs
# ================================================================

def comprimir_imagen_para_pdf(imagen_path_o_bytes, max_width=800, calidad=70):
    """Comprime imagen para reducir peso en PDFs (14MB → <2MB)"""
    try:
        # Cargar imagen
        if isinstance(imagen_path_o_bytes, (str, Path)):
            img = Image.open(imagen_path_o_bytes)
        else:
            img = Image.open(io.BytesIO(imagen_path_o_bytes))
        
        # Convertir a RGB si es necesario
        if img.mode in ('RGBA', 'LA', 'P'):
            # Crear fondo blanco para transparencias
            background = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'RGBA' or img.mode == 'LA':
                background.paste(img, mask=img.split()[-1])
                img = background
            else:
                img = img.convert('RGB')
        
        # Redimensionar si es muy grande
        if img.width > max_width:
            ratio = max_width / img.width
            nuevo_alto = int(img.height * ratio)
            img = img.resize((max_width, nuevo_alto), Image.LANCZOS)
        
        # Guardar con compresión
        output = io.BytesIO()
        img.save(output, format='JPEG', quality=calidad, optimize=True)
        output.seek(0)
        return output
    except Exception:
        # Si falla, retornar original
        if isinstance(imagen_path_o_bytes, bytes):
            return io.BytesIO(imagen_path_o_bytes)
        elif isinstance(imagen_path_o_bytes, (str, Path)):
            with open(imagen_path_o_bytes, 'rb') as f:
                return io.BytesIO(f.read())
        else:
            return imagen_path_o_bytes


# ================================================================
# FERIADOS OFICIALES DE PERÚ
# ================================================================

FERIADOS_PERU = {
    (1, 1): "Año Nuevo",
    (5, 1): "Día del Trabajo",
    (6, 7): "Batalla de Arica",
    (6, 29): "San Pedro y San Pablo",
    (7, 23): "Fuerza Aérea del Perú",
    (7, 28): "Fiestas Patrias",
    (7, 29): "Fiestas Patrias",
    (8, 6): "Batalla de Junín",
    (8, 30): "Santa Rosa de Lima",
    (10, 8): "Combate de Angamos",
    (11, 1): "Día de Todos los Santos",
    (12, 8): "Inmaculada Concepción",
    (12, 9): "Batalla de Ayacucho",
    (12, 25): "Navidad",
}


def dias_habiles_mes(anio, mes):
    dias = []
    _, ndays = calendar.monthrange(anio, mes)
    for d in range(1, ndays + 1):
        dt = date(anio, mes, d)
        if dt.weekday() < 5 and (mes, d) not in FERIADOS_PERU:
            dias.append(d)
    return dias


def feriados_del_mes(mes):
    resultado = []
    for (m, d), nombre in FERIADOS_PERU.items():
        if m == mes:
            resultado.append(f"{d} - {nombre}")
    return resultado


# ================================================================
# LINKS INSTITUCIONALES
# ================================================================

LINK_SIAGIE = "https://sistemas10.minedu.gob.pe/siagie3/"
LINK_GOOGLE = ("https://accounts.google.com/v3/signin/identifier?"
               "continue=https%3A%2F%2Fmail.google.com%2Fmail%2F"
               "&hd=ieyachay.org&osid=1&sacu=1&service=mail"
               "&flowName=GlifWebSignIn&flowEntry=AddSession"
               "&dsh=S386112432%3A1698624419248117&theme=glif")

# ================================================================
# SISTEMA DE USUARIOS — DINÁMICO (archivo JSON)
# ================================================================

ARCHIVO_USUARIOS = "usuarios.json"

USUARIOS_DEFAULT = {
    "administrador": {
        "password": "306020",
        "rol": "admin",
        "label": "Administrador",
        "docente_info": None
    },
}


def cargar_usuarios():
    """Carga usuarios combinando Google Sheets + local + defaults"""
    gs = _gs()
    usuarios_final = {}
    
    # 1. Empezar con defaults
    for uname, datos in USUARIOS_DEFAULT.items():
        usuarios_final[uname] = datos.copy()
    
    # 2. Cargar de Google Sheets
    if gs:
        usuarios_gs = gs.leer_usuarios()
        if usuarios_gs:
            for uname, datos in usuarios_gs.items():
                # SIEMPRE convertir password a string (GS convierte números)
                datos['password'] = str(datos.get('password', '')).strip()
                
                # Reconstruir docente_info si tiene grado
                if 'docente_info' not in datos and datos.get('grado'):
                    nombre_gs = datos.get('nombre', '').strip()
                    label_gs = datos.get('label', '').strip()
                    nombre_final = nombre_gs if (nombre_gs and nombre_gs != uname) else (
                        label_gs if (label_gs and label_gs != uname) else uname
                    )
                    datos['docente_info'] = {
                        'label': nombre_final,
                        'grado': datos.get('grado', ''),
                        'nivel': datos.get('nivel', ''),
                    }
                
                # Reconstruir label si falta o es el username
                nombre_gs = datos.get('nombre', '').strip()
                label_gs = datos.get('label', '').strip()
                if nombre_gs and ' ' in nombre_gs and len(nombre_gs) > 5:
                    datos['label'] = nombre_gs
                elif not label_gs or label_gs == uname.replace('.', ' ').title():
                    datos['label'] = nombre_gs if nombre_gs else uname.replace('.', ' ').title()
                
                if uname in USUARIOS_DEFAULT:
                    # Para admin: mantener password default, actualizar el resto
                    usuarios_final[uname]['docente_info'] = datos.get('docente_info')
                else:
                    # Usuarios creados dinámicamente: usar datos de GS
                    usuarios_final[uname] = datos
    
    # 3. Fallback: archivo local
    if Path(ARCHIVO_USUARIOS).exists():
        try:
            with open(ARCHIVO_USUARIOS, 'r', encoding='utf-8') as f:
                usuarios_local = json.load(f)
            for uname, datos in usuarios_local.items():
                if uname not in usuarios_final:
                    datos['password'] = str(datos.get('password', '')).strip()
                    usuarios_final[uname] = datos
        except Exception:
            pass
    
    guardar_usuarios_local(usuarios_final)
    return usuarios_final


def guardar_usuarios_local(usuarios):
    """Solo guarda localmente (sin Google Sheets)"""
    with open(ARCHIVO_USUARIOS, 'w', encoding='utf-8') as f:
        json.dump(usuarios, f, indent=2, ensure_ascii=False)


def guardar_usuarios(usuarios):
    with open(ARCHIVO_USUARIOS, 'w', encoding='utf-8') as f:
        json.dump(usuarios, f, indent=2, ensure_ascii=False)
    # Sincronizar con Google Sheets
    gs = _gs()
    if gs:
        gs.sync_usuarios_completo(usuarios)


# ================================================================
# CONSTANTES EDUCATIVAS
# ================================================================

NIVELES_GRADOS = {
    "INICIAL": ["Inicial 3 años", "Inicial 4 años", "Inicial 5 años"],
    "PRIMARIA": [
        "1° Primaria", "2° Primaria", "3° Primaria",
        "4° Primaria", "5° Primaria", "6° Primaria"
    ],
    "SECUNDARIA": [
        "1° Secundaria", "2° Secundaria", "3° Secundaria",
        "4° Secundaria", "5° Secundaria"
    ],
    "PREUNIVERSITARIO": [
        "GRUPO AB — CEPRE UNSAAC", "GRUPO CD — CEPRE UNSAAC",
        "Ciclo Verano", "Ciclo Regular", "Ciclo Intensivo",
        "Reforzamiento Primaria"
    ]
}

# Áreas CEPRE UNSAAC por grupo
AREAS_CEPRE_UNSAAC = {
    'GRUPO AB': [
        'Aritmética', 'Álgebra', 'Geometría', 'Trigonometría',
        'Física', 'Química', 'Biología',
        'Competencia Comunicativa',
    ],
    'GRUPO CD': [
        'Aritmética', 'Álgebra', 'Competencia Comunicativa',
        'Historia', 'Geografía', 'Educación Cívica',
        'Economía', 'Filosofía y Lógica',
    ],
}

SECCIONES = ["Única", "A", "B"]

TODOS_LOS_GRADOS = []
for nk, gl in NIVELES_GRADOS.items():
    for gi in gl:
        TODOS_LOS_GRADOS.append(gi)

NIVELES_LIST = list(NIVELES_GRADOS.keys())
GRADOS_OPCIONES = TODOS_LOS_GRADOS.copy()

MESES_ESCOLARES = {
    1: "Enero", 2: "Febrero",
    3: "Marzo", 4: "Abril", 5: "Mayo", 6: "Junio",
    7: "Julio", 8: "Agosto", 9: "Septiembre",
    10: "Octubre", 11: "Noviembre", 12: "Diciembre"
}

MESES_ESP = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
]

BIMESTRES = {
    "Bimestre 1": [3, 4, 5],
    "Bimestre 2": [5, 6, 7],
    "Bimestre 3": [8, 9, 10],
    "Bimestre 4": [10, 11, 12]
}

ARCHIVO_BD = "base_datos.xlsx"
ARCHIVO_MATRICULA = "matricula.xlsx"
ARCHIVO_DOCENTES = "docentes.xlsx"
ARCHIVO_ASISTENCIAS = "asistencias.json"
ARCHIVO_RESULTADOS = "resultados_examenes.json"



def verificar_acceso_docente(nivel):
    """Verifica y devuelve las áreas accesibles según nivel"""
    if nivel == "INICIAL":
        return ["INICIAL"], AREAS_INICIAL
    elif nivel == "PRIMARIA":
        grados = ["1° PRIMARIA", "2° PRIMARIA", "3° PRIMARIA", 
                  "4° PRIMARIA", "5° PRIMARIA", "6° PRIMARIA"]
        return grados, AREAS_PRIMARIA
    elif nivel in ["SECUNDARIA", "PREUNIVERSITARIO"]:
        grados = ["1° SECUNDARIA", "2° SECUNDARIA", "3° SECUNDARIA",
                  "4° SECUNDARIA", "5° SECUNDARIA",
                  "6° PREUNIVERSITARIO"]
        return grados, AREAS_SECUNDARIA + AREAS_PREUNIVERSITARIO
    return [], []


# ================================================================
# SESSION STATE
# ================================================================

def init_session_state():
    defaults = {
        'rol': None,
        'docente_info': None,
        'usuario_actual': '',
        'alumno': '',
        'dni': '',
        'grado': '',
        'apoderado': '',
        'dni_apo': '',
        'tipo_asistencia': 'Entrada',
        'activar_camara_asist': False,
        'areas_examen': [],
        'modulo_activo': None,
        'cola_asistencia': [],
        'wa_enviados': set(),
        'evaluaciones_guardadas': {},
        'ultimo_pdf_incidencia': None,
        'ultimo_codigo_incidencia': '',
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


init_session_state()


# ================================================================
# ESTILOS CSS + ANIMACIONES + SONIDO
# ================================================================

st.markdown("""
<style>
/* === ANIMACIÓN DE ENTRADA === */
@keyframes fadeInUp {
    from { opacity: 0; transform: translateY(20px); }
    to { opacity: 1; transform: translateY(0); }
}
@keyframes pulse {
    0%, 100% { transform: scale(1); }
    50% { transform: scale(1.05); }
}
@keyframes shimmer {
    0% { background-position: -200% 0; }
    100% { background-position: 200% 0; }
}
@keyframes slideIn {
    from { opacity: 0; transform: translateX(-30px); }
    to { opacity: 1; transform: translateX(0); }
}
@keyframes glow {
    0%, 100% { box-shadow: 0 0 5px rgba(26,86,219,0.3); }
    50% { box-shadow: 0 0 20px rgba(26,86,219,0.6); }
}
@keyframes gradient {
    0% { background-position: 0% 50%; }
    50% { background-position: 100% 50%; }
    100% { background-position: 0% 50%; }
}
@keyframes float {
    0%, 100% { transform: translateY(0px); }
    50% { transform: translateY(-8px); }
}
@keyframes rainbow {
    0% { filter: hue-rotate(0deg); }
    100% { filter: hue-rotate(360deg); }
}

/* === HEADER PRINCIPAL === */
.main-header {
    text-align: center; padding: 2rem;
    background: linear-gradient(270deg, #FF6B6B, #4ECDC4, #45B7D1, #F7B731, #5F27CD);
    background-size: 400% 400%;
    color: white; border-radius: 15px; margin-bottom: 2rem;
    box-shadow: 0 8px 25px rgba(0,30,124,0.35);
    animation: gradient 8s ease infinite, fadeInUp 0.6s ease-out;
}

/* === TABS ANIMADOS === */
.stTabs [data-baseweb="tab-list"] {
    gap: 8px;
}
.stTabs [data-baseweb="tab"] {
    border-radius: 10px 10px 0 0;
    padding: 10px 20px;
    transition: all 0.3s ease;
    font-weight: 600;
}
.stTabs [data-baseweb="tab"]:hover {
    background: rgba(26,86,219,0.1);
    transform: translateY(-2px) scale(1.02);
    animation: float 2s ease-in-out infinite;
}
.stTabs [aria-selected="true"] {
    background: linear-gradient(135deg, #1a56db, #0052cc) !important;
    color: white !important;
    box-shadow: 0 4px 12px rgba(26,86,219,0.3);
    animation: glow 2s ease-in-out infinite;
}

/* === BOTONES CON EFECTO === */
.stButton > button {
    transition: all 0.3s ease !important;
    border-radius: 10px !important;
    font-weight: 600 !important;
}
.stButton > button:hover {
    transform: translateY(-2px) scale(1.03) !important;
    box-shadow: 0 0 20px rgba(26,86,219,0.6), 0 0 40px rgba(26,86,219,0.3) !important;
    animation: float 1.5s ease-in-out infinite;
}
.stButton > button:active {
    transform: translateY(0) scale(0.98) !important;
}

/* === CARDS DE ESTADÍSTICAS === */
.stat-card {
    background: linear-gradient(135deg, #f8fafc, #e2e8f0);
    border-radius: 12px; padding: 1.2rem;
    border-left: 4px solid #1a56db;
    box-shadow: 0 2px 10px rgba(0,0,0,0.08);
    animation: slideIn 0.5s ease-out;
    transition: all 0.3s;
}
.stat-card:hover { 
    transform: translateY(-3px) scale(1.02); 
    box-shadow: 0 8px 25px rgba(26,86,219,0.15);
}
.stat-card h3 { margin: 0; color: #1a56db; font-size: 2rem; }
.stat-card p { margin: 0; color: #64748b; font-size: 0.9rem; }

/* === ASISTENCIA REGISTRADA === */
.asist-ok {
    background: linear-gradient(135deg, #dcfce7, #bbf7d0);
    border-radius: 10px; padding: 12px 16px;
    border-left: 4px solid #16a34a;
    animation: fadeInUp 0.4s ease-out;
    margin: 4px 0;
}
.asist-salida {
    background: linear-gradient(135deg, #fef3c7, #fde68a);
    border-radius: 10px; padding: 12px 16px;
    border-left: 4px solid #f59e0b;
    animation: fadeInUp 0.4s ease-out;
    margin: 4px 0;
}

/* === GOOGLE SHEETS STATUS === */
.gs-connected {
    background: linear-gradient(135deg, #dcfce7, #bbf7d0);
    border-radius: 8px; padding: 8px 12px;
    text-align: center; font-weight: 600;
    color: #166534; font-size: 0.85rem;
    animation: pulse 2s infinite;
}
.gs-offline {
    background: #fef3c7; border-radius: 8px;
    padding: 8px 12px; text-align: center;
    color: #92400e; font-size: 0.85rem;
}

/* === RANKING CON ANIMACIÓN === */
.ranking-gold {
    background: linear-gradient(135deg, #FFD700, #FFA500);
    background-size: 200% auto;
    animation: shimmer 3s linear infinite;
    color: #000; padding: 14px; border-radius: 10px;
    font-weight: bold; text-align: center; margin: 5px 0;
    box-shadow: 0 4px 15px rgba(255,215,0,0.4);
}
.ranking-silver {
    background: linear-gradient(135deg, #C0C0C0, #E8E8E8, #C0C0C0);
    background-size: 200% auto;
    animation: shimmer 3s linear infinite;
    color: #000; padding: 14px; border-radius: 10px;
    font-weight: bold; text-align: center; margin: 5px 0;
    box-shadow: 0 4px 12px rgba(192,192,192,0.4);
}
.ranking-bronze {
    background: linear-gradient(135deg, #CD7F32, #E8A849, #CD7F32);
    background-size: 200% auto;
    animation: shimmer 3s linear infinite;
    color: #fff; padding: 14px; border-radius: 10px;
    font-weight: bold; text-align: center; margin: 5px 0;
    box-shadow: 0 4px 12px rgba(205,127,50,0.4);
}

/* === WHATSAPP / LINKS === */
.wa-btn {
    background: linear-gradient(135deg, #25D366, #128C7E); color: white !important;
    padding: 10px 20px; border: none; border-radius: 10px;
    font-size: 15px; width: 100%; text-decoration: none;
    display: block; text-align: center; margin: 4px 0;
    transition: all 0.3s; font-weight: 600;
}
.wa-btn:hover { transform: translateY(-2px) scale(1.02); box-shadow: 0 4px 15px rgba(37,211,102,0.4); }
.link-btn {
    background: linear-gradient(135deg, #4285F4, #356AC3); color: white !important;
    padding: 8px 16px; border: none; border-radius: 10px;
    font-size: 14px; width: 100%; text-decoration: none;
    display: block; text-align: center; margin: 4px 0;
    transition: all 0.3s;
}
.link-btn:hover { transform: translateY(-2px) scale(1.02); box-shadow: 0 4px 12px rgba(66,133,244,0.4); }
.siagie-btn {
    background: linear-gradient(135deg, #E91E63, #C2185B); color: white !important;
    padding: 8px 16px; border: none; border-radius: 10px;
    font-size: 14px; width: 100%; text-decoration: none;
    display: block; text-align: center; margin: 4px 0;
    transition: all 0.3s;
}
.siagie-btn:hover { transform: translateY(-2px) scale(1.02); box-shadow: 0 4px 12px rgba(233,30,99,0.4); }

/* === EXPANDER MEJORADO === */
.streamlit-expanderHeader {
    font-weight: 600 !important;
    border-radius: 8px !important;
}

/* === SIDEBAR === */
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #f8fafc 0%, #e2e8f0 100%);
}

/* === SUCCESS/ERROR MEJORADOS === */
.stSuccess { animation: fadeInUp 0.4s ease-out; border-radius: 10px !important; }
.stError { animation: fadeInUp 0.4s ease-out; border-radius: 10px !important; }
.stInfo { animation: fadeInUp 0.4s ease-out; border-radius: 10px !important; }

/* === DASHBOARD GRID === */
.stButton > button[kind="secondary"] {
    min-height: 100px !important;
    font-size: 1.1rem !important;
    border-radius: 16px !important;
    border: 2px solid #e2e8f0 !important;
    background: linear-gradient(135deg, #ffffff 0%, #f1f5f9 100%) !important;
    transition: all 0.3s ease !important;
    box-shadow: 0 2px 8px rgba(0,0,0,0.06) !important;
}
.stButton > button[kind="secondary"]:hover {
    transform: translateY(-4px) scale(1.02) !important;
    box-shadow: 0 8px 25px rgba(26,86,219,0.15) !important;
    border-color: #1a56db !important;
    background: linear-gradient(135deg, #eff6ff 0%, #dbeafe 100%) !important;
}

/* === NÚMERO ANIMADO === */
@keyframes countUp { from { opacity: 0; transform: translateY(10px); } to { opacity: 1; transform: translateY(0); } }
.stMetric { animation: countUp 0.5s ease-out; }

/* === INPUTS MEJORADOS === */
.stTextInput > div > div > input { border-radius: 10px !important; transition: all 0.3s; }
.stTextInput > div > div > input:focus { box-shadow: 0 0 0 3px rgba(26,86,219,0.2) !important; border-color: #1a56db !important; }
.stSelectbox > div > div { border-radius: 10px !important; }

/* === DATAFRAME === */
.stDataFrame { border-radius: 12px !important; overflow: hidden; box-shadow: 0 2px 10px rgba(0,0,0,0.08); }

/* === SEMÁFORO COLORES === */
.semaforo-ad { color: #16a34a; font-weight: bold; }
.semaforo-a { color: #2563eb; font-weight: bold; }
.semaforo-b { color: #f59e0b; font-weight: bold; }
.semaforo-c { color: #dc2626; font-weight: bold; }

/* === LOADING SPINNER === */
.stSpinner > div { border-color: #1a56db transparent transparent transparent !important; }

/* === DASHBOARD MÓDULOS GRID === */
@keyframes cardFloat {
    0%, 100% { transform: translateY(0); }
    50% { transform: translateY(-5px); }
}
@keyframes glow {
    0%, 100% { box-shadow: 0 4px 15px rgba(0,0,0,0.1); }
    50% { box-shadow: 0 8px 30px rgba(26,86,219,0.25); }
}
@keyframes borderPulse {
    0%, 100% { border-color: transparent; }
    50% { border-color: #1a56db; }
}

/* === SIDEBAR MEJORADO === */
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #f8fafc 0%, #e2e8f0 100%) !important;
}
section[data-testid="stSidebar"] .stMarkdown h1 {
    background: linear-gradient(135deg, #1a56db, #7c3aed);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    font-size: 1.5rem !important;
}

/* === EXPANDERS ANIMADOS === */
.streamlit-expanderHeader {
    border-radius: 10px !important;
    transition: all 0.3s ease !important;
    font-weight: 600 !important;
}
.streamlit-expanderHeader:hover {
    background: rgba(26,86,219,0.08) !important;
}

/* === TABS CON GRADIENTE === */
.stTabs [aria-selected="true"] {
    background: linear-gradient(135deg, #1a56db, #2563eb) !important;
    color: white !important;
    border-radius: 10px 10px 0 0 !important;
    box-shadow: 0 4px 15px rgba(26,86,219,0.3) !important;
    transition: all 0.3s ease !important;
}

/* === ALERTAS BONITAS === */
.stAlert { border-radius: 12px !important; animation: fadeInUp 0.4s ease-out; }

/* === ÉXITO CON BRILLO === */
.stSuccess {
    animation: fadeInUp 0.4s ease-out;
    border-radius: 12px !important;
}

/* === BOTÓN PRIMARIO PREMIUM === */
.stButton > button[kind="primary"] {
    background: linear-gradient(135deg, #1a56db 0%, #2563eb 50%, #3b82f6 100%) !important;
    border: none !important;
    box-shadow: 0 4px 15px rgba(26,86,219,0.3) !important;
    letter-spacing: 0.5px !important;
}
.stButton > button[kind="primary"]:hover {
    background: linear-gradient(135deg, #1e40af 0%, #1a56db 50%, #2563eb 100%) !important;
    box-shadow: 0 8px 25px rgba(26,86,219,0.4) !important;
    transform: translateY(-3px) !important;
}

/* === DOWNLOAD BUTTON === */
.stDownloadButton > button {
    background: linear-gradient(135deg, #059669, #10b981) !important;
    color: white !important;
    border: none !important;
    border-radius: 10px !important;
    transition: all 0.3s ease !important;
}
.stDownloadButton > button:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 6px 20px rgba(5,150,105,0.3) !important;
}

/* === RADIO BUTTONS === */
.stRadio > div { gap: 8px; }
.stRadio [role="radiogroup"] > label {
    border-radius: 10px !important;
    transition: all 0.2s ease !important;
    padding: 8px 16px !important;
}

/* === NOTIFICACIONES FLOTANTES === */
@keyframes slideInRight {
    from { opacity: 0; transform: translateX(50px); }
    to { opacity: 1; transform: translateX(0); }
}
.stToast { animation: slideInRight 0.4s ease-out !important; }

/* === SELECTBOX HOVER === */
.stSelectbox > div > div:hover {
    border-color: #1a56db !important;
    box-shadow: 0 0 0 2px rgba(26,86,219,0.1) !important;
}
</style>
""", unsafe_allow_html=True)


def reproducir_sonido_asistencia():
    """Genera un beep/sonido cuando se registra asistencia"""
    st.markdown("""
    <audio autoplay>
        <source src="data:audio/wav;base64,UklGRl4FAABXQVZFZm10IBAAAAABAAEARKwAAIhYAQACABAAZGF0YToFAACAgICAgICAgICAkJigoKiouMDI0Njg6PD4+Pj48PDo4NjQyMC4sKignp6WjoaAgICAgICAgICAgJCYoKCorLjAyNDY4Ojw+Pj4+PDw6ODY0MjAuLCooJ6elo6GgICAgICAgICA" type="audio/wav">
    </audio>
    """, unsafe_allow_html=True)


def reproducir_beep_exitoso():
    """Sonido agradable tipo chime para registros exitosos"""
    st.markdown("""
    <script>
    (function() {
        try {
            var ctx = new (window.AudioContext || window.webkitAudioContext)();
            function nota(freq, inicio, dur) {
                var o = ctx.createOscillator();
                var g = ctx.createGain();
                o.connect(g); g.connect(ctx.destination);
                o.frequency.value = freq;
                o.type = 'sine';
                g.gain.setValueAtTime(0.25, ctx.currentTime + inicio);
                g.gain.exponentialRampToValueAtTime(0.001, ctx.currentTime + inicio + dur);
                o.start(ctx.currentTime + inicio);
                o.stop(ctx.currentTime + inicio + dur);
            }
            nota(523, 0, 0.15);
            nota(659, 0.1, 0.15);
            nota(784, 0.2, 0.3);
        } catch(e) {}
    })();
    </script>
    """, unsafe_allow_html=True)


def reproducir_beep_error():
    """Sonido de error"""
    st.markdown("""
    <script>
    (function() {
        try {
            var ctx = new (window.AudioContext || window.webkitAudioContext)();
            var osc = ctx.createOscillator();
            var gain = ctx.createGain();
            osc.connect(gain);
            gain.connect(ctx.destination);
            osc.frequency.value = 300;
            osc.type = 'square';
            gain.gain.setValueAtTime(0.2, ctx.currentTime);
            gain.gain.exponentialRampToValueAtTime(0.001, ctx.currentTime + 0.5);
            osc.start(ctx.currentTime);
            osc.stop(ctx.currentTime + 0.5);
        } catch(e) {}
    })();
    </script>
    """, unsafe_allow_html=True)


# ================================================================
# IMPORTACIONES OPCIONALES
# ================================================================

try:
    from barcode import Code128
    from barcode.writer import ImageWriter
    HAS_BARCODE = True
except ImportError:
    HAS_BARCODE = False

try:
    import cv2
    HAS_CV2 = True
except ImportError:
    HAS_CV2 = False

try:
    from pyzbar.pyzbar import decode as pyzbar_decode
    HAS_PYZBAR = True
except ImportError:
    HAS_PYZBAR = False


# ================================================================
# FUENTES
# ================================================================

class RecursoManager:
    @staticmethod
    def obtener_fuente(nombre, tamanio, bold=False):
        tam = int(tamanio)
        rutas = [
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
            "/usr/share/fonts/truetype/freefont/FreeSansBold.ttf" if bold else "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
            "/usr/share/fonts/truetype/ubuntu/Ubuntu-Bold.ttf" if bold else "/usr/share/fonts/truetype/ubuntu/Ubuntu-R.ttf",
            "/usr/share/fonts/truetype/noto/NotoSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
        ]
        for ruta in rutas:
            try:
                if Path(ruta).exists():
                    return ImageFont.truetype(ruta, tam)
            except Exception:
                continue
        # Fallback con tamanio (Pillow 10+)
        try:
            return ImageFont.load_default(size=tam)
        except TypeError:
            return ImageFont.load_default()


# ================================================================
# PERMISOS — SOLO ADMIN PUEDE BORRAR
# ================================================================

def puede_borrar():
    """Solo el admin puede borrar datos del sistema"""
    return st.session_state.rol == "admin"


# ================================================================
# BASE DE DATOS — ALUMNOS Y DOCENTES
# ================================================================

class BaseDatos:

    @staticmethod
    def cargar_matricula():
        # Después de escribir, forzar lectura local para evitar datos viejos de GS
        forzar_local = st.session_state.get('_forzar_local', False)
        if forzar_local:
            st.session_state['_forzar_local'] = False
            try:
                if Path(ARCHIVO_MATRICULA).exists():
                    df = pd.read_excel(ARCHIVO_MATRICULA, dtype=str, engine='openpyxl')
                    df.columns = df.columns.str.strip()
                    return df
            except Exception:
                pass
        # Intentar Google Sheets primero
        gs = _gs()
        if gs:
            try:
                df_gs = gs.leer_matricula()
                if not df_gs.empty:
                    col_map = {'nombre': 'Nombre', 'dni': 'DNI', 'nivel': 'Nivel',
                               'grado': 'Grado', 'seccion': 'Seccion',
                               'apoderado': 'Apoderado', 'dni_apoderado': 'DNI_Apoderado',
                               'celular_apoderado': 'Celular_Apoderado'}
                    df_gs = df_gs.rename(columns=col_map)
                    for col in df_gs.columns:
                        df_gs[col] = df_gs[col].astype(str).replace('nan', '').replace('None', '')
                    # ── PROTECCIÓN: combinar con local para no perder datos ──────
                    try:
                        if Path(ARCHIVO_MATRICULA).exists():
                            df_local = pd.read_excel(ARCHIVO_MATRICULA, dtype=str, engine='openpyxl')
                            df_local.columns = df_local.columns.str.strip()
                            if not df_local.empty and 'DNI' in df_local.columns and 'DNI' in df_gs.columns:
                                # Agregar al GS los que están en local pero no en GS
                                dnis_gs = set(df_gs['DNI'].astype(str).str.strip())
                                df_solo_local = df_local[~df_local['DNI'].astype(str).str.strip().isin(dnis_gs)]
                                if not df_solo_local.empty:
                                    df_gs = pd.concat([df_gs, df_solo_local], ignore_index=True)
                    except Exception:
                        pass
                    return df_gs
            except Exception:
                pass
        # Fallback: leer local
        try:
            if Path(ARCHIVO_MATRICULA).exists():
                df = pd.read_excel(ARCHIVO_MATRICULA, dtype=str, engine='openpyxl')
                df.columns = df.columns.str.strip()
                return df
        except Exception:
            pass
        return pd.DataFrame(columns=[
            'Nombre', 'DNI', 'Nivel', 'Grado', 'Seccion',
            'Apoderado', 'DNI_Apoderado', 'Celular_Apoderado'
        ])

    @staticmethod
    def guardar_matricula(df):
        try:
            df.to_excel(ARCHIVO_MATRICULA, index=False, engine='openpyxl')
        except Exception:
            # Fallback: guardar como CSV si openpyxl falla
            df.to_csv(ARCHIVO_MATRICULA.replace('.xlsx', '.csv'), index=False)
        # Forzar lectura local en el próximo cargar (GS puede tener datos viejos)
        st.session_state['_forzar_local'] = True
        # Invalidar índice DNI para que se reconstruya con datos nuevos
        st.session_state.pop('_indice_dni', None)
        st.session_state.pop('_indice_dni_ts', None)
        # Sincronizar con Google Sheets
        gs = _gs()
        if gs:
            try:
                col_map = {'Nombre': 'nombre', 'DNI': 'dni', 'Nivel': 'nivel',
                           'Grado': 'grado', 'Seccion': 'seccion',
                           'Apoderado': 'apoderado', 'DNI_Apoderado': 'dni_apoderado',
                           'Celular_Apoderado': 'celular_apoderado'}
                df_gs = df.rename(columns=col_map).copy()
                if 'fecha_matricula' not in df_gs.columns:
                    df_gs['fecha_matricula'] = fecha_peru_str()
                gs.sync_matricula_completa(df_gs)
            except Exception:
                pass

    @staticmethod
    def registrar_estudiante(datos):
        df = BaseDatos.cargar_matricula()
        # --- MATRÍCULA PROVISIONAL: generar ID temporal si no hay DNI ---
        if not datos.get('DNI') or str(datos.get('DNI', '')).strip() == '':
            # Generar ID provisional único: PROV + número correlativo
            existing_provs = []
            if not df.empty and 'DNI' in df.columns:
                for v in df['DNI'].astype(str):
                    if v.startswith('PROV'):
                        try:
                            existing_provs.append(int(v[4:]))
                        except Exception:
                            pass
            next_num = max(existing_provs, default=0) + 1
            datos['DNI'] = f'PROV{next_num:04d}'
            datos['_provisional'] = 'SI'
        else:
            datos['_provisional'] = datos.get('_provisional', 'NO')
        # --- Actualizar si ya existe, sino agregar ---
        if not df.empty and 'DNI' in df.columns and datos['DNI'] in df['DNI'].values:
            idx = df[df['DNI'] == datos['DNI']].index[0]
            for k, v in datos.items():
                df.at[idx, k] = v
        else:
            df = pd.concat([df, pd.DataFrame([datos])], ignore_index=True)
        BaseDatos.guardar_matricula(df)

    @staticmethod
    def buscar_por_dni(dni):
        """Búsqueda RÁPIDA: índice memoria (1ms) → local (50ms) → GS (último recurso)"""
        dni_str = str(dni).strip()

        def _buscar_en_df(df, tipo):
            if df is not None and not df.empty and 'DNI' in df.columns:
                df = df.copy()
                df['DNI'] = df['DNI'].astype(str).str.strip()
                res = df[df['DNI'] == dni_str]
                if not res.empty:
                    r = res.iloc[0].to_dict()
                    r['_tipo'] = tipo
                    return r
            return None

        # 1. ÍNDICE EN MEMORIA — instantáneo (<1ms)
        indice = st.session_state.get('_indice_dni')
        if indice and dni_str in indice:
            return indice[dni_str].copy()

        # 2. Reconstruir índice si no existe o es viejo (>3 min)
        ultima = st.session_state.get('_indice_dni_ts', 0)
        if not indice or (time.time() - ultima > 180):
            _construir_indice_dni()
            indice = st.session_state.get('_indice_dni', {})
            if dni_str in indice:
                return indice[dni_str].copy()

        # 3. Archivo local directo (sin GS, ~50ms)
        try:
            if Path(ARCHIVO_MATRICULA).exists():
                df_local = pd.read_excel(ARCHIVO_MATRICULA, dtype=str, engine='openpyxl')
                df_local.columns = df_local.columns.str.strip()
                found = _buscar_en_df(df_local, 'alumno')
                if found:
                    if '_indice_dni' not in st.session_state:
                        st.session_state['_indice_dni'] = {}
                    st.session_state['_indice_dni'][dni_str] = found
                    return found
        except Exception:
            pass

        # 4. Docentes local
        try:
            if Path(ARCHIVO_DOCENTES).exists():
                df_d = pd.read_excel(ARCHIVO_DOCENTES, dtype=str, engine='openpyxl')
                df_d.columns = df_d.columns.str.strip()
                found = _buscar_en_df(df_d, 'docente')
                if found:
                    if '_indice_dni' not in st.session_state:
                        st.session_state['_indice_dni'] = {}
                    st.session_state['_indice_dni'][dni_str] = found
                    return found
        except Exception:
            pass

        # 5. GS como último recurso (lento)
        try:
            df_m = BaseDatos.cargar_matricula()
            found = _buscar_en_df(df_m, 'alumno')
            if found:
                return found
            df_d = BaseDatos.cargar_docentes()
            found = _buscar_en_df(df_d, 'docente')
            if found:
                return found
        except Exception:
            pass

        # 6. Fallback archivo BD antiguo
        try:
            if Path(ARCHIVO_BD).exists():
                df2 = pd.read_excel(ARCHIVO_BD, dtype=str, engine='openpyxl')
                df2.columns = df2.columns.str.strip().str.title()
                if 'Dni' in df2.columns:
                    df2['Dni'] = df2['Dni'].astype(str).str.strip()
                    res3 = df2[df2['Dni'] == dni_str]
                    if not res3.empty:
                        row = res3.iloc[0].to_dict()
                        return {
                            'Nombre': row.get('Alumno', row.get('Nombre', '')),
                            'DNI': row.get('Dni', ''),
                            'Grado': row.get('Grado', ''),
                            'Nivel': row.get('Nivel', ''),
                            'Seccion': row.get('Seccion', ''),
                            'Apoderado': row.get('Apoderado', ''),
                            'DNI_Apoderado': row.get('Dni_Apoderado', ''),
                            'Celular_Apoderado': row.get('Celular', ''),
                            '_tipo': 'alumno'
                        }
        except Exception:
            pass
        return None

    @staticmethod
    def eliminar_estudiante(dni):
        dni_str = str(dni).strip()
        # 1. Quitar de matrícula
        df = BaseDatos.cargar_matricula()
        df['DNI'] = df['DNI'].astype(str).str.strip()
        df = df[df['DNI'] != dni_str]
        BaseDatos.guardar_matricula(df)
        # 2. Borrar notas en historial_evaluaciones.json
        try:
            if Path('historial_evaluaciones.json').exists():
                with open('historial_evaluaciones.json', 'r', encoding='utf-8') as f:
                    hist = json.load(f)
                # Eliminar todas las claves que contengan este DNI
                hist = {k: v for k, v in hist.items() if dni_str not in k}
                with open('historial_evaluaciones.json', 'w', encoding='utf-8') as f:
                    json.dump(hist, f, indent=2, ensure_ascii=False)
        except Exception:
            pass
        # 3. Borrar de resultados.json
        try:
            if Path('resultados.json').exists():
                with open('resultados.json', 'r', encoding='utf-8') as f:
                    res = json.load(f)
                res = {k: v for k, v in res.items() if dni_str not in k}
                with open('resultados.json', 'w', encoding='utf-8') as f:
                    json.dump(res, f, indent=2, ensure_ascii=False)
        except Exception:
            pass
        # 4. Borrar de resultados_examenes.json
        try:
            if Path('resultados_examenes.json').exists():
                with open('resultados_examenes.json', 'r', encoding='utf-8') as f:
                    rex = json.load(f)
                rex = {k: v for k, v in rex.items() if dni_str not in k}
                with open('resultados_examenes.json', 'w', encoding='utf-8') as f:
                    json.dump(rex, f, indent=2, ensure_ascii=False)
        except Exception:
            pass
        # Limpiar notas del estudiante al eliminar
        BaseDatos.eliminar_notas_por_dni(dni)

    @staticmethod
    def eliminar_notas_por_dni(dni):
        """Borra todas las notas y evaluaciones de un DNI específico"""
        dni_str = str(dni).strip()
        # historial_evaluaciones.json
        try:
            if Path('historial_evaluaciones.json').exists():
                with open('historial_evaluaciones.json', 'r', encoding='utf-8') as f:
                    hist = json.load(f)
                hist.pop(dni_str, None)
                with open('historial_evaluaciones.json', 'w', encoding='utf-8') as f:
                    json.dump(hist, f, indent=2, ensure_ascii=False)
        except Exception: pass
        # resultados.json y ARCHIVO_RESULTADOS
        for archivo in ['resultados.json', ARCHIVO_RESULTADOS]:
            try:
                if Path(archivo).exists():
                    with open(archivo, 'r', encoding='utf-8') as f:
                        res = json.load(f)
                    res = [r for r in res if str(r.get('dni', '')) != dni_str]
                    with open(archivo, 'w', encoding='utf-8') as f:
                        json.dump(res, f, indent=2, ensure_ascii=False)
            except Exception: pass

    @staticmethod
    def obtener_estudiantes_grado(grado, seccion=None):
        df = BaseDatos.cargar_matricula()
        if df.empty:
            return df
        if grado in ('ALL_NIVELES',):
            pass  # No filtrar — todos los grados/niveles
        elif grado in ('ALL_SECUNDARIA',):
            if 'Nivel' in df.columns:
                df = df[df['Nivel'] == "SECUNDARIA"]
        elif grado in ('ALL_SEC_PREU',):
            if 'Nivel' in df.columns:
                df = df[df['Nivel'].isin(['SECUNDARIA', 'PREUNIVERSITARIO'])]
        elif 'Grado' in df.columns:
            # Comparación flexible: strip y case-insensitive
            grado_norm = str(grado).strip().lower()
            df = df[df['Grado'].astype(str).str.strip().str.lower() == grado_norm]
        if seccion and seccion not in ["Todas", "Única"] and 'Seccion' in df.columns:
            df = df[df['Seccion'] == seccion]
        if 'Nombre' in df.columns:
            df = df.sort_values('Nombre', ascending=True).reset_index(drop=True)
        return df

    @staticmethod
    def cargar_docentes():
        # Después de escribir, forzar lectura local
        forzar_local = st.session_state.get('_forzar_local_doc', False)
        if forzar_local:
            st.session_state['_forzar_local_doc'] = False
            try:
                if Path(ARCHIVO_DOCENTES).exists():
                    df = pd.read_excel(ARCHIVO_DOCENTES, dtype=str, engine='openpyxl')
                    df.columns = df.columns.str.strip()
                    return df
            except Exception:
                pass
        # Intentar Google Sheets primero
        gs = _gs()
        if gs:
            try:
                df_gs = gs.leer_docentes()
                if not df_gs.empty:
                    col_map = {'nombre': 'Nombre', 'dni': 'DNI', 'cargo': 'Cargo',
                               'especialidad': 'Especialidad', 'celular': 'Celular',
                               'grado_asignado': 'Grado_Asignado'}
                    df_gs = df_gs.rename(columns=col_map)
                    for col in df_gs.columns:
                        df_gs[col] = df_gs[col].astype(str).replace('nan', '').replace('None', '')
                    return df_gs
            except Exception:
                pass
        try:
            if Path(ARCHIVO_DOCENTES).exists():
                df = pd.read_excel(ARCHIVO_DOCENTES, dtype=str, engine='openpyxl')
                df.columns = df.columns.str.strip()
                return df
        except Exception:
            pass
        return pd.DataFrame(columns=[
            'Nombre', 'DNI', 'Cargo', 'Especialidad', 'Celular', 'Grado_Asignado'
        ])

    @staticmethod
    def guardar_docentes(df):
        try:
            df.to_excel(ARCHIVO_DOCENTES, index=False, engine='openpyxl')
        except Exception:
            df.to_csv(ARCHIVO_DOCENTES.replace('.xlsx', '.csv'), index=False)
        # Forzar lectura local en el próximo cargar
        st.session_state['_forzar_local_doc'] = True
        # Invalidar índice DNI
        st.session_state.pop('_indice_dni', None)
        st.session_state.pop('_indice_dni_ts', None)
        gs = _gs()
        if gs:
            try:
                col_map = {'Nombre': 'nombre', 'DNI': 'dni', 'Cargo': 'cargo',
                           'Especialidad': 'especialidad', 'Celular': 'celular',
                           'Grado_Asignado': 'grado_asignado'}
                df_gs = df.rename(columns=col_map).copy()
                if 'fecha_registro' not in df_gs.columns:
                    df_gs['fecha_registro'] = fecha_peru_str()
                gs.sync_docentes_completo(df_gs)
            except Exception:
                pass

    @staticmethod
    def registrar_docente(datos):
        df = BaseDatos.cargar_docentes()
        if not df.empty and 'DNI' in df.columns and datos['DNI'] in df['DNI'].values:
            idx = df[df['DNI'] == datos['DNI']].index[0]
            for k, v in datos.items():
                df.at[idx, k] = v
        else:
            df = pd.concat([df, pd.DataFrame([datos])], ignore_index=True)
        BaseDatos.guardar_docentes(df)

    @staticmethod
    def eliminar_docente(dni):
        df = BaseDatos.cargar_docentes()
        df['DNI'] = df['DNI'].astype(str).str.strip()
        df = df[df['DNI'] != str(dni).strip()]
        BaseDatos.guardar_docentes(df)

    @staticmethod
    def guardar_asistencia(dni, nombre, tipo, hora, es_docente=False):
        fecha_hoy = fecha_peru_str()
        asistencias = {}
        if Path(ARCHIVO_ASISTENCIAS).exists():
            with open(ARCHIVO_ASISTENCIAS, 'r', encoding='utf-8') as f:
                asistencias = json.load(f)
        if fecha_hoy not in asistencias:
            asistencias[fecha_hoy] = {}
        if dni not in asistencias[fecha_hoy]:
            asistencias[fecha_hoy][dni] = {
                'nombre': nombre, 'entrada': '', 'salida': '',
                'tardanza': '', 'entrada_tarde': '', 'salida_tarde': '',
                'es_docente': es_docente
            }
        # Mapear tipos a campos
        campo = tipo.lower().replace(' ', '_')
        if campo in ('entrada', 'salida', 'tardanza', 'entrada_tarde', 'salida_tarde'):
            asistencias[fecha_hoy][dni][campo] = hora
        asistencias[fecha_hoy][dni]['nombre'] = nombre
        with open(ARCHIVO_ASISTENCIAS, 'w', encoding='utf-8') as f:
            json.dump(asistencias, f, indent=2, ensure_ascii=False)
        # Sincronizar con Google Sheets en silencio
        try:
            gs = _gs()
            if gs:
                grado = ''
                nivel = ''
                # Usar índice DNI (rápido) para obtener grado/nivel
                indice = st.session_state.get('_indice_dni', {})
                if dni in indice:
                    grado = str(indice[dni].get('Grado', indice[dni].get('grado', '')))
                    nivel = str(indice[dni].get('Nivel', indice[dni].get('nivel', '')))
                else:
                    # Fallback: buscar en archivo local
                    try:
                        if Path(ARCHIVO_MATRICULA).exists():
                            df_m = pd.read_excel(ARCHIVO_MATRICULA, dtype=str, engine='openpyxl')
                            df_m.columns = df_m.columns.str.strip()
                            if 'DNI' in df_m.columns:
                                est = df_m[df_m['DNI'].astype(str).str.strip() == str(dni).strip()]
                                if not est.empty:
                                    grado = str(est.iloc[0].get('Grado', ''))
                                    nivel = str(est.iloc[0].get('Nivel', ''))
                    except Exception:
                        pass
                reg = asistencias[fecha_hoy][dni]
                gs.guardar_asistencia({
                    'fecha': fecha_hoy,
                    'dni': str(dni),
                    'nombre': nombre,
                    'tipo_persona': 'docente' if es_docente else 'alumno',
                    'hora_entrada': reg.get('entrada', ''),
                    'hora_salida': reg.get('salida', ''),
                    'hora_entrada_tarde': reg.get('entrada_tarde', ''),
                    'hora_salida_tarde': reg.get('salida_tarde', ''),
                    'grado': grado,
                    'nivel': nivel,
                })
        except Exception:
            pass  # Error silencioso — asistencia ya guardada localmente

    @staticmethod
    def obtener_asistencias_hoy():
        fecha_hoy = fecha_peru_str()
        if Path(ARCHIVO_ASISTENCIAS).exists():
            with open(ARCHIVO_ASISTENCIAS, 'r', encoding='utf-8') as f:
                return json.load(f).get(fecha_hoy, {})
        return {}

    @staticmethod
    def borrar_asistencias_hoy():
        fecha_hoy = fecha_peru_str()
        if Path(ARCHIVO_ASISTENCIAS).exists():
            with open(ARCHIVO_ASISTENCIAS, 'r', encoding='utf-8') as f:
                a = json.load(f)
            if fecha_hoy in a:
                del a[fecha_hoy]
            with open(ARCHIVO_ASISTENCIAS, 'w', encoding='utf-8') as f:
                json.dump(a, f, indent=2, ensure_ascii=False)

    @staticmethod
    def obtener_estadisticas():
        df = BaseDatos.cargar_matricula()
        df_d = BaseDatos.cargar_docentes()
        return {
            'total_alumnos': len(df) if not df.empty else 0,
            'total_docentes': len(df_d) if not df_d.empty else 0,
            'grados': df['Grado'].nunique() if not df.empty and 'Grado' in df.columns else 0
        }

    # ---- RESULTADOS POR DOCENTE (separados por usuario) ----

    @staticmethod
    def guardar_resultados_examen(resultado, usuario_docente):
        """Guarda resultado asociado al usuario docente"""
        datos = {}
        if Path(ARCHIVO_RESULTADOS).exists():
            try:
                with open(ARCHIVO_RESULTADOS, 'r', encoding='utf-8') as f:
                    raw = json.load(f)
                if isinstance(raw, list):
                    datos = {"migrado": raw}
                elif isinstance(raw, dict):
                    datos = raw
                else:
                    datos = {}
            except Exception:
                datos = {}
        if usuario_docente not in datos:
            datos[usuario_docente] = []
        datos[usuario_docente].append(resultado)
        with open(ARCHIVO_RESULTADOS, 'w', encoding='utf-8') as f:
            json.dump(datos, f, indent=2, ensure_ascii=False)
        # Sincronizar con Google Sheets
        gs = _gs()
        if gs:
            try:
                import uuid
                eval_id = str(uuid.uuid4())[:8]
                titulo = resultado.get('titulo', 'Evaluación')
                fecha = resultado.get('fecha', fecha_peru_str())
                grado = resultado.get('grado', '')
                areas_info = resultado.get('areas', [])
                alumnos = resultado.get('alumnos', [])
                gs.guardar_resultados_examen(
                    eval_id, titulo, fecha, usuario_docente,
                    grado, areas_info, alumnos
                )
            except Exception:
                pass

    @staticmethod
    def cargar_resultados_examen(usuario_docente):
        """Carga solo los resultados del docente específico"""
        if Path(ARCHIVO_RESULTADOS).exists():
            try:
                with open(ARCHIVO_RESULTADOS, 'r', encoding='utf-8') as f:
                    datos = json.load(f)
                # Si es formato viejo (lista), retornar la lista completa
                if isinstance(datos, list):
                    return datos
                elif isinstance(datos, dict):
                    return datos.get(usuario_docente, [])
            except Exception:
                pass
        return []

    @staticmethod
    def limpiar_resultados_examen(usuario_docente):
        """Limpia solo los resultados del docente"""
        if Path(ARCHIVO_RESULTADOS).exists():
            try:
                with open(ARCHIVO_RESULTADOS, 'r', encoding='utf-8') as f:
                    datos = json.load(f)
                if isinstance(datos, list):
                    # Formato viejo, limpiar todo
                    datos = {}
                elif isinstance(datos, dict) and usuario_docente in datos:
                    datos[usuario_docente] = []
                with open(ARCHIVO_RESULTADOS, 'w', encoding='utf-8') as f:
                    json.dump(datos, f, indent=2, ensure_ascii=False)
            except Exception:
                pass

    @staticmethod
    def cargar_todos_resultados():
        """Carga todos los resultados (para admin)"""
        if Path(ARCHIVO_RESULTADOS).exists():
            try:
                with open(ARCHIVO_RESULTADOS, 'r', encoding='utf-8') as f:
                    datos = json.load(f)
                todos = []
                if isinstance(datos, list):
                    # Formato viejo
                    for r in datos:
                        r['_docente'] = 'migrado'
                        todos.append(r)
                elif isinstance(datos, dict):
                    for usr, lista in datos.items():
                        if isinstance(lista, list):
                            for r in lista:
                                r['_docente'] = usr
                                todos.append(r)
                return todos
            except Exception:
                pass
        return []

    @staticmethod
    def promover_grados():
        """Avanza todos los estudiantes al siguiente grado del año escolar.
        Retorna dict con resumen: {'promovidos': n, 'egresados': n, 'detalle': [...]}
        """
        # Cadena completa de promoción
        CADENA_PROMOCION = {
            # INICIAL
            "Inicial 3 años":  ("Inicial 4 años",  "INICIAL"),
            "Inicial 4 años":  ("Inicial 5 años",  "INICIAL"),
            "Inicial 5 años":  ("1° Primaria",      "PRIMARIA"),
            # PRIMARIA
            "1° Primaria":     ("2° Primaria",      "PRIMARIA"),
            "2° Primaria":     ("3° Primaria",      "PRIMARIA"),
            "3° Primaria":     ("4° Primaria",      "PRIMARIA"),
            "4° Primaria":     ("5° Primaria",      "PRIMARIA"),
            "5° Primaria":     ("6° Primaria",      "PRIMARIA"),
            "6° Primaria":     ("1° Secundaria",    "SECUNDARIA"),
            # SECUNDARIA
            "1° Secundaria":   ("2° Secundaria",    "SECUNDARIA"),
            "2° Secundaria":   ("3° Secundaria",    "SECUNDARIA"),
            "3° Secundaria":   ("4° Secundaria",    "SECUNDARIA"),
            "4° Secundaria":   ("5° Secundaria",    "SECUNDARIA"),
            "5° Secundaria":   ("EGRESADO",         "EGRESADO"),
        }

        df = BaseDatos.cargar_matricula()
        if df.empty or 'Grado' not in df.columns:
            return {'promovidos': 0, 'egresados': 0, 'sin_cambio': 0, 'detalle': []}

        promovidos = 0
        egresados = 0
        sin_cambio = 0
        detalle = []

        for idx, row in df.iterrows():
            grado_actual = str(row.get('Grado', '')).strip()
            if grado_actual in CADENA_PROMOCION:
                nuevo_grado, nuevo_nivel = CADENA_PROMOCION[grado_actual]
                if nuevo_grado == "EGRESADO":
                    df.at[idx, 'Grado'] = "EGRESADO 5° Sec"
                    df.at[idx, 'Nivel'] = "EGRESADO"
                    egresados += 1
                    detalle.append(f"🎓 {row.get('Nombre','')} — {grado_actual} → EGRESADO")
                else:
                    df.at[idx, 'Grado'] = nuevo_grado
                    df.at[idx, 'Nivel'] = nuevo_nivel
                    promovidos += 1
                    detalle.append(f"✅ {row.get('Nombre','')} — {grado_actual} → {nuevo_grado}")
            else:
                # PreUniversitario u otros ciclos especiales → no se tocan
                sin_cambio += 1

        BaseDatos.guardar_matricula(df)
        return {
            'promovidos': promovidos,
            'egresados': egresados,
            'sin_cambio': sin_cambio,
            'detalle': detalle
        }

    @staticmethod
    def previsualizar_promocion():
        """Retorna resumen de cómo quedarían los grados SIN guardar."""
        CADENA_PROMOCION = {
            "Inicial 3 años":  "Inicial 4 años",
            "Inicial 4 años":  "Inicial 5 años",
            "Inicial 5 años":  "1° Primaria",
            "1° Primaria":     "2° Primaria",
            "2° Primaria":     "3° Primaria",
            "3° Primaria":     "4° Primaria",
            "4° Primaria":     "5° Primaria",
            "5° Primaria":     "6° Primaria",
            "6° Primaria":     "1° Secundaria",
            "1° Secundaria":   "2° Secundaria",
            "2° Secundaria":   "3° Secundaria",
            "3° Secundaria":   "4° Secundaria",
            "4° Secundaria":   "5° Secundaria",
            "5° Secundaria":   "EGRESADO",
        }
        df = BaseDatos.cargar_matricula()
        resumen = {}
        if df.empty or 'Grado' not in df.columns:
            return resumen
        for grado_actual, nuevo in CADENA_PROMOCION.items():
            n = int((df['Grado'].astype(str).str.strip() == grado_actual).sum())
            if n > 0:
                resumen[grado_actual] = {'nuevo': nuevo, 'cantidad': n}
        return resumen

    @staticmethod
    def corregir_secciones_vacias():
        """Asigna sección 'A' a estudiantes sin sección o con 'Única' (excepto INICIAL)"""
        df = BaseDatos.cargar_matricula()
        if df.empty or 'Seccion' not in df.columns:
            return 0

        df['Seccion'] = df['Seccion'].fillna('').astype(str).str.strip()
        nivel_col = df['Nivel'].astype(str).str.strip().str.upper() if 'Nivel' in df.columns else None

        # Condición 1: sección vacía o valores nulos
        sin_seccion = df['Seccion'].isin(['', 'nan', 'None', 'N/A'])

        # Condición 2: sección "Única" en niveles que NO son INICIAL
        if nivel_col is not None:
            es_unica = df['Seccion'].str.upper().isin(['ÚNICA', 'UNICA', 'única', 'unica'])
            no_es_inicial = ~nivel_col.isin(['INICIAL'])
            unica_a_cambiar = es_unica & no_es_inicial
        else:
            unica_a_cambiar = df['Seccion'].str.upper().isin(['ÚNICA', 'UNICA', 'única', 'unica'])

        mask_total = sin_seccion | unica_a_cambiar
        cantidad = int(mask_total.sum())

        if cantidad > 0:
            df.loc[mask_total, 'Seccion'] = 'A'
            BaseDatos.guardar_matricula(df)
        return cantidad


def _construir_indice_dni():
    """Construye índice DNI→persona en memoria para búsqueda instantánea.
    Se refresca cada 3 minutos automáticamente."""
    indice = {}
    # 1. Archivo local matricula (rápido)
    try:
        if Path(ARCHIVO_MATRICULA).exists():
            df = pd.read_excel(ARCHIVO_MATRICULA, dtype=str, engine='openpyxl')
            df.columns = df.columns.str.strip()
            if not df.empty and 'DNI' in df.columns:
                for _, row in df.iterrows():
                    dni = str(row.get('DNI', '')).strip()
                    if dni and len(dni) >= 7:
                        r = row.to_dict()
                        r['_tipo'] = 'alumno'
                        indice[dni] = r
    except Exception:
        pass
    # 2. Archivo local docentes
    try:
        if Path(ARCHIVO_DOCENTES).exists():
            df_d = pd.read_excel(ARCHIVO_DOCENTES, dtype=str, engine='openpyxl')
            df_d.columns = df_d.columns.str.strip()
            if not df_d.empty and 'DNI' in df_d.columns:
                for _, row in df_d.iterrows():
                    dni = str(row.get('DNI', '')).strip()
                    if dni and len(dni) >= 7:
                        r = row.to_dict()
                        r['_tipo'] = 'docente'
                        indice[dni] = r
    except Exception:
        pass
    # 3. GS complementa (sin bloquear si falla)
    try:
        gs = _gs()
        if gs:
            df_gs = gs.leer_matricula()
            if not df_gs.empty:
                col_map = {'nombre': 'Nombre', 'dni': 'DNI', 'nivel': 'Nivel',
                           'grado': 'Grado', 'seccion': 'Seccion',
                           'apoderado': 'Apoderado', 'dni_apoderado': 'DNI_Apoderado',
                           'celular_apoderado': 'Celular_Apoderado'}
                df_gs = df_gs.rename(columns=col_map)
                for _, row in df_gs.iterrows():
                    dni = str(row.get('DNI', '')).strip()
                    if dni and len(dni) >= 7 and dni not in indice:
                        r = row.to_dict()
                        r['_tipo'] = 'alumno'
                        indice[dni] = r
    except Exception:
        pass
    st.session_state['_indice_dni'] = indice
    st.session_state['_indice_dni_ts'] = time.time()


def _nombre_completo_docente():
    """Resuelve el nombre COMPLETO del docente actual.
    Busca en: docente_info DNI → Docentes Excel/GS → match inteligente → label → usuario.
    Evita mostrar el nombre de cuenta (ej: 'aucordova') en PDFs y registros."""
    usuario = st.session_state.get('usuario_actual', '')
    info = st.session_state.get('docente_info', {}) or {}
    label_info = str(info.get('label', '')).strip()

    # 1. Si label parece nombre completo (tiene espacio y >8 chars), usarlo
    if label_info and ' ' in label_info and len(label_info) > 8:
        return label_info

    # 2. Buscar por DNI del docente (más confiable)
    dni_doc = str(info.get('dni', '')).strip()
    try:
        df_d = BaseDatos.cargar_docentes()
        if not df_d.empty and 'Nombre' in df_d.columns:
            # 2a. Búsqueda exacta por DNI
            if dni_doc and len(dni_doc) >= 7 and 'DNI' in df_d.columns:
                df_d['DNI'] = df_d['DNI'].astype(str).str.strip()
                match_dni = df_d[df_d['DNI'] == dni_doc]
                if not match_dni.empty:
                    nombre_d = str(match_dni.iloc[0]['Nombre']).strip()
                    if nombre_d and nombre_d.upper() not in ('NAN', 'NONE', ''):
                        _actualizar_label_docente(info, nombre_d)
                        return nombre_d

            # 2b. Búsqueda inteligente por nombre de usuario
            usuario_lower = usuario.replace('.', ' ').strip().lower()
            mejor_match = None
            mejor_score = 0
            for _, row in df_d.iterrows():
                nombre_d = str(row.get('Nombre', '')).strip()
                if not nombre_d or len(nombre_d) < 3:
                    continue
                nombre_lower = nombre_d.lower()
                # Verificar si alguna parte del nombre (>3 chars) aparece en el usuario
                partes_nombre = [p for p in nombre_lower.split() if len(p) > 3]
                score = 0
                for parte in partes_nombre:
                    if parte in usuario_lower:
                        score += len(parte)
                # También verificar si partes del usuario aparecen en el nombre
                partes_usuario = usuario_lower.replace('.', ' ').split()
                for parte_u in partes_usuario:
                    if len(parte_u) > 2 and parte_u in nombre_lower:
                        score += len(parte_u)
                if score > mejor_score:
                    mejor_score = score
                    mejor_match = nombre_d
            
            # Solo aceptar si el score es significativo (>=5 chars coinciden)
            if mejor_match and mejor_score >= 5:
                _actualizar_label_docente(info, mejor_match)
                return mejor_match
    except Exception:
        pass

    # 3. Buscar en usuarios dict
    try:
        usuarios = cargar_usuarios()
        datos_u = usuarios.get(usuario, {})
        lbl = str(datos_u.get('label', '')).strip()
        if lbl and ' ' in lbl and len(lbl) > 8:
            return lbl
    except Exception:
        pass

    # 4. Fallback: label o usuario formateado
    if label_info:
        return label_info
    return usuario.replace('.', ' ').title()


def _actualizar_label_docente(info, nombre_completo):
    """Actualiza docente_info con el nombre completo para futuras consultas."""
    if info:
        info['label'] = nombre_completo
        info['nombre'] = nombre_completo
        st.session_state['docente_info'] = info


# ================================================================
# GENERADOR PDF — DOCUMENTOS (6 tipos)
# CORREGIDO: "Se expide a solicitud del padre/madre/apoderado"
# ================================================================

class GeneradorPDF:
    def __init__(self, config):
        self.config = config
        self.buffer = io.BytesIO()
        self.canvas = canvas.Canvas(self.buffer, pagesize=A4)
        self.width, self.height = A4
        self.styles = getSampleStyleSheet()

    def _fondo(self):
        if Path("fondo.png").exists():
            try:
                self.canvas.drawImage("fondo.png", 0, 0,
                                       width=self.width, height=self.height)
            except Exception:
                pass

    def _marca_agua(self):
        if Path("escudo_upload.png").exists():
            try:
                self.canvas.saveState()
                self.canvas.setFillAlpha(0.35)
                self.canvas.drawImage("escudo_upload.png",
                                       self.width / 2 - 120, self.height / 2 - 120,
                                       240, 240, mask='auto')
                self.canvas.restoreState()
            except Exception:
                pass

    def _encabezado(self, titulo):
        self.canvas.setFont("Helvetica-Oblique", 11)
        self.canvas.drawCentredString(self.width / 2, self.config['y_frase'],
                                       f'"{self.config["frase"]}"')
        hoy = hora_peru()
        self.canvas.setFont("Helvetica", 11)
        self.canvas.drawRightString(
            self.width - 60, self.config['y_frase'] - 25,
            f"Chinchero, {hoy.day} de {MESES_ESP[hoy.month - 1]} de {self.config['anio']}"
        )
        self.canvas.setFont("Helvetica-Bold", 16)
        self.canvas.drawCentredString(self.width / 2, self.config['y_titulo'], titulo)
        self.canvas.line(100, self.config['y_titulo'] - 5,
                         self.width - 100, self.config['y_titulo'] - 5)

    def _parrafo(self, texto, x, y, ancho, estilo):
        p = Paragraph(texto, estilo)
        w, h = p.wrap(ancho, 600)
        p.drawOn(self.canvas, x, y - h)
        return y - h - 15

    def _qr(self, datos, tipo):
        data = (f"YACHAY|{tipo}|{datos.get('alumno', datos.get('Nombre', ''))}|"
                f"{datos.get('dni', datos.get('DNI', ''))}|"
                f"{hora_peru().strftime('%d/%m/%Y')}")
        q = qrcode.QRCode(box_size=10, border=1)
        q.add_data(data)
        q.make(fit=True)
        img = q.make_image(fill_color="black", back_color="white")
        tmp = "tmp_qr.png"
        img.save(tmp)
        self.canvas.drawImage(tmp, self.config['qr_x'], self.config['qr_y'], 70, 70)
        self.canvas.setFont("Helvetica", 6)
        self.canvas.drawCentredString(self.config['qr_x'] + 35,
                                       self.config['qr_y'] - 5, "VERIFICACIÓN")
        try:
            os.remove(tmp)
        except Exception:
            pass

    def _solicitante(self, datos, y):
        """CORREGIDO: Se expide a solicitud del padre/madre/apoderado"""
        apoderado = datos.get('apoderado', datos.get('Apoderado', '')).upper()
        dni_apo = datos.get('dni_apo', datos.get('DNI_Apoderado', ''))
        alumno = datos.get('alumno', datos.get('Nombre', '')).upper()
        e = ParagraphStyle('S', parent=self.styles['Normal'],
                            fontSize=10, leading=14, alignment=TA_JUSTIFY)
        if apoderado and apoderado.strip():
            texto = (f"Se expide el presente documento a solicitud del "
                     f"padre/madre/apoderado(a) <b>{apoderado}</b>, "
                     f"identificado(a) con DNI N° <b>{dni_apo}</b>, "
                     f"en representación del/la estudiante <b>{alumno}</b>.")
        else:
            texto = (f"Se expide el presente documento a solicitud de parte "
                     f"interesada, para los fines que estime conveniente.")
        return self._parrafo(texto, 60, y, self.width - 120, e)

    def _firmas(self):
        yf = 110
        self.canvas.line(200, yf, 395, yf)
        self.canvas.setFont("Helvetica-Bold", 11)
        self.canvas.drawCentredString(self.width / 2, yf - 15,
                                       self.config['directora'].upper())
        self.canvas.setFont("Helvetica", 9)
        self.canvas.drawCentredString(self.width / 2, yf - 28, "DIRECTORA")

    def _fin(self):
        self.canvas.save()
        self.buffer.seek(0)
        return self.buffer

    def generar_constancia_vacante(self, d):
        self._fondo()
        self._encabezado("CONSTANCIA DE VACANTE")
        y = self.config['y_titulo'] - 50
        mx, an = 60, self.width - 120
        e = ParagraphStyle('N', parent=self.styles['Normal'],
                            fontSize=11, leading=15, alignment=TA_JUSTIFY)
        el = ParagraphStyle('L', parent=e, leftIndent=25)
        al = d.get('alumno', d.get('Nombre', '')).upper()
        dni = d.get('dni', d.get('DNI', ''))
        gr = d.get('grado', d.get('Grado', '')).upper()
        y = self._parrafo(
            "La Dirección de la I.E.P. ALTERNATIVO YACHAY de Chinchero, "
            "debidamente representada por su Directora, certifica:", mx, y, an, e
        )
        y = self._parrafo(
            f"Que la I.E. cuenta con <b>VACANTE DISPONIBLE</b> en <b>{gr}</b> "
            f"para el/la estudiante <b>{al}</b>, DNI N° <b>{dni}</b>, "
            f"año escolar <b>{self.config['anio']}</b>.", mx, y, an, e
        )
        y = self._parrafo("Para formalizar la matrícula, presentar:", mx, y, an, e)
        for r in [
            "• Certificado Oficial de Estudios del SIAGIE (original).",
            "• Resolución Directoral de Traslado de Matrícula.",
            "• Libreta de Notas del Sistema SIAGIE.",
            "• Ficha Única de Matrícula del Sistema SIAGIE.",
            "• Copia del DNI del estudiante.",
            "• Constancia de No Adeudo de la IE de procedencia.",
            "• Folder o mica transparente."
        ]:
            y = self._parrafo(r, mx, y, an, el)
        y = self._solicitante(d, y)
        self._firmas()
        self._qr(d, "VACANTE")
        return self._fin()

    def generar_constancia_no_deudor(self, d):
        self._fondo()
        self._encabezado("CONSTANCIA DE NO ADEUDO")
        y = self.config['y_titulo'] - 50
        mx, an = 60, self.width - 120
        e = ParagraphStyle('N', parent=self.styles['Normal'],
                            fontSize=11, leading=15, alignment=TA_JUSTIFY)
        al = d.get('alumno', d.get('Nombre', '')).upper()
        dni = d.get('dni', d.get('DNI', ''))
        y = self._parrafo(
            "La Dirección de la I.E.P. ALTERNATIVO YACHAY:", mx, y, an, e
        )
        y = self._parrafo(
            f"Que el/la estudiante <b>{al}</b>, DNI N° <b>{dni}</b>, "
            f"ha cumplido con todas sus obligaciones económicas, "
            f"no registrando deuda alguna.", mx, y, an, e
        )
        y = self._solicitante(d, y)
        self._firmas()
        self._qr(d, "NO ADEUDO")
        return self._fin()

    def generar_constancia_estudios(self, d):
        self._fondo()
        self._encabezado("CONSTANCIA DE ESTUDIOS")
        y = self.config['y_titulo'] - 50
        mx, an = 60, self.width - 120
        e = ParagraphStyle('N', parent=self.styles['Normal'],
                            fontSize=11, leading=15, alignment=TA_JUSTIFY)
        al = d.get('alumno', d.get('Nombre', '')).upper()
        dni = d.get('dni', d.get('DNI', ''))
        gr = d.get('grado', d.get('Grado', '')).upper()
        y = self._parrafo(
            "La Dirección de la I.E.P. ALTERNATIVO YACHAY:", mx, y, an, e
        )
        y = self._parrafo(
            f"Que <b>{al}</b>, DNI N° <b>{dni}</b>, se encuentra "
            f"<b>MATRICULADO(A)</b> año <b>{self.config['anio']}</b>, "
            f"cursando <b>{gr}</b>, conforme consta en registros oficiales "
            f"y el Sistema SIAGIE.", mx, y, an, e
        )
        y = self._solicitante(d, y)
        self._firmas()
        self._qr(d, "ESTUDIOS")
        return self._fin()

    def generar_constancia_conducta(self, d):
        self._fondo()
        self._encabezado("CONSTANCIA DE CONDUCTA")
        y = self.config['y_titulo'] - 50
        mx, an = 60, self.width - 120
        e = ParagraphStyle('N', parent=self.styles['Normal'],
                            fontSize=10, leading=14, alignment=TA_JUSTIFY)
        al = d.get('alumno', d.get('Nombre', '')).upper()
        dni = d.get('dni', d.get('DNI', ''))
        y = self._parrafo(
            f"Que <b>{al}</b>, DNI N° <b>{dni}</b>, obtuvo en CONDUCTA:",
            mx, y, an, e
        )
        y -= 15
        tx = self.width / 2 - 200
        self.canvas.setFont("Helvetica-Bold", 10)
        self.canvas.drawString(tx, y, "GRADO")
        self.canvas.drawString(tx + 120, y, "AÑO")
        self.canvas.drawString(tx + 280, y, "CALIFICACIÓN")
        y -= 5
        self.canvas.line(tx - 10, y, tx + 380, y)
        y -= 20
        self.canvas.setFont("Helvetica", 9)
        ab = int(self.config['anio']) - 5
        for i, g in enumerate(["PRIMERO", "SEGUNDO", "TERCERO", "CUARTO", "QUINTO"]):
            n = d.get(f'nota_conducta_{i+1}', 'AD')
            self.canvas.drawString(tx, y, g)
            self.canvas.drawString(tx + 120, y, str(ab + i + 1))
            self.canvas.drawString(tx + 280, y, n)
            y -= 18
        y -= 10
        y = self._solicitante(d, y)
        self._firmas()
        self._qr(d, "CONDUCTA")
        return self._fin()

    def generar_carta_compromiso(self, d):
        self._fondo()
        self._encabezado("CARTA DE COMPROMISO")
        y = self.config['y_titulo'] - 40
        mx, an = 50, self.width - 100
        e = ParagraphStyle('C', parent=self.styles['Normal'],
                            fontSize=8.5, leading=11, alignment=TA_JUSTIFY)
        ei = ParagraphStyle('I', parent=e, leftIndent=10)
        apo = d.get('apoderado', d.get('Apoderado', '')).upper()
        dapo = d.get('dni_apo', d.get('DNI_Apoderado', ''))
        al = d.get('alumno', d.get('Nombre', '')).upper()
        gr = d.get('grado', d.get('Grado', '')).upper()
        y = self._parrafo(
            f"Yo, <b>{apo}</b>, DNI N° <b>{dapo}</b>, "
            f"padre/madre/apoderado(a) de <b>{al}</b>, del <b>{gr}</b>, "
            f"me comprometo a:", mx, y, an, e
        )
        for c in [
            "1. Velar por la asistencia puntual de mi hijo(a).",
            "2. Supervisar el cumplimiento de tareas.",
            "3. Asegurar asistencia uniformado(a).",
            "4. Inculcar respeto hacia docentes y compañeros.",
            "5. Participar en actividades del comité de aula.",
            "6. Ejercer crianza positiva, libre de violencia.",
            "7. Atender problemas de conducta oportunamente.",
            "8. Asumir responsabilidad por daños materiales.",
            "9. Vigilar vocabulario apropiado.",
            "10. Acudir cuando sea requerido(a).",
            "11. Asistir puntualmente a reuniones.",
            "12. Justificar inasistencias en 24 horas.",
            "13. Cumplir pagos de pensiones.",
            "14. Respetar la autonomía pedagógica."
        ]:
            y = self._parrafo(c, mx, y, an, ei)
            y += 2
        y = 120
        self.canvas.line(80, y, 200, y)
        self.canvas.line(220, y, 340, y)
        self.canvas.line(360, y, 480, y)
        y -= 10
        self.canvas.setFont("Helvetica-Bold", 7)
        self.canvas.drawCentredString(140, y, "FIRMA APODERADO")
        self.canvas.drawCentredString(280, y, self.config['directora'].upper())
        self.canvas.drawCentredString(280, y - 10, "DIRECTORA")
        self.canvas.drawCentredString(420, y, self.config['promotor'].upper())
        self.canvas.drawCentredString(420, y - 10, "PROMOTOR")
        return self._fin()

    def generar_resolucion_traslado(self, d):
        self._fondo()
        self.canvas.setFont("Helvetica-Oblique", 11)
        self.canvas.drawCentredString(self.width / 2, 700,
                                       f'"{self.config["frase"]}"')
        self.canvas.setFont("Helvetica-Bold", 14)
        self.canvas.drawCentredString(self.width / 2, 670,
                                       f"RESOLUCIÓN DIRECTORAL N° {d.get('num_resolucion', '')}")
        self.canvas.setFont("Helvetica", 11)
        self.canvas.drawCentredString(self.width / 2, 640,
                                       d.get('fecha_resolucion', ''))
        mx, an = 60, self.width - 120
        e = ParagraphStyle('N', parent=self.styles['Normal'],
                            fontSize=11, leading=15, alignment=TA_JUSTIFY)
        al = d.get('alumno', d.get('Nombre', '')).upper()
        niv = d.get('nivel', '').upper()
        y = 600
        self.canvas.setFont("Helvetica-Bold", 11)
        self.canvas.drawString(mx, y, "SE RESUELVE:")
        y -= 20
        t = Table([
            ['ALUMNO', al],
            ['NIVEL', niv],
            ['IE PROCEDENCIA', 'IEP ALTERNATIVO YACHAY'],
            ['IE DESTINO', d.get('ie_destino', '').upper()]
        ], colWidths=[200, 280])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        t.wrapOn(self.canvas, an, 200)
        t.drawOn(self.canvas, mx, y - 80)
        self._firmas()
        self._qr(d, "TRASLADO")
        return self._fin()


# ================================================================
# REGISTRO AUXILIAR PDF — 3 Cursos × 4 Competencias × 3 Desempeños
# ================================================================

def generar_registro_auxiliar_pdf(grado, seccion, anio, bimestre,
                                  estudiantes_df, cursos=None):
    if cursos is None:
        cursos = ["Matemática", "Comunicación", "Ciencia y Tec."]
    nc = len(cursos)
    dp = 3  # desempeños por competencia
    cp = 4  # competencias por curso
    total_d = nc * cp * dp
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=landscape(A4))
    w, h = landscape(A4)
    if Path("escudo_upload.png").exists():
        try:
            c.saveState()
            c.setFillAlpha(0.35)
            c.drawImage("escudo_upload.png", w / 2 - 100, h / 2 - 100,
                        200, 200, mask='auto')
            c.restoreState()
        except Exception:
            pass
    c.setFont("Helvetica-Bold", 11)
    c.drawCentredString(w / 2, h - 22,
                        "I.E.P. ALTERNATIVO YACHAY - REGISTRO AUXILIAR DE EVALUACIÓN")
    c.setFont("Helvetica", 8)
    c.drawCentredString(w / 2, h - 35,
                        f"Grado: {grado} | Sección: {seccion} | {bimestre} | Año: {anio}")
    c.setFont("Helvetica-Oblique", 7)
    c.drawCentredString(w / 2, h - 47,
                        '"Educar para la Vida — Pioneros en la Educación de Calidad"')

    cols_per_c = cp * dp
    r0 = ["N°", "APELLIDOS Y NOMBRES"]
    for curso in cursos:
        r0.append(curso.upper())
        r0.extend([""] * (cols_per_c - 1))
    r1 = ["", ""]
    for _ in range(nc):
        for ci in range(1, cp + 1):
            r1.append(f"C{ci}")
            r1.extend([""] * (dp - 1))
    r2 = ["", ""]
    for _ in range(nc):
        for _ in range(cp):
            for di in range(1, dp + 1):
                r2.append(f"D{di}")

    if not estudiantes_df.empty:
        est = estudiantes_df.sort_values('Nombre').reset_index(drop=True)
    else:
        est = pd.DataFrame()
    data = [r0, r1, r2]
    ne = len(est) if not est.empty else 25
    for idx in range(ne):
        nm = est.iloc[idx].get('Nombre', '') if idx < len(est) else ""
        if len(nm) > 28:
            nm = nm[:28] + "."
        data.append([str(idx + 1), nm] + [""] * total_d)

    avail = w - 30
    wn = 16
    wname = 115
    wd = max(16, min(25, (avail - wn - wname) / total_d))
    cw = [wn, wname] + [wd] * total_d
    # Calcular alturas para que todo quepa en la hoja sin desbordar
    ESPACIO_HEADER_AUX = 58   # título + subtítulos
    ESPACIO_PIE_AUX   = 20   # leyenda al pie
    altura_disp = h - ESPACIO_HEADER_AUX - ESPACIO_PIE_AUX
    total_filas_aux = 3 + ne   # 3 cabeceras + alumnos
    fila_h_aux = max(9, min(16, altura_disp / total_filas_aux))
    cab_h_aux  = min(18, fila_h_aux * 1.3)
    row_heights_aux = [cab_h_aux, cab_h_aux, cab_h_aux] + [fila_h_aux] * ne
    tabla = Table(data, colWidths=cw, rowHeights=row_heights_aux, repeatRows=3)
    sl = [
        ('FONTNAME', (0, 0), (-1, 2), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 2), 5),
        ('FONTSIZE', (0, 3), (-1, -1), 5.5),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('ALIGN', (1, 3), (1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('BACKGROUND', (0, 0), (1, 2), colors.Color(0.1, 0.1, 0.35)),
        ('TEXTCOLOR', (0, 0), (1, 2), colors.white),
        ('ROWBACKGROUNDS', (0, 3), (-1, -1),
         [colors.white, colors.Color(0.95, 0.95, 1)]),
    ]
    colores_c = [
        colors.Color(0, 0.2, 0.5),
        colors.Color(0.15, 0.35, 0.15),
        colors.Color(0.4, 0.15, 0.15)
    ]
    for ci, curso in enumerate(cursos):
        cs = 2 + ci * cols_per_c
        ce = cs + cols_per_c - 1
        sl.append(('SPAN', (cs, 0), (ce, 0)))
        bg = colores_c[ci % len(colores_c)]
        sl.append(('BACKGROUND', (cs, 0), (ce, 0), bg))
        sl.append(('TEXTCOLOR', (cs, 0), (ce, 0), colors.white))
        for ki in range(cp):
            s = cs + ki * dp
            e = s + dp - 1
            sl.append(('SPAN', (s, 1), (e, 1)))
            bg2 = colors.Color(min(bg.red + 0.1, 1),
                               min(bg.green + 0.1, 1),
                               min(bg.blue + 0.1, 1))
            sl.append(('BACKGROUND', (s, 1), (e, 1), bg2))
            sl.append(('TEXTCOLOR', (s, 1), (e, 1), colors.white))
            sl.append(('BACKGROUND', (s, 2), (e, 2), bg2))
            sl.append(('TEXTCOLOR', (s, 2), (e, 2), colors.white))
    tabla.setStyle(TableStyle(sl))
    tw, th = tabla.wrap(w - 20, h - 70)
    tabla.drawOn(c, 10, ESPACIO_PIE_AUX)
    c.setFont("Helvetica", 5)
    c.drawString(10, 12,
                 f"C=Competencia | D=Desempeño | AD(18-20) A(14-17) "
                 f"B(11-13) C(0-10) | {bimestre} | YACHAY PRO — {anio}")
    c.save()
    buffer.seek(0)
    return buffer


# ================================================================
# REGISTRO ASISTENCIA PDF (sin sáb/dom, sin feriados + pie feriados)
# ================================================================

def generar_registro_auxiliar_docx(grado, seccion, anio, bimestre, estudiantes_df, cursos=None):
    """Genera el registro auxiliar en formato Word (.docx)"""
    try:
        from docx import Document as DocxDoc
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy
    except ImportError:
        return None

    if cursos is None:
        cursos = ["Matemática", "Comunicación", "Ciencia y Tec."]

    nc = len(cursos)
    dp = 3  # desempeños por competencia
    cp = 4  # competencias por curso
    cols_per_c = cp * dp

    doc = DocxDoc()
    section = doc.sections[0]
    # Hoja A4 horizontal
    section.page_width  = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = section.right_margin = Cm(1.0)
    section.top_margin  = section.bottom_margin = Cm(1.0)

    # Título
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = titulo.add_run("I.E.P. ALTERNATIVO YACHAY — REGISTRO AUXILIAR DE EVALUACIÓN")
    run_t.bold = True; run_t.font.size = Pt(10)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Grado: {grado} | Sección: {seccion} | {bimestre} | Año: {anio}").font.size = Pt(8)

    # Preparar cabeceras
    r0 = ["N°", "APELLIDOS Y NOMBRES"]
    for curso in cursos:
        r0.append(curso.upper())
        r0.extend([""] * (cols_per_c - 1))
    r1 = ["", ""]
    for _ in range(nc):
        for ci in range(1, cp + 1):
            r1.append(f"C{ci}"); r1.extend([""] * (dp - 1))
    r2 = ["", ""]
    for _ in range(nc):
        for _ in range(cp):
            for di in range(1, dp + 1):
                r2.append(f"D{di}")

    if not estudiantes_df.empty:
        est = estudiantes_df.sort_values('Nombre').reset_index(drop=True)
    else:
        est = pd.DataFrame()
    ne = len(est) if not est.empty else 25

    total_cols = 2 + nc * cols_per_c
    table = doc.add_table(rows=3 + ne, cols=total_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Anchos de columna (en twips, 1cm=567)
    col_widths_cm = [0.7, 4.5] + [0.55] * (nc * cols_per_c)

    def set_cell(cell, text, bold=False, size=5, bg_hex=None, align_center=True):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(str(text))
        run.bold = bold; run.font.size = Pt(size)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if bg_hex:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), bg_hex)
            tcPr.append(shd)

    COLORES_CURSOS = ['1A237E', '1B5E20', '7B1FA2']

    # Fila 0 — cursos
    row0 = table.rows[0]
    set_cell(row0.cells[0], "N°", bold=True, size=5, bg_hex='1A1A5C')
    set_cell(row0.cells[1], "APELLIDOS Y NOMBRES", bold=True, size=5, bg_hex='1A1A5C')
    for ci, curso in enumerate(cursos):
        cs = 2 + ci * cols_per_c
        color_c = COLORES_CURSOS[ci % len(COLORES_CURSOS)]
        # Merge cells for curso header
        cell_start = row0.cells[cs]
        cell_end   = row0.cells[cs + cols_per_c - 1]
        cell_start.merge(cell_end)
        set_cell(cell_start, curso.upper(), bold=True, size=5, bg_hex=color_c)

    # Fila 1 — competencias
    row1 = table.rows[1]
    set_cell(row1.cells[0], "", bg_hex='1A1A5C'); set_cell(row1.cells[1], "", bg_hex='1A1A5C')
    for ci in range(nc):
        color_c = COLORES_CURSOS[ci % len(COLORES_CURSOS)]
        cs = 2 + ci * cols_per_c
        for ki in range(cp):
            s = cs + ki * dp
            cell_s = row1.cells[s]; cell_e = row1.cells[s + dp - 1]
            cell_s.merge(cell_e)
            set_cell(cell_s, f"C{ki+1}", bold=True, size=5, bg_hex=color_c)

    # Fila 2 — desempeños
    row2 = table.rows[2]
    set_cell(row2.cells[0], "", bg_hex='1A1A5C'); set_cell(row2.cells[1], "", bg_hex='1A1A5C')
    for ci in range(nc):
        color_c = COLORES_CURSOS[ci % len(COLORES_CURSOS)]
        cs = 2 + ci * cols_per_c
        for ki in range(cp):
            for di in range(dp):
                set_cell(row2.cells[cs + ki*dp + di], f"D{di+1}", bold=True, size=5, bg_hex=color_c)

    # Filas de alumnos
    for idx in range(ne):
        rw = table.rows[3 + idx]
        nm = est.iloc[idx].get('Nombre', '') if idx < len(est) else ""
        if len(nm) > 30: nm = nm[:30] + "."
        bg = 'FFFFFF' if idx % 2 == 0 else 'EEEEEE'
        set_cell(rw.cells[0], str(idx+1), size=5, bg_hex=bg)
        set_cell(rw.cells[1], nm, size=5, bg_hex=bg, align_center=False)
        for col in range(2, total_cols):
            set_cell(rw.cells[col], "", size=5, bg_hex=bg)

    # Ancho de columnas
    from docx.oxml import OxmlElement as OE
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OE('w:tblPr'); tbl.insert(0, tblPr)
    tblW = OE('w:tblW'); tblW.set(qn('w:w'), '0'); tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = OE('w:tcW')
            w_val = int(col_widths_cm[j] * 567) if j < len(col_widths_cm) else int(0.55 * 567)
            tcW.set(qn('w:w'), str(w_val)); tcW.set(qn('w:type'), 'dxa')
            existing = tcPr.find(qn('w:tcW'))
            if existing is not None: tcPr.remove(existing)
            tcPr.insert(0, tcW)
        trPr = row._tr.get_or_add_trPr()
        trH = OE('w:trHeight'); trH.set(qn('w:val'), '200'); trH.set(qn('w:hRule'), 'atLeast')
        existing_h = trPr.find(qn('w:trHeight'))
        if existing_h is not None: trPr.remove(existing_h)
        trPr.append(trH)

    doc.add_paragraph().add_run(
        f"C=Competencia | D=Desempeño | AD(18-20) A(14-17) B(11-13) C(0-10) | {bimestre} | YACHAY PRO — {anio}"
    ).font.size = Pt(6)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generar_registro_asistencia_pdf(grado, seccion, anio, estudiantes_df,
                                     meses_sel, docente=""):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=landscape(A4))
    w, h = landscape(A4)
    if not estudiantes_df.empty:
        est = estudiantes_df.sort_values('Nombre').reset_index(drop=True)
    else:
        est = pd.DataFrame()
    ds = {0: "L", 1: "M", 2: "Mi", 3: "J", 4: "V"}

    # Colores pastel por semana (suaves para llenar a mano)
    COLORES_SEMANA = [
        colors.Color(0.88, 0.95, 1.0),    # Sem 1: celeste pastel
        colors.Color(0.93, 1.0, 0.88),     # Sem 2: verde menta pastel
        colors.Color(1.0, 0.95, 0.85),     # Sem 3: durazno pastel
        colors.Color(0.95, 0.90, 1.0),     # Sem 4: lavanda pastel
        colors.Color(1.0, 0.92, 0.92),     # Sem 5: rosa pastel
    ]
    COLORES_SEMANA_HEADER = [
        colors.Color(0.70, 0.85, 0.95),    # Sem 1: celeste claro
        colors.Color(0.72, 0.92, 0.72),    # Sem 2: verde claro
        colors.Color(0.95, 0.82, 0.55),    # Sem 3: naranja claro
        colors.Color(0.82, 0.75, 0.95),    # Sem 4: lavanda claro
        colors.Color(0.95, 0.72, 0.72),    # Sem 5: rosa claro
    ]

    for mi, mn in enumerate(meses_sel):
        if mi > 0:
            c.showPage()
        mnm = MESES_ESCOLARES.get(mn, f"Mes {mn}")
        if Path("escudo_upload.png").exists():
            try:
                c.saveState()
                c.setFillAlpha(0.35)
                c.drawImage("escudo_upload.png", w / 2 - 100, h / 2 - 100,
                            200, 200, mask='auto')
                c.restoreState()
            except Exception:
                pass
        c.setFont("Helvetica-Bold", 11)
        c.drawCentredString(w / 2, h - 22,
                            "I.E.P. ALTERNATIVO YACHAY - REGISTRO DE ASISTENCIA")
        c.setFont("Helvetica", 8)
        info_line = (f"Grado: {grado} | Sección: {seccion} | "
                     f"Mes: {mnm} | Año: {anio}")
        if docente:
            info_line += f" | Docente: {docente}"
        c.drawCentredString(w / 2, h - 35, info_line)
        dias = dias_habiles_mes(int(anio), mn)
        nd = len(dias)

        # Calcular semana de cada día (1-5 dentro del mes)
        semana_de_dia = []
        for d in dias:
            semana_idx = (d - 1) // 7  # 0-based: days 1-7→0, 8-14→1, etc.
            if semana_idx > 4:
                semana_idx = 4
            semana_de_dia.append(semana_idx)

        header = ["N°", "APELLIDOS Y NOMBRES"]
        for d in dias:
            dt = date(int(anio), mn, d)
            header.append(f"{d}\n{ds[dt.weekday()]}")
        header.extend(["A", "T", "F", "J"])
        data = [header]
        ne = len(est) if not est.empty else 25
        for idx in range(ne):
            nm = est.iloc[idx].get('Nombre', '') if idx < len(est) else ""
            if len(nm) > 32:
                nm = nm[:32] + "."
            data.append([str(idx + 1), nm] + [""] * nd + ["", "", "", ""])
        dw = max(15, min(22, (w - 18 - 140 - 72 - 30) / max(nd, 1)))
        cw = [18, 140] + [dw] * nd + [18, 18, 18, 18]

        # Calcular altura disponible para la tabla (reservar espacio header + leyenda)
        ESPACIO_HEADER = 48   # título + subtítulo arriba
        ESPACIO_LEYENDA = 30  # leyenda de semanas abajo
        altura_disponible = h - ESPACIO_HEADER - ESPACIO_LEYENDA

        # Filas totales: 1 cabecera + alumnos
        total_filas = 1 + ne
        # Altura de fila: que quepan todas, mínimo 10pt máximo 18pt
        fila_h_auto = altura_disponible / total_filas
        fila_h = max(10, min(18, fila_h_auto))
        cabecera_h = min(22, fila_h * 1.4)

        row_heights = [cabecera_h] + [fila_h] * ne
        t = Table(data, colWidths=cw, rowHeights=row_heights, repeatRows=1)

        # Estilos base
        estilos = [
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 5),
            ('FONTSIZE', (0, 1), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 0.4, colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('ALIGN', (1, 1), (1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            # N° y Nombres header
            ('BACKGROUND', (0, 0), (1, 0), colors.Color(0, 0.3, 0.15)),
            ('TEXTCOLOR', (0, 0), (1, 0), colors.white),
            # Columnas A/T/F/J header
            ('BACKGROUND', (-4, 0), (-1, 0), colors.Color(0.6, 0, 0)),
            ('TEXTCOLOR', (-4, 0), (-1, 0), colors.white),
            # Columnas A/T/F/J data (gris claro)
            ('BACKGROUND', (-4, 1), (-1, -1), colors.Color(0.96, 0.96, 0.96)),
        ]

        # Aplicar colores por semana a cada columna de día
        for di, sem_idx in enumerate(semana_de_dia):
            col = di + 2  # +2 por N° y Nombre
            # Header: color fuerte de la semana
            estilos.append(('BACKGROUND', (col, 0), (col, 0),
                           COLORES_SEMANA_HEADER[sem_idx]))
            estilos.append(('TEXTCOLOR', (col, 0), (col, 0), colors.black))
            # Data: color pastel suave de la semana
            estilos.append(('BACKGROUND', (col, 1), (col, -1),
                           COLORES_SEMANA[sem_idx]))

        t.setStyle(TableStyle(estilos))
        tw, th2 = t.wrap(w - 20, h - 60)
        # Siempre dibujar la tabla justo encima de la leyenda (30pt desde abajo)
        tabla_y = ESPACIO_LEYENDA
        t.drawOn(c, 10, tabla_y)

        # Leyenda de semanas
        fer = feriados_del_mes(mn)
        c.setFont("Helvetica", 5)
        # Leyenda de colores
        ley_y = 18
        c.drawString(10, ley_y, "SEMANAS: ")
        ley_x = 55
        for si in range(5):
            # Ver si hay días de esta semana
            if si in semana_de_dia:
                c.setFillColor(COLORES_SEMANA_HEADER[si])
                c.rect(ley_x, ley_y - 1, 8, 7, fill=1, stroke=0)
                c.setFillColor(colors.black)
                c.drawString(ley_x + 10, ley_y, f"Sem {si+1}")
                ley_x += 40
        c.setFillColor(colors.black)
        pie = (" | A=Asistió | T=Tardanza | F=Falta | J=Justificada | "
               "Sin sáb/dom/feriados")
        if fer:
            pie += f" | FERIADOS: {', '.join(fer)}"
        c.drawString(ley_x + 5, ley_y, pie)
    c.save()
    buffer.seek(0)
    return buffer


# ================================================================
# RANKING PDF — COLUMNAS FIJAS + COLORES POR ÁREA
# ================================================================

def generar_ranking_pdf(resultados, anio):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    w, h = A4

    # ── Marca de agua UNA SOLA — zona inferior vacía ─────────────────────
    if Path("escudo_upload.png").exists():
        try:
            from PIL import Image as PILImage
            img = PILImage.open("escudo_upload.png")
            iw, ih = img.size
            mw = 220; mh = mw / (iw/ih)
            c.saveState()
            c.setFillAlpha(0.20)
            c.drawImage("escudo_upload.png", w/2-mw/2, 30, mw, mh, mask='auto')
            c.restoreState()
        except Exception:
            pass

    # Barra azul
    c.setFillColor(colors.HexColor("#001e7c"))
    c.rect(0, h-15, w, 15, fill=1, stroke=0)

    # ── Escudos — 60px izq, 60px der, separados del texto central ────────
    ALTO_ESC = 60
    esc_izq = "escudo_upload.png"
    esc_der = "escudo2_upload.png" if Path("escudo2_upload.png").exists() else "escudo_upload.png"
    try:
        from PIL import Image as PILImage
        if Path(esc_izq).exists():
            img = PILImage.open(esc_izq)
            iw, ih = img.size
            aw = ALTO_ESC * (iw/ih)
            c.drawImage(esc_izq, 18, h-12-ALTO_ESC, aw, ALTO_ESC, mask='auto')
        if Path(esc_der).exists():
            img2 = PILImage.open(esc_der)
            iw2, ih2 = img2.size
            aw2 = ALTO_ESC * (iw2/ih2)
            c.drawImage(esc_der, w-18-aw2, h-12-ALTO_ESC, aw2, ALTO_ESC, mask='auto')
    except Exception:
        pass

    # ── Textos institucionales centrados ─────────────────────────────────
    c.setFillColor(colors.HexColor("#001e7c"))
    c.setFont("Helvetica-Bold", 8)
    c.drawCentredString(w/2, h-28, "MINISTERIO DE EDUCACIÓN — DRE CUSCO — PIONEROS EN LA EDUCACION DE CALIDAD")
    c.setFont("Helvetica-Bold", 12)
    c.drawCentredString(w/2, h-43, "I.E.P. YACHAY — CHINCHERO")
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(w/2, h-60, f"RANKING DE RESULTADOS — {anio}")
    c.setFont("Helvetica", 9)
    c.setFillColor(colors.HexColor("#6b7280"))
    c.drawCentredString(w/2, h-75, f"Generado: {hora_peru().strftime('%d/%m/%Y %H:%M')}")
    c.setFillColor(colors.black)

    rk = sorted(resultados,
                key=lambda r: r.get('promedio_general', 0), reverse=True)
    all_a = set()
    for r in rk:
        for a in r.get('areas', []):
            all_a.add(a['nombre'])
    all_a = sorted(all_a)

    header = ["#", "APELLIDOS Y NOMBRES", "DNI"]
    header.extend(all_a)
    header.append("PROM.")
    data = [header]
    for idx, r in enumerate(rk):
        nm = r.get('nombre', '')
        if len(nm) > 30:
            nm = nm[:30] + "."
        fila = [str(idx + 1), nm, r.get('dni', '')]
        an_map = {a['nombre']: str(a['nota']) for a in r.get('areas', [])}
        for a in all_a:
            fila.append(an_map.get(a, '-'))
        fila.append(str(r.get('promedio_general', 0)))
        data.append(fila)

    na = len(all_a)
    # Anchos fijos: #=20, Nombre=150, DNI=55, Áreas=50 cada una, Prom=45
    cw = [20, 150, 55] + [50] * na + [45]
    t = Table(data, colWidths=cw, repeatRows=1)
    st_l = [
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 7),
        ('FONTSIZE', (0, 1), (-1, -1), 7),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('ALIGN', (1, 1), (1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('BACKGROUND', (0, 0), (2, 0), colors.Color(0.1, 0.1, 0.4)),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('BACKGROUND', (-1, 0), (-1, 0), colors.Color(0.3, 0, 0.3)),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1),
         [colors.white, colors.Color(0.95, 0.95, 1)]),
    ]
    # Colores diferentes por área
    colores_areas = [
        colors.Color(0, 0.3, 0.6),       # Azul
        colors.Color(0.2, 0.5, 0.1),      # Verde
        colors.Color(0.6, 0.2, 0),         # Naranja
        colors.Color(0.4, 0, 0.4),         # Morado
        colors.Color(0, 0.4, 0.4),         # Teal
        colors.Color(0.5, 0.3, 0),         # Marrón
        colors.Color(0.3, 0.1, 0.5),       # Índigo
        colors.Color(0.6, 0, 0.2),         # Rosa oscuro
    ]
    for i in range(na):
        col_idx = 3 + i
        bg = colores_areas[i % len(colores_areas)]
        st_l.append(('BACKGROUND', (col_idx, 0), (col_idx, 0), bg))
    # Top 3
    bg_top = [
        colors.Color(1, 0.84, 0),
        colors.Color(0.75, 0.75, 0.75),
        colors.Color(0.8, 0.5, 0.2),
    ]
    for i in range(min(3, len(rk))):
        st_l.append(('BACKGROUND', (0, i + 1), (-1, i + 1), bg_top[i]))
    t.setStyle(TableStyle(st_l))
    tw, th2 = t.wrap(w - 40, h - 150)
    t.drawOn(c, 20, h - 120 - th2)
    c.setFont("Helvetica", 7)
    c.drawCentredString(w / 2, 30,
                        f"YACHAY PRO — {hora_peru().strftime('%d/%m/%Y %H:%M')}")
    c.save()
    buffer.seek(0)
    return buffer


# ================================================================
# GENERADOR DE CARNETS
# ================================================================

class GeneradorCarnet:
    WIDTH = 1012
    HEIGHT = 638
    AZUL = (0, 30, 120)
    DORADO = (255, 215, 0)

    def __init__(self, datos, anio, foto_bytes=None, es_docente=False):
        self.datos = datos
        self.anio = anio
        self.foto_bytes = foto_bytes
        self.es_docente = es_docente
        self.img = Image.new('RGB', (self.WIDTH, self.HEIGHT), 'white')
        self.draw = ImageDraw.Draw(self.img)

    def _escudo_fondo(self):
        if Path("escudo_upload.png").exists():
            try:
                esc = Image.open("escudo_upload.png").convert("RGBA")
                esc = esc.resize((280, 280), Image.LANCZOS)
                capa = Image.new('RGBA', (self.WIDTH, self.HEIGHT), (255, 255, 255, 0))
                capa.paste(esc, ((self.WIDTH - 280) // 2, (self.HEIGHT - 280) // 2))
                px = [(d[0], d[1], d[2], min(d[3], 28)) for d in capa.getdata()]
                capa.putdata(px)
                self.img = Image.alpha_composite(
                    self.img.convert('RGBA'), capa
                ).convert('RGB')
                self.draw = ImageDraw.Draw(self.img)
            except Exception:
                pass

    def _barras(self):
        self.draw.rectangle([(0, 0), (self.WIDTH, 210)], fill=self.AZUL)
        self.draw.rectangle([(0, 207), (self.WIDTH, 213)], fill=self.DORADO)
        self.draw.rectangle([(0, self.HEIGHT - 65), (self.WIDTH, self.HEIGHT)],
                            fill=self.AZUL)
        self.draw.rectangle([(0, self.HEIGHT - 68), (self.WIDTH, self.HEIGHT - 63)],
                            fill=self.DORADO)

    def _textos(self):
        fh = RecursoManager.obtener_fuente("", 54, True)
        fm = RecursoManager.obtener_fuente("", 26, True)
        fc = RecursoManager.obtener_fuente("", 18, True)
        fp = RecursoManager.obtener_fuente("", 22, True)
        self.draw.text((self.WIDTH // 2, 65), "I.E. ALTERNATIVO YACHAY",  # Cambiado texto
                       font=fh, fill="white", anchor="mm")
        self.draw.text((self.WIDTH // 2, 115), '"EDUCAR PARA LA VIDA"',
                       font=fm, fill=self.DORADO, anchor="mm")
        tt = "CARNET DOCENTE" if self.es_docente else "CARNET ESCOLAR"
        self.draw.text((self.WIDTH // 2, 160), f"{tt} {self.anio}",
                       font=fc, fill="white", anchor="mm")
        self.draw.text((self.WIDTH // 2, self.HEIGHT - 35),
                       "PIONEROS EN LA EDUCACIÓN DE CALIDAD",
                       font=fp, fill=self.DORADO, anchor="mm")

    def _foto(self):
        x, y, wf, hf = 40, 228, 220, 280
        if self.foto_bytes:
            try:
                f = Image.open(self.foto_bytes).convert("RGB")
                self.img.paste(f.resize((wf, hf), Image.LANCZOS), (x, y))
            except Exception:
                self._ph(x, y, wf, hf)
        else:
            self._ph(x, y, wf, hf)
        self.draw.rectangle([(x - 3, y - 3), (x + wf + 3, y + hf + 3)],
                            outline=self.DORADO, width=4)

    def _ph(self, x, y, w, h):
        """Avatar por defecto según sexo del estudiante"""
        sexo = self.datos.get('Sexo', 'Masculino')
        if sexo == 'Femenino':
            bg_color = "#fce4ec"
            icon_color = "#e91e63"
            text_icon = "👩"
        else:
            bg_color = "#e3f2fd"
            icon_color = "#1565c0"
            text_icon = "👨"
        self.draw.rectangle([(x, y), (x + w, y + h)], fill=bg_color)
        # Silueta simple
        cx, cy = x + w // 2, y + h // 2
        # Cabeza
        head_r = min(w, h) // 6
        self.draw.ellipse([(cx - head_r, cy - head_r - 30),
                           (cx + head_r, cy + head_r - 30)],
                          fill=icon_color)
        # Cuerpo
        body_w = min(w, h) // 3
        self.draw.ellipse([(cx - body_w, cy + head_r - 10),
                           (cx + body_w, cy + head_r + body_w + 20)],
                          fill=icon_color)
        # Texto
        try:
            fn = RecursoManager.obtener_fuente("", 11)
            self.draw.text((cx, y + h - 15), "FOTO PENDIENTE",
                           font=fn, fill="#666", anchor="mm")
        except Exception:
            pass

    def _datos(self):
        xt = 290
        nm = self.datos.get('Nombre', self.datos.get('alumno', '')).upper()
        dni = str(self.datos.get('DNI', self.datos.get('dni', '')))
        fn = RecursoManager.obtener_fuente("", 56 if len(nm) > 25 else 68, True)
        fl = RecursoManager.obtener_fuente("", 42, True)
        fd = RecursoManager.obtener_fuente("", 42)
        yc = 240
        if len(nm) > 28:
            for l in textwrap.TextWrapper(width=28).wrap(nm)[:3]:
                self.draw.text((xt, yc), l, font=fn, fill="black")
                yc += 62
        else:
            self.draw.text((xt, yc), nm, font=fn, fill="black")
            yc += 70
        yc += 8
        # DNI prominente
        self.draw.text((xt, yc), "DNI:", font=fl, fill=self.AZUL)
        fd_dni = RecursoManager.obtener_fuente("", 52, True)
        self.draw.text((xt + 100, yc), dni, font=fd_dni, fill="black")
        yc += 56
        if self.es_docente:
            cg = self.datos.get('Cargo', 'DOCENTE').upper()
            self.draw.text((xt, yc), "CARGO:", font=fl, fill="black")
            self.draw.text((xt + 130, yc), cg, font=fd, fill="black")
            yc += 50
            esp = self.datos.get('Especialidad', '').upper()
            if esp:
                self.draw.text((xt, yc), "ESPEC.:", font=fl, fill="black")
                self.draw.text((xt + 140, yc), esp[:20], font=fd, fill="black")
                yc += 50
        else:
            gr = self.datos.get('Grado', self.datos.get('grado', '')).upper()
            sc = self.datos.get('Seccion', self.datos.get('seccion', ''))
            self.draw.text((xt, yc), "GRADO:", font=fl, fill="black")
            self.draw.text((xt + 130, yc), gr, font=fd, fill="black")
            yc += 50
            if sc:
                self.draw.text((xt, yc), "SECCIÓN:", font=fl, fill="black")
                self.draw.text((xt + 155, yc), str(sc), font=fd, fill="black")
                yc += 50
        self.draw.text((xt, yc), "VIGENCIA:", font=fl, fill="black")
        self.draw.text((xt + 160, yc), str(self.anio), font=fd, fill="black")

    def _qr(self):
        try:
            dni = str(self.datos.get('DNI', self.datos.get('dni', '')))
            q = qrcode.QRCode(box_size=16, border=1)
            q.add_data(dni)
            q.make(fit=True)
            iq = q.make_image(fill_color="black", back_color="white")
            iq = iq.resize((310, 310), Image.LANCZOS)
            self.img.paste(iq, (self.WIDTH - 345, 195))
            fs = RecursoManager.obtener_fuente("", 20, True)
            self.draw.text((self.WIDTH - 190, 510), "ESCANEAR QR",
                           font=fs, fill="black", anchor="mm")
        except Exception:
            pass

    def _barcode(self):
        if not HAS_BARCODE:
            return
        try:
            dni = str(self.datos.get('DNI', self.datos.get('dni', '')))
            buf2 = io.BytesIO()
            Code128(dni, writer=ImageWriter()).write(buf2, options={
                'write_text': False, 'module_width': 0.70,
                'module_height': 16, 'quiet_zone': 2
            })
            buf2.seek(0)
            ib = Image.open(buf2)
            ib = ib.crop(ib.getbbox())
            ib = ib.resize((420, 80), Image.LANCZOS)
            xb = (self.WIDTH - 420) // 2
            yb = self.HEIGHT - 140
            self.img.paste(ib, (xb, yb))
            fbc = RecursoManager.obtener_fuente("", 22, True)
            self.draw.text((self.WIDTH // 2, yb + 85), f"DNI: {dni}",
                           font=fbc, fill="black", anchor="mm")
        except Exception:
            pass

    def generar(self):
        self._escudo_fondo()
        self._barras()
        self._textos()
        self._foto()
        self._datos()
        self._qr()
        self._barcode()
        out = io.BytesIO()
        self.img.save(out, format='PNG', optimize=True, quality=95)
        out.seek(0)
        return out


# ================================================================
# CARNETS LOTE PDF — 8 POR HOJA (fotocheck)
# ================================================================

def generar_carnets_lote_pdf(lista_datos, anio, es_docente=False):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    w, h = A4
    mx = 8 * mm
    my = 6 * mm
    gx = 4 * mm
    gy = 3 * mm
    cw2 = (w - 2 * mx - gx) / 2             # 2 columnas
    ch2 = cw2 * 638 / 1012                   # Mantener proporcion del carnet
    pp = 6                                     # 6 por pagina (2x3)
    total = len(lista_datos)
    np2 = (total + pp - 1) // pp
    for pag in range(np2):
        if pag > 0:
            c.showPage()
        ini = pag * pp
        fin = min(ini + pp, total)
        for idx in range(ini, fin):
            pos = idx - ini
            col = pos % 2
            fila = pos // 2
            x = mx + col * (cw2 + gx)
            y = h - my - 8 - (fila + 1) * ch2 - fila * gy
            gen = GeneradorCarnet(lista_datos[idx], anio, es_docente=es_docente)
            ib = gen.generar()
            tmp = f"tmp_c_{idx}.png"
            with open(tmp, 'wb') as f:
                f.write(ib.getvalue())
            try:
                c.drawImage(tmp, x, y, width=cw2, height=ch2,
                            preserveAspectRatio=True)
                c.setStrokeColor(colors.grey)
                c.setDash(3, 3)
                c.setLineWidth(0.3)
                c.rect(x, y, cw2, ch2)
                c.setDash()
            except Exception:
                pass
            try:
                os.remove(tmp)
            except Exception:
                pass
        c.setFont("Helvetica", 6)
        c.setFillColor(colors.grey)
        c.drawCentredString(w / 2, 10,
                            f"YACHAY — Carnets {anio} — Pág {pag + 1}/{np2} — "
                            f"Cortar por líneas punteadas")
        c.setFillColor(colors.black)
    c.save()
    buffer.seek(0)
    return buffer


# ================================================================
# UTILIDADES
# ================================================================

def generar_link_whatsapp(tel, msg):
    t = str(tel).strip()
    if '.' in t:
        t = t.split('.')[0]
    t = t.replace("+", "").replace(" ", "").replace("-", "").replace("(", "").replace(")", "")
    t = ''.join(c for c in t if c.isdigit())
    if len(t) == 9:
        t = "51" + t
    elif not t.startswith("51"):
        t = "51" + t
    msg_encoded = urllib.parse.quote(msg.encode('utf-8'), safe=b'')
    # whatsapp:// abre directamente la app de escritorio
    return f"whatsapp://send?phone={t}&text={msg_encoded}"


FRASES_MOTIVACIONALES = [
    "🌟 La puntualidad es la cortesía de los reyes y la obligación de los caballeros.",
    "📚 Educar es sembrar semillas de futuro. ¡Gracias por confiar en YACHAY!",
    "🎯 El éxito es la suma de pequeños esfuerzos repetidos día tras día.",
    "💪 Cada día de clases es una oportunidad para crecer y aprender.",
    "🌈 La educación es el arma más poderosa para cambiar el mundo. — Nelson Mandela",
    "⭐ Un niño puntual hoy será un adulto responsable mañana.",
    "📖 Leer es soñar con los ojos abiertos. ¡Motivemos la lectura!",
    "🏆 El talento gana juegos, pero el trabajo en equipo gana campeonatos.",
    "🌱 Cada estudiante es una semilla; con amor y educación, florecerá.",
    "🔑 La disciplina es el puente entre las metas y los logros.",
    "💡 No hay atajos para ningún lugar que valga la pena ir.",
    "🎓 La mejor inversión es la educación de nuestros hijos.",
    "🌻 Con esfuerzo y dedicación, todo es posible. ¡Vamos YACHAY!",
    "📝 El hábito de estudiar hoy construye el profesional del mañana.",
    "🤝 Familia y escuela juntos: la fórmula del éxito educativo.",
    "⏰ La puntualidad es un valor que se enseña desde casa.",
    "🎒 Cada día es una nueva página en el libro de la vida.",
    "🏫 YACHAY significa aprender. ¡Aprendamos juntos!",
    "✨ El futuro pertenece a quienes creen en la belleza de sus sueños.",
    "🌟 Educar para la Vida — Pioneros en la Educación de Calidad.",
]

import random as _random


def generar_mensaje_asistencia(nombre, tipo, hora):
    saludo = "Buenos días" if int(hora.split(':')[0]) < 12 else "Buenas tardes"
    if tipo == "entrada":
        em = "✅ ENTRADA"
    elif tipo == "tardanza":
        em = "⏰ TARDANZA"
    else:
        em = "🏁 SALIDA"
    frase = _random.choice(FRASES_MOTIVACIONALES)
    return (f"{saludo}\n🏫 I.E. ALTERNATIVO YACHAY informa:\n"
            f"{em} registrada\n👤 {nombre}\n🕒 Hora: {hora}\n\n"
            f"{frase}")


def decodificar_qr_imagen(ib):
    if not HAS_PYZBAR:
        return None
    try:
        img = Image.open(io.BytesIO(ib))
        cod = pyzbar_decode(img)
        if cod:
            return cod[0].data.decode('utf-8')
    except Exception:
        pass
    if HAS_CV2:
        try:
            np2 = np.frombuffer(ib, np.uint8)
            ic = cv2.imdecode(np2, cv2.IMREAD_COLOR)
            gr = cv2.cvtColor(ic, cv2.COLOR_BGR2GRAY)
            for m in [cv2.THRESH_BINARY, cv2.THRESH_BINARY_INV]:
                _, th = cv2.threshold(gr, 127, 255, m)
                cod = pyzbar_decode(Image.fromarray(th))
                if cod:
                    return cod[0].data.decode('utf-8')
        except Exception:
            pass
    return None




# ================================================================
# HOJA DE RESPUESTAS + ESCÁNER OMR PROFESIONAL
# Sistema basado en posición con marcadores de alineación
# ================================================================

# Constantes de la hoja VERTICAL (compartidas entre generador y escáner)
HOJA_W = 2480       # Ancho A4 PORTRAIT 300dpi
HOJA_H = 3508       # Alto A4 PORTRAIT 300dpi
HOJA_MARKER_SIZE = 100   # Tamaño marcadores esquina
HOJA_MARKER_PAD = 40     # Padding de marcadores desde borde
HOJA_BUBBLE_R = 34       # Radio de burbuja
HOJA_Y_START = 950       # Y donde empiezan las burbujas
HOJA_X_START = 340       # X donde empieza la primera opción
HOJA_SP_Y = 108          # Espacio vertical entre preguntas
HOJA_SP_X = 155          # Espacio horizontal entre opciones A,B,C,D
HOJA_COL_SP = 750        # Espacio entre columnas de preguntas
HOJA_PPC = 20            # Preguntas por columna


def _posicion_burbuja(pregunta_idx, opcion_idx):
    """Calcula posición exacta (cx, cy) de una burbuja en la hoja"""
    col = pregunta_idx // HOJA_PPC
    fila = pregunta_idx % HOJA_PPC
    cx = HOJA_X_START + col * HOJA_COL_SP + opcion_idx * HOJA_SP_X
    cy = HOJA_Y_START + fila * HOJA_SP_Y
    return cx, cy


def generar_hoja_respuestas(np_, titulo):
    """Genera hoja de respuestas VERTICAL para escaneo OMR"""
    img = Image.new('RGB', (HOJA_W, HOJA_H), 'white')
    draw = ImageDraw.Draw(img)
    try:
        ft = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 70)
        fs = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 45)
        fn = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 42)
        fl = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 32)
        fb = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 30)
        fi = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 22)
    except Exception:
        ft = fs = fn = fl = fb = fi = ImageFont.load_default()

    # ===== 4 MARCADORES DE ESQUINA =====
    ms = HOJA_MARKER_SIZE
    mp = HOJA_MARKER_PAD
    draw.rectangle([(mp, mp), (mp + ms, mp + ms)], fill="black")
    draw.rectangle([(HOJA_W - mp - ms, mp), (HOJA_W - mp, mp + ms)], fill="black")
    draw.rectangle([(mp, HOJA_H - mp - ms), (mp + ms, HOJA_H - mp)], fill="black")
    draw.rectangle([(HOJA_W - mp - ms, HOJA_H - mp - ms),
                    (HOJA_W - mp, HOJA_H - mp)], fill="black")
    draw.rectangle([(mp, mp + ms + 10), (mp + ms, mp + ms + 30)], fill="black")

    # ===== ENCABEZADO =====
    draw.text((HOJA_W // 2, 200), "I.E.P. ALTERNATIVO YACHAY",
              font=ft, fill="black", anchor="mm")
    draw.text((HOJA_W // 2, 290), f"HOJA DE RESPUESTAS — {titulo.upper()}",
              font=fs, fill="black", anchor="mm")

    # ===== DATOS DEL ALUMNO =====
    draw.text((220, 400), "Nombre: _____________________________________________",
              font=fs, fill="black")
    draw.text((220, 480), "DNI: __________________  Grado: __________________",
              font=fs, fill="black")
    draw.text((220, 560), f"Fecha: __________________  Total: {np_} preguntas",
              font=fs, fill="black")

    # ===== INSTRUCCIONES =====
    draw.text((220, 660), "RELLENE COMPLETAMENTE el círculo de su respuesta",
              font=fb, fill="red")
    ex_y = 720
    draw.text((220, ex_y), "Correcto:", font=fl, fill="gray")
    draw.ellipse([(430, ex_y - 5), (490, ex_y + 55)], fill="black")
    draw.text((530, ex_y), "Incorrecto:", font=fl, fill="gray")
    draw.ellipse([(770, ex_y - 5), (830, ex_y + 55)], outline="black", width=3)
    draw.text((870, ex_y), "Use lápiz 2B o bolígrafo negro", font=fl, fill="gray")

    # Línea separadora
    draw.line([(100, 820), (HOJA_W - 100, 820)], fill="black", width=4)

    # ===== BURBUJAS =====
    for i in range(np_):
        col = i // HOJA_PPC
        fila = i % HOJA_PPC

        # Número de pregunta
        num_x = HOJA_X_START + col * HOJA_COL_SP - 120
        num_y = HOJA_Y_START + fila * HOJA_SP_Y
        draw.text((num_x, num_y), f"{i + 1}.",
                  font=fn, fill="black", anchor="rm")

        # 4 opciones: A, B, C, D
        for j, letra in enumerate(['A', 'B', 'C', 'D']):
            cx, cy = _posicion_burbuja(i, j)
            r = HOJA_BUBBLE_R
            # Círculo bien definido con borde grueso
            draw.ellipse([(cx - r, cy - r), (cx + r, cy + r)],
                         outline="black", width=5)
            # Letra pequeña dentro
            draw.text((cx, cy), letra, font=fl, fill=(100, 100, 100), anchor="mm")

    # ===== PIE DE PÁGINA =====
    draw.line([(100, HOJA_H - 250), (HOJA_W - 100, HOJA_H - 250)],
              fill="black", width=2)

    frases_seguridad = [
        "DOCUMENTO OFICIAL — CUALQUIER ALTERACIÓN INVALIDA ESTE EXAMEN",
        "I.E.P. ALTERNATIVO YACHAY — LECTURA ÓPTICA AUTOMATIZADA",
        "Use SOLO lápiz 2B o bolígrafo negro — Rellene completamente cada círculo",
    ]
    y_pie = HOJA_H - 230
    for frase in frases_seguridad:
        draw.text((HOJA_W // 2, y_pie), frase,
                  font=fb, fill="gray", anchor="mm")
        y_pie += 30

    codigo_seg = hashlib.md5(f"{titulo}{datetime.now().isoformat()}".encode()).hexdigest()[:12].upper()
    draw.text((HOJA_W // 2, HOJA_H - 60),
              f"Código: {codigo_seg} | YACHAY PRO {datetime.now().year}",
              font=fb, fill="black", anchor="mm")

    # Marca de agua diagonal
    try:
        marca_font = ImageFont.truetype(
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 60)
    except Exception:
        marca_font = fb
    marca_img = Image.new('RGBA', img.size, (255, 255, 255, 0))
    marca_draw = ImageDraw.Draw(marca_img)
    for yy in range(200, HOJA_H - 200, 400):
        for xx in range(-200, HOJA_W, 600):
            marca_draw.text((xx, yy), "YACHAY PRO",
                           font=marca_font, fill=(200, 200, 200, 35))
    img = Image.alpha_composite(img.convert('RGBA'), marca_img).convert('RGB')

    out = io.BytesIO()
    img.save(out, format='PNG', quality=95)
    out.seek(0)
    return out


# ================================================================
# ESCÁNER OMR — DETECCIÓN POR POSICIÓN
# ================================================================

def _encontrar_marcadores(gray):
    """
    Encuentra los 4 marcadores de esquina (cuadrados negros grandes).
    Retorna las coordenadas ordenadas: [TL, TR, BL, BR] o None.
    """
    alto, ancho = gray.shape[:2]
    resultados = []

    # Probar múltiples umbrales para robustez
    for metodo in range(3):
        if metodo == 0:
            blur = cv2.GaussianBlur(gray, (5, 5), 0)
            _, thresh = cv2.threshold(blur, 0, 255,
                                       cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
        elif metodo == 1:
            blur = cv2.GaussianBlur(gray, (7, 7), 0)
            thresh = cv2.adaptiveThreshold(blur, 255,
                                            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                            cv2.THRESH_BINARY_INV, 21, 5)
        else:
            blur = cv2.medianBlur(gray, 5)
            _, thresh = cv2.threshold(blur, 80, 255, cv2.THRESH_BINARY_INV)

        # Probar ambos modos de contorno para mayor robustez
        for retr_mode in [cv2.RETR_EXTERNAL, cv2.RETR_LIST]:
            contours, _ = cv2.findContours(thresh, retr_mode,
                                            cv2.CHAIN_APPROX_SIMPLE)

            # Buscar contornos grandes y cuadrados (los marcadores)
            candidatos = []
            min_size = min(ancho, alto) * 0.02  # Al menos 2% del tamaño
            max_size = min(ancho, alto) * 0.12  # Máximo 12%

            for ct in contours:
                x, y, w, h = cv2.boundingRect(ct)
                if w < min_size or h < min_size:
                    continue
                if w > max_size or h > max_size:
                    continue

                aspect = w / float(h) if h > 0 else 0
                if not (0.6 <= aspect <= 1.6):
                    continue

                area = cv2.contourArea(ct)
                rect_area = w * h
                solidez = area / rect_area if rect_area > 0 else 0
                if solidez < 0.6:
                    continue

                # Centro del contorno
                cx = x + w // 2
                cy = y + h // 2
                candidatos.append((cx, cy, w * h, x, y, w, h))

            if len(candidatos) < 4:
                continue

            # Ordenar por tamaño y tomar los más grandes
            candidatos = sorted(candidatos, key=lambda c: c[2], reverse=True)

            if len(candidatos) >= 4:
                top = candidatos[:min(12, len(candidatos))]
                mejor = _seleccionar_esquinas(top, ancho, alto)
                if mejor is not None:
                    resultados.append(mejor)
                    break  # Encontrado, no seguir probando modos

    if not resultados:
        return None

    # Retornar el primer resultado exitoso
    return resultados[0]


def _seleccionar_esquinas(candidatos, ancho, alto):
    """
    De una lista de candidatos, selecciona 4 que forman las esquinas
    de la hoja. Retorna [TL, TR, BL, BR] como arrays de coordenadas.
    """
    puntos = [(c[0], c[1]) for c in candidatos]

    # Clasificar por cuadrante
    cx_medio = ancho / 2
    cy_medio = alto / 2

    tl_cands = [(x, y) for x, y in puntos if x < cx_medio and y < cy_medio]
    tr_cands = [(x, y) for x, y in puntos if x > cx_medio and y < cy_medio]
    bl_cands = [(x, y) for x, y in puntos if x < cx_medio and y > cy_medio]
    br_cands = [(x, y) for x, y in puntos if x > cx_medio and y > cy_medio]

    if not (tl_cands and tr_cands and bl_cands and br_cands):
        return None

    # Tomar el más cercano a cada esquina
    tl = min(tl_cands, key=lambda p: p[0]**2 + p[1]**2)
    tr = min(tr_cands, key=lambda p: (ancho - p[0])**2 + p[1]**2)
    bl = min(bl_cands, key=lambda p: p[0]**2 + (alto - p[1])**2)
    br = min(br_cands, key=lambda p: (ancho - p[0])**2 + (alto - p[1])**2)

    return [list(tl), list(tr), list(bl), list(br)]


def _corregir_perspectiva(gray, esquinas):
    """
    Aplica transformación de perspectiva para alinear la hoja.
    esquinas = [TL, TR, BL, BR]
    Retorna imagen corregida de tamaño HOJA_W x HOJA_H
    """
    tl, tr, bl, br = esquinas

    # Puntos origen (de la foto)
    src = np.array([tl, tr, bl, br], dtype="float32")

    # Puntos destino (hoja perfecta) — ajustados a los centros de marcadores
    mp = HOJA_MARKER_PAD + HOJA_MARKER_SIZE // 2
    dst = np.array([
        [mp, mp],
        [HOJA_W - mp, mp],
        [mp, HOJA_H - mp],
        [HOJA_W - mp, HOJA_H - mp]
    ], dtype="float32")

    # Calcular y aplicar transformación
    M = cv2.getPerspectiveTransform(src, dst)
    warped = cv2.warpPerspective(gray, M, (HOJA_W, HOJA_H))
    return warped


def _leer_burbujas(warped_gray, num_preguntas):
    """
    Lee las respuestas de la imagen ya corregida/alineada.
    MEJORADO: Lógica estricta anti-falsos positivos.
    - Pre-procesamiento con GaussianBlur + OTSU
    - Erosión para eliminar ruido/sombras
    - Umbral de relleno mínimo 45%
    - Comparación relativa: la más marcada debe ser >1.4x la segunda
    - Si no cumple condiciones → '?' (indeterminado)
    """
    # Pre-procesamiento robusto
    blur = cv2.GaussianBlur(warped_gray, (5, 5), 0)
    _, thresh = cv2.threshold(blur, 0, 255,
                               cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)

    # Erosión para eliminar ruido, trazos débiles y sombras
    kernel = np.ones((2, 2), np.uint8)
    thresh = cv2.erode(thresh, kernel, iterations=1)

    respuestas = []
    radio_muestra = int(HOJA_BUBBLE_R * 0.60)
    UMBRAL_RELLENO_MINIMO = 0.45   # Mínimo 45% del círculo relleno
    RATIO_DIFERENCIA = 1.4          # La más marcada debe ser 1.4x la segunda

    for i in range(num_preguntas):
        intensidades = []
        for j in range(4):  # A, B, C, D
            cx, cy = _posicion_burbuja(i, j)

            # Verificar límites
            if (cy - radio_muestra < 0 or cy + radio_muestra >= HOJA_H or
                    cx - radio_muestra < 0 or cx + radio_muestra >= HOJA_W):
                intensidades.append(0.0)
                continue

            # Crear máscara circular localizada (más eficiente)
            y1 = max(0, cy - radio_muestra - 5)
            y2 = min(HOJA_H, cy + radio_muestra + 5)
            x1 = max(0, cx - radio_muestra - 5)
            x2 = min(HOJA_W, cx + radio_muestra + 5)

            roi = thresh[y1:y2, x1:x2]
            mask_local = np.zeros_like(roi, dtype="uint8")
            cv2.circle(mask_local,
                       (cx - x1, cy - y1),
                       radio_muestra, 255, -1)

            masked = cv2.bitwise_and(roi, roi, mask=mask_local)
            total = cv2.countNonZero(mask_local)
            filled = cv2.countNonZero(masked)
            ratio = filled / total if total > 0 else 0.0
            intensidades.append(ratio)

        if not intensidades:
            respuestas.append('-')  # Sin datos = en blanco = 0 puntos
            continue

        max_val = max(intensidades)
        max_idx = intensidades.index(max_val)

        # Condición 1: Relleno mínimo
        if max_val < UMBRAL_RELLENO_MINIMO:
            respuestas.append('-')  # En blanco = 0 puntos (no marcó nada)
            continue

        # Condición 2: Diferencia significativa con la segunda opción
        sorted_vals = sorted(intensidades, reverse=True)
        segunda = sorted_vals[1] if len(sorted_vals) >= 2 else 0

        if segunda > 0 and max_val / segunda < RATIO_DIFERENCIA:
            respuestas.append('?')  # Ambiguo — corregir manualmente
            continue

        # Respuesta clara
        respuestas.append(['A', 'B', 'C', 'D'][max_idx])

    return respuestas


def _leer_sin_perspectiva(gray, num_preguntas):
    """
    Método alternativo cuando no se detectan marcadores.
    Intenta detectar la región de burbujas directamente.
    Busca patrones de filas de 4 elementos oscuros.
    """
    alto, ancho = gray.shape[:2]

    # Redimensionar a tamaño estándar para posiciones conocidas
    resized = cv2.resize(gray, (HOJA_W, HOJA_H), interpolation=cv2.INTER_LINEAR)

    _, thresh = cv2.threshold(resized, 0, 255,
                               cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)

    # Intentar leer directamente asumiendo que la imagen ya está alineada
    respuestas = _leer_burbujas(resized, num_preguntas)

    # Verificar calidad: si más del 70% son '?', falló
    preguntas_detectadas = sum(1 for r in respuestas if r != '?')
    if preguntas_detectadas < num_preguntas * 0.3:
        return None

    return respuestas


def procesar_examen(image_bytes, num_preguntas):
    """
    ESCÁNER OMR PROFESIONAL - Basado en posición.
    
    Método principal:
    1. Detecta 4 marcadores de esquina
    2. Corrige perspectiva (la foto se vuelve una hoja plana)
    3. Lee cada burbuja en su posición exacta
    
    Método alternativo (sin marcadores):
    - Redimensiona la imagen al tamaño de la hoja
    - Intenta leer las posiciones directamente
    
    Retorna lista de respuestas ['A','B','C','D','?'] o None si falla
    """
    if not HAS_CV2:
        return None

    try:
        # Decodificar imagen
        nparr = np.frombuffer(image_bytes, np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        if img is None:
            return None

        # Escalar si es muy grande (>4000px)
        h_orig, w_orig = img.shape[:2]
        escala = 1.0
        if max(h_orig, w_orig) > 4000:
            escala = 4000 / max(h_orig, w_orig)
            img = cv2.resize(img, (int(w_orig * escala), int(h_orig * escala)),
                             interpolation=cv2.INTER_AREA)

        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

        # === MÉTODO 1: Con marcadores (el más preciso) ===
        esquinas = _encontrar_marcadores(gray)
        if esquinas is not None:
            warped = _corregir_perspectiva(gray, esquinas)
            respuestas = _leer_burbujas(warped, num_preguntas)
            detectadas = sum(1 for r in respuestas if r != '?')
            if detectadas >= num_preguntas * 0.3:
                return respuestas

        # === MÉTODO 2: Redimensionar directo (sin marcadores) ===
        respuestas = _leer_sin_perspectiva(gray, num_preguntas)
        if respuestas:
            return respuestas

        # === MÉTODO 3: Mejorar contraste y reintentar ===
        clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(gray)
        esquinas2 = _encontrar_marcadores(enhanced)
        if esquinas2 is not None:
            warped2 = _corregir_perspectiva(enhanced, esquinas2)
            respuestas2 = _leer_burbujas(warped2, num_preguntas)
            detectadas2 = sum(1 for r in respuestas2 if r != '?')
            if detectadas2 >= num_preguntas * 0.3:
                return respuestas2

        # === MÉTODO 4: Umbral manual y reintentar ===
        for umbral in [100, 120, 140, 160]:
            _, manual_thresh = cv2.threshold(gray, umbral, 255, cv2.THRESH_BINARY)
            esquinas3 = _encontrar_marcadores(manual_thresh)
            if esquinas3 is not None:
                warped3 = _corregir_perspectiva(gray, esquinas3)
                respuestas3 = _leer_burbujas(warped3, num_preguntas)
                detectadas3 = sum(1 for r in respuestas3 if r != '?')
                if detectadas3 >= num_preguntas * 0.3:
                    return respuestas3

        return None

    except Exception:
        return None

# ================================================================
# PANTALLA DE LOGIN (Usuario + Contraseña — SEGURO)
# ================================================================

def pantalla_login():
    # CSS especial para login
    st.markdown("""
    <style>
    .login-container {
        max-width: 480px;
        margin: 0 auto;
        padding: 2rem;
        background: white;
        border-radius: 20px;
        box-shadow: 0 20px 60px rgba(0,0,0,0.15);
        animation: fadeInUp 0.6s ease-out;
    }
    .login-header {
        text-align: center;
        padding: 2.5rem 1.5rem;
        background: linear-gradient(135deg, #001e7c 0%, #0044cc 40%, #0066ff 80%, #3b82f6 100%);
        color: white;
        border-radius: 20px;
        margin-bottom: 2rem;
        box-shadow: 0 12px 35px rgba(0,30,124,0.4);
        position: relative;
        overflow: hidden;
    }
    .login-header::before {
        content: '';
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: radial-gradient(circle, rgba(255,255,255,0.1) 0%, transparent 60%);
        animation: shimmer 3s infinite;
    }
    .login-title {
        font-size: 2rem;
        font-weight: 800;
        margin: 0;
        text-shadow: 0 2px 10px rgba(0,0,0,0.3);
        letter-spacing: 1px;
    }
    .login-subtitle {
        color: #b8d4ff;
        margin: 8px 0 4px;
        font-size: 0.95rem;
    }
    .login-motto {
        color: #FFD700;
        font-style: italic;
        font-size: 1.15rem;
        margin: 10px 0 5px;
        font-weight: 600;
    }
    .login-slogan {
        color: #FFD700;
        font-size: 0.85rem;
        letter-spacing: 2px;
        text-transform: uppercase;
    }
    .login-location {
        color: #8fb8f0;
        font-size: 0.8rem;
        margin-top: 12px;
    }
    .login-divider {
        border: none;
        border-top: 1px solid rgba(255,215,0,0.4);
        margin: 12px 40px;
    }
    .login-footer {
        text-align: center;
        color: #94a3b8;
        font-size: 0.75rem;
        margin-top: 1.5rem;
        padding-top: 1rem;
        border-top: 1px solid #e2e8f0;
    }
    @keyframes shimmer {
        0% { transform: rotate(0deg); }
        100% { transform: rotate(360deg); }
    }
    @keyframes fadeInUp {
        from { opacity: 0; transform: translateY(20px); }
        to { opacity: 1; transform: translateY(0); }
    }
    @keyframes pulseGlow {
        0%, 100% { filter: drop-shadow(0 0 8px rgba(26,86,219,0.4)); }
        50% { filter: drop-shadow(0 0 25px rgba(26,86,219,0.8)); }
    }
    .escudo-login img {
        animation: pulseGlow 3s ease-in-out infinite;
        border-radius: 50%;
    }
    [data-testid="stTextInput"] input {
        border-radius: 12px !important;
        padding: 12px 16px !important;
        font-size: 1rem !important;
        border: 2px solid #e2e8f0 !important;
        transition: border-color 0.3s !important;
    }
    [data-testid="stTextInput"] input:focus {
        border-color: #1a56db !important;
        box-shadow: 0 0 0 3px rgba(26,86,219,0.15) !important;
    }
    </style>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        # ═══ DOS ESCUDOS (base64 para render confiable) ═══
        esc_izq = Path("escudo_upload.png").exists()
        esc_der = Path("escudo2_upload.png").exists()
        if esc_izq or esc_der:
            esc_html = '<div style="display:flex;justify-content:center;align-items:center;gap:20px;margin-bottom:10px;">'
            if esc_izq:
                with open("escudo_upload.png", "rb") as ef:
                    b64_izq = base64.b64encode(ef.read()).decode()
                esc_html += f'<img src="data:image/png;base64,{b64_izq}" style="width:120px;height:auto;border-radius:50%;filter:drop-shadow(0 0 12px rgba(26,86,219,0.5));" />'
            esc_html += '<span style="font-size:2.2rem;">🎓</span>'
            if esc_der:
                with open("escudo2_upload.png", "rb") as ef:
                    b64_der = base64.b64encode(ef.read()).decode()
                esc_html += f'<img src="data:image/png;base64,{b64_der}" style="width:120px;height:auto;border-radius:50%;filter:drop-shadow(0 0 12px rgba(26,86,219,0.5));" />'
            esc_html += '</div>'
            st.markdown(esc_html, unsafe_allow_html=True)
        
        st.markdown("""
        <div class='login-header'>
            <p class='login-title'>🎓 SISTEMA YACHAY PRO</p>
            <p class='login-subtitle'>Sistema Integral de Gestión Educativa</p>
            <p class='login-motto'>"Educar para la Vida"</p>
            <p class='login-slogan'>Pioneros en la Educación de Calidad</p>
            <hr class='login-divider'>
            <p class='login-location'>📍 Chinchero, Cusco — Perú</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("")
        usuario = st.text_input("👤 Nombre de usuario:", key="login_user",
                                placeholder="Ingrese su usuario")
        contrasena = st.text_input("🔑 Contraseña:", type="password",
                                    key="login_pwd",
                                    placeholder="Ingrese su contraseña")
        
        if st.button("🔐 INGRESAR AL SISTEMA", use_container_width=True,
                     type="primary"):
            usuarios = cargar_usuarios()
            usuario_lower = usuario.strip().lower()
            if usuario_lower in usuarios:
                datos_u = usuarios[usuario_lower]
                pwd_guardado = str(datos_u.get('password', '')).strip()
                # Limpiar .0 si GS lo convirtió
                if pwd_guardado.endswith('.0'):
                    pwd_guardado = pwd_guardado[:-2]
                
                if str(contrasena).strip() == pwd_guardado:
                    rol = datos_u.get('rol', 'docente')
                    # Directivos/Promotor/Coordinador → acceso como directivo
                    if rol in ['directivo', 'promotor', 'coordinador']:
                        st.session_state.rol = 'directivo'
                    else:
                        st.session_state.rol = rol
                    st.session_state.docente_info = datos_u.get('docente_info')
                    st.session_state.usuario_actual = usuario_lower
                    st.toast(f"✅ Bienvenido, {datos_u.get('label', usuario_lower)}")
                    st.rerun()
                else:
                    st.error("⛔ Contraseña incorrecta")
            else:
                st.error("⛔ Usuario no encontrado")
        
        st.markdown("""
        <div class='login-footer'>
            💡 Ingrese usuario y contraseña asignados por el administrador<br>
            © 2026 YACHAY PRO — Todos los derechos reservados
        </div>
        """, unsafe_allow_html=True)

        # Libro de reclamaciones
        st.markdown("---")
        with st.expander("📕 Libro de Reclamaciones Virtual"):
            st.markdown("*Según normativa MINEDU*")
            with st.form("form_reclamo_login", clear_on_submit=True):
                r_nombre = st.text_input("Nombre completo:", key="rl_nombre")
                r_dni = st.text_input("DNI:", key="rl_dni")
                r_cel = st.text_input("Celular:", key="rl_cel")
                r_tipo = st.selectbox("Tipo:", ["Queja", "Reclamo", "Sugerencia"], key="rl_tipo")
                r_detalle = st.text_area("Detalle:", key="rl_detalle")
                if st.form_submit_button("📩 ENVIAR", type="primary",
                                          use_container_width=True):
                    if r_nombre and r_dni and r_detalle:
                        gs = _gs()
                        if gs:
                            try:
                                ws = gs._get_hoja('config')
                                if ws:
                                    codigo_rec = f"REC-{hora_peru().year}-{int(time.time()) % 10000:04d}"
                                    ws.append_row([
                                        f"reclamo_{codigo_rec}",
                                        json.dumps({
                                            'codigo': codigo_rec, 'nombre': r_nombre,
                                            'dni': r_dni, 'celular': r_cel,
                                            'tipo': r_tipo, 'detalle': r_detalle,
                                            'fecha': fecha_peru_str(), 'hora': hora_peru_str(),
                                            'estado': 'Pendiente'
                                        }, ensure_ascii=False)
                                    ])
                                    st.success(f"✅ Reclamo registrado. Código: **{codigo_rec}**")
                            except Exception:
                                st.error("Error al enviar. Intente más tarde.")
                        else:
                            st.warning("Sistema en modo local.")
                    else:
                        st.error("Complete todos los campos.")


# ================================================================
# SIDEBAR — Con links SIAGIE y Google Institucional
# ================================================================

# ================================================================
# SISTEMA DE BACKUP Y RESTAURACIÓN
# ================================================================

ARCHIVOS_BACKUP = [
    ARCHIVO_MATRICULA,    # matricula.xlsx
    ARCHIVO_DOCENTES,     # docentes.xlsx
    ARCHIVO_BD,           # base_datos.xlsx
    ARCHIVO_ASISTENCIAS,  # asistencias.json
    ARCHIVO_RESULTADOS,   # resultados_examenes.json
    ARCHIVO_USUARIOS,     # usuarios.json
    "escudo_upload.png",
    "fondo.png",
    "materiales_docente.json",   # Aula Virtual
    "examenes_semanales.json",   # Exámenes Semanales
    "notas.json",                # Notas registradas
    "diagnostico_data.json",      # Exámenes de diagnóstico
    "historial_evaluaciones.json", # Historial evaluaciones
]


def crear_backup():
    """Crea un ZIP con TODOS los datos del sistema"""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for archivo in ARCHIVOS_BACKUP:
            if Path(archivo).exists():
                zf.write(archivo, archivo)
        # Agregar un manifiesto con info del backup
        info = {
            "fecha": hora_peru().strftime('%d/%m/%Y %H:%M:%S'),
            "version": "YACHAY PRO v4.0",
            "archivos": [a for a in ARCHIVOS_BACKUP if Path(a).exists()],
            "total_alumnos": len(BaseDatos.cargar_matricula()),
            "total_docentes": len(BaseDatos.cargar_docentes()),
        }
        zf.writestr("_backup_info.json", json.dumps(info, indent=2, ensure_ascii=False))
    buf.seek(0)
    return buf


def restaurar_backup(zip_bytes):
    """Restaura datos desde un ZIP de backup"""
    errores = []
    restaurados = []
    try:
        with zipfile.ZipFile(io.BytesIO(zip_bytes), 'r') as zf:
            nombres = zf.namelist()
            for archivo in nombres:
                if archivo.startswith("_backup_"):
                    continue  # Saltar manifiesto
                try:
                    zf.extract(archivo, '.')
                    restaurados.append(archivo)
                except Exception as e:
                    errores.append(f"{archivo}: {str(e)}")
    except Exception as e:
        errores.append(f"Error ZIP: {str(e)}")
    return restaurados, errores


def configurar_sidebar():
    with st.sidebar:
        # Escudo
        if Path("escudo_upload.png").exists():
            st.image("escudo_upload.png", width=80)
        st.title("🎓 YACHAY PRO")
        roles_nombres = {
            "admin": "⚙️ Administrador",
            "directivo": "📋 Directivo",
            "auxiliar": "👤 Auxiliar",
            "docente": "👨‍🏫 Docente"
        }
        label = roles_nombres.get(st.session_state.rol, '')
        if st.session_state.rol == "docente" and st.session_state.docente_info:
            label += f" — {_nombre_completo_docente()}"
        st.info(f"**{label}**")
        st.caption(f"🕒 {hora_peru().strftime('%H:%M:%S')} | "
                   f"📅 {hora_peru().strftime('%d/%m/%Y')}")

        # Estado Google Sheets
        gs = _gs()
        if gs:
            st.markdown('<div class="gs-connected">☁️ Google Sheets: Conectado ✅</div>',
                       unsafe_allow_html=True)
        else:
            st.markdown('<div class="gs-offline">💾 Modo local (sin Google Sheets)</div>',
                       unsafe_allow_html=True)

        # Links institucionales para directivo y docentes
        if st.session_state.rol in ["directivo", "docente"]:
            st.markdown("---")
            c1, c2 = st.columns(2)
            with c1:
                st.markdown(
                    f'<a href="{LINK_SIAGIE}" target="_blank" class="siagie-btn">'
                    f'📚 Ir a SIAGIE</a>', unsafe_allow_html=True)
            with c2:
                st.markdown(
                    f'<a href="{LINK_GOOGLE}" target="_blank" class="link-btn">'
                    f'📧 Cuenta Google</a>', unsafe_allow_html=True)

        st.markdown("---")
        directora = "Prof. Ana María CUSI INCA"
        promotor = "Prof. Leandro CORDOVA TOCRE"
        frase = "Año de la Esperanza y el Fortalecimiento de la Democracia"

        if st.session_state.rol == "admin":
            with st.expander("📂 Archivos"):
                ub = st.file_uploader("📊 Base Datos", type=["xlsx"], key="ub")
                if ub:
                    with open(ARCHIVO_BD, "wb") as f:
                        f.write(ub.getbuffer())
                    st.success("🎉")
                    st.rerun()
                uf = st.file_uploader("🖼️ Fondo docs", type=["png"], key="uf")
                if uf:
                    with open("fondo.png", "wb") as f:
                        f.write(uf.getbuffer())
                    _guardar_archivo_binario_gs("bin_fondo", "fondo.png")
                    st.success("🎉 Fondo guardado y sincronizado")
                ue = st.file_uploader("🛡️ Escudo Izquierda", type=["png"], key="ue")
                if ue:
                    with open("escudo_upload.png", "wb") as f:
                        f.write(ue.getbuffer())
                    _guardar_archivo_binario_gs("bin_escudo_izq", "escudo_upload.png")
                    st.success("🎉 Escudo izquierda guardado y sincronizado")
                ue2 = st.file_uploader("🛡️ Escudo Derecha", type=["png"], key="ue2")
                if ue2:
                    with open("escudo2_upload.png", "wb") as f:
                        f.write(ue2.getbuffer())
                    _guardar_archivo_binario_gs("bin_escudo_der", "escudo2_upload.png")
                    st.success("🎉 Escudo derecho guardado y sincronizado")
            with st.expander("👥 Autoridades"):
                directora = st.text_input("Directora:", directora, key="di")
                promotor = st.text_input("Promotor:", promotor, key="pi")
            with st.expander("🎯 Título del Año"):
                frase = st.text_input("Frase:", frase, key="fi")
            with st.expander("🔐 Gestionar Usuarios"):
                _gestion_usuarios_admin()
            with st.expander("💾 BACKUP / RESTAURAR", expanded=False):
                st.caption("⚠️ **IMPORTANTE:** Streamlit Cloud puede borrar "
                           "tus datos. Haz backup frecuentemente.")
                st.markdown("---")
                st.markdown("**📥 DESCARGAR BACKUP:**")
                if st.button("💾 CREAR BACKUP AHORA", type="primary",
                             use_container_width=True, key="btn_backup"):
                    with st.spinner("📦 Empaquetando datos..."):
                        backup_zip = crear_backup()
                    fecha_bk = hora_peru().strftime('%Y%m%d_%H%M')
                    st.download_button(
                        f"⬇️ Descargar backup_{fecha_bk}.zip",
                        backup_zip,
                        f"backup_yachay_{fecha_bk}.zip",
                        "application/zip",
                        use_container_width=True,
                        key="dl_backup"
                    )
                    st.success("🎉 Backup listo. ¡Guárdalo en tu PC!")
                st.markdown("---")
                st.markdown("**📤 RESTAURAR DESDE BACKUP:**")
                uploaded_backup = st.file_uploader(
                    "Subir archivo .zip de backup:",
                    type=["zip"], key="upload_backup"
                )
                if uploaded_backup:
                    st.warning("⚠️ Esto REEMPLAZARÁ todos los datos actuales "
                               "con los del backup.")
                    if st.button("🔄 RESTAURAR DATOS", type="primary",
                                 use_container_width=True, key="btn_restaurar"):
                        with st.spinner("🔄 Restaurando..."):
                            rest, errs = restaurar_backup(
                                uploaded_backup.getvalue()
                            )
                        if rest:
                            st.success(f"✅ Restaurados {len(rest)} archivos:\n"
                                       f"{', '.join(rest)}")
                        if errs:
                            st.error(f"❌ Errores: {', '.join(errs)}")
                        if rest:
                            st.balloons()
                            time.sleep(1)
                            st.rerun()
            
            with st.expander("🔧 Herramientas"):
                st.markdown("### 📝 Corregir Secciones")
                st.caption("Cambia a 'A' los estudiantes sin sección o con 'Única' (Primaria, Secundaria, PreU)")
                if st.button("🔄 Corregir Secciones", type="primary", 
                           use_container_width=True, key="btn_corr_sec"):
                    cantidad = BaseDatos.corregir_secciones_vacias()
                    if cantidad > 0:
                        st.success(f"✅ Se asignó sección 'A' a {cantidad} estudiante(s)")
                        st.balloons()
                    else:
                        st.info("✅ Todos los estudiantes ya tienen sección")

                st.markdown("---")
                st.markdown("### 🎓 Promoción de Año Escolar")
                st.caption("Avanza todos los estudiantes al siguiente grado (Inicial→Primaria→Secundaria)")

                # Previsualización
                if st.button("🔍 Ver previsualización", key="btn_prev_promo", use_container_width=True):
                    prev = BaseDatos.previsualizar_promocion()
                    if prev:
                        st.session_state['_prev_promo'] = prev
                    else:
                        st.warning("No hay estudiantes para promover.")

                if st.session_state.get('_prev_promo'):
                    prev = st.session_state['_prev_promo']
                    st.markdown("**Vista previa de cambios:**")
                    total_prev = sum(v['cantidad'] for v in prev.values())
                    for grado_act, info in prev.items():
                        emoji = "🎓" if info['nuevo'] == "EGRESADO" else "➡️"
                        st.markdown(
                            f"- **{grado_act}** {emoji} **{info['nuevo']}** "
                            f"&nbsp;&nbsp;`{info['cantidad']} alumno(s)`"
                        )
                    st.warning(f"⚠️ Se modificarán **{total_prev} estudiantes**. ¿Confirmar?")
                    c_si, c_no = st.columns(2)
                    with c_si:
                        if st.button("✅ SÍ, PROMOVER AHORA", type="primary",
                                     use_container_width=True, key="btn_confirmar_promo"):
                            with st.spinner("🔄 Promoviendo grados..."):
                                resultado = BaseDatos.promover_grados()
                            st.session_state.pop('_prev_promo', None)
                            st.success(
                                f"🎉 ¡Promoción completada!\n\n"
                                f"✅ Promovidos: **{resultado['promovidos']}**\n\n"
                                f"🎓 Egresados: **{resultado['egresados']}**\n\n"
                                f"⏭️ Sin cambio (PreU/otros): **{resultado['sin_cambio']}**"
                            )
                            st.balloons()
                            time.sleep(1)
                            st.rerun()
                    with c_no:
                        if st.button("❌ Cancelar", use_container_width=True, key="btn_cancel_promo"):
                            st.session_state.pop('_prev_promo', None)
                            st.rerun()

                st.markdown("---")
                st.markdown("### 🗑️ Resetear TODAS las Notas")
                st.caption("⚠️ Borra todos los registros de notas y evaluaciones del sistema.")
                _chk_all = st.checkbox("Confirmo que deseo borrar TODAS las notas", key="chk_reset_all")
                if _chk_all and st.button("🗑️ BORRAR TODAS LAS NOTAS", type="primary",
                                           use_container_width=True, key="btn_reset_all"):
                    for _archivo in ['historial_evaluaciones.json', 'resultados.json',
                                     'resultados_examenes.json', ARCHIVO_RESULTADOS]:
                        try:
                            with open(_archivo, 'w', encoding='utf-8') as _f:
                                json.dump({}, _f)
                        except Exception: pass
                    # Limpiar también en Google Sheets
                    try:
                        _gs_inst = _gs()
                        if _gs_inst:
                            ws_cfg = _gs_inst._get_hoja('config')
                            if ws_cfg:
                                registros = ws_cfg.get_all_records()
                                # Borrar desde el final para no desplazar índices
                                for i in range(len(registros), 0, -1):
                                    clave = str(registros[i-1].get('clave', ''))
                                    if clave.startswith('nota_') or clave.startswith('resultado_'):
                                        ws_cfg.delete_rows(i + 1)
                    except Exception: pass
                    st.success("✅ Todas las notas eliminadas del sistema y GS")
                    st.rerun()

        st.markdown("---")
        anio = st.number_input("📅 Año:", 2024, 2040, 2026, key="ai")
        
        # Solo admin y directivo ven estadísticas
        if st.session_state.rol in ['admin', 'directivo']:
            stats = BaseDatos.obtener_estadisticas()
            c1, c2 = st.columns(2)
            with c1:
                st.metric("📚 Alumnos", stats['total_alumnos'])
            with c2:
                st.metric("👨‍🏫 Docentes", stats['total_docentes'])
        
        # Mensaje de guardado para todos
        st.markdown("""<div style="background: #dcfce7; border-radius: 8px; 
                    padding: 8px; text-align: center; font-size: 0.8rem; color: #166534;">
                    💾 Todo se guarda automáticamente en la nube
                    </div>""", unsafe_allow_html=True)
        
        st.markdown("---")
        
        if st.button("🚪 CERRAR SESIÓN", use_container_width=True, key="btn_logout_sidebar", type="primary"):
            for k in list(st.session_state.keys()):
                del st.session_state[k]
            st.rerun()

    return {
        'anio': anio, 'directora': directora, 'promotor': promotor,
        'frase': frase, 'y_frase': 700, 'y_titulo': 630,
        'qr_x': 435, 'qr_y': 47
    }


# ================================================================
# GESTIÓN DE USUARIOS DESDE ADMIN
# ================================================================

def _gestion_usuarios_admin():
    """Admin puede editar/eliminar usuarios. Crear cuentas = Registrar Docente en Matrícula."""
    usuarios = cargar_usuarios()
    
    # NUEVO: Detectar y borrar cuentas obsoletas
    cuentas_viejas = [u for u in usuarios.keys() 
                      if (u.startswith('profe') or u.startswith('prof.')) 
                      and u not in ['profesional', 'profesor']]
    
    if cuentas_viejas:
        st.warning(f"⚠️ Se detectaron {len(cuentas_viejas)} cuentas obsoletas")
        
        # Mostrar lista sin expander (para evitar anidamiento)
        st.markdown("**👀 Cuentas obsoletas detectadas:**")
        for cv in cuentas_viejas:
            st.caption(f"🗑️ {cv} → {usuarios[cv].get('label', 'Sin nombre')}")
        
        if st.button("🗑️ ELIMINAR TODAS LAS CUENTAS OBSOLETAS", 
                   type="primary", key="btn_del_obsoletas"):
            for cv in cuentas_viejas:
                del usuarios[cv]
            guardar_usuarios(usuarios)
            st.success(f"✅ {len(cuentas_viejas)} cuentas obsoletas eliminadas")
            st.balloons()
            time.sleep(1)
            st.rerun()
        st.markdown("---")
    
    st.caption(f"**{len(usuarios)} cuentas de acceso:**")
    for usr, datos in usuarios.items():
        rol_emoji = {"admin": "⚙️", "directivo": "📋", "auxiliar": "👤", "docente": "👨‍🏫"}.get(datos.get('rol', ''), '•')
        grado_txt = ""
        if datos.get('docente_info') and datos['docente_info'].get('grado'):
            grado_txt = f" — {datos['docente_info']['grado']}"
        st.caption(f"{rol_emoji} **{usr}** → {datos.get('label', datos['rol'])}{grado_txt}")

    st.info("💡 Para crear nuevas cuentas: vaya a **Matrícula > Registrar Docente**")

    st.markdown("---")
    st.markdown("**✏️ Editar cuenta:**")
    editable = [u for u in usuarios.keys() if u != "administrador"]
    if not editable:
        st.caption("No hay cuentas editables.")
        return
    edit_usr = st.selectbox("Seleccionar usuario:", editable, key="edit_usr")
    if edit_usr:
        datos_edit = usuarios[edit_usr]
        ne_label = st.text_input("Nombre completo:", value=datos_edit.get('label', ''), key="ne_label")
        ne_pass = st.text_input("Contraseña:", value=datos_edit.get('password', ''), key="ne_pass")
        ne_rol = st.selectbox("Rol:", ["docente", "directivo", "auxiliar"],
                               index=["docente", "directivo", "auxiliar"].index(datos_edit.get('rol', 'docente')),
                               key="ne_rol")
        
        # Solo docentes necesitan grado y nivel
        ne_nivel = "PRIMARIA"
        ne_grado = "N/A"
        if ne_rol == "docente":
            niveles_opciones = ["INICIAL", "PRIMARIA", "SECUNDARIA", "PREUNIVERSITARIO", "TODOS"]
            nivel_actual = datos_edit.get('docente_info', {}).get('nivel', 'PRIMARIA') if datos_edit.get('docente_info') else 'PRIMARIA'
            if nivel_actual not in niveles_opciones:
                nivel_actual = "PRIMARIA"
            ne_nivel = st.selectbox("Nivel:", niveles_opciones,
                                     index=niveles_opciones.index(nivel_actual),
                                     key="ne_nivel")
            if ne_nivel == "TODOS":
                grados_opts = ["ALL_NIVELES"]
                st.success("🔓 Acceso a TODOS los grados (Inicial, Primaria, Secundaria, Pre-U)")
            else:
                grados_opts = ["N/A"] + NIVELES_GRADOS.get(ne_nivel, []) + ["ALL_SECUNDARIA"]
            ne_grado = st.selectbox("Grado asignado:", grados_opts, key="ne_grado")
        else:
            st.caption(f"🔓 **{ne_rol.title()}** tiene acceso completo (sin grado específico)")
        
        c1, c2 = st.columns(2)
        with c1:
            if st.button("💾 GUARDAR", type="primary", key="btn_edit_usr"):
                usuarios[edit_usr]['label'] = ne_label
                usuarios[edit_usr]['password'] = ne_pass
                usuarios[edit_usr]['rol'] = ne_rol
                if ne_rol == "docente":
                    # Preservar DNI existente o intentar resolverlo
                    old_di = datos_edit.get('docente_info') or {}
                    dni_doc = old_di.get('dni', '')
                    di = {"label": ne_label, "grado": ne_grado, "nivel": ne_nivel, "dni": dni_doc}
                    usuarios[edit_usr]['docente_info'] = di
                else:
                    usuarios[edit_usr]['docente_info'] = None
                guardar_usuarios(usuarios)
                st.success(f"✅ {edit_usr} actualizado")
                st.rerun()
        with c2:
            if st.button("🗑️ Eliminar", key="btn_del_usr", type="primary"):
                if edit_usr != "administrador":
                    del usuarios[edit_usr]
                    guardar_usuarios(usuarios)
                    st.success(f"✅ {edit_usr} eliminado")
                    st.rerun()


# ================================================================
# TAB: MATRÍCULA (Alumnos + Docentes)
# ================================================================

def tab_matricula(config):
    st.header("📝 Matrícula")
    tab_est, tab_doc, tab_lista, tab_pdf = st.tabs([
        "➕ Registrar Alumno", "👨‍🏫 Registrar Docente",
        "📋 Listas", "⬇️ Registros PDF"
    ])

    with tab_est:
        st.subheader("📝 Matrícula de Estudiante")

        # --- MODO DE REGISTRO ---
        modo_reg = st.radio(
            "Tipo de registro:",
            ["✅ Completo (con DNI)", "⏳ Provisional (solo nombre)"],
            horizontal=True, key="modo_matricula"
        )
        es_provisional = "Provisional" in modo_reg

        if es_provisional:
            st.info("💡 **Modo provisional**: solo ingrese apellidos y nombres. Puede completar DNI, apoderado y celular después en **Base de Datos → Completar Datos**.")

        c1, c2 = st.columns(2)
        with c1:
            mn = st.text_input("Apellidos y Nombres: *", key="mn",
                               placeholder="APELLIDO APELLIDO, Nombre")
            if not es_provisional:
                md = st.text_input("DNI: *", key="md", max_chars=8,
                                   placeholder="12345678")
            else:
                md = ""
            mnv = st.selectbox("Nivel:", list(NIVELES_GRADOS.keys()), key="mnv")
            mg = st.selectbox("Grado:", NIVELES_GRADOS[mnv], key="mg")
            # Secundaria/Preuniversitario → sección A por defecto
            if mnv in ("SECUNDARIA", "PREUNIVERSITARIO"):
                ms = st.selectbox("Sección:", SECCIONES, index=SECCIONES.index("A"), key="ms")
            else:
                ms = st.selectbox("Sección:", SECCIONES, key="ms")
        with c2:
            msexo = st.selectbox("Sexo:", ["Masculino", "Femenino"], key="msexo")
            if not es_provisional:
                ma = st.text_input("Apoderado (Padre/Madre):", key="ma")
                mda = st.text_input("DNI Apoderado:", key="mda", max_chars=8)
                mc = st.text_input("Celular Apoderado:", key="mc", max_chars=9,
                                   placeholder="987654321")
            else:
                ma = ""
                mda = ""
                mc = ""
                st.caption("👆 DNI, apoderado y celular se completan luego en Base de Datos")

        if st.button("✅ MATRICULAR", type="primary", use_container_width=True,
                     key="bm"):
            if mn:
                if es_provisional:
                    # Registro provisional — sin DNI
                    with st.spinner("💾 Guardando matrícula provisional..."):
                        BaseDatos.registrar_estudiante({
                            'Nombre': mn.strip().upper(), 'DNI': '',
                            'Nivel': mnv, 'Grado': mg, 'Seccion': ms,
                            'Sexo': msexo, 'Apoderado': '', 'DNI_Apoderado': '',
                            'Celular_Apoderado': '', '_provisional': 'SI'
                        })
                        time.sleep(1)
                    avatar = "👦" if msexo == "Masculino" else "👧"
                    st.success(f"⏳ **MATRÍCULA PROVISIONAL REGISTRADA**")
                    st.markdown(f"""
                    <div class="asist-ok">
                        <strong>📋 Matrícula Provisional</strong><br>
                        {avatar} {mn.strip().upper()}<br>
                        🎓 {mg} — Sección {ms}<br>
                        📅 {fecha_peru_str()}<br>
                        <span style="color:orange;font-weight:bold;">⚠️ PENDIENTE: completar DNI y celular en Base de Datos</span>
                    </div>
                    """, unsafe_allow_html=True)
                    reproducir_beep_exitoso()
                else:
                    # Registro completo — requiere DNI
                    if not md:
                        st.error("⚠️ En modo Completo el DNI es obligatorio")
                    else:
                        md_clean = ''.join(c for c in md.strip() if c.isdigit())
                        if len(md_clean) != 8:
                            st.error(f"⚠️ El DNI debe tener 8 dígitos ({len(md_clean)} encontrados)")
                        else:
                            with st.spinner("💾 Guardando matrícula..."):
                                BaseDatos.registrar_estudiante({
                                    'Nombre': mn.strip().upper(), 'DNI': md_clean,
                                    'Nivel': mnv, 'Grado': mg, 'Seccion': ms,
                                    'Sexo': msexo, 'Apoderado': ma.strip(),
                                    'DNI_Apoderado': mda.strip(),
                                    'Celular_Apoderado': mc.strip(),
                                    '_provisional': 'NO'
                                })
                                time.sleep(2)
                            verificar = BaseDatos.buscar_por_dni(md_clean)
                            if verificar:
                                avatar = "👦" if msexo == "Masculino" else "👧"
                                st.success(f"✅ **MATRICULADO CORRECTAMENTE** ☁️ Guardado en la nube")
                                st.markdown(f"""
                                <div class="asist-ok">
                                    <strong>📋 Confirmación de Matrícula</strong><br>
                                    {avatar} {mn.strip().upper()}<br>
                                    🆔 DNI: {md_clean}<br>
                                    🎓 {mg} — Sección {ms}<br>
                                    📅 {fecha_peru_str()}<br>
                                    <span style="color:green;font-weight:bold;">☑️ VERIFICADO EN BASE DE DATOS</span>
                                </div>
                                """, unsafe_allow_html=True)
                                reproducir_beep_exitoso()
                                st.balloons()
                            else:
                                st.warning("⚠️ Se intentó guardar pero no se pudo verificar. Revise en la lista.")
            else:
                st.error("⚠️ El nombre es obligatorio")

    with tab_doc:
        st.subheader("👨‍🏫 Registro de Docente / Personal")
        c1, c2 = st.columns(2)
        with c1:
            dn_n = st.text_input("👤 Apellidos y Nombres:", key="dn_nom")
            dn_d = st.text_input("🆔 DNI:", key="dn_dni", max_chars=8)
            dn_c = st.selectbox("💼 Cargo:", [
                "Docente", "Directora", "Auxiliar", "Coordinador",
                "Secretaria", "Personal de Limpieza", "Otro"
            ], key="dn_cargo")
            dn_e = st.text_input("📚 Especialidad:", key="dn_esp",
                                  placeholder="Ej: Educación Primaria")
        with c2:
            dn_t = st.text_input("📱 Celular:", key="dn_cel", max_chars=9,
                                  placeholder="987654321")
            
            # Solo Docente/Coordinador necesitan nivel y grado
            dn_areas_sel = ""
            if dn_c in ["Docente", "Coordinador"]:
                dn_nivel = st.selectbox("🏫 Nivel:", 
                                         ["INICIAL", "PRIMARIA", "SECUNDARIA", "PREUNIVERSITARIO",
                                          "TODOS LOS NIVELES"],
                                         key="dn_nivel_reg")
                if dn_nivel == "TODOS LOS NIVELES":
                    # Ed. Física, Inglés, etc. — acceso a TODOS los grados
                    dn_g = "ALL_NIVELES"
                    dn_nivel = "TODOS"
                    st.success(f"✅ Acceso a TODOS los grados: Inicial ({len(NIVELES_GRADOS['INICIAL'])}), "
                               f"Primaria ({len(NIVELES_GRADOS['PRIMARIA'])}), "
                               f"Secundaria ({len(NIVELES_GRADOS['SECUNDARIA'])}), "
                               f"Pre-U ({len(NIVELES_GRADOS['PREUNIVERSITARIO'])})")
                    st.caption("Grados: " + ", ".join(TODOS_LOS_GRADOS))
                    dn_areas_sel = ""
                elif dn_nivel in ["INICIAL", "PRIMARIA"]:
                    dn_g = st.selectbox("🎓 Grado Asignado:",
                                         ["N/A"] + NIVELES_GRADOS.get(dn_nivel, []),
                                         key="dn_grado")
                else:
                    # SECUNDARIA y PREUNIVERSITARIO: acceso a TODOS los grados de ambos niveles
                    dn_g = "ALL_SEC_PREU"
                    grados_sec = NIVELES_GRADOS.get('SECUNDARIA', [])
                    grados_preu = NIVELES_GRADOS.get('PREUNIVERSITARIO', [])
                    st.success(f"✅ Acceso automático a TODOS los grados de Secundaria ({len(grados_sec)}) y Pre-U ({len(grados_preu)})")
                    st.caption("Grados: " + ", ".join(grados_sec + grados_preu))
                    # Todas las áreas SEC + PREU combinadas
                    todas_areas = list(AREAS_MINEDU.get('SECUNDARIA', []))
                    for a in AREAS_CEPRE_UNSAAC.get('GRUPO AB', []):
                        if a not in todas_areas:
                            todas_areas.append(a)
                    for a in AREAS_CEPRE_UNSAAC.get('GRUPO CD', []):
                        if a not in todas_areas:
                            todas_areas.append(a)
                    st.info(f"📚 {len(todas_areas)} áreas disponibles (Secundaria + CEPRE UNSAAC)")
                    dn_areas_sel = todas_areas  # Todas seleccionadas por defecto
            else:
                # Directora, Auxiliar, etc. — acceso completo sin grado
                dn_nivel = "PRIMARIA"
                dn_g = "N/A"
                st.caption(f"🔓 {dn_c}: acceso completo (sin grado específico)")
            dn_email = st.text_input("📧 Email:", key="dn_email",
                                      placeholder="nombre@ieyachay.org")
            dn_foto = st.file_uploader("📸 Foto:", type=['jpg', 'png', 'jpeg'],
                                        key="dn_foto")
            if dn_foto:
                st.image(dn_foto, width=120)
            # Opción para crear cuenta de acceso
            crear_cuenta = st.checkbox("🔐 Crear cuenta de acceso al sistema", value=True, key="crear_cuenta_doc")
            if crear_cuenta:
                cc1, cc2 = st.columns(2)
                with cc1:
                    dn_usuario = st.text_input("👤 Usuario:", 
                                                value=dn_n.strip().lower().replace(' ', '.').split('.')[0] if dn_n else '',
                                                key="dn_usuario_auto",
                                                placeholder="ej: prof.matematica")
                with cc2:
                    dn_password = st.text_input("🔑 Contraseña:", 
                                                 value=dn_d.strip() if dn_d else '',
                                                 key="dn_pass_auto",
                                                 placeholder="DNI por defecto")

        if st.button("✅ REGISTRAR DOCENTE", type="primary",
                     use_container_width=True, key="bd"):
            if dn_n and dn_d:
                if dn_foto:
                    foto_path = f"foto_doc_{dn_d.strip()}.png"
                    with open(foto_path, "wb") as fout:
                        fout.write(dn_foto.getbuffer())
                areas_txt = ', '.join(dn_areas_sel) if dn_areas_sel else dn_e.strip()
                BaseDatos.registrar_docente({
                    'Nombre': dn_n.strip().upper(), 'DNI': dn_d.strip(),
                    'Cargo': dn_c, 'Especialidad': areas_txt,
                    'Celular': dn_t.strip(), 'Grado_Asignado': dn_g,
                    'Email': dn_email.strip(), 'Nivel': dn_nivel,
                    'Areas': areas_txt
                })
                st.success(f"✅ {dn_n} registrado como {dn_c}")
                
                # Auto-crear cuenta de usuario
                if crear_cuenta and dn_usuario and dn_password:
                    usuarios = cargar_usuarios()
                    u_key = dn_usuario.strip().lower()
                    rol_auto = "docente" if dn_c == "Docente" else ("auxiliar" if dn_c == "Auxiliar" else "directivo")
                    
                    # Solo docentes tienen grado/nivel
                    if rol_auto == "docente":
                        di = {"label": dn_n.strip().upper(), "grado": dn_g, "nivel": dn_nivel, "dni": dn_d.strip()}
                    else:
                        di = None  # Directivos y auxiliares no necesitan grado
                    
                    usuarios[u_key] = {
                        "password": dn_password,
                        "rol": rol_auto,
                        "label": dn_n.strip().upper(),
                        "docente_info": di,
                    }
                    guardar_usuarios(usuarios)
                    st.success(f"🔐 Cuenta creada: **{u_key}** / contraseña: **{dn_password}** / rol: **{rol_auto}**")
                
                if dn_areas_sel:
                    st.info(f"📚 Áreas: {areas_txt}")
                reproducir_beep_exitoso()
                st.balloons()
            else:
                st.error("⚠️ Nombre y DNI requeridos")

    with tab_lista:
        st.subheader("📚 Alumnos Matriculados")
        df = BaseDatos.cargar_matricula()
        if not df.empty:
            c1, c2, c3 = st.columns(3)
            with c1:
                fn = st.selectbox("Nivel:", ["Todos"] + list(NIVELES_GRADOS.keys()),
                                  key="fn")
            with c2:
                go = ["Todos"] + (NIVELES_GRADOS[fn] if fn != "Todos"
                                  else TODOS_LOS_GRADOS)
                fg = st.selectbox("Grado:", go, key="fg")
            with c3:
                bq = st.text_input("🔍 Buscar:", key="bq")
            d = df.copy()
            if fn != "Todos" and 'Nivel' in d.columns:
                d = d[d['Nivel'] == fn]
            if fg != "Todos" and 'Grado' in d.columns:
                d = d[d['Grado'] == fg]
            if bq:
                d = d[d.apply(lambda r: bq.lower() in str(r).lower(), axis=1)]
            if 'Nombre' in d.columns:
                d = d.sort_values('Nombre')
            st.metric("Resultados", len(d))
            st.dataframe(d, use_container_width=True, hide_index=True, height=400)
            buf = io.BytesIO()
            d.to_excel(buf, index=False, engine='openpyxl')
            buf.seek(0)
            st.download_button("⬇️ Excel", buf,
                               f"Matricula_{config['anio']}.xlsx", key="dme")
            # Solo admin puede eliminar
            if puede_borrar():
                with st.expander("🗑️ Eliminar Alumno"):
                    deld = st.text_input("DNI a eliminar:", key="dd")
                    if st.button("🗑️ Eliminar", key="bdel", type="primary"):
                        if deld:
                            BaseDatos.eliminar_estudiante(deld)
                            st.rerun()
        else:
            st.info("📝 Sin alumnos matriculados.")

        st.markdown("---")
        st.subheader("👨‍🏫 Docentes Registrados")
        df_doc = BaseDatos.cargar_docentes()
        if not df_doc.empty:
            if 'Nombre' in df_doc.columns:
                df_doc = df_doc.sort_values('Nombre')
            st.dataframe(df_doc, use_container_width=True, hide_index=True)
            buf2 = io.BytesIO()
            df_doc.to_excel(buf2, index=False, engine='openpyxl')
            buf2.seek(0)
            st.download_button("⬇️ Excel Docentes", buf2,
                               "docentes.xlsx", key="dmedoc")
            if puede_borrar():
                with st.expander("🗑️ Eliminar Docente"):
                    deld2 = st.text_input("DNI:", key="dddoc")
                    if st.button("🗑️ Eliminar", key="bdeldoc", type="primary"):
                        if deld2:
                            BaseDatos.eliminar_docente(deld2)
                            st.rerun()
        else:
            st.info("📝 Sin docentes registrados.")

    with tab_pdf:
        _seccion_registros_pdf(config)


def _seccion_registros_pdf(config):
    df = BaseDatos.cargar_matricula()
    if df.empty:
        st.info("📝 Registra estudiantes primero.")
        return
    c1, c2 = st.columns(2)
    with c1:
        np_ = st.selectbox("Nivel:", list(NIVELES_GRADOS.keys()), key="pn")
        gp = st.selectbox("Grado:", NIVELES_GRADOS[np_], key="pg")
    with c2:
        sp = st.selectbox("Sección:", ["Todas"] + SECCIONES, key="ps")
    dg = BaseDatos.obtener_estudiantes_grado(gp, sp)
    st.info(f"📊 {len(dg)} estudiantes (orden alfabético)")

    st.markdown("---")
    st.markdown("**📝 Registro Auxiliar (Cursos × Competencias × Desempeños)**")
    bim = st.selectbox("📅 Periodo:", list(BIMESTRES.keys()), key="bim_sel")
    st.markdown("**Cursos (hasta 3 por hoja):**")
    c1, c2, c3 = st.columns(3)
    with c1:
        curso1 = st.text_input("Curso 1:", "Matemática", key="c1")
    with c2:
        curso2 = st.text_input("Curso 2:", "Comunicación", key="c2")
    with c3:
        curso3 = st.text_input("Curso 3:", "Ciencia y Tec.", key="c3")
    cursos = [c for c in [curso1, curso2, curso3] if c.strip()]
    st.caption(f"{len(cursos)} cursos × 4 competencias × 3 desempeños")
    c1a, c2a = st.columns(2)
    with c1a:
        if st.button("📝 Generar PDF Auxiliar", type="primary",
                     use_container_width=True, key="gra"):
            sl = sp if sp != "Todas" else "Todas"
            with st.spinner("Generando PDF..."):
                pdf = generar_registro_auxiliar_pdf(gp, sl, config['anio'], bim, dg, cursos)
            st.session_state['_aux_pdf'] = pdf
            st.session_state['_aux_key'] = f"RegAux_{gp}_{bim}"
    with c2a:
        if st.button("📄 Generar Word Auxiliar", use_container_width=True, key="gra_docx"):
            sl = sp if sp != "Todas" else "Todas"
            with st.spinner("Generando Word..."):
                docx_aux = generar_registro_auxiliar_docx(gp, sl, config['anio'], bim, dg, cursos)
            if docx_aux:
                st.session_state['_aux_docx'] = docx_aux
                st.session_state['_aux_key'] = f"RegAux_{gp}_{bim}"
    if st.session_state.get('_aux_pdf'):
        st.download_button("⬇️ Descargar PDF", st.session_state['_aux_pdf'],
                           f"{st.session_state.get('_aux_key','RegAux')}.pdf",
                           "application/pdf", key="dra")
    if st.session_state.get('_aux_docx'):
        st.download_button("⬇️ Descargar Word (.docx)", st.session_state['_aux_docx'],
                           f"{st.session_state.get('_aux_key','RegAux')}.docx",
                           "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           key="dra_docx")

    st.markdown("---")
    st.markdown("**📋 Registro Asistencia (sin sáb/dom, sin feriados)**")
    meses_opts = list(MESES_ESCOLARES.items())
    meses_sel = st.multiselect(
        "Meses:",
        [f"{v} ({k})" for k, v in meses_opts],
        default=[f"{v} ({k})" for k, v in meses_opts[:3]],
        key="msel"
    )
    meses_nums = [int(m.split('(')[1].replace(')', '')) for m in meses_sel]
    if st.button("📋 Generar Registro Asistencia PDF", type="primary",
                 use_container_width=True, key="gras"):
        if meses_nums:
            sl = sp if sp != "Todas" else "Todas"
            pdf = generar_registro_asistencia_pdf(
                gp, sl, config['anio'], dg, meses_nums
            )
            st.download_button("⬇️ Descargar", pdf,
                               f"RegAsist_{gp}.pdf", "application/pdf", key="dras")


# ================================================================
# TAB: DOCUMENTOS
# ================================================================

def tab_documentos(config):
    st.header("📄 Documentos")
    c1, c2 = st.columns([1, 2])
    with c1:
        td = st.selectbox("📑 Tipo:", [
            "CONSTANCIA DE VACANTE", "CONSTANCIA DE NO DEUDOR",
            "CONSTANCIA DE ESTUDIOS", "CONSTANCIA DE CONDUCTA",
            "CARTA COMPROMISO", "RESOLUCIÓN DE TRASLADO"
        ], key="td")
        st.markdown("---")
        db = st.text_input("🔍 Buscar DNI:", key="db")
        if st.button("🔎 Buscar", use_container_width=True, key="bb", type="primary"):
            r = BaseDatos.buscar_por_dni(db)
            if r:
                st.session_state.alumno = r.get('Nombre', '')
                st.session_state.dni = r.get('DNI', '')
                st.session_state.grado = r.get('Grado', '')
                st.session_state.apoderado = r.get('Apoderado', '')
                st.session_state.dni_apo = r.get('DNI_Apoderado', '')
                st.success("🎉 Datos cargados")
                st.rerun()
            else:
                st.error("❌ No encontrado")
    with c2:
        with st.container(border=True):
            nm = st.text_input("👤 Estudiante:", key="alumno")
            dn = st.text_input("🆔 DNI Estudiante:", key="dni")
            gr = st.text_input("📚 Grado:", key="grado")
            ap = st.text_input("👨‍👩‍👧 Padre/Madre/Apoderado:", key="apoderado")
            da = st.text_input("🆔 DNI Padre/Madre/Apoderado:", key="dni_apo")
            nc = {}
            if td == "CONSTANCIA DE CONDUCTA":
                cols = st.columns(5)
                for i, col in enumerate(cols):
                    with col:
                        nc[f'nota_conducta_{i+1}'] = st.selectbox(
                            f"{i+1}°", ["AD", "A", "B", "C"], key=f"n{i}")
            ex = {}
            if td == "RESOLUCIÓN DE TRASLADO":
                ex['num_resolucion'] = st.text_input("N° Resolución:", key="nr")
                ex['fecha_resolucion'] = st.text_input("Fecha:", key="fr2")
                ex['nivel'] = st.selectbox("Nivel:",
                                           ["INICIAL", "PRIMARIA", "SECUNDARIA"],
                                           key="nl")
                ex['ie_destino'] = st.text_input("IE Destino:", key="ie")
        if st.button("✨ GENERAR DOCUMENTO", type="primary",
                     use_container_width=True, key="gd"):
            if nm and dn:
                d = {'alumno': nm, 'dni': dn, 'grado': gr,
                     'apoderado': ap, 'dni_apo': da, **nc, **ex}
                g = GeneradorPDF(config)
                metodos = {
                    "CONSTANCIA DE VACANTE": g.generar_constancia_vacante,
                    "CONSTANCIA DE NO DEUDOR": g.generar_constancia_no_deudor,
                    "CONSTANCIA DE ESTUDIOS": g.generar_constancia_estudios,
                    "CONSTANCIA DE CONDUCTA": g.generar_constancia_conducta,
                    "CARTA COMPROMISO": g.generar_carta_compromiso,
                    "RESOLUCIÓN DE TRASLADO": g.generar_resolucion_traslado,
                }
                pdf = metodos[td](d)
                st.success("🎉 Documento generado")
                st.download_button("⬇️ Descargar PDF", pdf,
                                   f"{nm}_{td}.pdf", "application/pdf",
                                   use_container_width=True, key="dd2")


# ================================================================
# TAB: CARNETS (Individual, Matrícula, Lote Alumnos PDF, Lote Docentes PDF)
# ================================================================

def tab_carnets(config):
    st.header("🪪 Centro de Carnetización")
    t1, t2, t3, t4 = st.tabs([
        "⚡ Individual", "📋 Desde Matrícula",
        "📦 Lote Alumnos (PDF)", "👨‍🏫 Lote Docentes (PDF)"
    ])

    with t1:
        c1, c2 = st.columns(2)
        with c1:
            c_tipo = st.radio("Tipo de carnet:", ["🎓 Alumno", "👨‍🏫 Docente"],
                              horizontal=True, key="c_tipo")
            es_doc_ind = "Docente" in c_tipo
            cn = st.text_input("👤 Nombre:", key="cn")
            cd = st.text_input("🆔 DNI:", key="cd")
            if es_doc_ind:
                c_cargo = st.selectbox("💼 Cargo:", ["Docente", "Directora", "Coordinador",
                                                      "Auxiliar", "Administrativo"], key="c_cargo")
                c_esp = st.text_input("📚 Especialidad:", key="c_esp",
                                      placeholder="Ej: Matemática")
            else:
                cg = st.selectbox("📚 Grado:", TODOS_LOS_GRADOS, key="cg")
                cs = st.selectbox("📂 Sección:", SECCIONES, key="cs")
        with c2:
            cf = st.file_uploader("📸 Foto:", type=['jpg', 'png', 'jpeg'], key="cf")
            if cf:
                st.image(cf, width=180)
        if st.button("🪪 GENERAR CARNET", type="primary",
                     use_container_width=True, key="gc"):
            if cn and cd:
                fi = io.BytesIO(cf.getvalue()) if cf else None
                if es_doc_ind:
                    datos_c = {'Nombre': cn, 'DNI': cd, 'Cargo': c_cargo,
                               'Especialidad': c_esp, 'Grado': ''}
                else:
                    datos_c = {'Nombre': cn, 'DNI': cd, 'Grado': cg, 'Seccion': cs}
                cr = GeneradorCarnet(
                    datos_c, config['anio'], fi, es_docente=es_doc_ind
                ).generar()
                st.image(cr, use_container_width=True)
                st.download_button("⬇️ Descargar", cr,
                                   f"Carnet_{cn.replace(' ', '_')}.png",
                                   "image/png", use_container_width=True, key="dc")

    with t2:
        dbs = st.text_input("🔍 DNI:", key="cbd")
        if st.button("🔎 Buscar", key="cbb", type="primary"):
            a = BaseDatos.buscar_por_dni(dbs)
            if a:
                st.session_state['ce'] = a
                st.success(f"✅ {a.get('Nombre', '')}")
            else:
                st.error("❌ No encontrado")
        if st.session_state.get('ce') and isinstance(st.session_state['ce'], dict):
            a = st.session_state['ce']
            es_d = a.get('_tipo', '') == 'docente'
            tt = "DOCENTE" if es_d else "ALUMNO"
            st.markdown(f"**[{tt}]** {a.get('Nombre', '')} | DNI: {a.get('DNI', '')}")
            fm = st.file_uploader("📸 Foto:", type=['jpg', 'png', 'jpeg'], key="cfm")
            if st.button("🪪 GENERAR", type="primary",
                         use_container_width=True, key="gcm"):
                fi = io.BytesIO(fm.getvalue()) if fm else None
                cr = GeneradorCarnet(a, config['anio'], fi, es_docente=es_d).generar()
                st.image(cr, use_container_width=True)
                st.download_button("⬇️", cr, "Carnet.png", "image/png",
                                   use_container_width=True, key="dcm")

    with t3:
        st.subheader("📦 Carnets Alumnos — PDF (8 por hoja)")
        st.caption("Tamaño fotocheck con líneas de corte para plastificar")
        df = BaseDatos.cargar_matricula()
        if not df.empty:
            nl = st.selectbox("Nivel:", ["Todos"] + list(NIVELES_GRADOS.keys()),
                              key="ln")
            d = df.copy()
            if nl != "Todos" and 'Nivel' in d.columns:
                d = d[d['Nivel'] == nl]
            if 'Nombre' in d.columns:
                d = d.sort_values('Nombre')
            st.info(f"📊 {len(d)} carnets de alumnos")
            if st.button("🚀 GENERAR PDF CARNETS", type="primary",
                         use_container_width=True, key="gl"):
                progreso = st.progress(0)
                lista = d.to_dict('records')
                pdf = generar_carnets_lote_pdf(lista, config['anio'], es_docente=False)
                progreso.progress(100)
                st.balloons()
                st.download_button("⬇️ DESCARGAR PDF", pdf,
                                   f"Carnets_Alumnos_{config['anio']}.pdf",
                                   "application/pdf", use_container_width=True,
                                   key="dlz")
        else:
            st.info("📝 Registra estudiantes.")

    with t4:
        st.subheader("👨‍🏫 Carnets Docentes — PDF (8 por hoja)")
        st.caption("Tamaño fotocheck con líneas de corte para plastificar")
        df_doc = BaseDatos.cargar_docentes()
        if not df_doc.empty:
            if 'Nombre' in df_doc.columns:
                df_doc = df_doc.sort_values('Nombre')
            st.info(f"📊 {len(df_doc)} carnets de docentes")
            st.dataframe(df_doc[['Nombre', 'DNI', 'Cargo']],
                         use_container_width=True, hide_index=True)
            if st.button("🚀 GENERAR PDF CARNETS DOCENTES", type="primary",
                         use_container_width=True, key="gld"):
                lista = df_doc.to_dict('records')
                pdf = generar_carnets_lote_pdf(lista, config['anio'], es_docente=True)
                st.balloons()
                st.download_button("⬇️ DESCARGAR PDF", pdf,
                                   f"Carnets_Docentes_{config['anio']}.pdf",
                                   "application/pdf", use_container_width=True,
                                   key="dlzd")
        else:
            st.info("📝 Registra docentes en Matrícula.")


# ================================================================
# TAB: ASISTENCIAS (Alumnos + Docentes)
# ================================================================

def tab_asistencias():
    st.header("📋 Control de Asistencia")
    st.caption(f"🕒 **{hora_peru().strftime('%H:%M:%S')}** | "
               f"📅 {hora_peru().strftime('%d/%m/%Y')}")

    # Inicializar tracking de WhatsApp enviados
    if 'wa_enviados' not in st.session_state:
        st.session_state.wa_enviados = set()

    # ── Horario y Modo ──────────────────────────────────────────
    col_h, col_modo = st.columns([2, 3])
    with col_h:
        horario_sel = st.radio("⏰ Horario:", ['normal', 'invierno'],
                                format_func=lambda x: HORARIOS[x]['nombre'],
                                horizontal=True, key="horario_radio",
                                index=0 if _horario_activo() == 'normal' else 1)
        _guardar_horario(horario_sel)
        limite = HORARIOS[horario_sel]['limite']
        st.caption(f"Límite puntualidad: **{limite}** — después = tardanza automática")

    with col_modo:
        c1, c2 = st.columns(2)
        with c1:
            if st.button("🟢 ENTRADA", use_container_width=True, key="be", type="primary"):
                st.session_state.tipo_asistencia = "Entrada"
                st.rerun()
        with c2:
            if st.button("🔵 SALIDA", use_container_width=True, key="bs", type="primary"):
                st.session_state.tipo_asistencia = "Salida"
                st.rerun()
        st.caption("💡 El sistema detecta automáticamente si es turno mañana o tarde")

    _color_modo = {"Entrada": "#16a34a", "Salida": "#2563eb"}
    _modo = st.session_state.get('tipo_asistencia', 'Entrada')
    modo_label = _modo.replace('_', ' ')
    st.markdown(f"<div style='background:{_color_modo.get(_modo,'#2563eb')};color:white;padding:8px 14px;border-radius:8px;font-weight:bold;'>📌 Modo: {modo_label} | Horario: {HORARIOS[horario_sel]['nombre']} — Tardanza auto después de {limite}</div>", unsafe_allow_html=True)
    st.markdown("---")

    # ===== ZONA DE REGISTRO RÁPIDO =====
    cc, cm = st.columns(2)
    with cc:
        st.markdown("### 📸 Escanear QR / Código")
        act = st.checkbox("📷 Activar cámara", key="chkc",
                          value=st.session_state.get('activar_camara_asist', False))
        st.session_state.activar_camara_asist = act
        if act:
            foto = st.camera_input("Apunta al QR:", key="ca")
            if foto:
                d = decodificar_qr_imagen(foto.getvalue())
                if d:
                    _registrar_asistencia_rapida(d)
                else:
                    st.warning("⚠️ QR no detectado.")
        else:
            st.info("💡 Activa la cámara para escanear.")
    with cm:
        st.markdown("### ✏️ Registro Manual / Lector de Código de Barras")
        st.caption("💡 Con lector de barras: apunte al carnet y se registra automáticamente")

        # Callback que se ejecuta al cambiar el campo (Enter o scanner)
        def _on_dni_submit():
            val = st.session_state.get('dm_input', '').strip()
            dni_limpio = ''.join(c for c in val if c.isdigit())
            if len(dni_limpio) >= 7:  # Aceptar 7-8 dígitos
                # Guardar DNI pendiente para procesarlo fuera del callback
                st.session_state['_dni_pendiente'] = dni_limpio[:8]
            # Limpiar campo inmediatamente
            st.session_state['dm_input'] = ''

        dm = st.text_input("DNI:", key="dm_input",
                           placeholder="Escanee código de barras o escriba DNI + Enter",
                           on_change=_on_dni_submit)

        # Procesar DNI pendiente FUERA del callback (más estable)
        _dni_pend = st.session_state.pop('_dni_pendiente', None)
        if _dni_pend:
            _registrar_asistencia_rapida(_dni_pend)
        
        # También procesar si escribieron DNI y no activó on_change
        if dm:
            dni_directo = ''.join(c for c in dm.strip() if c.isdigit())
            if len(dni_directo) == 8:
                _registrar_asistencia_rapida(dni_directo)
                st.session_state['dm_input'] = ''

        # Sonido/vibración via JS después de registrar
        if not dm:  # Campo fue limpiado = se registró
            st.markdown("""
            <script>
            (function() {
                if ('vibrate' in navigator) { navigator.vibrate([200, 100, 200]); }
                try {
                    var ctx = new (window.AudioContext || window.webkitAudioContext)();
                    var o = ctx.createOscillator();
                    o.type = 'sine'; o.frequency.value = 800;
                    o.connect(ctx.destination);
                    o.start(); setTimeout(function(){ o.stop(); }, 200);
                } catch(e) {}
                // Devolver foco al campo para siguiente escaneo
                setTimeout(function() {
                    var inputs = window.parent.document.querySelectorAll('input[placeholder*="DNI"]');
                    if (inputs.length > 0) inputs[0].focus();
                }, 100);
            })();
            </script>
            """, unsafe_allow_html=True)

    # ===== LISTA DE ASISTENCIA DE HOY =====
    st.markdown("---")
    st.subheader("📊 Registros de Hoy")
    asis = BaseDatos.obtener_asistencias_hoy()
    if asis:
        # Separar alumnos y docentes
        alumnos_h = []
        docentes_h = []
        for dk, v in asis.items():
            reg = {'DNI': dk, 'Nombre': v['nombre'],
                   'Entrada': v.get('entrada', '—'),
                   'Tardanza': '⏰' if _es_tardanza(v.get('entrada', '') or v.get('tardanza', '')) and (v.get('entrada') or v.get('tardanza')) else '—',
                   'Salida': v.get('salida', '—'),
                   'Ent.Tarde': v.get('entrada_tarde', '—'),
                   'Sal.Tarde': v.get('salida_tarde', '—'),
                   'es_docente': v.get('es_docente', False)}
            if v.get('es_docente', False):
                docentes_h.append(reg)
            else:
                alumnos_h.append(reg)

        # Métricas rápidas
        c1, c2, c3, c4, c5 = st.columns(5)
        with c1:
            st.metric("📚 Alumnos", len(alumnos_h))
        with c2:
            st.metric("👨‍🏫 Docentes", len(docentes_h))
        with c3:
            entradas = sum(1 for v in asis.values() if v.get('entrada'))
            st.metric("🌅 Entradas", entradas)
        with c4:
            tardanzas = sum(1 for v in asis.values()
                          if _es_tardanza(v.get('entrada', '') or v.get('tardanza', ''))
                          and (v.get('entrada') or v.get('tardanza')))
            st.metric("⏰ Tardanzas", tardanzas)
        with c5:
            salidas = sum(1 for v in asis.values() if v.get('salida'))
            st.metric("🌙 Salidas", salidas)

        if alumnos_h:
            st.markdown("**📚 Alumnos registrados:**")
            st.dataframe(pd.DataFrame(alumnos_h).drop(columns=['es_docente']),
                         use_container_width=True, hide_index=True)
        if docentes_h:
            st.markdown("**👨‍🏫 Docentes registrados:**")
            st.dataframe(pd.DataFrame(docentes_h).drop(columns=['es_docente']),
                         use_container_width=True, hide_index=True)

        # ===== ZONA WHATSAPP — TABS ENTRADA / SALIDA =====
        st.markdown("---")
        st.subheader("📱 Enviar Notificaciones WhatsApp")
        st.caption("Toque cada botón para enviar. Al marcar ✅ desaparece de la lista.")

        tab_ent, tab_tard, tab_sal = st.tabs(["🌅 Entrada", "⏰ Tardanza", "🌙 Salida"])

        def _render_wa_tab(tipo_tab):
            asis_fresh = BaseDatos.obtener_asistencias_hoy()
            pendientes = 0
            enviados = 0
            sin_celular = []
            links_pendientes = []  # Para el botón enviar todo

            for dk, dat in asis_fresh.items():
                hora_reg = dat.get(tipo_tab, '')
                if not hora_reg:
                    continue

                clave_envio = f"{dk}_{tipo_tab}_{fecha_peru_str()}"
                ya_enviado = clave_envio in st.session_state.wa_enviados
                if ya_enviado:
                    enviados += 1
                    continue

                nombre = dat['nombre']
                es_doc = dat.get('es_docente', False)
                tipo_icon = "👨‍🏫" if es_doc else "📚"

                cel = ''
                # ── Buscar celular en AMBAS tablas siempre ───────────────
                # 1) DOCENTES — por DNI y por nombre
                df_doc_wa = st.session_state.get('_cache_docentes_wa', pd.DataFrame())
                if df_doc_wa.empty:
                    df_doc_wa = BaseDatos.cargar_docentes()
                    st.session_state['_cache_docentes_wa'] = df_doc_wa
                if not df_doc_wa.empty and 'Celular' in df_doc_wa.columns:
                    # Por DNI
                    if 'DNI' in df_doc_wa.columns:
                        fila_d = df_doc_wa[df_doc_wa['DNI'].astype(str).str.strip() == str(dk).strip()]
                        if not fila_d.empty:
                            cv = str(fila_d.iloc[0].get('Celular', '')).strip()
                            if cv and cv not in ('nan', 'None', '', 'NaN'):
                                cel = cv
                    # Por nombre (fallback si DNI no coincide)
                    if not cel and 'Nombre' in df_doc_wa.columns:
                        nm_up = nombre.strip().upper()
                        fila_n = df_doc_wa[df_doc_wa['Nombre'].astype(str).str.strip().str.upper() == nm_up]
                        if not fila_n.empty:
                            cv = str(fila_n.iloc[0].get('Celular', '')).strip()
                            if cv and cv not in ('nan', 'None', '', 'NaN'):
                                cel = cv

                # 2) MATRÍCULA — alumnos por DNI
                if not cel:
                    df_m = st.session_state.get('_cache_matricula', pd.DataFrame())
                    if df_m.empty:
                        df_m = BaseDatos.cargar_matricula()
                        st.session_state['_cache_matricula'] = df_m
                    if not df_m.empty and 'DNI' in df_m.columns:
                        fila = df_m[df_m['DNI'].astype(str).str.strip() == str(dk).strip()]
                        if not fila.empty:
                            cv = str(fila.iloc[0].get('Celular_Apoderado',
                                     fila.iloc[0].get('Celular', ''))).strip()
                            if cv and cv not in ('nan', 'None', '', 'NaN'):
                                cel = cv

                # Limpiar número
                if cel:
                    if '.' in cel:
                        cel = cel.split('.')[0]
                    cel = ''.join(c for c in cel if c.isdigit())
                    cel = '' if len(cel) < 7 else cel

                if not cel:
                    sin_celular.append(f"{tipo_icon} {nombre}")
                    continue

                pendientes += 1
                msg = generar_mensaje_asistencia(nombre, tipo_tab, hora_reg)
                link = generar_link_whatsapp(cel, msg)
                links_pendientes.append({'link': link, 'clave': clave_envio, 'nombre': nombre, 'hora': hora_reg, 'cel': cel, 'icon': tipo_icon})

                col_btn, col_check = st.columns([4, 1])
                with col_btn:
                    st.markdown(
                        f'<a href="{link}" target="_blank" class="wa-btn">'
                        f'📱 {tipo_icon} {nombre} — 🕒 {hora_reg} → {cel}</a>',
                        unsafe_allow_html=True)
                with col_check:
                    if st.button("✅", key=f"wa_{dk}_{tipo_tab}",
                                 help="Marcar como enviado y quitar de lista", type="primary"):
                        st.session_state.wa_enviados.add(clave_envio)
                        st.rerun()

            if links_pendientes:
                for item in links_pendientes:
                    col_btn, col_check = st.columns([4, 1])
                    with col_btn:
                        st.markdown(
                            f'<a href="{item["link"]}" target="_blank" class="wa-btn">'
                            f'📱 {item["icon"]} {item["nombre"]} — 🕒 {item["hora"]} → {item["cel"]}</a>',
                            unsafe_allow_html=True)
                    with col_check:
                        if st.button("✅", key=f"next_wa_{tipo_tab}_{item['clave']}", type="primary",
                                     help="Enviado — quitar de lista"):
                            st.session_state.wa_enviados.add(item['clave'])
                            st.rerun()

            if sin_celular:
                with st.expander(f"⚠️ {len(sin_celular)} sin celular registrado"):
                    for s in sin_celular:
                        st.caption(f"• {s}")

            _total_enviados = len(st.session_state.wa_enviados)
            if pendientes == 0 and _total_enviados > 0:
                st.success(f"🎉 ¡Todos enviados! ({_total_enviados} mensajes)")
            elif pendientes == 0 and _total_enviados == 0:
                st.info("No hay registros de este tipo aún.")

        with tab_ent:
            _render_wa_tab("entrada")
        with tab_tard:
            _render_wa_tab("tardanza")
        with tab_sal:
            _render_wa_tab("salida")

        # Botón para resetear marcas de enviado
        if st.session_state.wa_enviados:
            if st.button("🔄 Resetear marcas de enviado", key="reset_wa", type="primary"):
                st.session_state.wa_enviados = set()
                st.rerun()

        st.markdown("---")
        # Solo admin puede borrar
        if puede_borrar():
            if st.button("🗑️ BORRAR ASISTENCIAS DEL DÍA", type="secondary",
                         use_container_width=True, key="borrar_asist"):
                BaseDatos.borrar_asistencias_hoy()
                st.session_state.wa_enviados = set()
                st.success("🎉 Eliminadas")
                st.rerun()
    else:
        st.info("📝 No hay registros hoy. Escanee QR o ingrese DNI para registrar.")


def _registrar_asistencia_rapida(dni):
    """Registra asistencia — AUTO-DETECTA mañana/tarde, tardanza, horas docente"""
    persona = BaseDatos.buscar_por_dni(dni)
    if persona:
        hora = hora_peru_str()
        modo = st.session_state.get('tipo_asistencia', 'Entrada')
        es_d = persona.get('_tipo', '') == 'docente'

        # Obtener nombre
        if es_d:
            df_doc = BaseDatos.cargar_docentes()
            if not df_doc.empty and 'DNI' in df_doc.columns:
                df_doc['DNI'] = df_doc['DNI'].astype(str).str.strip()
                doc_e = df_doc[df_doc['DNI'] == str(dni).strip()]
                nombre = doc_e.iloc[0]['Nombre'] if not doc_e.empty else persona.get('Nombre', '')
            else:
                nombre = persona.get('Nombre', '')
        else:
            nombre = persona.get('Nombre', '')

        tp = "👨‍🏫 DOCENTE" if es_d else "📚 ALUMNO"
        limite_txt = HORARIOS[_horario_activo()]['limite']

        # ── AUTO-DETECTAR TURNO mañana/tarde ─────────────────────────
        asis_hoy = BaseDatos.obtener_asistencias_hoy()
        reg_hoy = asis_hoy.get(str(dni).strip(), {})

        if modo == "Entrada":
            tiene_entrada = reg_hoy.get('entrada') or reg_hoy.get('tardanza')
            tiene_salida = reg_hoy.get('salida')
            tiene_ent_tarde = reg_hoy.get('entrada_tarde')

            if tiene_entrada and tiene_salida and not tiene_ent_tarde:
                # Mañana completa → auto ENTRADA TARDE
                tipo = 'entrada_tarde'
                emoji_tipo = "🌤️"
                msg_extra = " 📌 TURNO TARDE (auto-detectado)"
            elif tiene_entrada and not tiene_salida:
                st.warning(f"⚠️ **{nombre}** ya registró entrada ({reg_hoy.get('entrada') or reg_hoy.get('tardanza')}). Registre salida primero.")
                return
            elif tiene_ent_tarde:
                st.warning(f"⚠️ **{nombre}** ya tiene entrada tarde ({tiene_ent_tarde}). Ya registrado.")
                return
            else:
                # Primera entrada del día
                if _es_tardanza(hora):
                    tipo = 'tardanza'
                    emoji_tipo = "🟡"
                    msg_extra = f" ⏰ TARDANZA (después de {limite_txt})"
                else:
                    tipo = 'entrada'
                    emoji_tipo = "🟢"
                    msg_extra = " ✅ PUNTUAL"

        elif modo == "Salida":
            tiene_ent_tarde = reg_hoy.get('entrada_tarde')
            tiene_sal_tarde = reg_hoy.get('salida_tarde')
            tiene_salida = reg_hoy.get('salida')

            if tiene_ent_tarde and not tiene_sal_tarde:
                # Tarde abierta → auto SALIDA TARDE
                tipo = 'salida_tarde'
                emoji_tipo = "🌙"
                msg_extra = " 📌 SALIDA TARDE (auto-detectado)"
            elif tiene_sal_tarde:
                st.warning(f"⚠️ **{nombre}** ya completó ambos turnos hoy.")
                return
            elif tiene_salida and not tiene_ent_tarde:
                st.info(f"ℹ️ **{nombre}** ya salió de mañana ({tiene_salida}). Presione Entrada para turno tarde.")
                return
            else:
                tipo = 'salida'
                emoji_tipo = "🔵"
                msg_extra = ""
        else:
            tipo = modo.lower()
            emoji_tipo = "⚪"
            msg_extra = ""

        # ── Calcular horas trabajadas para docentes ──────────────────
        horas_info = ""
        if es_d and tipo in ('salida', 'salida_tarde'):
            try:
                ent_key = 'entrada_tarde' if tipo == 'salida_tarde' else 'entrada'
                hora_ent = reg_hoy.get(ent_key) or reg_hoy.get('tardanza', '')
                if hora_ent:
                    h1, m1 = hora_ent.split(':')[:2]
                    h2, m2 = hora.split(':')[:2]
                    mins = (int(h2)*60+int(m2)) - (int(h1)*60+int(m1))
                    if mins > 0:
                        horas_info = f" | ⏱️ {mins//60}h{mins%60:02d}m"
            except Exception:
                pass

        try:
            BaseDatos.guardar_asistencia(dni, nombre, tipo, hora, es_docente=es_d)
        except Exception as e:
            st.error(f"❌ Error al guardar: {e}")
            return

        label = tipo.replace('_', ' ').title()
        color_div = "ok" if 'entrada' in tipo else "salida"
        st.toast(f"{emoji_tipo} {tp} {nombre} — {label}: {hora}{msg_extra}", icon="✅")
        st.markdown(f"""<div class="asist-{color_div}">
            {emoji_tipo} <strong>[{tp}] {nombre}</strong> — {label}: <strong>{hora}</strong>{msg_extra}{horas_info}
        </div>""", unsafe_allow_html=True)
        reproducir_beep_exitoso()
    else:
        st.warning(f"⚠️ DNI **{dni}** no está en matrícula. Puede registrarlo manualmente:")
        nombre_manual = st.text_input("Nombre completo:", key=f"nombre_manual_{dni}",
                                      placeholder="Ej: FLORES QUISPE JUAN")
        if nombre_manual and st.button("✅ Registrar de todas formas", key=f"reg_manual_{dni}", type="primary"):
            hora = hora_peru_str()
            tipo = st.session_state.get('tipo_asistencia', 'Entrada').lower()
            BaseDatos.guardar_asistencia(dni, nombre_manual.upper().strip(), tipo, hora, es_docente=False)
            st.success(f"✅ Registrado: {nombre_manual.upper()} — {hora}")
            reproducir_beep_exitoso()
            st.info("💡 Recuerda matricular a este estudiante para que aparezca normalmente.")




def nota_a_letra(nota):
    if nota >= 18: return 'AD'
    elif nota >= 14: return 'A'
    elif nota >= 11: return 'B'
    else: return 'C'

# Escala de calificación MINEDU — global para uso en toda la app
ESCALA_MINEDU = {
    'AD': {'nombre': 'Logro Destacado',  'rango': '18-20', 'color': '#15803d', 'desc': 'El estudiante evidencia un nivel superior a lo esperado respecto a la competencia.'},
    'A':  {'nombre': 'Logro Esperado',   'rango': '14-17', 'color': '#2563eb', 'desc': 'El estudiante evidencia el nivel esperado respecto a la competencia.'},
    'B':  {'nombre': 'En Proceso',        'rango': '11-13', 'color': '#d97706', 'desc': 'El estudiante está próximo al nivel esperado, requiere acompañamiento.'},
    'C':  {'nombre': 'En Inicio',         'rango': '0-10',  'color': '#dc2626', 'desc': 'El estudiante muestra un progreso mínimo. Requiere mayor tiempo e intervención.'},
}

def color_semaforo(letra):
    return ESCALA_MINEDU.get(letra, {}).get('color', '#888')

def generar_reporte_estudiante_pdf(nombre, dni, grado, resultados_hist, config):
    """PDF individual del estudiante con semáforo AD/A/B/C y recomendaciones"""
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    w, h = A4

    # ── Marca de agua ────────────────────────────────────────────────────
    if Path("escudo_upload.png").exists():
        try:
            from PIL import Image as PILImage
            img = PILImage.open("escudo_upload.png")
            iw, ih = img.size
            ratio = iw / ih
            mw = 420; mh = mw / ratio
            c.saveState()
            c.setFillAlpha(0.35)
            c.drawImage("escudo_upload.png", w/2-mw/2, h/2-mh/2, mw, mh, mask='auto')
            c.restoreState()
        except Exception:
            pass

    # ── Barra azul superior ──────────────────────────────────────────────
    c.setFillColor(colors.HexColor("#001e7c"))
    c.rect(0, h-15, w, 15, fill=1, stroke=0)

    # ── Escudo izquierda y derecha (proporción correcta) ─────────────────
    ALTO_ESC = 65
    esc_izq = "escudo_upload.png"
    esc_der = "escudo2_upload.png" if Path("escudo2_upload.png").exists() else "escudo_upload.png"
    try:
        from PIL import Image as PILImage
        if Path(esc_izq).exists():
            img = PILImage.open(esc_izq)
            iw, ih = img.size
            aw = ALTO_ESC * (iw/ih)
            c.drawImage(esc_izq, 18, h-12-ALTO_ESC, aw, ALTO_ESC, mask='auto')
        if Path(esc_der).exists():
            img2 = PILImage.open(esc_der)
            iw2, ih2 = img2.size
            aw2 = ALTO_ESC * (iw2/ih2)
            _alto_der2 = 80
            _aw_der2 = _alto_der2 * (iw2/ih2)
            c.drawImage(esc_der, w-18-_aw_der2, h-12-_alto_der2, _aw_der2, _alto_der2, mask='auto')
    except Exception:
        pass

    # ── Textos institucionales ───────────────────────────────────────────
    c.setFillColor(colors.HexColor("#001e7c"))
    c.setFont("Helvetica-Bold", 7.5)
    c.drawCentredString(w/2, h-28, "MINISTERIO DE EDUCACIÓN — DRE CUSCO — PIONEROS EN LA EDUCACION DE CALIDAD")
    c.setFont("Helvetica-Bold", 11)
    c.drawCentredString(w/2, h-43, "I.E.P. YACHAY — CHINCHERO")
    c.setFont("Helvetica-Bold", 13)
    c.drawCentredString(w/2, h-60, "INFORME ACADÉMICO INDIVIDUAL")
    c.setFont("Helvetica", 9)
    c.setFillColor(colors.HexColor("#6b7280"))
    c.drawCentredString(w/2, h-73, f"Año Escolar {config.get('anio', 2026)}")
    c.setFillColor(colors.black)

    # ── Datos del estudiante ─────────────────────────────────────────────
    c.setStrokeColor(colors.HexColor("#1a56db"))
    c.setLineWidth(1.5)
    c.roundRect(25, h-148, w-50, 60, 8, fill=0)
    c.setFont("Helvetica-Bold", 10)
    y = h - 105
    c.drawString(35, y, f"Estudiante: {nombre}")
    c.drawRightString(w-35, y, f"DNI: {dni}")
    y -= 18
    c.drawString(35, y, f"Grado: {grado}")
    c.drawRightString(w-35, y, f"Fecha: {fecha_peru_str()}")
    c.setStrokeColor(colors.HexColor("#1a56db"))
    c.setLineWidth(2)
    c.line(25, h-152, w-25, h-152)
    y = h - 175
    
    if not resultados_hist:
        c.setFont("Helvetica", 12)
        c.drawString(50, y, "Sin evaluaciones registradas.")
        c.save()
        buf.seek(0)
        return buf
    
    # Tabla de evaluaciones
    c.setFont("Helvetica-Bold", 11)
    c.drawString(50, y, "HISTORIAL DE EVALUACIONES")
    y -= 20
    
    # Headers
    c.setFont("Helvetica-Bold", 8)
    c.setFillColor(colors.HexColor("#1a56db"))
    c.rect(50, y-2, w-100, 16, fill=True)
    c.setFillColor(colors.white)
    cols_x = [55, 140, 250, 340, 400, 460, 520]
    headers = ["Fecha", "Evaluación", "Área", "Nota", "Literal", "Estado", ""]
    for i, hd in enumerate(headers[:6]):
        c.drawString(cols_x[i], y+2, hd)
    y -= 18
    
    c.setFillColor(colors.black)
    c.setFont("Helvetica", 7)
    promedios_areas = {}
    total_general = []
    
    for r in resultados_hist:
        for area in r.get('areas', []):
            nota = area.get('nota', 0)
            letra = nota_a_letra(nota)
            col = color_semaforo(letra)
            nombre_area = area.get('nombre', '')
            
            # Acumular para estadísticas
            if nombre_area not in promedios_areas:
                promedios_areas[nombre_area] = []
            promedios_areas[nombre_area].append(nota)
            total_general.append(nota)
            
            c.drawString(cols_x[0], y, str(r.get('fecha', ''))[:10])
            c.drawString(cols_x[1], y, str(r.get('titulo', 'Evaluación'))[:18])
            c.drawString(cols_x[2], y, nombre_area[:15])
            c.drawString(cols_x[3], y, f"{nota}/20")
            c.drawString(cols_x[4], y, letra)
            
            # Semáforo de color
            c.setFillColor(colors.HexColor(col))
            c.circle(cols_x[5]+15, y+3, 5, fill=True)
            c.setFillColor(colors.black)
            
            y -= 14
            if y < 120:
                c.showPage()
                y = h - 60
                c.setFont("Helvetica", 7)
    
    # Resumen estadístico
    y -= 15
    if y < 200:
        c.showPage()
        y = h - 60
    
    c.setFont("Helvetica-Bold", 11)
    c.drawString(50, y, "RESUMEN POR ÁREAS")
    y -= 20
    
    for area_nombre, notas in promedios_areas.items():
        if not notas:
            continue
        prom = round(sum(notas) / len(notas), 1)
        letra = nota_a_letra(prom)
        col = color_semaforo(letra)
        
        c.setFont("Helvetica-Bold", 9)
        c.drawString(55, y, f"{area_nombre}:")
        c.drawString(200, y, f"Promedio: {prom}/20")
        c.drawString(310, y, f"({letra} — {ESCALA_MINEDU[letra]['nombre']})")
        
        # Barra visual
        c.setFillColor(colors.HexColor("#e2e8f0"))
        c.rect(420, y-2, 120, 12, fill=True)
        c.setFillColor(colors.HexColor(col))
        c.rect(420, y-2, (prom/20)*120, 12, fill=True)
        c.setFillColor(colors.black)
        y -= 18
    
    # Promedio General
    if total_general:
        prom_gen = round(sum(total_general) / len(total_general), 1)
        letra_gen = nota_a_letra(prom_gen)
        y -= 10
        c.setFont("Helvetica-Bold", 13)
        c.drawString(55, y, f"PROMEDIO GENERAL: {prom_gen}/20 ({letra_gen})")
        
        # Semáforo grande
        col_gen = color_semaforo(letra_gen)
        c.setFillColor(colors.HexColor(col_gen))
        c.circle(430, y+5, 12, fill=True)
        c.setFillColor(colors.white)
        c.setFont("Helvetica-Bold", 10)
        c.drawCentredString(430, y+1, letra_gen)
        c.setFillColor(colors.black)
    
    # Recomendaciones pedagógicas
    y -= 35
    if y < 180:
        c.showPage()
        y = h - 60
    c.setFont("Helvetica-Bold", 11)
    c.drawString(50, y, "RECOMENDACIONES PEDAGÓGICAS Y PSICOLÓGICAS")
    y -= 18
    c.setFont("Helvetica", 8)
    
    if total_general:
        letra_gen = nota_a_letra(prom_gen)
        desc = ESCALA_MINEDU[letra_gen]['desc']
        c.drawString(55, y, f"• Nivel actual: {desc}")
        y -= 14
        
        if letra_gen == 'AD':
            recs = [
                "Excelente desempeño. Mantener el ritmo y motivar con retos académicos mayores.",
                "Se recomienda participación en concursos académicos y olimpiadas.",
                "Fomentar el liderazgo y tutoría entre pares.",
            ]
        elif letra_gen == 'A':
            recs = [
                "Buen rendimiento. Reforzar áreas con menor puntaje para alcanzar el nivel destacado.",
                "Establecer metas semanales de estudio con apoyo familiar.",
                "Incentivar hábitos de lectura diaria de 30 minutos.",
            ]
        elif letra_gen == 'B':
            recs = [
                "En proceso de logro. Requiere acompañamiento adicional del docente y familia.",
                "Se sugiere sesiones de refuerzo en las áreas con menor calificación.",
                "Establecer un horario de estudio fijo en casa con supervisión.",
                "Diálogo constante entre padres y docentes sobre avances.",
            ]
        else:
            recs = [
                "Necesita apoyo inmediato. Coordinar con el docente un plan de recuperación.",
                "Se recomienda evaluación psicopedagógica para identificar dificultades.",
                "Sesiones de refuerzo diarias con material adaptado a su ritmo.",
                "Reunión urgente con padres para establecer compromisos.",
                "Considerar apoyo emocional si hay factores externos que afectan el aprendizaje.",
            ]
        for rec in recs:
            c.drawString(55, y, f"• {rec}")
            y -= 12
    
    # Escala de calificación
    y -= 20
    c.setFont("Helvetica-Bold", 9)
    c.drawString(50, y, "ESCALA DE CALIFICACIÓN MINEDU:")
    y -= 14
    c.setFont("Helvetica", 7)
    for sigla, info in ESCALA_MINEDU.items():
        c.setFillColor(colors.HexColor(info['color']))
        c.circle(60, y+3, 4, fill=True)
        c.setFillColor(colors.black)
        c.drawString(70, y, f"{sigla} ({info['min']}-{info['max']}): {info['nombre']}")
        y -= 12
    
    # Pie
    c.setFont("Helvetica-Oblique", 7)
    c.drawCentredString(w/2, 30, f"YACHAY PRO — Sistema de Gestión Educativa © {hora_peru().year}")
    
    c.save()
    buf.seek(0)
    return buf


def tab_calificacion_yachay(config):
    st.header("📝 Sistema de Calificación YACHAY")
    usuario_actual = st.session_state.usuario_actual

    tabs_cal = st.tabs([
        "🔑 Crear Claves", "📄 Hoja de Respuestas",
        "✅ Calificar", "🏆 Ranking", "📊 Historial"
    ])

    titulo_eval = "Evaluación"  # Default

    # ===== TAB: CREAR CLAVES (Grid estilo ZipGrade) =====
    with tabs_cal[0]:
        st.subheader("🔑 Crear Claves de Evaluación")
        st.markdown("Marque la alternativa correcta para cada pregunta:")

        ec1, ec2 = st.columns(2)
        with ec1:
            titulo_eval = st.text_input("📝 Nombre de la evaluación:",
                                         "Evaluación Bimestral", key="tit_eval")
        with ec2:
            num_areas = st.number_input("Número de áreas:", 1, 6, 1, key="num_areas_grid")

        areas_grid = []
        total_preguntas = 0
        for a_idx in range(int(num_areas)):
            st.markdown(f"---")
            ac1, ac2 = st.columns([2, 1])
            with ac1:
                area_nom = st.text_input(f"Área {a_idx+1}:", key=f"area_nom_{a_idx}",
                                          value=["Matemática", "Comunicación",
                                                 "Ciencia y Tec.", "Personal Social",
                                                 "Arte y Cultura", "Ed. Física"][a_idx]
                                          if a_idx < 6 else f"Área {a_idx+1}")
            with ac2:
                area_num = st.selectbox(f"Preguntas:",
                                         [5, 10, 15, 20, 25],
                                         index=1, key=f"area_num_{a_idx}")

            # Grid de alternativas
            claves_area = []
            cols_header = st.columns([1, 1, 1, 1, 1])
            with cols_header[0]:
                st.markdown("**#**")
            for opt_idx, opt in enumerate(['A', 'B', 'C', 'D']):
                with cols_header[opt_idx + 1]:
                    st.markdown(f"**{opt}**")

            for p in range(int(area_num)):
                p_global = total_preguntas + p + 1
                resp = st.radio(
                    f"P{p_global}",
                    ['A', 'B', 'C', 'D'],
                    horizontal=True,
                    key=f"grid_{a_idx}_{p}",
                    label_visibility="collapsed" if p > 0 else "visible"
                )
                claves_area.append(resp)

            areas_grid.append({
                'nombre': area_nom,
                'num': int(area_num),
                'claves': claves_area
            })
            total_preguntas += int(area_num)

        st.markdown("---")
        st.info(f"📊 Total: **{total_preguntas} preguntas** en **{len(areas_grid)} áreas**")

        # Resumen visual de claves
        if areas_grid:
            resumen = ""
            for ag in areas_grid:
                resumen += f"**{ag['nombre']}:** {''.join(ag['claves'])}\n\n"
            st.markdown(resumen)

        # Guardar evaluación
        if st.button("💾 GUARDAR EVALUACIÓN", type="primary",
                     use_container_width=True, key="guardar_eval"):
            if titulo_eval:
                eval_data = {
                    'titulo': titulo_eval,
                    'fecha': fecha_peru_str(),
                    'hora': hora_peru_str(),
                    'usuario': usuario_actual,
                    'areas': areas_grid,
                    'total_preguntas': total_preguntas
                }
                # Guardar en session_state
                if 'evaluaciones_guardadas' not in st.session_state:
                    st.session_state.evaluaciones_guardadas = {}
                eval_key = f"{titulo_eval}_{fecha_peru_str()}"
                st.session_state.evaluaciones_guardadas[eval_key] = eval_data

                # Guardar en Google Sheets
                gs = _gs()
                if gs:
                    try:
                        ws = gs._get_hoja('config')
                        if ws:
                            ws.append_row([
                                f"eval_{eval_key}",
                                json.dumps(eval_data, ensure_ascii=False, default=str)
                            ])
                    except Exception:
                        pass

                st.success(f"✅ Evaluación **'{titulo_eval}'** guardada exitosamente")
                st.markdown(f"**Claves:** {total_preguntas} preguntas en {len(areas_grid)} áreas")
                reproducir_beep_exitoso()
            else:
                st.error("⚠️ Ingrese un nombre para la evaluación")

    # ===== TAB: HOJA DE RESPUESTAS =====
    with tabs_cal[1]:
        st.subheader("📄 Hoja de Respuestas")
        c1, c2 = st.columns(2)
        with c1:
            npg = st.selectbox("Preguntas:", [10, 20, 30, 40, 50],
                               index=1, key="npg")
        with c2:
            th = st.text_input("Título:", "EVALUACIÓN BIMESTRAL", key="th")
        
        if st.button("📄 GENERAR HOJA DE RESPUESTAS PDF", type="primary",
                     use_container_width=True, key="gh"):
            hoja_bio = generar_hoja_respuestas(npg, th)
            hoja_bytes = hoja_bio.getvalue()
            
            # Vista previa
            st.image(hoja_bytes, use_container_width=True)
            
            # PDF LANDSCAPE: 2 hojas verticales lado a lado
            try:
                from PIL import Image as PILImage
                img_pil = PILImage.open(io.BytesIO(hoja_bytes))
                img_w, img_h = img_pil.size
                
                pdf_buf = io.BytesIO()
                # A4 Landscape: 841.89 x 595.27 puntos
                pw, ph = 841.89, 595.27
                c_pdf = canvas.Canvas(pdf_buf, pagesize=(pw, ph))
                
                img_path = "/tmp/hoja_temp.png"
                img_pil.save(img_path)
                
                # 2 hojas verticales lado a lado en página horizontal
                half_w = pw / 2
                margin = 5
                # Escalar cada hoja para que quepa en media página
                scale = min((half_w - margin * 2) / img_w, (ph - margin * 2) / img_h) * 0.95
                draw_w = img_w * scale
                draw_h = img_h * scale
                
                # Hoja izquierda
                x_left = (half_w - draw_w) / 2
                y_bot = (ph - draw_h) / 2
                c_pdf.drawImage(img_path, x_left, y_bot,
                                width=draw_w, height=draw_h)
                
                # Línea de corte vertical al centro
                c_pdf.setStrokeColor(colors.gray)
                c_pdf.setLineWidth(0.5)
                c_pdf.setDash(6, 3)
                c_pdf.line(half_w, 10, half_w, ph - 10)
                c_pdf.setFont("Helvetica", 5)
                c_pdf.saveState()
                c_pdf.translate(half_w + 4, ph / 2)
                c_pdf.rotate(90)
                c_pdf.drawCentredString(0, 0, "- - - CORTAR AQUI - - -")
                c_pdf.restoreState()
                c_pdf.setDash()
                
                # Hoja derecha
                x_right = half_w + (half_w - draw_w) / 2
                c_pdf.drawImage(img_path, x_right, y_bot,
                                width=draw_w, height=draw_h)
                
                c_pdf.save()
                pdf_buf.seek(0)
                st.download_button("📥 DESCARGAR PDF (2 hojas lado a lado)",
                                   pdf_buf.getvalue(),
                                   f"Hojas_Respuesta_{npg}p.pdf",
                                   "application/pdf",
                                   use_container_width=True, key="dh_pdf")
                st.success("🎉 PDF listo — página horizontal con 2 hojas verticales")
            except Exception as e:
                st.error(f"Error PDF: {e}")
                st.download_button("⬇️ Descargar PNG", hoja_bytes,
                                   f"Hoja_{npg}p.png", "image/png", key="dh_png")

    # ===== TAB: CALIFICAR =====
    with tabs_cal[2]:
        st.subheader("✅ Calificar Examen")

        # Cargar evaluación guardada o crear nueva
        modo_cal = st.radio("Modo:", [
            "📂 Evaluación Guardada",
            "✏️ Claves Manuales",
            "⚡ Evaluación Rápida (solo nombres)"
        ], key="modo_cal")

        ia = []
        tc_ = []
        tp = 0

        if modo_cal == "📂 Evaluación Guardada":
            # Cargar de Google Sheets
            evals_disp = {}
            gs = _gs()
            if gs:
                try:
                    ws = gs._get_hoja('config')
                    if ws:
                        data = ws.get_all_records()
                        for d in data:
                            clave = str(d.get('clave', ''))
                            if clave.startswith('eval_'):
                                try:
                                    evals_disp[clave[5:]] = json.loads(d.get('valor', '{}'))
                                except Exception:
                                    pass
                except Exception:
                    pass

            # Agregar de session_state
            for k, v in st.session_state.get('evaluaciones_guardadas', {}).items():
                if k not in evals_disp:
                    evals_disp[k] = v

            if evals_disp:
                sel_eval = st.selectbox("Seleccionar evaluación:",
                                         list(evals_disp.keys()), key="sel_eval_cal")
                if sel_eval:
                    ev = evals_disp[sel_eval]
                    st.success(f"📝 **{ev.get('titulo', sel_eval)}** — "
                              f"{ev.get('total_preguntas', 0)} preguntas")
                    for a in ev.get('areas', []):
                        claves_list = a.get('claves', [])
                        ia.append({
                            'nombre': a['nombre'],
                            'num': a['num'],
                            'claves': claves_list
                        })
                        tc_.extend(claves_list)
                        tp += a['num']
            else:
                st.warning("No hay evaluaciones guardadas. Cree una en la pestaña 🔑 Crear Claves.")

        elif modo_cal == "✏️ Claves Manuales":
            if 'areas_examen' not in st.session_state:
                st.session_state.areas_examen = []
            ca, cn_, cb = st.columns([2, 1, 1])
            with ca:
                na = st.text_input("Área:", key="na")
            with cn_:
                nn = st.selectbox("Preguntas:", [5, 10, 15, 20, 25, 30],
                                  index=1, key="nn")
            with cb:
                st.markdown("###")
                if st.button("➕ Agregar", key="aa", type="primary"):
                    if na:
                        st.session_state.areas_examen.append({
                            'nombre': na, 'num': nn, 'claves': ''})
                        st.rerun()

            for i, a in enumerate(st.session_state.areas_examen):
                with st.expander(f"📚 {a['nombre']} ({a['num']}p)", expanded=True):
                    cl = st.text_input("Claves (ej: ABCDABCDAB):",
                                       value=a.get('claves', ''),
                                       key=f"cl{i}", max_chars=a['num'])
                    st.session_state.areas_examen[i]['claves'] = cl.upper()
                    ia.append({'nombre': a['nombre'], 'num': a['num'],
                               'claves': list(cl.upper())})
                    tc_.extend(list(cl.upper()))
                    tp += a['num']
                    if len(st.session_state.areas_examen) > 1:
                        if st.button("🗑️ Quitar", key=f"d{i}", type="primary"):
                            st.session_state.areas_examen.pop(i)
                            st.rerun()

        else:  # Evaluación Rápida
            st.info("⚡ En este modo solo ingresa nombre del estudiante (sin DNI)")

            ca, cn_, cb = st.columns([2, 1, 1])
            with ca:
                na = st.text_input("Área:", key="na_r")
            with cn_:
                nn = st.selectbox("Preguntas:", [5, 10, 15, 20, 25, 30],
                                  index=1, key="nn_r")
            with cb:
                st.markdown("###")
                if st.button("➕ Agregar", key="aa_r", type="primary"):
                    if na:
                        st.session_state.areas_examen.append({
                            'nombre': na, 'num': nn, 'claves': ''})
                        st.rerun()

            for i, a in enumerate(st.session_state.get('areas_examen', [])):
                with st.expander(f"📚 {a['nombre']} ({a['num']}p)", expanded=True):
                    cl = st.text_input("Claves:", value=a.get('claves', ''),
                                       key=f"clr{i}", max_chars=a['num'])
                    st.session_state.areas_examen[i]['claves'] = cl.upper()
                    ia.append({'nombre': a['nombre'], 'num': a['num'],
                               'claves': list(cl.upper())})
                    tc_.extend(list(cl.upper()))
                    tp += a['num']

        if ia:
            st.info(f"📊 {tp} preguntas en {len(ia)} áreas")

        # Seleccionar alumno
        st.markdown("---")
        st.markdown("**👤 Seleccionar Alumno:**")

        de = ""
        nombre_sel = ""

        if modo_cal == "⚡ Evaluación Rápida (solo nombres)":
            nombre_sel = st.text_input("Nombre completo del estudiante:",
                                        key="nombre_rapido",
                                        placeholder="Ej: JUAN PEREZ QUISPE")
            de = ""
        else:
            metodo_sel = st.radio("Método:",
                                   ["📋 Lista de mi grado", "🔍 Buscar por DNI"],
                                   horizontal=True, key="metodo_sel")
            if metodo_sel == "📋 Lista de mi grado":
                # Usar el helper central que filtra según rol
                rol_act = st.session_state.get('rol', '')
                info_act = st.session_state.get('docente_info', {}) or {}
                nivel_act = str(info_act.get('nivel', '')).upper()
                grado_act = str(info_act.get('grado', ''))

                es_sec_act = ('SECUNDARIA' in nivel_act or 'PREUNIVERSITARIO' in nivel_act
                              or 'GRUPO' in grado_act or grado_act in ('ALL_NIVELES', 'ALL_SEC_PREU', 'ALL_SECUNDARIA'))

                if rol_act in ['admin', 'directivo']:
                    grado_doc = st.selectbox("🎓 Grado:", GRADOS_OPCIONES, key="grado_cal_sel")
                elif es_sec_act:
                    grados_d = _grados_del_docente()
                    grado_doc = st.selectbox("🎓 Grado:", grados_d, key="grado_cal_sel_sec")
                elif grado_act and grado_act != 'N/A':
                    grado_doc = grado_act
                    st.info(f"🎓 **{grado_doc}**")
                else:
                    grado_doc = None
                    st.warning("Sin grado asignado.")

                if grado_doc:
                    dg = BaseDatos.obtener_estudiantes_grado(grado_doc)
                    if not dg.empty and 'Nombre' in dg.columns:
                        opciones = [f"{row.get('Nombre', '')} — DNI: {row.get('DNI', '')}"
                                    for _, row in dg.iterrows()]
                        sel = st.selectbox("Estudiante:", opciones, key="sel_est")
                        if sel:
                            de = sel.split("DNI: ")[-1].strip()
                            nombre_sel = sel.split(" — ")[0].strip()
                    else:
                        st.warning("No hay estudiantes en este grado.")
                else:
                    de = st.text_input("DNI:", key="de_manual")
            else:
                de = st.text_input("DNI del alumno:", key="de")
                if de:
                    ae = BaseDatos.buscar_por_dni(de)
                    if ae:
                        nombre_sel = str(ae.get('Nombre', ''))
                        st.success(f"👤 {nombre_sel}")

        # Respuestas
        st.markdown("**📝 Respuestas:**")
        met = st.radio("Método:", ["✏️ Manual", "📸 Cámara/Foto"],
                       horizontal=True, key="met")
        ra = []
        if met == "✏️ Manual":
            for i, a in enumerate(ia):
                r = st.text_input(f"{a['nombre']} ({a['num']}):",
                                  key=f"r{i}", max_chars=a['num'],
                                  placeholder="Ej: ABCDABCDAB")
                ra.extend(list(r.upper()))
        else:
            st.info("📸 Use la hoja generada por el sistema. Buena luz, que se vean los 4 cuadrados negros.")
            src_img = st.radio("Fuente:",
                                ["📷 Cámara", "📁 Subir foto"],
                                horizontal=True, key="src_img")
            image_data = None
            if src_img == "📷 Cámara":
                ac = st.checkbox("📷 Activar cámara", key="chce")
                if ac:
                    fe = st.camera_input("Apunta a la hoja:", key="ce")
                    if fe:
                        image_data = fe.getvalue()
            else:
                fu = st.file_uploader("📁 Subir foto:", type=['jpg', 'jpeg', 'png'], key="fu_hoja")
                if fu:
                    image_data = fu.getvalue()

            if image_data:
                with st.spinner("🔍 Escaneando..."):
                    det = procesar_examen(image_data, tp)
                if det:
                    detectadas = sum(1 for x in det if x != '?')
                    if detectadas == len(det):
                        st.success(f"✅ {detectadas}/{len(det)} respuestas detectadas")
                    else:
                        st.warning(f"⚠️ {detectadas}/{len(det)} detectadas. Corrija las '?' abajo.")
                    det_str = ''.join(det)
                    corregido = st.text_input("Respuestas detectadas:", value=det_str,
                                              key="det_corr", max_chars=tp)
                    ra = list(corregido.upper())
                else:
                    st.error("❌ No se pudo leer. Intente con mejor luz o use modo Manual.")

        # CALIFICAR
        st.markdown("---")
        if st.button("📊 CALIFICAR", type="primary",
                     use_container_width=True, key="cal"):
            if tc_ and ra:
                ad = BaseDatos.buscar_por_dni(de) if de else None
                nm = nombre_sel if nombre_sel else (
                    str(ad.get('Nombre', '')) if ad else "Sin nombre")
                grado_est = str(ad.get('Grado', '')) if ad else ""
                res = {
                    'fecha': hora_peru().strftime('%d/%m/%Y %H:%M'),
                    'titulo': titulo_eval if modo_cal == "📂 Evaluación Guardada" else "Evaluación",
                    'dni': de, 'nombre': nm, 'grado': grado_est,
                    'areas': [], 'promedio_general': 0
                }
                idx = 0
                sn = 0
                mw = (f"📝 *RESULTADOS*\n🏫 YACHAY\n👤 {nm}\n"
                      f"📅 {hora_peru().strftime('%d/%m/%Y')}\n\n")
                for a in ia:
                    n = a['num']
                    ck = a['claves'][:n]
                    rk = ra[idx:idx + n]
                    ok = sum(1 for j in range(min(len(ck), len(rk)))
                             if ck[j] == rk[j])
                    nota = round((ok / n) * 20, 1) if n else 0
                    lt = nota_a_letra(nota)
                    detalle = []
                    for j in range(n):
                        cj = ck[j] if j < len(ck) else '?'
                        rj = rk[j] if j < len(rk) else '?'
                        detalle.append({
                            'p': idx + j + 1, 'c': cj, 'r': rj,
                            'ok': (j < len(ck) and j < len(rk) and ck[j] == rk[j])
                        })
                    res['areas'].append({
                        'nombre': a['nombre'], 'correctas': ok,
                        'total': n, 'nota': nota, 'letra': lt,
                        'detalle': detalle
                    })
                    sn += nota
                    mw += f"📚 *{a['nombre']}:* {nota}/20 ({lt})\n"
                    idx += n

                pm = round(sn / len(ia), 1) if ia else 0
                lp = nota_a_letra(pm)
                res['promedio_general'] = pm
                mw += f"\n📊 *PROMEDIO: {pm}/20 ({lp})*"
                BaseDatos.guardar_resultados_examen(res, usuario_actual)

                # Confirmación visual
                st.markdown("### 📊 Resultados")
                cols = st.columns(len(ia) + 1)
                for i, ar in enumerate(res['areas']):
                    with cols[i]:
                        st.metric(f"📚 {ar['nombre']}", f"{ar['nota']}/20",
                                  f"{ar['letra']}")
                with cols[-1]:
                    st.metric("📊 PROMEDIO", f"{pm}/20", lp)

                # Detalle por área
                for ar in res['areas']:
                    with st.expander(f"📋 {ar['nombre']}"):
                        st.dataframe(pd.DataFrame([
                            {'#': d['p'], 'Clave': d['c'], 'Resp': d['r'],
                             '': '✅' if d['ok'] else '❌'}
                            for d in ar['detalle']
                        ]), use_container_width=True, hide_index=True)

                # WhatsApp
                if ad:
                    cel = str(ad.get('Celular_Apoderado', '')).strip()
                    if cel and cel not in ('', 'None', 'nan'):
                        link = generar_link_whatsapp(cel, mw)
                        st.markdown(
                            f'<a href="{link}" target="_blank" class="wa-btn">'
                            f'📱 Enviar resultado → {cel}</a>',
                            unsafe_allow_html=True)

                # Reporte PDF individual
                if st.button("📥 Descargar Reporte PDF del Estudiante", key="dl_rep_est", type="primary"):
                    pdf = generar_reporte_estudiante_pdf(
                        nm, de, grado_est, [res], config)
                    st.download_button("⬇️ PDF", pdf,
                                       f"Reporte_{nm.replace(' ', '_')}.pdf",
                                       "application/pdf", key="dl_rep_est2")

                st.success("🎉 Resultado guardado correctamente en la base de datos")
                reproducir_beep_exitoso()
                st.balloons()
            else:
                st.error("⚠️ Configure claves y respuestas")

    # ===== TAB: RANKING =====
    with tabs_cal[3]:
        st.subheader("🏆 Ranking de Evaluación")

        if st.session_state.rol in ["admin", "directivo"]:
            grado_rank = st.selectbox("Ver grado:", ["TODOS"] + GRADOS_OPCIONES,
                                       key="grado_rank")
            rs = BaseDatos.cargar_todos_resultados()
            if grado_rank != "TODOS":
                # Pre-cargar matrícula para filtrar eficientemente
                df_mat = BaseDatos.cargar_matricula()
                grados_por_dni = {}
                if not df_mat.empty and 'DNI' in df_mat.columns and 'Grado' in df_mat.columns:
                    for _, row_m in df_mat.iterrows():
                        grados_por_dni[str(row_m.get('DNI', '')).strip()] = str(row_m.get('Grado', ''))
                rs = [r for r in rs if str(r.get('grado', '')) == grado_rank or
                      grados_por_dni.get(str(r.get('dni', '')).strip(), '') == grado_rank]
        else:
            rs = BaseDatos.cargar_resultados_examen(usuario_actual)

        if rs:
            df = pd.DataFrame([{
                'Fecha': r.get('fecha', ''),
                'Nombre': r.get('nombre', ''),
                'DNI': str(r.get('dni', '')),
                'Promedio': r.get('promedio_general', 0),
                'Literal': nota_a_letra(r.get('promedio_general', 0)),
                'Áreas': ', '.join([f"{a['nombre']}:{a['nota']}"
                                    for a in r.get('areas', [])])
            } for r in rs])
            df = df.sort_values('Promedio', ascending=False).reset_index(drop=True)
            df.insert(0, '#', range(1, len(df) + 1))

            st.dataframe(df, use_container_width=True, hide_index=True)

            # Podio
            if len(df) >= 1:
                cols = st.columns(min(3, len(df)))
                medallas = ["🥇", "🥈", "🥉"]
                estilos = ["ranking-gold", "ranking-silver", "ranking-bronze"]
                for i in range(min(3, len(df))):
                    with cols[i]:
                        r = df.iloc[i]
                        st.markdown(
                            f'<div class="{estilos[i]}">'
                            f'{medallas[i]} {r["Nombre"]}<br>'
                            f'{r["Promedio"]}/20 ({r["Literal"]})</div>',
                            unsafe_allow_html=True)

            st.markdown("---")
            bc1, bc2 = st.columns(2)
            with bc1:
                pdf = generar_ranking_pdf(rs, config['anio'])
                st.download_button("📥 RANKING PDF", pdf,
                                   f"Ranking_{config['anio']}.pdf",
                                   "application/pdf", key="drpdf",
                                   use_container_width=True, type="primary")
            with bc2:
                if st.button("📥 REPORTES INDIVIDUALES PDF", type="primary",
                             use_container_width=True, key="reps_ind"):
                    # Generar un PDF multi-página con todos los estudiantes
                    buf_all = io.BytesIO()
                    c_all = canvas.Canvas(buf_all, pagesize=A4)
                    w, h_page = A4
                    for r_item in rs:
                        nm = r_item.get('nombre', '')
                        dn = str(r_item.get('dni', ''))
                        gr = str(r_item.get('grado', ''))
                        
                        c_all.setFont("Helvetica-Bold", 14)
                        c_all.drawCentredString(w/2, h_page-40, f"REPORTE: {nm}")
                        c_all.setFont("Helvetica", 10)
                        y = h_page - 70
                        c_all.drawString(50, y, f"DNI: {dn} | Grado: {gr} | Fecha: {r_item.get('fecha', '')}")
                        y -= 25
                        
                        for area in r_item.get('areas', []):
                            nota = area.get('nota', 0)
                            letra = nota_a_letra(nota)
                            c_all.drawString(60, y, f"• {area['nombre']}: {nota}/20 ({letra})")
                            y -= 16
                        
                        pm = r_item.get('promedio_general', 0)
                        lp = nota_a_letra(pm)
                        y -= 10
                        c_all.setFont("Helvetica-Bold", 12)
                        c_all.drawString(60, y, f"PROMEDIO: {pm}/20 ({lp})")
                        c_all.showPage()
                    
                    c_all.save()
                    buf_all.seek(0)
                    st.download_button("⬇️ Reportes PDF", buf_all,
                                       "Reportes_Individuales.pdf",
                                       "application/pdf", key="dl_reps_all")

            # WhatsApp individual
            st.markdown("---")
            st.markdown("### 📱 Enviar por WhatsApp")
            for _, row in df.iterrows():
                al = BaseDatos.buscar_por_dni(row['DNI']) if row['DNI'] else None
                if al:
                    cel = str(al.get('Celular_Apoderado', '')).strip()
                    if cel and cel not in ('', 'None', 'nan'):
                        ro = next((r for r in rs if str(r.get('dni')) == str(row['DNI'])), None)
                        if ro:
                            msg = f"📝 *RANKING YACHAY*\n👤 {row['Nombre']}\n🏆 #{row['#']}°/{len(df)}\n"
                            for a in ro.get('areas', []):
                                msg += f"📚 {a['nombre']}: {a['nota']}/20\n"
                            msg += f"\n📊 *PROMEDIO: {row['Promedio']}/20 ({row['Literal']})*"
                            link = generar_link_whatsapp(cel, msg)
                            st.markdown(
                                f'<a href="{link}" target="_blank" class="wa-btn">'
                                f'📱 #{row["#"]} {row["Nombre"]} — {row["Promedio"]}/20</a>',
                                unsafe_allow_html=True)

            st.markdown("---")
            if st.button("🔄 NUEVA EVALUACIÓN", type="secondary",
                         use_container_width=True, key="nueva_eval"):
                BaseDatos.limpiar_resultados_examen(usuario_actual)
                st.session_state.areas_examen = []
                st.success("🎉 Resultados limpiados. Nueva evaluación lista.")
                st.rerun()
        else:
            st.info("📝 Califica exámenes para ver tu ranking.")

    # ===== TAB: HISTORIAL =====
    with tabs_cal[4]:
        st.subheader("📊 Historial de Evaluaciones")
        
        # NUEVO: Mostrar evaluaciones guardadas
        st.markdown("### 💾 Evaluaciones Guardadas")
        try:
            historial_file = 'historial_evaluaciones.json'
            if Path(historial_file).exists():
                with open(historial_file, 'r', encoding='utf-8') as f:
                    hist_data = json.load(f)
                
                if hist_data:
                    for clave, eval_data in sorted(hist_data.items(), reverse=True):
                        with st.expander(f"📝 {eval_data['grado']} - {eval_data['periodo']} ({eval_data['fecha']})"):
                            st.write(f"**Hora:** {eval_data.get('hora', 'N/A')}")
                            st.write(f"**Estudiantes evaluados:** {len(eval_data.get('ranking', []))}")
                            st.write(f"**Áreas:** {', '.join([a['nombre'] for a in eval_data.get('areas', [])] if isinstance(eval_data.get('areas', []), list) else eval_data.get('areas', []))}")
                            
                            col_ver, col_del = st.columns([3, 1])
                            with col_ver:
                                if st.button("📊 Ver Ranking", key=f"ver_rank_{clave}", type="primary"):
                                    df_hist = pd.DataFrame(eval_data.get('ranking', []))
                                    st.dataframe(df_hist, use_container_width=True)
                            with col_del:
                                if st.button("🗑️ Eliminar", key=f"del_eval_{clave}", type="primary"):
                                    del hist_data[clave]
                                    with open(historial_file, 'w', encoding='utf-8') as f:
                                        json.dump(hist_data, f, ensure_ascii=False, indent=2)
                                    st.success("✅ Evaluación eliminada")
                                    st.rerun()
                else:
                    st.info("No hay evaluaciones guardadas en historial")
            else:
                st.info("No hay historial disponible aún")
        except Exception as e:
            st.error(f"Error al cargar historial: {str(e)}")
        
        st.markdown("---")
        st.markdown("### 👤 Historial por Estudiante")

        if st.session_state.rol in ["admin", "directivo"]:
            grado_hist = st.selectbox("Grado:", GRADOS_OPCIONES, key="grado_hist")
            dg = BaseDatos.obtener_estudiantes_grado(grado_hist)
            if not dg.empty:
                est_sel = st.selectbox("Estudiante:",
                                        [f"{r['Nombre']} — {r['DNI']}"
                                         for _, r in dg.iterrows()],
                                        key="est_hist")
                if est_sel:
                    dni_hist = est_sel.split(" — ")[-1].strip()
                    nombre_hist = est_sel.split(" — ")[0].strip()
            else:
                st.warning("No hay estudiantes en este grado.")
                dni_hist = ""
                nombre_hist = ""
        else:
            dni_hist = st.text_input("DNI del estudiante:", key="dni_hist")
            al_h = BaseDatos.buscar_por_dni(dni_hist) if dni_hist else None
            nombre_hist = str(al_h.get('Nombre', '')) if al_h else ""

        if dni_hist or nombre_hist:
            # Buscar todos los resultados
            all_res = BaseDatos.cargar_todos_resultados()
            hist = [r for r in all_res if str(r.get('dni', '')) == str(dni_hist)
                    or (not dni_hist and r.get('nombre', '') == nombre_hist)]

            if hist:
                st.success(f"📋 {len(hist)} evaluaciones encontradas para **{nombre_hist}**")

                for h in hist:
                    with st.expander(f"📝 {h.get('titulo', 'Evaluación')} — {h.get('fecha', '')}"):
                        for a in h.get('areas', []):
                            st.write(f"**{a['nombre']}:** {a['nota']}/20 ({nota_a_letra(a['nota'])})")
                        st.write(f"**Promedio:** {h.get('promedio_general', 0)}/20")

                # Descargar reporte completo
                if st.button("📥 Descargar Reporte Completo PDF", key="dl_hist_pdf", type="primary"):
                    al_h = BaseDatos.buscar_por_dni(dni_hist)
                    grado_h = str(al_h.get('Grado', '')) if al_h else ""
                    pdf = generar_reporte_estudiante_pdf(
                        nombre_hist, dni_hist, grado_h, hist, config)
                    st.download_button("⬇️ PDF", pdf,
                                       f"Historial_{nombre_hist.replace(' ', '_')}.pdf",
                                       "application/pdf", key="dl_hist_pdf2")
            else:
                st.info("No hay evaluaciones registradas para este estudiante.")


# ================================================================
# TAB: BASE DE DATOS
# ================================================================

def generar_pdf_datos_pendientes(df_pendientes, config):
    """Genera un PDF imprimible con los datos faltantes de los estudiantes para llenar a mano."""
    buf = io.BytesIO()
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors as rl_colors
    from reportlab.platypus import Table, TableStyle

    W, H = A4
    c = rl_canvas.Canvas(buf, pagesize=A4)

    colegio = config.get('colegio', 'INSTITUCIÓN EDUCATIVA')
    anio = config.get('anio', '2025')

    def nueva_pagina(titulo_extra=""):
        c.setFillColorRGB(0.09, 0.39, 0.67)
        c.rect(0, H - 50, W, 50, fill=1, stroke=0)
        c.setFillColorRGB(1, 1, 1)
        c.setFont("Helvetica-Bold", 13)
        c.drawCentredString(W / 2, H - 30, f"{colegio} — FICHA DE DATOS PENDIENTES {anio}")
        if titulo_extra:
            c.setFont("Helvetica", 9)
            c.drawCentredString(W / 2, H - 44, titulo_extra)

    nueva_pagina()

    # Agrupar por grado
    grados = df_pendientes['Grado'].dropna().unique() if 'Grado' in df_pendientes.columns else ['Sin Grado']
    y = H - 70

    for grado in sorted(grados):
        if 'Grado' in df_pendientes.columns:
            sub = df_pendientes[df_pendientes['Grado'] == grado].copy()
        else:
            sub = df_pendientes.copy()

        if sub.empty:
            continue

        # Cabecera de grado
        if y < 140:
            c.showPage()
            nueva_pagina()
            y = H - 70

        c.setFillColorRGB(0.18, 0.55, 0.34)
        c.rect(30, y - 18, W - 60, 18, fill=1, stroke=0)
        c.setFillColorRGB(1, 1, 1)
        c.setFont("Helvetica-Bold", 10)
        c.drawString(35, y - 13, f"  GRADO: {grado}  —  {len(sub)} estudiante(s) con datos pendientes")
        y -= 22

        # Tabla por grado
        col_widths = [25, 165, 65, 65, 80, 80]
        headers = ["N°", "APELLIDOS Y NOMBRES", "DNI ALUMNO", "DNI APODERADO", "APODERADO", "CELULAR"]

        table_data = [headers]
        for i, (_, row) in enumerate(sub.iterrows(), 1):
            dni_val = str(row.get('DNI', '')).strip()
            es_prov = dni_val.startswith('PROV') or str(row.get('_provisional', '')).upper() == 'SI'
            dni_show = "___________" if es_prov else (dni_val if dni_val else "___________")
            apod = str(row.get('Apoderado', '')).strip()
            dni_apod = str(row.get('DNI_Apoderado', '')).strip()
            cel = str(row.get('Celular_Apoderado', '')).strip()
            table_data.append([
                str(i),
                str(row.get('Nombre', '')),
                dni_show if es_prov else (dni_val or "___________"),
                dni_apod if (dni_apod and dni_apod not in ('nan', '')) else "___________",
                apod if (apod and apod not in ('nan', '')) else "___________________________",
                cel if (cel and cel not in ('nan', '')) else "_________"
            ])

        rows_per_page = int((y - 60) / 16)
        chunks = [table_data[0:1] + table_data[1:][i:i + rows_per_page]
                  for i in range(0, len(table_data) - 1, max(rows_per_page, 1))]

        for chunk_idx, chunk in enumerate(chunks):
            if chunk_idx > 0:
                c.showPage()
                nueva_pagina(f"(continuación: {grado})")
                y = H - 70

            tbl = Table(chunk, colWidths=col_widths)
            style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), rl_colors.HexColor('#1565C0')),
                ('TEXTCOLOR', (0, 0), (-1, 0), rl_colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 7),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 7),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [rl_colors.white, rl_colors.HexColor('#EEF4FB')]),
                ('GRID', (0, 0), (-1, -1), 0.4, rl_colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('ALIGN', (0, 0), (0, -1), 'CENTER'),
                ('LEFTPADDING', (0, 0), (-1, -1), 3),
                ('RIGHTPADDING', (0, 0), (-1, -1), 3),
                ('TOPPADDING', (0, 0), (-1, -1), 3),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
            ])
            # Marcar en naranja las celdas vacías (con guiones)
            for row_i, row_data in enumerate(chunk[1:], 1):
                for col_i, cell in enumerate(row_data[2:], 2):
                    if '_' in str(cell):
                        style.add('BACKGROUND', (col_i, row_i), (col_i, row_i),
                                  rl_colors.HexColor('#FFF3CD'))
            tbl.setStyle(style)
            tbl_w, tbl_h = tbl.wrapOn(c, W - 60, y - 60)
            if y - tbl_h < 50:
                c.showPage()
                nueva_pagina(f"(continuación: {grado})")
                y = H - 70
            tbl.drawOn(c, 30, y - tbl_h)
            y = y - tbl_h - 12

    # Pie de página final
    c.setFont("Helvetica-Oblique", 8)
    c.setFillColorRGB(0.5, 0.5, 0.5)
    c.drawCentredString(W / 2, 30,
                        f"Sistema YACHAY PRO — Impreso {fecha_peru_str()} — Las celdas en amarillo están PENDIENTES de completar")
    c.save()
    buf.seek(0)
    return buf


def tab_base_datos():
    st.header("📊 Base de Datos")
    df = BaseDatos.cargar_matricula()
    df_doc = BaseDatos.cargar_docentes()

    # Contar provisionales
    n_prov = 0
    if not df.empty:
        if '_provisional' in df.columns:
            n_prov = (df['_provisional'].astype(str).str.upper() == 'SI').sum()
        elif 'DNI' in df.columns:
            n_prov = df['DNI'].astype(str).str.startswith('PROV').sum()

    c1, c2, c3, c4, c5 = st.columns(5)
    with c1:
        st.metric("📚 Alumnos", len(df) if not df.empty else 0)
    with c2:
        st.metric("👨‍🏫 Docentes", len(df_doc) if not df_doc.empty else 0)
    with c3:
        st.metric("🎓 Grados",
                   df['Grado'].nunique() if not df.empty and 'Grado' in df.columns
                   else 0)
    with c4:
        st.metric("📱 Con Celular",
                   df['Celular_Apoderado'].notna().sum()
                   if not df.empty and 'Celular_Apoderado' in df.columns else 0)
    with c5:
        st.metric("⏳ Provisionales", n_prov,
                  delta="pendientes" if n_prov > 0 else None,
                  delta_color="inverse" if n_prov > 0 else "off")

    tab_al, tab_completar, tab_dc = st.tabs(["📚 Alumnos", "✏️ Completar Datos", "👨‍🏫 Docentes"])
    with tab_al:
        if not df.empty:
            c1, c2 = st.columns(2)
            with c1:
                opts = ['Todos'] + (
                    sorted(df['Grado'].dropna().unique().tolist())
                    if 'Grado' in df.columns else [])
                fg = st.selectbox("Filtrar:", opts, key="fbd")
            with c2:
                bq = st.text_input("🔍", key="bbd")
            d = df.copy()
            if fg != 'Todos' and 'Grado' in d.columns:
                d = d[d['Grado'] == fg]
            if bq:
                d = d[d.apply(lambda r: bq.lower() in str(r).lower(), axis=1)]
            if 'Nombre' in d.columns:
                d = d.sort_values('Nombre')
            st.dataframe(d, use_container_width=True, hide_index=True, height=500)
            c1, c2, c3 = st.columns(3)
            with c1:
                st.download_button("⬇️ CSV",
                                   d.to_csv(index=False).encode('utf-8'),
                                   "alumnos.csv", "text/csv", key="dcsv")
            with c2:
                buf = io.BytesIO()
                d.to_excel(buf, index=False, engine='openpyxl')
                buf.seek(0)
                st.download_button("⬇️ Excel", buf, "alumnos.xlsx", key="dxlsx")
            with c3:
                st.markdown("")
            # Editar nombre de alumno
            with st.expander("✏️ Editar Nombre de Alumno", expanded=False):
                st.caption("Busque al alumno por DNI o ID provisional para corregir su nombre")
                enomb_busq = st.text_input("DNI o ID provisional:", key="enomb_busq", max_chars=10, placeholder="12345678 o PROV0001")
                if enomb_busq and len(enomb_busq.strip()) >= 4:
                    df_enomb = BaseDatos.cargar_matricula()
                    df_enomb["DNI"] = df_enomb["DNI"].astype(str).str.strip()
                    rows_found = df_enomb[df_enomb["DNI"] == enomb_busq.strip()]
                    if rows_found.empty:
                        # buscar por nombre parcial
                        rows_found = df_enomb[df_enomb.apply(lambda r: enomb_busq.strip().upper() in str(r.get("Nombre","")).upper(), axis=1)]
                    if not rows_found.empty:
                        alumno_enomb = rows_found.iloc[0]
                        st.info(f"**{alumno_enomb.get('Nombre','')}** — {alumno_enomb.get('Grado','')} {alumno_enomb.get('Seccion','')}")
                        nuevo_nomb = st.text_input("Nuevo nombre completo:", value=str(alumno_enomb.get('Nombre','')), key="enomb_nuevo")
                        if st.button("💾 GUARDAR NOMBRE", type="primary", key="btn_guardar_nomb"):
                            if nuevo_nomb.strip():
                                df_enomb.loc[rows_found.index, "Nombre"] = nuevo_nomb.strip().upper()
                                BaseDatos.guardar_matricula(df_enomb)
                                st.success(f"✅ Nombre actualizado a: **{nuevo_nomb.strip().upper()}**")
                                time.sleep(1)
                                st.rerun()
                            else:
                                st.error("El nombre no puede estar vacío")
                    else:
                        st.warning("No se encontró ningún alumno con ese DNI o nombre")
            # Cambiar Grado/Seccion de alumno
            with st.expander("✏️ Cambiar Grado / Sección de Alumno", expanded=False):
                st.caption("Busque al alumno por DNI para cambiar su grado o sección (ej: Ciclo Intensivo → Ciclo Ordinario)")
                edit_dni = st.text_input("DNI del alumno:", key="edit_dni_gs", max_chars=8, placeholder="12345678")
                if edit_dni and len(edit_dni.strip()) == 8:
                    alumno_edit = BaseDatos.buscar_por_dni(edit_dni.strip())
                    if alumno_edit:
                        st.info(f"**{alumno_edit.get('Nombre', '')}** — Actual: **{alumno_edit.get('Grado', '')}** | Sección: **{alumno_edit.get('Seccion', '')}**")
                        ce1, ce2 = st.columns(2)
                        with ce1:
                            nuevo_grado = st.selectbox("Nuevo Grado:", GRADOS_OPCIONES, key="edit_new_grado",
                                index=GRADOS_OPCIONES.index(alumno_edit.get('Grado','')) if alumno_edit.get('Grado','') in GRADOS_OPCIONES else 0)
                        with ce2:
                            nuevo_sec = st.selectbox("Nueva Sección:", SECCIONES, key="edit_new_sec",
                                index=SECCIONES.index(alumno_edit.get('Seccion','')) if alumno_edit.get('Seccion','') in SECCIONES else 0)
                        # Detectar nivel del nuevo grado
                        nuevo_nivel = ""
                        for nv, grados in NIVELES_GRADOS.items():
                            if nuevo_grado in grados:
                                nuevo_nivel = nv
                                break
                        cambio_grado = nuevo_grado != alumno_edit.get('Grado', '')
                        cambio_sec = nuevo_sec != alumno_edit.get('Seccion', '')
                        if cambio_grado or cambio_sec:
                            cambios = []
                            if cambio_grado:
                                cambios.append(f"Grado: {alumno_edit.get('Grado','')} → **{nuevo_grado}**")
                            if cambio_sec:
                                cambios.append(f"Sección: {alumno_edit.get('Seccion','')} → **{nuevo_sec}**")
                            st.warning(f"Cambios: {' | '.join(cambios)}")
                            if st.button("✅ APLICAR CAMBIO", type="primary", key="btn_apply_edit"):
                                df_edit = BaseDatos.cargar_matricula()
                                df_edit['DNI'] = df_edit['DNI'].astype(str).str.strip()
                                mask = df_edit['DNI'] == edit_dni.strip()
                                if mask.any():
                                    df_edit.loc[mask, 'Grado'] = nuevo_grado
                                    df_edit.loc[mask, 'Seccion'] = nuevo_sec
                                    if nuevo_nivel:
                                        df_edit.loc[mask, 'Nivel'] = nuevo_nivel
                                    BaseDatos.guardar_matricula(df_edit)
                                    st.success(f"✅ {alumno_edit.get('Nombre','')} actualizado a {nuevo_grado} - {nuevo_sec}")
                                    time.sleep(1)
                                    st.rerun()
                                else:
                                    st.error("No se encontró el alumno en la base de datos")
                        else:
                            st.success("Sin cambios — el grado y sección son los mismos")
                    else:
                        st.error("⚠️ No se encontró alumno con ese DNI")

            # Eliminar alumno
            with st.expander("🗑️ Eliminar Alumno", expanded=False):
                del_dni_a = st.text_input("DNI del alumno a eliminar:", key="del_dni_alum",
                                          max_chars=8, placeholder="12345678")
                if st.button("❌ ELIMINAR ALUMNO", type="primary", key="btn_del_alum"):
                    if del_dni_a and len(del_dni_a.strip()) == 8:
                        alumno = BaseDatos.buscar_por_dni(del_dni_a.strip())
                        if alumno:
                            BaseDatos.eliminar_estudiante(del_dni_a.strip())
                            st.success(f"✅ Alumno con DNI {del_dni_a} eliminado")
                            time.sleep(1)
                            st.rerun()
                        else:
                            st.error("⚠️ No se encontró alumno con ese DNI")
                    else:
                        st.error("⚠️ Ingrese un DNI válido de 8 dígitos")
        else:
            st.info("📝 Sin alumnos.")

    # ================================================================
    # PESTAÑA: COMPLETAR DATOS (nueva)
    # ================================================================
    with tab_completar:
        st.subheader("✏️ Completar Datos de Estudiantes")

        df_comp = BaseDatos.cargar_matricula()
        if df_comp.empty:
            st.warning("No hay alumnos registrados aún.")
        else:
            # --- Filtros superiores ---
            config_c = st.session_state.get("config", {"colegio": "IE YACHAY", "anio": "2025"})
            fc1, fc2, fc3, fc4 = st.columns([2, 2, 2, 2])
            with fc1:
                filtro_grado = st.selectbox("Grado:", ["Todos"] + sorted(df_comp["Grado"].dropna().unique().tolist()) if "Grado" in df_comp.columns else ["Todos"], key="fc_grado_comp")
            with fc2:
                filtro_estado = st.selectbox("Mostrar:", ["Solo incompletos", "Todos los alumnos"], key="fc_estado_comp")
            with fc3:
                filtro_busq = st.text_input("🔍 Buscar nombre:", key="fc_busq_comp", placeholder="Apellido o nombre...")
            with fc4:
                st.markdown("<br>", unsafe_allow_html=True)

            def _es_incompleto(row):
                dni = str(row.get("DNI", "")).strip()
                cel = str(row.get("Celular_Apoderado", "")).strip()
                prov = str(row.get("_provisional", "")).upper() == "SI"
                return prov or dni == "" or dni.startswith("PROV") or dni == "nan" or cel in ("", "nan")

            df_vista = df_comp.copy()
            if filtro_grado != "Todos" and "Grado" in df_vista.columns:
                df_vista = df_vista[df_vista["Grado"] == filtro_grado]
            if filtro_estado == "Solo incompletos":
                df_vista = df_vista[df_vista.apply(_es_incompleto, axis=1)]
            if filtro_busq:
                df_vista = df_vista[df_vista.apply(lambda r: filtro_busq.strip().upper() in str(r.get("Nombre","")).upper(), axis=1)]
            if "Nombre" in df_vista.columns:
                df_vista = df_vista.sort_values("Nombre").reset_index(drop=True)

            n_incomp = df_comp[df_comp.apply(_es_incompleto, axis=1)].shape[0]
            m1, m2, m3 = st.columns([2, 2, 2])
            with m1:
                st.metric("⚠️ Con datos incompletos", n_incomp)
            with m2:
                st.metric("✅ Completos", len(df_comp) - n_incomp)
            with m3:
                if n_incomp > 0:
                    df_pdf = df_comp[df_comp.apply(_es_incompleto, axis=1)].copy()
                    pdf_pend = generar_pdf_datos_pendientes(df_pdf, config_c)
                    st.download_button("🖨️ PDF para llenar a mano", pdf_pend,
                        f"datos_pendientes_{config_c.get('anio','2025')}.pdf",
                        mime="application/pdf", key="btn_pdf_pendientes")

            if df_vista.empty:
                st.success("✅ No hay estudiantes con datos pendientes en este filtro.")
            else:
                st.caption(f"📋 Mostrando {len(df_vista)} estudiante(s) — Haga clic en ✏️ para editar")
                st.markdown("---")

                # Mostrar cada estudiante como fila editable
                for i, (_, row) in enumerate(df_vista.iterrows()):
                    dni_val = str(row.get("DNI", "")).strip()
                    nombre = str(row.get("Nombre", "")).strip()
                    grado = str(row.get("Grado", "")).strip()
                    seccion = str(row.get("Seccion", "")).strip()
                    cel = str(row.get("Celular_Apoderado", "")).replace("nan", "").strip()
                    apod = str(row.get("Apoderado", "")).replace("nan", "").strip()
                    dni_apod = str(row.get("DNI_Apoderado", "")).replace("nan", "").strip()
                    es_prov = dni_val.startswith("PROV") or str(row.get("_provisional","")).upper() == "SI"

                    # Badge de estado
                    badge = "🟡 PROVISIONAL" if es_prov else ("🟠 Sin celular" if not cel else "✅ Completo")
                    icono = "👦" if str(row.get("Sexo","")).strip() == "Masculino" else "👧"

                    with st.expander(f"{icono} {nombre}   —   {grado} {seccion}   {badge}", expanded=False):
                        # Mostrar datos actuales
                        ca, cb, cc = st.columns(3)
                        with ca:
                            st.caption("DNI actual")
                            st.code(dni_val if not es_prov else f"⚠️ {dni_val} (provisional)")
                        with cb:
                            st.caption("Celular Apoderado")
                            st.code(cel if cel else "— vacío —")
                        with cc:
                            st.caption("Apoderado")
                            st.code(apod if apod else "— vacío —")

                        # Formulario de edición
                        with st.form(f"form_edit_{i}_{dni_val}"):
                            # Fila 0: Nombre y Sexo
                            en0, en1 = st.columns([3, 1])
                            with en0:
                                inp_nombre = st.text_input("✏️ Apellidos y Nombres:",
                                    value=nombre, key=f"inp_nombre_{i}",
                                    placeholder="APELLIDOS, Nombres")
                            with en1:
                                sexo_actual = str(row.get("Sexo", "Masculino")).strip()
                                inp_sexo = st.selectbox("Sexo:",
                                    ["Masculino", "Femenino"],
                                    index=0 if sexo_actual == "Masculino" else 1,
                                    key=f"inp_sexo_{i}")
                            e1, e2 = st.columns(2)
                            with e1:
                                inp_dni = st.text_input("DNI del alumno:",
                                    value="" if es_prov else (dni_val if not es_prov else ""),
                                    max_chars=8, placeholder="12345678",
                                    key=f"inp_dni_{i}")
                                inp_apod = st.text_input("Apoderado:",
                                    value=apod, key=f"inp_apod_{i}",
                                    placeholder="Apellidos y nombres del apoderado")
                            with e2:
                                inp_cel = st.text_input("📱 Celular apoderado:",
                                    value=cel, max_chars=9, placeholder="987654321",
                                    key=f"inp_cel_{i}")
                                inp_dni_apod = st.text_input("DNI apoderado:",
                                    value=dni_apod, max_chars=8, placeholder="12345678",
                                    key=f"inp_dni_apod_{i}")
                            guardar = st.form_submit_button(f"💾 GUARDAR — {nombre}", type="primary", use_container_width=True)

                        if guardar:
                            df_upd = BaseDatos.cargar_matricula()
                            df_upd["DNI"] = df_upd["DNI"].astype(str).str.strip()
                            mask_upd = df_upd["DNI"] == dni_val
                            if not mask_upd.any():
                                mask_upd = df_upd["Nombre"] == nombre
                            if mask_upd.any():
                                if inp_nombre.strip():
                                    df_upd.loc[mask_upd, "Nombre"] = inp_nombre.strip().upper()
                                df_upd.loc[mask_upd, "Sexo"] = inp_sexo
                                if inp_dni.strip() and len(inp_dni.strip()) == 8 and inp_dni.strip().isdigit():
                                    df_upd.loc[mask_upd, "DNI"] = inp_dni.strip()
                                    df_upd.loc[mask_upd, "_provisional"] = "NO"
                                if inp_apod.strip():
                                    df_upd.loc[mask_upd, "Apoderado"] = inp_apod.strip().upper()
                                if inp_dni_apod.strip():
                                    df_upd.loc[mask_upd, "DNI_Apoderado"] = inp_dni_apod.strip()
                                if inp_cel.strip():
                                    df_upd.loc[mask_upd, "Celular_Apoderado"] = inp_cel.strip()
                                BaseDatos.guardar_matricula(df_upd)
                                st.success(f"✅ **{inp_nombre.strip().upper()}** actualizado correctamente.")
                                if inp_cel.strip():
                                    st.info("📱 Celular guardado — ya puede generar QR y carnet.")
                                time.sleep(1)
                                st.rerun()
                            else:
                                st.error("⚠️ No se encontró el alumno.")

    with tab_dc:
        if not df_doc.empty:
            if 'Nombre' in df_doc.columns:
                df_doc = df_doc.sort_values('Nombre')
            st.dataframe(df_doc, use_container_width=True, hide_index=True)
            buf2 = io.BytesIO()
            df_doc.to_excel(buf2, index=False, engine='openpyxl')
            buf2.seek(0)
            st.download_button("⬇️ Excel", buf2,
                               "docentes_export.xlsx", key="dxlsxd")
            # Eliminar docente
            with st.expander("🗑️ Eliminar Docente", expanded=False):
                del_dni_d = st.text_input("DNI del docente a eliminar:", key="del_dni_doc",
                                          max_chars=8, placeholder="12345678")
                if st.button("❌ ELIMINAR DOCENTE", type="primary", key="btn_del_doc"):
                    if del_dni_d and len(del_dni_d.strip()) == 8:
                        BaseDatos.eliminar_docente(del_dni_d.strip())
                        st.success(f"✅ Docente con DNI {del_dni_d} eliminado")
                        time.sleep(1)
                        st.rerun()
                    else:
                        st.error("⚠️ Ingrese un DNI válido de 8 dígitos")
        else:
            st.info("📝 Sin docentes.")


# ================================================================
# VISTA DOCENTE — Con links Google e institucionales
# ================================================================

def vista_docente(config):
    info = st.session_state.docente_info
    usuario = st.session_state.get('usuario_actual', '')
    
    # Si no hay info, intentar reconstruir desde usuarios
    if not info or not isinstance(info, dict):
        usuarios = cargar_usuarios()
        user_data = usuarios.get(usuario, {})
        info = user_data.get('docente_info')
        if not info:
            # Intentar reconstruir desde datos del usuario
            info = {
                'grado': user_data.get('grado', ''),
                'label': user_data.get('label', usuario.replace('.', ' ').title()),
                'nivel': user_data.get('nivel', ''),
            }
        st.session_state.docente_info = info
    
    grado = str(info.get('grado', ''))
    label = _nombre_completo_docente()
    if grado == 'ALL_NIVELES':
        st.markdown(f"### 👨‍🏫 {label} — Todos los Niveles")
    elif grado in ('ALL_SEC_PREU', 'ALL_SECUNDARIA'):
        st.markdown(f"### 👨‍🏫 {label} — Secundaria / Pre-Universitario")
    elif grado:
        st.markdown(f"### 👨‍🏫 {label} — {grado}")
    else:
        st.markdown(f"### 👨‍🏫 {label}")
        st.info("💡 Pida al administrador que asigne su grado en 'Gestionar Usuarios'.")

    # Determinar nivel del docente
    nivel_doc = str(info.get('nivel', ''))
    es_multigrado = (grado == 'ALL_NIVELES'
                     or 'SECUNDARIA' in nivel_doc or 'PREUNIVERSITARIO' in nivel_doc
                     or 'GRUPO' in grado or 'Sec' in grado
                     or grado in ('ALL_NIVELES', 'ALL_SEC_PREU', 'ALL_SECUNDARIA'))
    
    # Para multigrado: mostrar info general
    if es_multigrado and grado == 'ALL_NIVELES':
        st.caption("🔓 Acceso a todos los grados: Inicial, Primaria, Secundaria y Pre-Universitario")
    elif es_multigrado and grado in ('ALL_SEC_PREU', 'ALL_SECUNDARIA'):
        st.caption("🔓 Acceso a todos los grados de Secundaria y Pre-Universitario")
    
    if es_multigrado:
        # SECUNDARIA/PREUNIVERSITARIO: Sin asistencia, acceso a todos los grados
        tabs = st.tabs([
            "📝 Registrar Notas", "📝 Registro Auxiliar",
            "📋 Registro de Asistencia", "📝 Registro Bimestral",
            "📄 Registrar Ficha",
            "📝 Exámenes", "📸 Calificación YACHAY"
        ])
        with tabs[0]:
            tab_registrar_notas(config)
        with tabs[1]:
            _tab_registro_auxiliar_docente(grado, config)
        with tabs[2]:
            _tab_registro_pdf_docente(grado, config)
        with tabs[3]:
            _tab_registro_bimestral_docente(grado, config)
        with tabs[4]:
            tab_material_docente(config)
        with tabs[5]:
            tab_examenes_semanales(config)
        with tabs[6]:
            tab_calificacion_yachay(config)
    else:
        # INICIAL/PRIMARIA: Sin asistencia (solo directivo/auxiliar la manejan)
        tabs = st.tabs([
            "📝 Registrar Notas", "📝 Registro Auxiliar",
            "📋 Registro de Asistencia", "📝 Registro Bimestral",
            "📄 Registrar Ficha",
            "📝 Exámenes", "📸 Calificación YACHAY"
        ])
        with tabs[0]:
            tab_registrar_notas(config)
        with tabs[1]:
            _tab_registro_auxiliar_docente(grado, config)
        with tabs[2]:
            _tab_registro_pdf_docente(grado, config)
        with tabs[3]:
            _tab_registro_bimestral_docente(grado, config)
        with tabs[4]:
            tab_material_docente(config)
        with tabs[5]:
            tab_examenes_semanales(config)
        with tabs[6]:
            tab_calificacion_yachay(config)


def _tab_registro_auxiliar_docente(grado, config):
    """Tab de registro auxiliar para docentes"""
    st.subheader("📝 Registro Auxiliar de Evaluación")
    
    # Sec/Preu: seleccionar grado
    info = st.session_state.get('docente_info', {}) or {}
    nivel_d = str(info.get('nivel', '')).upper()
    es_sec = ('SECUNDARIA' in nivel_d or 'PREUNIVERSITARIO' in nivel_d
              or str(grado) in ('ALL_NIVELES', 'ALL_SEC_PREU', 'ALL_SECUNDARIA')
              or 'GRUPO' in str(grado) or 'Sec' in str(grado))
    
    if es_sec:
        grados_disp = _grados_del_docente()
        grado_sel = st.selectbox("🎓 Grado:", grados_disp, key="reg_aux_grado")
    else:
        grado_sel = grado
    
    tipo_reg = st.radio("Tipo:", ["📄 En blanco", "📊 Con notas registradas"],
                        horizontal=True, key="tipo_reg_aux")
    sec = st.selectbox("Sección:", ["Todas"] + SECCIONES, key="ds")
    bim = st.selectbox("📅 Periodo:", list(BIMESTRES.keys()), key="dbim")
    
    # Determinar áreas según nivel del grado seleccionado
    grado_str = str(grado_sel)
    if any(x in grado_str for x in ['GRUPO', 'Ciclo', 'Reforzamiento']):
        todas_areas = list(set(AREAS_CEPRE_UNSAAC.get('GRUPO AB', []) + AREAS_CEPRE_UNSAAC.get('GRUPO CD', [])))
    elif any(x in grado_str for x in ['Sec']):
        todas_areas = list(AREAS_MINEDU.get('SECUNDARIA', []))
        for a in set(AREAS_CEPRE_UNSAAC.get('GRUPO AB', []) + AREAS_CEPRE_UNSAAC.get('GRUPO CD', [])):
            if a not in todas_areas:
                todas_areas.append(a)
    elif 'Inicial' in grado_str:
        todas_areas = AREAS_MINEDU.get('INICIAL', ['Comunicación', 'Matemática'])
    else:
        todas_areas = AREAS_MINEDU.get('PRIMARIA', ['Comunicación', 'Matemática'])
    
    if tipo_reg == "📄 En blanco":
        st.markdown("**Cursos:**")
        cursos_d = st.multiselect("Seleccione cursos:", todas_areas,
                                   default=todas_areas[:3], key="dc_cursos")
    else:
        # Mostrar cursos con notas registradas
        notas = {}
        if Path('notas.json').exists():
            with open('notas.json', 'r', encoding='utf-8') as f:
                notas = json.load(f)
        cursos_con_notas = {}
        for k, v in notas.items():
            if isinstance(v, dict) and v.get('grado') == grado_sel:
                area_n = v.get('area', '')
                if area_n not in cursos_con_notas:
                    cursos_con_notas[area_n] = 0
                cursos_con_notas[area_n] += 1
        # Mostrar info de notas registradas
        if cursos_con_notas:
            st.success(f"📊 Cursos con notas: {len(cursos_con_notas)}")
            for cn, cnt in sorted(cursos_con_notas.items()):
                st.caption(f"  📚 **{cn}** — {cnt} registro(s)")
        else:
            st.info("📭 No hay notas registradas aún para este grado")
        # Permitir seleccionar cursos también en este modo
        opciones_areas = list(cursos_con_notas.keys()) if cursos_con_notas else todas_areas
        for a in todas_areas:
            if a not in opciones_areas:
                opciones_areas.append(a)
        cursos_d = st.multiselect("📚 Seleccione cursos:", opciones_areas,
                                   default=list(cursos_con_notas.keys())[:3] if cursos_con_notas else opciones_areas[:3],
                                   key="dc_cursos_notas")
        
    dg = BaseDatos.obtener_estudiantes_grado(grado_sel, sec)
    st.info(f"📊 {len(dg)} estudiantes — {grado_sel}")
    if not dg.empty:
        st.dataframe(dg[['Nombre', 'DNI', 'Grado', 'Seccion']],
                     use_container_width=True, hide_index=True)
    cd1, cd2 = st.columns(2)
    with cd1:
        if st.button("📥 Generar PDF Auxiliar", type="primary",
                     use_container_width=True, key="ddra"):
            if not dg.empty:
                lg = grado if grado != "ALL_SECUNDARIA" else "Secundaria"
                sl = sec if sec != "Todas" else "Todas"
                with st.spinner("Generando PDF..."):
                    pdf = generar_registro_auxiliar_pdf(lg, sl, config['anio'], bim, dg, cursos_d)
                st.session_state['_aux_pdf_d'] = pdf
                st.session_state['_aux_key_d'] = f"RegAux_{lg}_{bim}"
    with cd2:
        if st.button("📄 Generar Word Auxiliar", use_container_width=True, key="ddra_docx_btn"):
            if not dg.empty:
                lg = grado if grado != "ALL_SECUNDARIA" else "Secundaria"
                sl = sec if sec != "Todas" else "Todas"
                with st.spinner("Generando Word..."):
                    docx_aux2 = generar_registro_auxiliar_docx(lg, sl, config['anio'], bim, dg, cursos_d)
                if docx_aux2:
                    st.session_state['_aux_docx_d'] = docx_aux2
                    st.session_state['_aux_key_d'] = f"RegAux_{lg}_{bim}"
    if st.session_state.get('_aux_pdf_d'):
        st.download_button("⬇️ Descargar PDF", st.session_state['_aux_pdf_d'],
                           f"{st.session_state.get('_aux_key_d','RegAux')}.pdf",
                           "application/pdf", key="ddra2")
    if st.session_state.get('_aux_docx_d'):
        st.download_button("⬇️ Descargar Word (.docx)", st.session_state['_aux_docx_d'],
                           f"{st.session_state.get('_aux_key_d','RegAux')}.docx",
                           "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           key="ddra2_docx")


def _tab_registro_pdf_docente(grado, config):
    """Tab de registro PDF para docentes — SOLO Asistencia"""
    st.subheader("📋 Registro de Asistencia PDF")
    
    # Resolver grado
    info = st.session_state.get('docente_info', {}) or {}
    nivel_d = str(info.get('nivel', '')).upper()
    es_sec = ('SECUNDARIA' in nivel_d or 'PREUNIVERSITARIO' in nivel_d
              or str(grado) in ('ALL_NIVELES', 'ALL_SEC_PREU', 'ALL_SECUNDARIA')
              or 'GRUPO' in str(grado) or 'Sec' in str(grado))

    if es_sec:
        grados_disp = _grados_del_docente()
        grado_sel = st.selectbox("🎓 Grado:", grados_disp, key="reg_pdf_grado")
    else:
        grado_sel = grado
    
    sec2 = st.selectbox("Sección:", ["Todas"] + SECCIONES, key="ds2")
    meses_opts = list(MESES_ESCOLARES.items())
    meses_sel = st.multiselect(
        "Meses:",
        [f"{v} ({k})" for k, v in meses_opts],
        default=[f"{v} ({k})" for k, v in meses_opts[:1]],
        key="dmsel")
    meses_nums = [int(m.split('(')[1].replace(')', '')) for m in meses_sel]
    dg2 = BaseDatos.obtener_estudiantes_grado(grado_sel, sec2)
    st.info(f"📊 {len(dg2)} estudiantes — {grado_sel}")
    if st.button("📥 Descargar Registro Asistencia PDF", type="primary",
                 use_container_width=True, key="ddas"):
        if not dg2.empty and meses_nums:
            lg = grado_sel if grado_sel not in ("ALL_SECUNDARIA", "ALL_SEC_PREU") else "Secundaria"
            sl = sec2 if sec2 != "Todas" else "Todas"
            pdf = generar_registro_asistencia_pdf(
                lg, sl, config['anio'], dg2, meses_nums,
                docente=_nombre_completo_docente())
            st.download_button("⬇️ PDF", pdf,
                               f"RegAsist_{lg}.pdf",
                               "application/pdf", key="ddas2")


def _tab_registro_bimestral_docente(grado, config):
    """Tab SEPARADA de registro bimestral de notas"""
    st.subheader("📝 Registro Bimestral de Notas")
    st.caption("Registro para llenar notas consolidadas por competencia del Currículo Nacional")

    info = st.session_state.get('docente_info', {}) or {}
    nivel_d = str(info.get('nivel', '')).upper()
    es_sec = ('SECUNDARIA' in nivel_d or 'PREUNIVERSITARIO' in nivel_d
              or str(grado) in ('ALL_NIVELES', 'ALL_SEC_PREU', 'ALL_SECUNDARIA')
              or 'GRUPO' in str(grado) or 'Sec' in str(grado))

    if es_sec:
        grados_disp = _grados_del_docente()
        grado_sel = st.selectbox("🎓 Grado:", grados_disp, key="rbim_grado")
    else:
        grado_sel = grado

    sec2 = st.selectbox("Sección:", ["Todas"] + SECCIONES, key="rbim_sec")

    bimestre = st.selectbox("📅 Bimestre:", [
        "I Bimestre", "II Bimestre", "III Bimestre", "IV Bimestre"
    ], key="rbim_bim")

    # Determinar nivel del grado seleccionado
    nivel_grado = "PRIMARIA"
    for niv, grados_niv in NIVELES_GRADOS.items():
        if grado_sel in grados_niv:
            nivel_grado = niv
            break

    # Áreas del Currículo Nacional para este nivel
    areas_cn = AREAS_MINEDU.get(nivel_grado, AREAS_MINEDU.get('PRIMARIA', []))
    areas_sel = st.multiselect(
        "📚 Áreas (Currículo Nacional):", areas_cn,
        default=areas_cn, key="rbim_areas")

    if areas_sel:
        with st.expander("👁️ Ver competencias por área", expanded=False):
            for area in areas_sel:
                comps = COMPETENCIAS_CN.get(area, ['Competencia general'])
                st.markdown(f"**{area}:** {', '.join(comps)}")

        dg2 = BaseDatos.obtener_estudiantes_grado(grado_sel, sec2)
        st.info(f"📊 {len(dg2)} estudiantes — {grado_sel} | "
                f"{len(areas_sel)} áreas | {bimestre}")

        if st.button("📥 Descargar Registro Bimestral PDF", type="primary",
                     use_container_width=True, key="rbim_dl"):
            if not dg2.empty:
                lg = grado_sel if grado_sel not in ("ALL_SECUNDARIA", "ALL_SEC_PREU") else "Secundaria"
                sl = sec2 if sec2 != "Todas" else "Todas"
                pdf = generar_registro_bimestral_pdf(
                    lg, sl, config['anio'], dg2,
                    bimestre, areas_sel, nivel_grado,
                    docente=_nombre_completo_docente())
                st.download_button("⬇️ PDF Bimestral", pdf,
                                   f"RegNotas_{bimestre.replace(' ','_')}_{lg}.pdf",
                                   "application/pdf", key="rbim_dl2")
            else:
                st.warning("No hay estudiantes para este grado/sección.")
    else:
        st.warning("Seleccione al menos un área.")


# ================================================================
# REGISTRO DE INCIDENCIAS
# ================================================================

TIPOS_INCIDENCIA = [
    'Conductual (Indisciplina)',
    'Académica (Plagio, falta de tareas)',
    'Convivencia (Conflicto entre pares)',
    'Presunto caso de Violencia Escolar (Bullying)',
    'Salud / Accidente',
    'Infraestructura / Daño a propiedad',
]

DERIVACIONES = [
    'No requiere derivación',
    'Psicología',
    'Dirección',
    'Tutoría',
    'Reporte portal SíseVe',
    'DEMUNA',
    'Otra',
]


def generar_incidencia_pdf(datos, config):
    """Genera PDF del registro de incidencia"""
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    w, h = A4

    # Encabezado
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(w / 2, h - 50, config.get('nombre_ie', 'I.E.P. ALTERNATIVO YACHAY'))
    c.setFont("Helvetica-Bold", 12)
    c.drawCentredString(w / 2, h - 70, "REGISTRO DE INCIDENCIAS")
    c.setFont("Helvetica", 9)
    c.drawCentredString(w / 2, h - 85, config.get('ubicacion', 'Chinchero, Cusco'))

    c.setStrokeColor(colors.HexColor("#1a56db"))
    c.setLineWidth(2)
    c.line(40, h - 95, w - 40, h - 95)

    y = h - 120
    c.setFont("Helvetica-Bold", 11)
    c.drawString(50, y, "I. INFORMACIÓN GENERAL")
    y -= 20

    campos = [
        ("Código de Incidencia:", datos.get('codigo', '')),
        ("Fecha y Hora:", f"{datos.get('fecha', '')} — {datos.get('hora', '')}"),
        ("Lugar:", datos.get('lugar', '')),
        ("Nivel:", datos.get('nivel', '')),
        ("Grado y Sección:", f"{datos.get('grado', '')} — {datos.get('seccion', '')}"),
    ]
    c.setFont("Helvetica", 10)
    for label, valor in campos:
        c.setFont("Helvetica-Bold", 10)
        c.drawString(60, y, label)
        c.setFont("Helvetica", 10)
        c.drawString(200, y, str(valor))
        y -= 18

    y -= 10
    c.setFont("Helvetica-Bold", 11)
    c.drawString(50, y, "II. IDENTIFICACIÓN DE INVOLUCRADOS")
    y -= 20
    for label_campo in ['Afectado(s)', 'Implicado(s)', 'Reportante']:
        c.setFont("Helvetica-Bold", 10)
        c.drawString(60, y, f"{label_campo}:")
        c.setFont("Helvetica", 10)
        c.drawString(160, y, str(datos.get(label_campo.lower().replace('(s)', '').strip(), '')))
        y -= 18

    y -= 10
    c.setFont("Helvetica-Bold", 11)
    c.drawString(50, y, "III. DESCRIPCIÓN DE LA INCIDENCIA")
    y -= 20
    c.setFont("Helvetica-Bold", 10)
    c.drawString(60, y, f"Tipo: {datos.get('tipo', '')}")
    y -= 20

    c.setFont("Helvetica-Bold", 10)
    c.drawString(60, y, "Relato de los hechos:")
    y -= 15
    c.setFont("Helvetica", 9)
    relato = str(datos.get('relato', ''))
    for linea in textwrap.wrap(relato, 85):
        c.drawString(70, y, linea)
        y -= 13
        if y < 100:
            c.showPage()
            y = h - 50

    y -= 15
    c.setFont("Helvetica-Bold", 11)
    c.drawString(50, y, "IV. MEDIDAS Y ACCIONES TOMADAS")
    y -= 20
    for label_accion, key in [("Acción Inmediata:", 'accion_inmediata'),
                               ("Compromisos:", 'compromisos'),
                               ("Derivación:", 'derivacion')]:
        c.setFont("Helvetica-Bold", 10)
        c.drawString(60, y, label_accion)
        y -= 15
        c.setFont("Helvetica", 9)
        for linea in textwrap.wrap(str(datos.get(key, '')), 85):
            c.drawString(70, y, linea)
            y -= 13

    y -= 30
    c.setFont("Helvetica-Bold", 11)
    c.drawString(50, y, "V. FIRMAS DE CONFORMIDAD")
    y -= 40
    firmas = ["Responsable del Registro", "Estudiante (si aplica)",
              "Padre de Familia"]
    for i, firma in enumerate(firmas):
        x = 60 + (i * 170)
        c.line(x, y, x + 140, y)
        c.setFont("Helvetica", 8)
        c.drawCentredString(x + 70, y - 12, firma)

    # Pie
    c.setFont("Helvetica", 7)
    c.drawCentredString(w / 2, 30, f"Generado por YACHAY PRO — {hora_peru_str()}")

    c.save()
    buf.seek(0)
    return buf.getvalue()


def tab_incidencias(config):
    """Tab de Registro de Incidencias"""
    st.subheader("📝 Registro de Incidencias")

    gs = _gs()

    col1, col2 = st.columns([2, 1])

    with col1:
        st.markdown("### Nueva Incidencia")

        # Generar código automático
        if gs:
            codigo = gs.generar_siguiente_codigo_incidencia()
        else:
            codigo = f"INC-{hora_peru().year}-{int(time.time()) % 1000:03d}"

        with st.form("form_incidencia", clear_on_submit=True):
            st.info(f"📌 Código: **{codigo}**")

            ci1, ci2 = st.columns(2)
            with ci1:
                fecha_inc = st.date_input("Fecha:", value=hora_peru().date(),
                                           key="fld_inc_fecha")
                nivel_inc = st.selectbox("Nivel:", NIVELES_LIST, key="fld_inc_nivel")
            with ci2:
                hora_inc = st.text_input("Hora:", value=hora_peru().strftime('%H:%M'),
                                          key="fld_inc_hora")
                grado_inc = st.selectbox("Grado:", GRADOS_OPCIONES, key="fld_inc_grado")

            lugar = st.text_input("Lugar:", placeholder="Ej: Aula, patio, alrededores",
                                  key="fld_inc_lugar")
            seccion_inc = st.selectbox("Sección:", SECCIONES, key="fld_inc_sec")

            tipo_inc = st.selectbox("Tipo de Incidencia:", TIPOS_INCIDENCIA,
                                    key="fld_inc_tipo")

            st.markdown("**Involucrados:**")
            afectados = st.text_area("Afectado(s) - Nombres, DNI:",
                                     key="fld_inc_afect")
            implicados = st.text_area("Implicado(s) - Nombres, DNI:",
                                      key="fld_inc_implic")
            reportante = st.text_input("Informante/Reportante:",
                                       key="fld_inc_report")

            relato = st.text_area("Relato de los hechos:",
                                  placeholder="Descripción objetiva...",
                                  key="fld_inc_relato")

            accion = st.text_area("Acción Inmediata:", key="fld_inc_accion")
            compromisos = st.text_area("Compromisos:", key="fld_inc_comp")
            derivacion = st.selectbox("Derivación:", DERIVACIONES, key="fld_inc_deriv")

            submitted = st.form_submit_button("💾 REGISTRAR INCIDENCIA",
                                               type="primary",
                                               use_container_width=True)
            if submitted:
                datos_inc = {
                    'codigo': codigo,
                    'fecha': str(fecha_inc),
                    'hora': str(hora_inc),
                    'lugar': lugar,
                    'nivel': nivel_inc,
                    'grado': grado_inc,
                    'seccion': seccion_inc,
                    'tipo': tipo_inc,
                    'afectados': afectados,
                    'implicados': implicados,
                    'reportante': reportante,
                    'dni_reportante': '',
                    'relato': relato,
                    'accion_inmediata': accion,
                    'compromisos': compromisos,
                    'derivacion': derivacion,
                    'registrado_por': st.session_state.get('usuario_actual', ''),
                }

                # Guardar en Google Sheets
                if gs:
                    gs.guardar_incidencia(datos_inc)
                    st.success(f"✅ Incidencia {codigo} registrada y guardada en Google Sheets")
                else:
                    st.success(f"✅ Incidencia {codigo} registrada")

                # Guardar PDF en session para descargar fuera del form
                pdf = generar_incidencia_pdf(datos_inc, config)
                st.session_state['ultimo_pdf_incidencia'] = pdf
                st.session_state['ultimo_codigo_incidencia'] = codigo

        # Botón de descarga FUERA del formulario
        if st.session_state.get('ultimo_pdf_incidencia'):
            cod = st.session_state.get('ultimo_codigo_incidencia', 'INC')
            st.download_button("📥 Descargar PDF de Incidencia", 
                               st.session_state['ultimo_pdf_incidencia'],
                               f"Incidencia_{cod}.pdf",
                               "application/pdf", key="dl_inc_outside")

    with col2:
        st.markdown("### 📋 Historial")
        if gs:
            incidencias = gs.leer_incidencias()
            if incidencias:
                for inc in reversed(incidencias[-20:]):
                    with st.expander(
                        f"📌 {inc.get('codigo', '?')} — {inc.get('fecha', '')}"):
                        st.write(f"**Tipo:** {inc.get('tipo', '')}")
                        st.write(f"**Grado:** {inc.get('grado', '')}")
                        st.write(f"**Afectados:** {inc.get('afectados', '')}")
                        st.write(f"**Relato:** {inc.get('relato', '')[:200]}")
                        try:
                            pdf_h = generar_incidencia_pdf(inc, config)
                            st.download_button("📥 PDF",
                                               pdf_h,
                                               f"Inc_{inc.get('codigo', '')}.pdf",
                                               "application/pdf",
                                               key=f"dl_hist_{inc.get('codigo', '')}_{id(inc)}")
                        except Exception:
                            pass
            else:
                st.info("Sin incidencias registradas")
        else:
            st.warning("⚠️ Conecta Google Sheets para ver historial")


# ================================================================
# REPORTES MENSUALES Y HISTORIAL
# ================================================================

def generar_reporte_asistencia_mensual_pdf(datos_mes, grado, mes, anio, config):
    """PDF de reporte mensual de asistencia por grado"""
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=landscape(A4))
    w, h = landscape(A4)

    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(w / 2, h - 40,
                        config.get('nombre_ie', 'I.E.P. ALTERNATIVO YACHAY'))
    c.setFont("Helvetica-Bold", 11)
    nombre_mes = ['', 'Enero', 'Febrero', 'Marzo', 'Abril', 'Mayo', 'Junio',
                  'Julio', 'Agosto', 'Septiembre', 'Octubre', 'Noviembre',
                  'Diciembre']
    c.drawCentredString(w / 2, h - 58,
                        f"REPORTE MENSUAL DE ASISTENCIA — {nombre_mes[mes]} {anio}")
    c.drawCentredString(w / 2, h - 73, f"Grado: {grado}")

    # Calcular días del mes
    import calendar as cal_mod
    dias_mes = cal_mod.monthrange(anio, mes)[1]

    y = h - 100
    c.setFont("Helvetica-Bold", 7)

    # Encabezados
    x_start = 30
    c.drawString(x_start, y, "#")
    c.drawString(x_start + 15, y, "Nombre")
    c.drawString(x_start + 180, y, "DNI")

    # Días como columnas
    x_dia = x_start + 225
    for d in range(1, dias_mes + 1):
        dia_semana = cal_mod.weekday(anio, mes, d)
        if dia_semana < 5:  # L-V
            c.drawCentredString(x_dia, y, str(d))
            x_dia += 18

    c.drawString(x_dia + 5, y, "Total")
    c.drawString(x_dia + 35, y, "%")

    y -= 3
    c.setLineWidth(0.5)
    c.line(x_start, y, w - 30, y)
    y -= 12

    c.setFont("Helvetica", 6)
    num = 0
    for nombre, info in sorted(datos_mes.items()):
        num += 1
        c.drawString(x_start, y, str(num))
        c.drawString(x_start + 15, y, nombre[:30])
        c.drawString(x_start + 180, y, str(info.get('dni', '')))

        x_dia = x_start + 225
        total_asist = 0
        total_dias_hab = 0
        for d in range(1, dias_mes + 1):
            dia_semana = cal_mod.weekday(anio, mes, d)
            if dia_semana < 5:
                total_dias_hab += 1
                fecha_str = f"{anio}-{mes:02d}-{d:02d}"
                if fecha_str in info.get('fechas', {}):
                    c.setFillColor(colors.HexColor("#16a34a"))
                    c.drawCentredString(x_dia, y, "✓")
                    c.setFillColor(colors.black)
                    total_asist += 1
                else:
                    c.setFillColor(colors.HexColor("#dc2626"))
                    c.drawCentredString(x_dia, y, "✗")
                    c.setFillColor(colors.black)
                x_dia += 18

        pct = (total_asist / total_dias_hab * 100) if total_dias_hab > 0 else 0
        c.drawString(x_dia + 5, y, str(total_asist))
        c.drawString(x_dia + 35, y, f"{pct:.0f}%")

        y -= 11
        if y < 40:
            c.showPage()
            y = h - 50
            c.setFont("Helvetica", 6)

    c.setFont("Helvetica", 7)
    c.drawCentredString(w / 2, 20,
                        f"YACHAY PRO — Generado: {hora_peru_str()}")
    c.save()
    buf.seek(0)
    return buf.getvalue()


def generar_reporte_examen_zipgrade(resultado, config):
    """Genera reporte estilo ZipGrade: verde=correcta, rojo=incorrecta, azul=no marcó pero era correcta"""
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    w, h = A4

    titulo = resultado.get('titulo', 'Evaluación')
    fecha = resultado.get('fecha', '')
    areas = resultado.get('areas', [])
    alumnos = resultado.get('alumnos', [])

    # COLOR DEFINITIONS
    COLOR_CORRECTO = colors.HexColor("#16a34a")   # Verde
    COLOR_INCORRECTO = colors.HexColor("#dc2626")  # Rojo
    COLOR_NO_MARCO = colors.HexColor("#2563eb")    # Azul

    pagina = 0
    for alumno in alumnos:
        if pagina > 0:
            c.showPage()
        pagina += 1

        # Encabezado
        c.setFont("Helvetica-Bold", 12)
        c.drawCentredString(w / 2, h - 40,
                            config.get('nombre_ie', 'I.E.P. ALTERNATIVO YACHAY'))
        c.setFont("Helvetica-Bold", 10)
        c.drawCentredString(w / 2, h - 58, f"REPORTE DE EVALUACIÓN — {titulo}")
        c.setFont("Helvetica", 9)
        c.drawCentredString(w / 2, h - 73, f"Fecha: {fecha}")

        c.setLineWidth(2)
        c.setStrokeColor(colors.HexColor("#1a56db"))
        c.line(40, h - 80, w - 40, h - 80)

        # Datos del alumno
        y = h - 100
        c.setFont("Helvetica-Bold", 11)
        c.drawString(50, y, f"Alumno: {alumno.get('nombre', '')}")
        c.setFont("Helvetica", 10)
        c.drawString(400, y, f"DNI: {alumno.get('dni', '')}")
        y -= 18
        prom = alumno.get('promedio', 0)
        nota_letra = 'AD' if prom >= 18 else 'A' if prom >= 14 else 'B' if prom >= 11 else 'C'
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y, f"PROMEDIO GENERAL: {prom:.1f}/20 ({nota_letra})")
        y -= 25

        # Leyenda de colores
        c.setFont("Helvetica-Bold", 8)
        c.setFillColor(COLOR_CORRECTO)
        c.rect(50, y, 12, 10, fill=1, stroke=0)
        c.setFillColor(colors.black)
        c.drawString(65, y + 2, "= Correcta")

        c.setFillColor(COLOR_INCORRECTO)
        c.rect(150, y, 12, 10, fill=1, stroke=0)
        c.setFillColor(colors.black)
        c.drawString(165, y + 2, "= Incorrecta")

        c.setFillColor(COLOR_NO_MARCO)
        c.rect(270, y, 12, 10, fill=1, stroke=0)
        c.setFillColor(colors.black)
        c.drawString(285, y + 2, "= No marcó (era correcta)")
        y -= 25

        # Por cada área
        notas_alumno = alumno.get('notas', [])
        for idx_area, nota_area in enumerate(notas_alumno):
            area_nombre = nota_area.get('area', f'Área {idx_area + 1}')
            nota_val = nota_area.get('nota', 0)
            correctas = nota_area.get('correctas', 0)
            total = nota_area.get('total', 10)
            respuestas = str(nota_area.get('respuestas', ''))
            claves = str(nota_area.get('claves', ''))

            c.setFont("Helvetica-Bold", 10)
            c.setFillColor(colors.HexColor("#1a56db"))
            c.drawString(50, y, f"📝 {area_nombre} — {nota_val:.1f}/20 ({correctas}/{total})")
            c.setFillColor(colors.black)
            y -= 18

            # Tabla de respuestas con colores
            opciones = ['A', 'B', 'C', 'D']
            c.setFont("Helvetica-Bold", 8)
            c.drawString(60, y, "Preg")
            for oi, op in enumerate(opciones):
                c.drawCentredString(120 + oi * 40, y, op)
            c.drawString(290, y, "Correcta")
            c.drawString(355, y, "Marcó")
            c.drawString(410, y, "Resultado")
            y -= 3
            c.line(55, y, 470, y)
            y -= 12

            c.setFont("Helvetica", 8)
            for p in range(total):
                clave_p = claves[p] if p < len(claves) else '?'
                resp_p = respuestas[p] if p < len(respuestas) else '?'
                es_correcta = resp_p == clave_p and resp_p != '?'
                no_marco = resp_p == '?'

                c.drawString(60, y, f"  {p + 1}")

                # Dibujar burbujas con colores
                for oi, op in enumerate(opciones):
                    cx = 120 + oi * 40
                    if op == clave_p and es_correcta:
                        c.setFillColor(COLOR_CORRECTO)
                        c.circle(cx, y + 3, 7, fill=1, stroke=0)
                        c.setFillColor(colors.white)
                        c.drawCentredString(cx, y + 0.5, op)
                    elif op == resp_p and not es_correcta and not no_marco:
                        c.setFillColor(COLOR_INCORRECTO)
                        c.circle(cx, y + 3, 7, fill=1, stroke=0)
                        c.setFillColor(colors.white)
                        c.drawCentredString(cx, y + 0.5, op)
                    elif op == clave_p and (not es_correcta):
                        c.setFillColor(COLOR_NO_MARCO)
                        c.circle(cx, y + 3, 7, fill=1, stroke=0)
                        c.setFillColor(colors.white)
                        c.drawCentredString(cx, y + 0.5, op)
                    else:
                        c.setStrokeColor(colors.HexColor("#94a3b8"))
                        c.setFillColor(colors.white)
                        c.circle(cx, y + 3, 7, fill=1, stroke=1)
                        c.setFillColor(colors.HexColor("#94a3b8"))
                        c.drawCentredString(cx, y + 0.5, op)

                c.setFillColor(colors.black)
                c.drawString(290, y, clave_p)
                c.drawString(355, y, resp_p)

                if es_correcta:
                    c.setFillColor(COLOR_CORRECTO)
                    c.drawString(410, y, "✓ Correcta")
                elif no_marco:
                    c.setFillColor(COLOR_NO_MARCO)
                    c.drawString(410, y, "— Sin marcar")
                else:
                    c.setFillColor(COLOR_INCORRECTO)
                    c.drawString(410, y, "✗ Incorrecta")

                c.setFillColor(colors.black)
                y -= 14

                if y < 60:
                    c.showPage()
                    y = h - 50
                    c.setFont("Helvetica", 8)

            y -= 10

    # Pie
    c.setFont("Helvetica", 7)
    c.drawCentredString(w / 2, 25, f"YACHAY PRO — Generado: {hora_peru_str()}")
    c.save()
    buf.seek(0)
    return buf.getvalue()


def tab_reportes(config):
    """Tab de reportes y historial — COMPLETO"""
    st.subheader("📊 Reportes e Historial")

    subtab = st.radio("Seleccionar:", [
        "📋 Asistencia Mensual", "👨‍🏫 Asistencia Docentes",
        "📊 Reporte Integral",
        "📄 Reporte ZipGrade", "🏆 Historial de Evaluaciones",
        "📁 Fichas Docentes"
    ], horizontal=True, key="rep_tipo")

    gs = _gs()

    if subtab == "🏆 Historial de Evaluaciones":
        st.markdown("### 🏆 Historial de Evaluaciones — Vista Director")
        hist = _cargar_historial_evaluaciones()
        if not hist:
            st.info("📭 No hay evaluaciones guardadas en el historial.")
            return

        # Filtros
        fc1, fc2 = st.columns(2)
        with fc1:
            grados_hist = sorted(set(v.get('grado','') for v in hist.values()))
            filtro_grado = st.selectbox("Filtrar por grado:", ["Todos"] + grados_hist, key="rep_hist_grado")
        with fc2:
            docentes_hist = sorted(set(v.get('docente','') for v in hist.values()))
            filtro_doc = st.selectbox("Filtrar por docente:", ["Todos"] + docentes_hist, key="rep_hist_doc")

        total_mostradas = 0
        for clave, ev in sorted(hist.items(), reverse=True):
            if filtro_grado != "Todos" and ev.get('grado') != filtro_grado:
                continue
            if filtro_doc != "Todos" and ev.get('docente') != filtro_doc:
                continue
            total_mostradas += 1
            areas_ev = ev.get('areas', [])
            if areas_ev and isinstance(areas_ev[0], dict):
                areas_nombres_ev = [a['nombre'] for a in areas_ev]
            else:
                areas_nombres_ev = list(areas_ev)
            titulo_ev = ev.get('titulo', '') or ''
            label_ev = f"📝 {ev.get('grado','')} | {ev.get('periodo','')} | {ev.get('fecha','')}"
            if titulo_ev:
                label_ev += f" — {titulo_ev}"
            with st.expander(label_ev):
                st.caption(f"👤 Docente: {ev.get('docente_nombre', ev.get('docente','—'))} | 📚 Áreas: {', '.join(areas_nombres_ev)} | 👥 Estudiantes: {len(ev.get('ranking',[]))}")
                ranking_ev = ev.get('ranking', [])
                if ranking_ev:
                    df_ev = pd.DataFrame(ranking_ev)
                    cols_ev = ['Puesto','Medalla','Nombre'] + areas_nombres_ev + ['Promedio']
                    cols_ev = [c for c in cols_ev if c in df_ev.columns]
                    st.dataframe(df_ev[cols_ev], use_container_width=True, hide_index=True)
                    if st.button("📥 PDF Ranking", key=f"rep_pdf_{clave}", type="primary"):
                        pdf_ev = _generar_ranking_pdf(ranking_ev, areas_nombres_ev,
                                                      ev.get('grado',''), ev.get('periodo',''), config)
                        st.download_button("⬇️ Descargar PDF", pdf_ev,
                                           f"Ranking_{ev.get('grado','')}_{ev.get('periodo','')}_{ev.get('fecha','')}.pdf",
                                           "application/pdf", key=f"dl_rep_{clave}")
        if total_mostradas == 0:
            st.info("No hay evaluaciones para los filtros seleccionados.")
        return

    if subtab == "👨‍🏫 Asistencia Docentes":
        st.markdown("""<div style='background:linear-gradient(135deg,#001e7c,#0044cc);color:white;
            padding:15px 20px;border-radius:12px;margin-bottom:15px;text-align:center;'>
            <h3 style='margin:0;color:white;'>👨‍🏫 I.E.P. ALTERNATIVO YACHAY</h3>
            <p style='margin:4px 0 0;color:#b8d4ff;'>Reporte de Asistencia y Puntualidad — Docentes</p>
            <p style='margin:2px 0 0;color:#FFD700;font-size:0.85rem;'>📍 Chinchero, Cusco — {hora_peru().year}</p>
        </div>""", unsafe_allow_html=True)

        # ── Cargar datos de asistencia ──────────────────────────────────
        asistencias = {}
        if Path(ARCHIVO_ASISTENCIAS).exists():
            try:
                with open(ARCHIVO_ASISTENCIAS, 'r', encoding='utf-8') as f:
                    asistencias = json.load(f)
            except Exception:
                pass

        # GS
        datos_gs_doc = {}
        if gs:
            try:
                ws = gs.sh.worksheet('asistencias')
                registros = ws.get_all_records()
                for r in registros:
                    if str(r.get('tipo_persona', '')).lower() == 'docente':
                        fecha = str(r.get('fecha', ''))
                        dni = str(r.get('dni', ''))
                        if fecha and dni:
                            if fecha not in datos_gs_doc:
                                datos_gs_doc[fecha] = {}
                            datos_gs_doc[fecha][dni] = {
                                'nombre': r.get('nombre', ''),
                                'entrada': r.get('hora_entrada', ''),
                                'salida': r.get('hora_salida', ''),
                                'es_docente': True
                            }
            except Exception:
                pass

        todas_fechas = set(asistencias.keys()) | set(datos_gs_doc.keys())

        # Extraer solo docentes
        docentes_asist = {}
        for fecha in sorted(todas_fechas):
            regs = {**datos_gs_doc.get(fecha, {}), **asistencias.get(fecha, {})}
            for dni, dat in regs.items():
                if dat.get('es_docente', False):
                    nm = dat.get('nombre', dni)
                    if nm not in docentes_asist:
                        docentes_asist[nm] = {}
                    entrada = dat.get('entrada', '') or dat.get('tardanza', '')
                    tardanza_auto = _es_tardanza_docente(entrada) if entrada else False
                    docentes_asist[nm][fecha] = {
                        'entrada': entrada,
                        'salida': dat.get('salida', ''),
                        'tardanza': tardanza_auto,
                        'entrada_tarde': dat.get('entrada_tarde', ''),
                        'salida_tarde': dat.get('salida_tarde', ''),
                        'dni': dni,
                    }

        if not docentes_asist:
            st.info("📭 No hay registros de asistencia de docentes aún.")
        else:
            df_doc_list = BaseDatos.cargar_docentes()

            modo = st.radio("Vista:", [
                "📅 Semanal por Mes", "📆 Resumen Mensual",
                "✏️ Editar Registros", "⏱️ Horas Sec/PreU",
                "📱 WhatsApp Docentes"
            ], horizontal=True, key="rep_doc_modo")

            # ── Meses escolares (marzo a diciembre) ──────────────────────
            meses_esc = [(m, n) for m, n in MESES_ESCOLARES.items() if m >= 3]

            if modo == "📅 Semanal por Mes":
                mes_sel = st.selectbox("📆 Mes:", meses_esc,
                                        format_func=lambda x: x[1],
                                        key="rep_doc_mes_sem")
                mes_num = mes_sel[0]
                mes_nombre = mes_sel[1]
                anio_sel = hora_peru().year

                semanas = _semanas_del_mes(mes_num, anio_sel)
                if not semanas:
                    st.info("Sin semanas para este mes.")
                else:
                    st.markdown(f"### 📅 {mes_nombre} {anio_sel} — Semanas {semanas[0][0]} a {semanas[-1][0]}")

                    for sem_num, lun, vie in semanas:
                        with st.expander(
                            f"📌 Semana {sem_num}: {lun.strftime('%d/%m')} – {vie.strftime('%d/%m')}",
                            expanded=(sem_num == _semana_escolar_actual())
                        ):
                            data_sem = []
                            for nm in sorted(docentes_asist.keys()):
                                dias_ok = 0
                                tardanzas = 0
                                faltas = 0
                                detalle = []
                                for d in range(5):
                                    dia = lun + timedelta(days=d)
                                    fecha_str = dia.strftime('%Y-%m-%d')
                                    reg = docentes_asist[nm].get(fecha_str, {})
                                    dia_nombre = ['Lun','Mar','Mié','Jue','Vie'][d]
                                    if reg.get('entrada'):
                                        dias_ok += 1
                                        if reg.get('tardanza'):
                                            tardanzas += 1
                                            detalle.append(f"{dia_nombre}:⏰{reg['entrada']}")
                                        else:
                                            detalle.append(f"{dia_nombre}:✅{reg['entrada']}")
                                    elif dia <= hora_peru().date():
                                        faltas += 1
                                        detalle.append(f"{dia_nombre}:❌")
                                    else:
                                        detalle.append(f"{dia_nombre}:—")
                                puntualidad = round((dias_ok - tardanzas) / max(dias_ok, 1) * 100)
                                data_sem.append({
                                    'Docente': nm, 'Asist.': dias_ok,
                                    'Tard.': tardanzas, 'Faltas': faltas,
                                    'Puntualidad': f"{puntualidad}%",
                                    'Detalle': ' | '.join(detalle)
                                })

                            if data_sem:
                                df_s = pd.DataFrame(data_sem)
                                st.dataframe(df_s, use_container_width=True, hide_index=True)

                    # ── Gráfico de barras del mes ──────────────────────────
                    st.markdown("---")
                    st.markdown(f"### 📊 Resumen Gráfico — {mes_nombre}")
                    chart_data = []
                    for nm in sorted(docentes_asist.keys()):
                        dias_total = 0
                        puntuales = 0
                        tardes = 0
                        for fecha_str, reg in docentes_asist[nm].items():
                            try:
                                fd = datetime.strptime(fecha_str, '%Y-%m-%d').date()
                                if fd.month == mes_num and fd.year == anio_sel:
                                    if reg.get('entrada'):
                                        dias_total += 1
                                        if reg.get('tardanza'):
                                            tardes += 1
                                        else:
                                            puntuales += 1
                            except Exception:
                                pass
                        # Nombre corto
                        nm_corto = nm.split()[-1] if ' ' in nm else nm
                        if len(nm_corto) > 12:
                            nm_corto = nm_corto[:10] + ".."
                        chart_data.append({'Docente': nm_corto, 'Puntuales': puntuales, 'Tardanzas': tardes})

                    if chart_data:
                        df_chart = pd.DataFrame(chart_data)
                        import altair as alt
                        df_melt = df_chart.melt(id_vars='Docente', value_vars=['Puntuales', 'Tardanzas'],
                                                 var_name='Tipo', value_name='Días')
                        chart = alt.Chart(df_melt).mark_bar().encode(
                            x=alt.X('Docente:N', sort='-y', title=''),
                            y=alt.Y('Días:Q', title='Días'),
                            color=alt.Color('Tipo:N', scale=alt.Scale(
                                domain=['Puntuales', 'Tardanzas'],
                                range=['#22c55e', '#f59e0b']
                            )),
                            xOffset='Tipo:N'
                        ).properties(height=300, title=f'Asistencia y Puntualidad — {mes_nombre} {anio_sel}')
                        st.altair_chart(chart, use_container_width=True)

            elif modo == "📆 Resumen Mensual":
                mes_sel2 = st.selectbox("📆 Mes:", meses_esc,
                                         format_func=lambda x: x[1],
                                         key="rep_doc_mes2")
                mes_num = mes_sel2[0]
                mes_nombre = mes_sel2[1]
                anio_sel = hora_peru().year

                st.markdown(f"### 📆 Resumen Mensual — {mes_nombre} {anio_sel}")

                data_tabla = []
                for nm in sorted(docentes_asist.keys()):
                    dias_mes = 0
                    puntuales = 0
                    tardanzas = 0
                    entradas = []
                    for fecha_str, reg in docentes_asist[nm].items():
                        try:
                            fd = datetime.strptime(fecha_str, '%Y-%m-%d').date()
                            if fd.month == mes_num and fd.year == anio_sel:
                                if reg.get('entrada'):
                                    dias_mes += 1
                                    entradas.append(reg['entrada'])
                                    if reg.get('tardanza'):
                                        tardanzas += 1
                                    else:
                                        puntuales += 1
                        except Exception:
                            pass
                    puntualidad = round(puntuales / max(dias_mes, 1) * 100)
                    data_tabla.append({
                        'Docente': nm, 'Días': dias_mes,
                        'Puntuales': puntuales, 'Tardanzas': tardanzas,
                        'Puntualidad': f"{puntualidad}%",
                        'Hora Prom.': sorted(entradas)[len(entradas)//2] if entradas else '—',
                    })

                if data_tabla:
                    df_mes = pd.DataFrame(data_tabla).sort_values('Días', ascending=False)
                    st.dataframe(df_mes, use_container_width=True, hide_index=True)

                    # Métricas
                    col_m1, col_m2, col_m3, col_m4 = st.columns(4)
                    col_m1.metric("👨‍🏫 Docentes", len(data_tabla))
                    col_m2.metric("📅 Prom. días", f"{sum(r['Días'] for r in data_tabla)/max(len(data_tabla),1):.1f}")
                    total_punt = sum(r['Puntuales'] for r in data_tabla)
                    total_tard = sum(r['Tardanzas'] for r in data_tabla)
                    col_m3.metric("✅ Puntuales", total_punt)
                    col_m4.metric("⏰ Tardanzas", total_tard)

                    # Gráfico
                    st.markdown("---")
                    chart_data = []
                    for r in data_tabla:
                        nm_c = r['Docente'].split()[-1] if ' ' in r['Docente'] else r['Docente']
                        if len(nm_c) > 12:
                            nm_c = nm_c[:10] + ".."
                        chart_data.append({'Docente': nm_c, 'Puntuales': r['Puntuales'], 'Tardanzas': r['Tardanzas']})

                    df_chart = pd.DataFrame(chart_data)
                    import altair as alt
                    df_melt = df_chart.melt(id_vars='Docente', value_vars=['Puntuales', 'Tardanzas'],
                                             var_name='Tipo', value_name='Días')
                    chart = alt.Chart(df_melt).mark_bar().encode(
                        x=alt.X('Docente:N', sort='-y', title=''),
                        y=alt.Y('Días:Q', title='Días'),
                        color=alt.Color('Tipo:N', scale=alt.Scale(
                            domain=['Puntuales', 'Tardanzas'],
                            range=['#22c55e', '#f59e0b']
                        )),
                        xOffset='Tipo:N'
                    ).properties(height=300, title=f'Asistencia y Puntualidad — {mes_nombre} {anio_sel}')
                    st.altair_chart(chart, use_container_width=True)

                    # PDF
                    if st.button("📥 Descargar PDF", type="primary", key="dl_asist_doc_pdf"):
                        buf = io.BytesIO()
                        c_p = canvas.Canvas(buf, pagesize=landscape(A4))
                        wp, hp = landscape(A4)
                        # Encabezado oficial
                        c_p.setFillColor(colors.HexColor("#001e7c"))
                        c_p.rect(0, hp - 55, wp, 55, fill=1, stroke=0)
                        c_p.setFillColor(colors.white)
                        c_p.setFont("Helvetica-Bold", 16)
                        c_p.drawCentredString(wp/2, hp-22, "I.E.P. ALTERNATIVO YACHAY")
                        c_p.setFont("Helvetica", 10)
                        c_p.drawCentredString(wp/2, hp-38,
                                              f"REPORTE DE ASISTENCIA DOCENTES — {mes_nombre.upper()} {anio_sel}")
                        c_p.setFillColor(colors.HexColor("#FFD700"))
                        c_p.setFont("Helvetica", 8)
                        c_p.drawCentredString(wp/2, hp-50, "Chinchero, Cusco — Perú")
                        c_p.setFillColor(colors.black)
                        # Tabla
                        headers = ["N°", "DOCENTE", "DÍAS", "PUNTUALES", "TARDANZAS", "PUNTUALIDAD", "HORA PROM."]
                        rows = [headers]
                        for i, r in enumerate(data_tabla):
                            rows.append([str(i+1), r['Docente'], str(r['Días']),
                                        str(r['Puntuales']), str(r['Tardanzas']),
                                        r['Puntualidad'], r['Hora Prom.']])
                        t_doc = Table(rows, colWidths=[25, 200, 50, 65, 65, 70, 65])
                        t_doc.setStyle(TableStyle([
                            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                            ('FONTSIZE', (0,0), (-1,-1), 7),
                            ('GRID', (0,0), (-1,-1), 0.5, colors.black),
                            ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#001e7c")),
                            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
                            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                            ('ALIGN', (1,1), (1,-1), 'LEFT'),
                            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.Color(0.95,0.95,1.0)]),
                        ]))
                        tw2, th2 = t_doc.wrap(wp-40, hp-80)
                        t_doc.drawOn(c_p, 20, hp-70-th2)
                        # Pie
                        c_p.setFont("Helvetica", 6)
                        c_p.drawString(20, 12, f"Generado: {fecha_peru_str()} | Puntualidad: entrada antes de 8:05am")
                        c_p.drawRightString(wp-20, 12, "Sistema YACHAY PRO")
                        c_p.save()
                        buf.seek(0)
                        st.download_button("⬇️ PDF", buf, f"Asistencia_Docentes_{mes_nombre}.pdf",
                                           "application/pdf", key="dl_doc_pdf2")
                else:
                    st.info("Sin datos para este mes.")

            elif modo == "✏️ Editar Registros":
                # ── Editar registros de docentes ─────────────────────────
                st.markdown("### ✏️ Editar Registro de Asistencia — Docente")
                st.caption("Para docentes con actividades especiales que justifican la hora de llegada")

                docente_edit = st.selectbox("👨‍🏫 Docente:",
                                            sorted(docentes_asist.keys()),
                                            key="edit_doc_sel")
                fecha_edit = st.date_input("📅 Fecha:", value=hora_peru().date(),
                                            key="edit_fecha")
                fecha_str_e = fecha_edit.strftime('%Y-%m-%d')

                reg_actual = docentes_asist.get(docente_edit, {}).get(fecha_str_e, {})
                st.info(f"Registro actual: Entrada: **{reg_actual.get('entrada', '—')}** | "
                        f"Salida: **{reg_actual.get('salida', '—')}** | "
                        f"Tardanza: **{'Sí' if reg_actual.get('tardanza') else 'No'}**")

                col_e1, col_e2, col_e3 = st.columns(3)
                with col_e1:
                    nueva_entrada = st.text_input("🕒 Entrada:", value=reg_actual.get('entrada', ''),
                                                   placeholder="07:45", key="edit_ent")
                with col_e2:
                    nueva_salida = st.text_input("🕒 Salida:", value=reg_actual.get('salida', ''),
                                                  placeholder="13:30", key="edit_sal")
                with col_e3:
                    quitar_tard = st.checkbox("✅ Marcar como PUNTUAL (quitar tardanza)",
                                              key="edit_puntual")

                motivo = st.text_input("📝 Motivo de la modificación:",
                                        placeholder="Ej: Actividad extracurricular autorizada",
                                        key="edit_motivo")

                if st.button("💾 GUARDAR CAMBIOS", type="primary", key="btn_edit_doc"):
                    if motivo:
                        try:
                            asist_data = {}
                            if Path(ARCHIVO_ASISTENCIAS).exists():
                                with open(ARCHIVO_ASISTENCIAS, 'r', encoding='utf-8') as f:
                                    asist_data = json.load(f)
                            # Buscar DNI del docente
                            dni_edit = None
                            if fecha_str_e in asist_data:
                                for dk, dv in asist_data[fecha_str_e].items():
                                    if dv.get('nombre', '').strip().upper() == docente_edit.strip().upper():
                                        dni_edit = dk
                                        break
                            if dni_edit and fecha_str_e in asist_data and dni_edit in asist_data[fecha_str_e]:
                                if nueva_entrada:
                                    asist_data[fecha_str_e][dni_edit]['entrada'] = nueva_entrada
                                if nueva_salida:
                                    asist_data[fecha_str_e][dni_edit]['salida'] = nueva_salida
                                if quitar_tard:
                                    asist_data[fecha_str_e][dni_edit]['tardanza'] = ''
                                    if not asist_data[fecha_str_e][dni_edit].get('entrada'):
                                        asist_data[fecha_str_e][dni_edit]['entrada'] = nueva_entrada or '07:30'
                                asist_data[fecha_str_e][dni_edit]['modificado'] = motivo
                                asist_data[fecha_str_e][dni_edit]['modificado_por'] = st.session_state.get('usuario_actual', '')
                                with open(ARCHIVO_ASISTENCIAS, 'w', encoding='utf-8') as f:
                                    json.dump(asist_data, f, indent=2, ensure_ascii=False)
                                st.success(f"✅ Registro de {docente_edit} modificado — {fecha_str_e}")
                                st.rerun()
                            else:
                                st.warning("No se encontró registro para esa fecha. Puede crear uno nuevo.")
                                # Crear registro
                                if fecha_str_e not in asist_data:
                                    asist_data[fecha_str_e] = {}
                                # Buscar DNI en docentes
                                df_doc_e = BaseDatos.cargar_docentes()
                                if not df_doc_e.empty:
                                    fd = df_doc_e[df_doc_e['Nombre'].astype(str).str.upper() == docente_edit.upper()]
                                    if not fd.empty:
                                        dn = str(fd.iloc[0].get('DNI', f'edit_{docente_edit[:10]}'))
                                        asist_data[fecha_str_e][dn] = {
                                            'nombre': docente_edit, 'entrada': nueva_entrada or '07:30',
                                            'salida': nueva_salida, 'tardanza': '' if quitar_tard else '',
                                            'es_docente': True, 'modificado': motivo,
                                            'entrada_tarde': '', 'salida_tarde': '',
                                        }
                                        with open(ARCHIVO_ASISTENCIAS, 'w', encoding='utf-8') as f:
                                            json.dump(asist_data, f, indent=2, ensure_ascii=False)
                                        st.success(f"✅ Registro creado para {docente_edit} — {fecha_str_e}")
                                        st.rerun()
                        except Exception as e:
                            st.error(f"❌ Error: {e}")
                    else:
                        st.warning("⚠️ Debe ingresar el motivo de la modificación.")

            elif modo == "⏱️ Horas Sec/PreU":
                # ── Horas de trabajo Sec/PreU ────────────────────────────
                st.markdown("### ⏱️ Control de Horas — Docentes Secundaria / PreUniversitario")
                st.caption("Docentes que trabajan por horas: registro de horas por día, semana y mes")

                mes_h = st.selectbox("📆 Mes:", meses_esc,
                                      format_func=lambda x: x[1],
                                      key="horas_mes")
                mes_num_h = mes_h[0]
                mes_nombre_h = mes_h[1]
                anio_h = hora_peru().year

                # Calcular horas por docente
                horas_data = []
                for nm in sorted(docentes_asist.keys()):
                    horas_dia = {}
                    total_horas_mes = 0.0
                    for fecha_str, reg in docentes_asist[nm].items():
                        try:
                            fd = datetime.strptime(fecha_str, '%Y-%m-%d').date()
                            if fd.month == mes_num_h and fd.year == anio_h:
                                ent = reg.get('entrada', '')
                                sal = reg.get('salida', '')
                                horas_d = 0.0
                                if ent and sal:
                                    try:
                                        h1, m1 = ent.split(':')[:2]
                                        h2, m2 = sal.split(':')[:2]
                                        min1 = int(h1) * 60 + int(m1)
                                        min2 = int(h2) * 60 + int(m2)
                                        horas_d = max(0, (min2 - min1) / 60)
                                    except Exception:
                                        pass
                                # Sumar turno tarde si existe
                                ent_t = reg.get('entrada_tarde', '')
                                sal_t = reg.get('salida_tarde', '')
                                if ent_t and sal_t:
                                    try:
                                        h1t, m1t = ent_t.split(':')[:2]
                                        h2t, m2t = sal_t.split(':')[:2]
                                        min1t = int(h1t) * 60 + int(m1t)
                                        min2t = int(h2t) * 60 + int(m2t)
                                        horas_d += max(0, (min2t - min1t) / 60)
                                    except Exception:
                                        pass
                                horas_dia[fecha_str] = round(horas_d, 1)
                                total_horas_mes += horas_d
                        except Exception:
                            pass

                    if horas_dia:
                        # Calcular por semana
                        semanas_h = _semanas_del_mes(mes_num_h, anio_h)
                        horas_por_sem = []
                        for sem_n, lun_s, vie_s in semanas_h:
                            h_sem = 0.0
                            for d in range(5):
                                dia_s = lun_s + timedelta(days=d)
                                h_sem += horas_dia.get(dia_s.strftime('%Y-%m-%d'), 0.0)
                            horas_por_sem.append(round(h_sem, 1))

                        horas_data.append({
                            'Docente': nm,
                            'Días': len(horas_dia),
                            'Horas Total': round(total_horas_mes, 1),
                            'Prom/Día': round(total_horas_mes / max(len(horas_dia), 1), 1),
                            **{f'Sem {s[0]}': h for s, h in zip(semanas_h, horas_por_sem)},
                        })

                if horas_data:
                    st.markdown(f"### {mes_nombre_h} {anio_h}")
                    df_h = pd.DataFrame(horas_data).sort_values('Horas Total', ascending=False)
                    st.dataframe(df_h, use_container_width=True, hide_index=True)

                    # Métricas
                    ch1, ch2, ch3 = st.columns(3)
                    ch1.metric("👨‍🏫 Docentes", len(horas_data))
                    ch2.metric("⏱️ Prom. horas/mes", f"{sum(r['Horas Total'] for r in horas_data)/max(len(horas_data),1):.1f}h")
                    ch3.metric("⏱️ Total horas", f"{sum(r['Horas Total'] for r in horas_data):.0f}h")

                    # Gráfico
                    import altair as alt
                    chart_h = []
                    for r in horas_data:
                        nm_c = r['Docente'].split()[-1] if ' ' in r['Docente'] else r['Docente']
                        if len(nm_c) > 12:
                            nm_c = nm_c[:10] + ".."
                        chart_h.append({'Docente': nm_c, 'Horas': r['Horas Total']})
                    df_ch = pd.DataFrame(chart_h)
                    bar = alt.Chart(df_ch).mark_bar(color='#3b82f6').encode(
                        x=alt.X('Docente:N', sort='-y', title=''),
                        y=alt.Y('Horas:Q', title='Horas trabajadas'),
                    ).properties(height=280, title=f'Horas Trabajadas — {mes_nombre_h} {anio_h}')
                    st.altair_chart(bar, use_container_width=True)
                else:
                    st.info("Sin datos de horas para este mes. Se necesita entrada Y salida registradas.")

            elif modo == "📱 WhatsApp Docentes":
                # ── WhatsApp Docentes ──────────────────────────────────────
                st.markdown("### 📱 Enviar WhatsApp a Docentes")
                st.caption("Enviar mensajes de asistencia o comunicados a docentes")

                if not df_doc_list.empty and 'Celular' in df_doc_list.columns:
                    msg_tipo = st.selectbox("Tipo de mensaje:", [
                        "Recordatorio de asistencia",
                        "Felicitación por puntualidad",
                        "Comunicado general",
                        "Mensaje personalizado"
                    ], key="wa_doc_tipo")

                    msg_custom = ""
                    if msg_tipo == "Mensaje personalizado":
                        msg_custom = st.text_area("Mensaje:", placeholder="Escriba su mensaje...",
                                                   key="wa_doc_msg")

                    st.markdown("---")
                    enviados_count = 0
                    for _, row in df_doc_list.iterrows():
                        nm = str(row.get('Nombre', '')).strip()
                        cel = str(row.get('Celular', '')).strip()
                        if cel and cel not in ('nan', 'None', ''):
                            if '.' in cel:
                                cel = cel.split('.')[0]
                            cel = ''.join(c for c in cel if c.isdigit())
                            if len(cel) >= 7:
                                if msg_tipo == "Recordatorio de asistencia":
                                    msg = (f"Buenos días Prof. {nm}. Le recordamos registrar "
                                           f"su asistencia en el sistema YACHAY. Gracias.")
                                elif msg_tipo == "Felicitación por puntualidad":
                                    msg = (f"Estimado(a) Prof. {nm}, felicitamos su puntualidad "
                                           f"y compromiso con la I.E.P. YACHAY. ¡Siga así!")
                                elif msg_tipo == "Comunicado general":
                                    msg = (f"Estimado(a) Prof. {nm}, se le comunica que "
                                           f"hay una reunión de coordinación. "
                                           f"Por favor revise el sistema YACHAY. Gracias.")
                                else:
                                    msg = f"Prof. {nm}: {msg_custom}" if msg_custom else f"Mensaje para {nm}"

                                link = generar_link_whatsapp(cel, msg)
                                st.markdown(
                                    f'<a href="{link}" target="_blank" class="wa-btn">'
                                    f'📱 👨‍🏫 {nm} → {cel}</a>',
                                    unsafe_allow_html=True)
                                enviados_count += 1
                    if enviados_count == 0:
                        st.warning("No hay docentes con celular registrado.")
                    else:
                        st.success(f"📱 {enviados_count} docentes con WhatsApp disponible")
                else:
                    st.warning("No hay docentes registrados o sin campo Celular.")

    if subtab == "📋 Asistencia Mensual":
        st.markdown("### 📋 Reporte Mensual de Asistencia por Grado")
        if gs:
            c1, c2, c3 = st.columns(3)
            with c1:
                grado_rep = st.selectbox("Grado:", GRADOS_OPCIONES, key="rep_gr")
            with c2:
                mes_rep = st.selectbox("Mes:", list(range(1, 13)),
                                        format_func=lambda x: ['', 'Enero', 'Febrero',
                                        'Marzo', 'Abril', 'Mayo', 'Junio', 'Julio',
                                        'Agosto', 'Septiembre', 'Octubre', 'Noviembre',
                                        'Diciembre'][x], key="rep_mes")
            with c3:
                anio_rep = st.number_input("Año:", value=hora_peru().year,
                                            key="rep_anio")

            if st.button("📊 Generar Reporte", type="primary", key="btn_rep_asist"):
                datos = gs.reporte_asistencia_mensual(grado_rep, mes_rep, int(anio_rep))
                if datos:
                    st.success(f"✅ {len(datos)} estudiantes encontrados")
                    for nombre, info_a in sorted(datos.items()):
                        total = len(info_a.get('fechas', {}))
                        st.write(f"**{nombre}** — {total} días asistidos")
                    pdf = generar_reporte_asistencia_mensual_pdf(
                        datos, grado_rep, mes_rep, int(anio_rep), config)
                    st.download_button("📥 PDF Asistencia Mensual", pdf,
                                       f"Asistencia_{grado_rep}_{mes_rep}.pdf",
                                       "application/pdf", key="dl_rep_asist")
                else:
                    st.warning("No hay datos para este período")
        else:
            st.warning("⚠️ Conecta Google Sheets")

    elif subtab == "📊 Reporte Integral":
        st.markdown("### 📊 Reporte Integral del Estudiante")
        st.markdown("*Incluye: Notas + Asistencia + Semáforo + Recomendaciones*")

        rc1, rc2 = st.columns(2)
        with rc1:
            grado_ri = st.selectbox("Grado:", GRADOS_OPCIONES, key="ri_grado")
        with rc2:
            modo_ri = st.radio("Generar:", ["Un estudiante", "Todo el grado"],
                               horizontal=True, key="ri_modo")

        dg = BaseDatos.obtener_estudiantes_grado(grado_ri)
        if dg.empty:
            st.warning("Sin estudiantes en este grado")
            return

        if modo_ri == "Un estudiante":
            opciones = [f"{r['Nombre']} — {r['DNI']}" for _, r in dg.iterrows()]
            sel = st.selectbox("Estudiante:", opciones, key="ri_est")
            if sel:
                dni_ri = sel.split(" — ")[-1].strip()
                nombre_ri = sel.split(" — ")[0].strip()
                
                # EDITAR NOTAS (Solo Director/Admin/Promotor)
                if st.session_state.rol in ['admin', 'directivo']:
                    st.markdown("---")
                    with st.expander("✏️ Editar Notas del Estudiante", expanded=False):
                        st.caption("🔐 Solo Director, Administrador y Promotor")
                        
                        notas_edit = []
                        if gs:
                            try:
                                ws = gs._get_hoja('config')
                                if ws:
                                    data = ws.get_all_records()
                                    for d in data:
                                        clave = str(d.get('clave', ''))
                                        if clave.startswith(f'nota_{dni_ri}'):
                                            try:
                                                nota_data = json.loads(d.get('valor', '{}'))
                                                nota_data['_clave'] = clave
                                                notas_edit.append(nota_data)
                                            except Exception:
                                                pass
                            except Exception:
                                pass
                        
                        if notas_edit:
                            for idx, nota in enumerate(notas_edit):
                                col1, col2, col3 = st.columns([3, 1, 1])
                                with col1:
                                    st.text(f"{nota.get('area', 'N/A')} - {nota.get('bimestre', 'N/A')}")
                                with col2:
                                    nueva = st.number_input("Nota:", 0.0, 20.0, 
                                                          float(nota.get('nota', 0)), 0.5,
                                                          key=f"ne_{idx}_{dni_ri}")
                                with col3:
                                    if st.button("💾", key=f"sv_{idx}_{dni_ri}", type="primary"):
                                        if gs and '_clave' in nota:
                                            try:
                                                nota['nota'] = nueva
                                                nota['literal'] = nota_a_letra(nueva)
                                                ws = gs._get_hoja('config')
                                                if ws:
                                                    all_data = ws.get_all_records()
                                                    for row_idx, row in enumerate(all_data, start=2):
                                                        if str(row.get('clave', '')) == nota['_clave']:
                                                            nota_copy = nota.copy()
                                                            nota_copy.pop('_clave', None)
                                                            ws.update_cell(row_idx, 2, 
                                                                         json.dumps(nota_copy, ensure_ascii=False))
                                                            break
                                                st.success(f"✅ Actualizado: {nueva}/20")
                                                time.sleep(1)
                                                st.rerun()
                                            except Exception as e:
                                                st.error(f"Error: {str(e)[:50]}")
                        else:
                            st.info("Sin notas para editar")
                    st.markdown("---")
        else:
            dni_ri = None
            nombre_ri = None

        if st.button("📥 GENERAR REPORTE INTEGRAL", type="primary",
                     use_container_width=True, key="btn_ri"):
            with st.spinner("Generando reporte..."):
                if modo_ri == "Un estudiante" and dni_ri:
                    # Cargar notas del estudiante
                    notas_est = []
                    asist_est = {}

                    # Cargar asistencia LOCAL primero (fuente principal)
                    try:
                        if Path(ARCHIVO_ASISTENCIAS).exists():
                            with open(ARCHIVO_ASISTENCIAS, 'r', encoding='utf-8') as f:
                                todas_asis = json.load(f)
                            for fecha_a, registros in todas_asis.items():
                                if str(dni_ri) in registros:
                                    asist_est[fecha_a] = registros[str(dni_ri)]
                    except Exception:
                        pass

                    # SIEMPRE complementar con GS (puede tener datos que local perdió)
                    if gs:
                        try:
                            asist_gs = gs.historial_asistencia_estudiante(dni_ri)
                            if asist_gs:
                                for fecha_g, datos_g in asist_gs.items():
                                    if fecha_g not in asist_est:
                                        asist_est[fecha_g] = datos_g
                                    else:
                                        # Completar campos vacíos
                                        if not asist_est[fecha_g].get('entrada') and datos_g.get('entrada'):
                                            asist_est[fecha_g]['entrada'] = datos_g['entrada']
                                        if not asist_est[fecha_g].get('salida') and datos_g.get('salida'):
                                            asist_est[fecha_g]['salida'] = datos_g['salida']
                        except Exception:
                            pass

                    if gs:
                        try:
                            ws = gs._get_hoja('config')
                            if ws:
                                data = ws.get_all_records()
                                for d in data:
                                    clave = str(d.get('clave', ''))
                                    if clave.startswith(f'nota_{dni_ri}'):
                                        try:
                                            notas_est.append(json.loads(d.get('valor', '{}')))
                                        except Exception:
                                            pass
                        except Exception:
                            pass

                    # También cargar de resultados de examen
                    _anio_rep = str(config.get('anio', 2026))
                    all_res = BaseDatos.cargar_todos_resultados()
                    for r in all_res:
                        if str(r.get('dni', '')) == str(dni_ri):
                            _fecha_r = str(r.get('fecha', ''))
                            if not _fecha_r.startswith(_anio_rep):
                                continue
                            for area in r.get('areas', []):
                                notas_est.append({
                                    'area': area['nombre'],
                                    'nota': area['nota'],
                                    'literal': nota_a_letra(area['nota']),
                                    'bimestre': r.get('titulo', 'Evaluación'),
                                    'fecha': r.get('fecha', ''),
                                    'tipo': 'examen'
                                })

                    # Cargar notas del historial de evaluaciones (Registrar Notas)
                    hist_eval = _cargar_historial_evaluaciones()
                    for clave_h, ev_h in hist_eval.items():
                        _fecha_h = str(ev_h.get('fecha', ''))
                        if not _fecha_h.startswith(_anio_rep):
                            continue
                        for fila_h in ev_h.get('ranking', []):
                            if str(fila_h.get('DNI', '')) == str(dni_ri):
                                areas_h = ev_h.get('areas', [])
                                areas_nombres_h = [a['nombre'] for a in areas_h] if areas_h and isinstance(areas_h[0], dict) else list(areas_h)
                                for a_n in areas_nombres_h:
                                    nota_v = fila_h.get(a_n, 0)
                                    if nota_v and float(nota_v) > 0:
                                        notas_est.append({
                                            'area': a_n,
                                            'nota': float(nota_v),
                                            'literal': nota_a_letra(float(nota_v)),
                                            'bimestre': ev_h.get('periodo', ''),
                                            'fecha': ev_h.get('fecha', ''),
                                            'titulo': ev_h.get('titulo', ''),
                                            'tipo': 'registro_notas'
                                        })

                    al = BaseDatos.buscar_por_dni(dni_ri)
                    grado_est = str(al.get('Grado', grado_ri)) if al else grado_ri

                    pdf = generar_reporte_integral_pdf(
                        nombre_ri, dni_ri, grado_est, notas_est, asist_est, config)
                    st.download_button("⬇️ Descargar PDF", pdf,
                                       f"Reporte_{nombre_ri.replace(' ', '_')}.pdf",
                                       "application/pdf", key="dl_ri")
                    st.success(f"✅ Reporte de {nombre_ri} generado")

                else:
                    # Todo el grado - un PDF multi-página
                    buf_all = io.BytesIO()
                    c_pdf = canvas.Canvas(buf_all, pagesize=A4)
                    w_page, h_page = A4

                    for _, row in dg.iterrows():
                        n_est = str(row.get('Nombre', ''))
                        d_est = str(row.get('DNI', ''))

                        # Cargar notas
                        notas_est = []
                        if gs:
                            try:
                                ws = gs._get_hoja('config')
                                if ws:
                                    data = ws.get_all_records()
                                    for d in data:
                                        clave = str(d.get('clave', ''))
                                        if clave.startswith(f'nota_{d_est}'):
                                            try:
                                                notas_est.append(json.loads(d.get('valor', '{}')))
                                            except Exception:
                                                pass
                            except Exception:
                                pass

                        # De exámenes también
                        _anio_rep_g = str(config.get('anio', 2026))
                        all_res = BaseDatos.cargar_todos_resultados()
                        for r in all_res:
                            if str(r.get('dni', '')) == d_est:
                                if not str(r.get('fecha', '')).startswith(_anio_rep_g):
                                    continue
                                for area in r.get('areas', []):
                                    notas_est.append({
                                        'area': area['nombre'],
                                        'nota': area['nota'],
                                        'fecha': r.get('fecha', ''),
                                    })

                        # De historial de evaluaciones (Registrar Notas)
                        hist_eval_g = _cargar_historial_evaluaciones()
                        for clave_hg, ev_hg in hist_eval_g.items():
                            if not str(ev_hg.get('fecha', '')).startswith(_anio_rep_g):
                                continue
                            for fila_hg in ev_hg.get('ranking', []):
                                if str(fila_hg.get('DNI', '')) == str(d_est):
                                    areas_hg = ev_hg.get('areas', [])
                                    areas_noms_hg = [a['nombre'] for a in areas_hg] if areas_hg and isinstance(areas_hg[0], dict) else list(areas_hg)
                                    for a_ng in areas_noms_hg:
                                        nota_vg = fila_hg.get(a_ng, 0)
                                        if nota_vg and float(nota_vg) > 0:
                                            notas_est.append({
                                                'area': a_ng,
                                                'nota': float(nota_vg),
                                                'fecha': ev_hg.get('fecha', ''),
                                            })

                        # Página del estudiante
                        c_pdf.setFont("Helvetica-Bold", 14)
                        c_pdf.drawCentredString(w_page/2, h_page-40, f"REPORTE: {n_est}")
                        c_pdf.setFont("Helvetica", 10)
                        y = h_page-65
                        c_pdf.drawString(50, y, f"DNI: {d_est} | Grado: {grado_ri}")
                        y -= 25

                        if notas_est:
                            for n in notas_est:
                                nota_v = float(n.get('nota', 0))
                                lit = nota_a_letra(nota_v)
                                c_pdf.drawString(60, y,
                                    f"• {n.get('area', '')}: {nota_v}/20 ({lit}) — {n.get('fecha', '')}")
                                y -= 14
                                if y < 80:
                                    break

                            # Promedio
                            promedios = [float(n.get('nota', 0)) for n in notas_est if float(n.get('nota', 0)) > 0]
                            if promedios:
                                prom = round(sum(promedios)/len(promedios), 1)
                                lit_p = nota_a_letra(prom)
                                y -= 10
                                c_pdf.setFont("Helvetica-Bold", 12)
                                c_pdf.drawString(60, y, f"PROMEDIO: {prom}/20 ({lit_p})")
                        else:
                            c_pdf.drawString(60, y, "Sin calificaciones registradas.")

                        c_pdf.showPage()

                    c_pdf.save()
                    buf_all.seek(0)
                    st.download_button("⬇️ Reportes Todo el Grado", buf_all,
                                       f"Reportes_{grado_ri}.pdf",
                                       "application/pdf", key="dl_ri_all")
                    st.success(f"✅ Reportes de {len(dg)} estudiantes generados")

    elif subtab == "📄 Reporte ZipGrade":
        st.markdown("### 📄 Reporte estilo ZipGrade")
        usuario = st.session_state.get('usuario_actual', '')
        resultados = BaseDatos.cargar_resultados_examen(usuario)
        if st.session_state.rol in ['admin', 'directivo']:
            resultados = BaseDatos.cargar_todos_resultados()

        if resultados:
            opciones_eval = [
                f"{r.get('nombre', '?')} — {r.get('fecha', '')}"
                for r in resultados
            ]
            sel_eval = st.selectbox("Evaluación:", opciones_eval, key="zg_eval")
            idx_eval = opciones_eval.index(sel_eval)
            eval_sel = resultados[idx_eval]

            # Mostrar detalles
            for area in eval_sel.get('areas', []):
                nota = area.get('nota', 0)
                lit = nota_a_letra(nota)
                col = color_semaforo(lit)
                st.markdown(f"**{area['nombre']}:** <span style='color:{col};'>{nota}/20 ({lit})</span>",
                           unsafe_allow_html=True)

            if st.button("📥 PDF ZipGrade", type="primary", key="btn_zg"):
                pdf = generar_reporte_examen_zipgrade(eval_sel, config)
                st.download_button("⬇️ PDF", pdf,
                                   f"ZipGrade_{sel_eval[:20]}.pdf",
                                   "application/pdf", key="dl_zg")
        else:
            st.info("Sin evaluaciones. Califica exámenes primero.")

    elif subtab == "📁 Fichas Docentes":
        st.markdown("### 📁 Fichas de Trabajo — Registro por Docente")
        st.caption("Vista consolidada de todas las fichas subidas por los docentes")

        fichas = _cargar_fichas_registro()
        if not fichas:
            st.info("📭 No hay fichas registradas aún. Los docentes pueden subirlas desde 'Registrar Ficha'.")
            return

        # ── Filtros ──
        fc1, fc2, fc3, fc4 = st.columns(4)
        with fc1:
            grados_f = sorted(set(f.get('grado', '') for f in fichas if f.get('grado')))
            filtro_grado = st.selectbox("🎓 Grado:", ["Todos"] + grados_f, key="fichas_f_grado")
        with fc2:
            meses_f = sorted(set(f.get('mes', '') for f in fichas if f.get('mes')))
            filtro_mes = st.selectbox("📅 Mes:", ["Todos"] + meses_f, key="fichas_f_mes")
        with fc3:
            semanas_f = sorted(set(str(f.get('semana', '')) for f in fichas if f.get('semana')))
            filtro_semana = st.selectbox("📆 Semana:", ["Todas"] + semanas_f, key="fichas_f_sem")
        with fc4:
            docentes_f = sorted(set(f.get('docente_nombre', f.get('docente', ''))
                                    for f in fichas if f.get('docente')))
            filtro_doc = st.selectbox("👨‍🏫 Docente:", ["Todos"] + docentes_f, key="fichas_f_doc")

        # ── Aplicar filtros ──
        fichas_filtradas = fichas.copy()
        if filtro_grado != "Todos":
            fichas_filtradas = [f for f in fichas_filtradas if f.get('grado') == filtro_grado]
        if filtro_mes != "Todos":
            fichas_filtradas = [f for f in fichas_filtradas if f.get('mes') == filtro_mes]
        if filtro_semana != "Todas":
            fichas_filtradas = [f for f in fichas_filtradas if str(f.get('semana')) == filtro_semana]
        if filtro_doc != "Todos":
            fichas_filtradas = [f for f in fichas_filtradas
                                if f.get('docente_nombre') == filtro_doc or f.get('docente') == filtro_doc]

        # ── Métricas ──
        mc1, mc2, mc3, mc4 = st.columns(4)
        with mc1:
            st.metric("📄 Total fichas", len(fichas_filtradas))
        with mc2:
            docs_unicos = len(set(f.get('docente', '') for f in fichas_filtradas))
            st.metric("👨‍🏫 Docentes", docs_unicos)
        with mc3:
            grados_unicos = len(set(f.get('grado', '') for f in fichas_filtradas))
            st.metric("🎓 Grados", grados_unicos)
        with mc4:
            areas_unicas = len(set(f.get('area', '') for f in fichas_filtradas))
            st.metric("📚 Áreas", areas_unicas)

        st.markdown("---")

        # ── Tabla de fichas ──
        if fichas_filtradas:
            for ficha in sorted(fichas_filtradas,
                                key=lambda x: x.get('fecha_subida', ''), reverse=True):
                doc_nombre = ficha.get('docente_nombre', ficha.get('docente', ''))
                titulo_f = ficha.get('titulo', 'Sin título')
                area_f = ficha.get('area', '')
                grado_f = ficha.get('grado', '')
                sem_f = ficha.get('semana', '')
                mes_f = ficha.get('mes', '')
                fecha_f = ficha.get('fecha_subida', '')
                tipo_f = ficha.get('tipo', '')
                archivo_f = ficha.get('archivo', '')

                with st.expander(f"📄 **{titulo_f}** — {area_f} | {grado_f} | Sem {sem_f} | {mes_f}"):
                    ec1, ec2 = st.columns([3, 1])
                    with ec1:
                        st.markdown(f"**Docente:** {doc_nombre}")
                        st.markdown(f"**Área:** {area_f} | **Grado:** {grado_f} | "
                                    f"**Semana:** {sem_f} | **Mes:** {mes_f}")
                        st.caption(f"📅 Subido: {fecha_f} | Tipo: {tipo_f} | Archivo: {archivo_f}")
                    with ec2:
                        # Intentar ofrecer descarga si el archivo existe localmente
                        ruta_ficha = Path("fichas") / archivo_f if archivo_f else None
                        if ruta_ficha and ruta_ficha.exists():
                            with open(ruta_ficha, 'rb') as ff:
                                st.download_button("⬇️ Descargar", ff.read(),
                                                   archivo_f, "application/pdf",
                                                   key=f"dl_rep_{ficha.get('id', '')}")
                        else:
                            st.caption("📁 Solo en servidor")
        else:
            st.info("No hay fichas que coincidan con los filtros seleccionados.")


# ================================================================
# ÁREAS DEL CURRÍCULO NACIONAL MINEDU — Por Nivel
# ================================================================

AREAS_MINEDU = {
    'INICIAL': [
        'Personal Social', 'Psicomotriz', 'Comunicación',
        'Castellano como segunda lengua', 'Matemática',
        'Ciencia y Tecnología', 'Educación Física', 'Inglés'
    ],
    'PRIMARIA': [
        'Personal Social', 'Educación Física', 'Comunicación',
        'Arte y Cultura', 'Castellano como segunda lengua',
        'Inglés', 'Matemática', 'Ciencia y Tecnología',
        'Educación Religiosa'
    ],
    'SECUNDARIA': [
        'Desarrollo Personal, Ciudadanía y Cívica', 'Ciencias Sociales',
        'Educación para el Trabajo', 'Educación Física', 'Comunicación',
        'Arte y Cultura', 'Castellano como segunda lengua', 'Inglés',
        'Matemática', 'Ciencia y Tecnología', 'Educación Religiosa'
    ],
    'PREUNIVERSITARIO': [
        'Razonamiento Matemático', 'Aritmética', 'Álgebra', 'Geometría',
        'Trigonometría', 'Lenguaje', 'Literatura', 'Razonamiento Verbal',
        'Historia del Perú', 'Historia Universal', 'Geografía', 'Economía',
        'Filosofía y Lógica', 'Psicología', 'Educación Cívica',
        'Biología', 'Química', 'Física', 'Anatomía'
    ]
}

PERIODOS_EVALUACION = [
    'Semana 1', 'Semana 2', 'Semana 3', 'Semana 4',
    'Semana 5', 'Semana 6', 'Semana 7', 'Semana 8',
    'Quincenal 1', 'Quincenal 2',
    'I Bimestre', 'II Bimestre', 'III Bimestre', 'IV Bimestre',
    'Evaluación Parcial', 'Evaluación Final', 'Práctica Calificada',
    'Ciclo Verano', 'Ciclo Regular', 'Ciclo Intensivo',
    'Reforzamiento Pre-U',
]
BIMESTRES_LISTA = PERIODOS_EVALUACION  # Alias

# ================================================================
# COMPETENCIAS DEL CURRÍCULO NACIONAL DEL PERÚ (MINEDU)
# ================================================================
COMPETENCIAS_CN = {
    'Personal Social': [
        'Construye su identidad',
        'Convive y participa democráticamente'
    ],
    'Psicomotriz': [
        'Se desenvuelve de manera autónoma a través de su motricidad'
    ],
    'Comunicación': [
        'Se comunica oralmente en su lengua materna',
        'Lee diversos tipos de textos escritos',
        'Escribe diversos tipos de textos'
    ],
    'Matemática': [
        'Resuelve problemas de cantidad',
        'Resuelve problemas de regularidad, equivalencia y cambio',
        'Resuelve problemas de forma, movimiento y localización',
        'Resuelve problemas de gestión de datos e incertidumbre'
    ],
    'Ciencia y Tecnología': [
        'Indaga mediante métodos científicos',
        'Explica el mundo físico basándose en conocimientos científicos',
        'Diseña y construye soluciones tecnológicas'
    ],
    'Educación Física': [
        'Se desenvuelve de manera autónoma a través de su motricidad',
        'Asume una vida saludable',
        'Interactúa a través de sus habilidades sociomotrices'
    ],
    'Inglés': [
        'Se comunica oralmente en inglés como lengua extranjera',
        'Lee diversos tipos de textos en inglés',
        'Escribe diversos tipos de textos en inglés'
    ],
    'Arte y Cultura': [
        'Aprecia de manera crítica manifestaciones artístico-culturales',
        'Crea proyectos desde los lenguajes artísticos'
    ],
    'Educación Religiosa': [
        'Construye su identidad como persona humana, amada por Dios',
        'Asume la experiencia del encuentro personal y comunitario con Dios'
    ],
    'Castellano como segunda lengua': [
        'Se comunica oralmente en castellano como segunda lengua',
        'Lee diversos tipos de textos en castellano como segunda lengua',
        'Escribe diversos tipos de textos en castellano como segunda lengua'
    ],
    'Ciencias Sociales': [
        'Construye interpretaciones históricas',
        'Gestiona responsablemente el espacio y el ambiente',
        'Gestiona responsablemente los recursos económicos'
    ],
    'Desarrollo Personal, Ciudadanía y Cívica': [
        'Construye su identidad',
        'Convive y participa democráticamente'
    ],
    'Educación para el Trabajo': [
        'Gestiona proyectos de emprendimiento económico o social'
    ],
}


def generar_registro_bimestral_pdf(grado, seccion, anio, estudiantes_df,
                                    bimestre, areas_sel, nivel="PRIMARIA",
                                    docente=""):
    """Genera PDF de registro bimestral — hasta 3 áreas por hoja, proporcional."""
    buffer = io.BytesIO()
    c_pdf = canvas.Canvas(buffer, pagesize=landscape(A4))
    w, h = landscape(A4)

    if not estudiantes_df.empty:
        est = estudiantes_df.sort_values('Nombre').reset_index(drop=True)
    else:
        est = pd.DataFrame()

    ne = len(est) if not est.empty else 25

    # Agrupar áreas: hasta 3 por página
    AREAS_POR_PAG = 3
    grupos = []
    for i in range(0, len(areas_sel), AREAS_POR_PAG):
        grupos.append(areas_sel[i:i + AREAS_POR_PAG])

    COLORES_AREA = [
        colors.Color(0.85, 0.92, 1.0),    # celeste
        colors.Color(0.88, 1.0, 0.88),     # verde
        colors.Color(1.0, 0.92, 0.85),     # durazno
        colors.Color(0.92, 0.88, 1.0),     # lavanda
        colors.Color(1.0, 0.90, 0.90),     # rosa
        colors.Color(0.92, 0.98, 0.88),    # lima
        colors.Color(0.90, 0.90, 1.0),     # azul
        colors.Color(1.0, 0.92, 0.88),     # salmon
        colors.Color(0.88, 1.0, 0.95),     # menta
    ]
    COLORES_HDR = [
        colors.Color(0.50, 0.70, 0.90),
        colors.Color(0.50, 0.82, 0.50),
        colors.Color(0.90, 0.70, 0.40),
        colors.Color(0.68, 0.55, 0.85),
        colors.Color(0.88, 0.50, 0.50),
        colors.Color(0.65, 0.82, 0.50),
        colors.Color(0.55, 0.55, 0.88),
        colors.Color(0.88, 0.65, 0.50),
        colors.Color(0.50, 0.85, 0.72),
    ]

    area_idx_global = 0
    for pg_idx, grupo in enumerate(grupos):
        if pg_idx > 0:
            c_pdf.showPage()

        # Encabezado
        c_pdf.setFont("Helvetica-Bold", 11)
        c_pdf.drawCentredString(w / 2, h - 18,
                                "I.E.P. ALTERNATIVO YACHAY — REGISTRO DE NOTAS BIMESTRAL")
        c_pdf.setFont("Helvetica", 7)
        info_line = (f"Grado: {grado} | Sección: {seccion} | {bimestre} | "
                     f"Año: {anio}")
        if docente:
            info_line += f" | Docente: {docente}"
        c_pdf.drawCentredString(w / 2, h - 30, info_line)

        # Calcular todas las competencias de este grupo
        all_comps = []  # [(area_name, comp_name, color_idx)]
        for gi, area in enumerate(grupo):
            comps = COMPETENCIAS_CN.get(area, [f'Competencia {area}'])
            for comp in comps:
                all_comps.append((area, comp, area_idx_global + gi))

        total_comp_cols = len(all_comps)

        # Construir DOBLE header: Fila 0 = ÁREAS (merged), Fila 1 = Competencias
        # Fila 0: N° | NOMBRES | [ÁREA spanning comps] ... | PROM | OBS
        header_areas = ["N°", "APELLIDOS Y NOMBRES"]
        header_comps = ["", ""]
        for gi, area in enumerate(grupo):
            comps = COMPETENCIAS_CN.get(area, [f'Competencia {area}'])
            header_areas.append(area)
            header_comps.append(comps[0][:30] if comps else "")
            for ci in range(1, len(comps)):
                header_areas.append("")  # merged visually
                header_comps.append(comps[ci][:30])
        header_areas.extend(["PROM.", "OBS."])
        header_comps.extend(["", ""])

        # Data rows
        data = [header_areas, header_comps]
        for idx in range(ne):
            nm = est.iloc[idx].get('Nombre', '') if idx < len(est) else ""
            if len(nm) > 28:
                nm = nm[:28] + "."
            fila = [str(idx + 1), nm] + [""] * total_comp_cols + ["", ""]
            data.append(fila)

        # Column widths — proporcional
        avail_w = w - 20  # 10px margins
        num_w = 16
        name_w = 130
        prom_w = 30
        obs_w = 35
        fixed = num_w + name_w + prom_w + obs_w
        comp_w = max(28, (avail_w - fixed) / max(total_comp_cols, 1))
        cw = [num_w, name_w] + [comp_w] * total_comp_cols + [prom_w, obs_w]

        t = Table(data, colWidths=cw, repeatRows=2)

        estilos = [
            ('FONTNAME', (0, 0), (-1, 1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 6),       # Áreas
            ('FONTSIZE', (0, 1), (-1, 1), 4.5),     # Competencias
            ('FONTSIZE', (0, 2), (-1, -1), 6),       # Datos
            ('GRID', (0, 0), (-1, -1), 0.4, colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('ALIGN', (1, 2), (1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ROWHEIGHTS', (0, 0), (-1, 0), 16),    # Fila áreas
            ('ROWHEIGHTS', (0, 1), (-1, 1), 22),    # Fila competencias
            # N° y Nombres header
            ('BACKGROUND', (0, 0), (1, 1), colors.Color(0, 0.30, 0.15)),
            ('TEXTCOLOR', (0, 0), (1, 1), colors.white),
            ('SPAN', (0, 0), (0, 1)),  # N° merge vertical
            ('SPAN', (1, 0), (1, 1)),  # Nombres merge vertical
            # PROM y OBS header
            ('BACKGROUND', (-2, 0), (-1, 1), colors.Color(0.50, 0, 0.15)),
            ('TEXTCOLOR', (-2, 0), (-1, 1), colors.white),
            ('SPAN', (-2, 0), (-2, 1)),  # PROM merge vertical
            ('SPAN', (-1, 0), (-1, 1)),  # OBS merge vertical
        ]

        # Merge ÁREA cells y colorear
        col_offset = 2
        for gi, area in enumerate(grupo):
            comps = COMPETENCIAS_CN.get(area, [f'Competencia {area}'])
            nc = len(comps)
            ci_global = area_idx_global + gi
            c_hdr = COLORES_HDR[ci_global % len(COLORES_HDR)]
            c_data = COLORES_AREA[ci_global % len(COLORES_AREA)]

            # Merge área row
            if nc > 1:
                estilos.append(('SPAN', (col_offset, 0), (col_offset + nc - 1, 0)))
            # Color área header
            estilos.append(('BACKGROUND', (col_offset, 0), (col_offset + nc - 1, 0), c_hdr))
            estilos.append(('TEXTCOLOR', (col_offset, 0), (col_offset + nc - 1, 0), colors.white))
            # Color competencias header
            estilos.append(('BACKGROUND', (col_offset, 1), (col_offset + nc - 1, 1), c_hdr))
            estilos.append(('TEXTCOLOR', (col_offset, 1), (col_offset + nc - 1, 1), colors.white))
            # Data pastel
            estilos.append(('BACKGROUND', (col_offset, 2), (col_offset + nc - 1, -1), c_data))

            col_offset += nc

        t.setStyle(TableStyle(estilos))
        tw, th2 = t.wrap(w - 20, h - 50)
        t.drawOn(c_pdf, 10, h - 40 - th2)

        # Pie
        c_pdf.setFont("Helvetica", 5)
        areas_texto = " | ".join(grupo)
        c_pdf.drawString(10, 8, f"Áreas: {areas_texto}")
        c_pdf.drawRightString(w - 10, 8,
                              f"Currículo Nacional MINEDU — {bimestre} {anio} — "
                              f"Pág {pg_idx + 1}/{len(grupos)}")

        area_idx_global += len(grupo)

    c_pdf.save()
    buffer.seek(0)
    return buffer

# ================================================================
# TAB: REGISTRAR NOTAS (Manual — Para todos los docentes)
# ================================================================

def _sync_resultados_a_gs():
    """Sincroniza resultados.json a Google Sheets"""
    try:
        gs = _gs()
        if not gs:
            return
        if Path('resultados.json').exists():
            with open('resultados.json', 'r', encoding='utf-8') as f:
                data = json.load(f)
            ws = gs._get_hoja('config')
            if ws:
                data_str = json.dumps(data, ensure_ascii=False, default=str)
                all_vals = ws.get_all_values()
                found = False
                for idx, row in enumerate(all_vals):
                    if row and row[0] == 'resultados_json':
                        ws.update_cell(idx + 1, 2, data_str)
                        found = True
                        break
                if not found:
                    ws.append_row(['resultados_json', data_str])
    except Exception:
        pass

def _sync_horario_a_gs():
    """Sincroniza config_horario.json a Google Sheets"""
    try:
        gs = _gs()
        if not gs:
            return
        if Path('config_horario.json').exists():
            with open('config_horario.json', 'r', encoding='utf-8') as f:
                data = json.load(f)
            ws = gs._get_hoja('config')
            if ws:
                data_str = json.dumps(data, ensure_ascii=False, default=str)
                all_vals = ws.get_all_values()
                found = False
                for idx, row in enumerate(all_vals):
                    if row and row[0] == 'config_horario':
                        ws.update_cell(idx + 1, 2, data_str)
                        found = True
                        break
                if not found:
                    ws.append_row(['config_horario', data_str])
    except Exception:
        pass

def _guardar_archivo_binario_gs(nombre_clave, filepath):
    """Guarda un archivo binario (imagen, mp3) como base64 en Google Sheets hoja config"""
    try:
        if not Path(filepath).exists():
            return False
        import base64 as b64mod
        with open(filepath, "rb") as fbin:
            data_b64 = b64mod.b64encode(fbin.read()).decode('utf-8')
        gs = _gs()
        if not gs:
            return False
        ws = gs._get_hoja('config')
        if not ws:
            return False
        all_vals = ws.get_all_values()
        found = False
        for idx, row in enumerate(all_vals):
            if row and row[0] == nombre_clave:
                ws.update_cell(idx + 1, 2, data_b64)
                found = True
                break
        if not found:
            ws.append_row([nombre_clave, data_b64])
        return True
    except Exception:
        return False

def _restaurar_archivo_binario_gs(nombre_clave, filepath):
    """Restaura un archivo binario desde base64 en Google Sheets si no existe localmente"""
    try:
        if Path(filepath).exists():
            return True  # ya existe, no sobreescribir
        import base64 as b64mod
        gs = _gs()
        if not gs:
            return False
        ws = gs._get_hoja('config')
        if not ws:
            return False
        all_vals = ws.get_all_values()
        for row in all_vals:
            if row and row[0] == nombre_clave and len(row) > 1 and row[1]:
                data = b64mod.b64decode(row[1])
                with open(filepath, "wb") as fbin:
                    fbin.write(data)
                return True
        return False
    except Exception:
        return False

def _pausa_guardar_mp3(modelo_id, audio_bytes, extension="mp3"):
    """Guarda MP3 de pausa activa en disco + sincroniza a GSheets como base64"""
    import base64 as b64mod
    path = f"pausa_mp3_{modelo_id}.{extension}"
    with open(path, "wb") as f:
        f.write(audio_bytes)
    try:
        _guardar_archivo_binario_gs(f"bin_pausa_mp3_{modelo_id}", path)
    except Exception:
        pass
    return path

def _pausa_cargar_mp3_b64(modelo_id):
    """Carga MP3 de pausa activa como base64 para reproduccion HTML"""
    import base64 as b64mod
    for ext in ["mp3", "ogg", "wav"]:
        path = f"pausa_mp3_{modelo_id}.{ext}"
        if Path(path).exists():
            with open(path, "rb") as f:
                return b64mod.b64encode(f.read()).decode("utf-8"), ext
    return None, None

def _restaurar_todos_archivos_binarios():
    """Restaura todos los archivos binarios (escudos, fondo, mp3 pausas, mp3 qaway) desde GSheets"""
    archivos_binarios = [
        ("bin_escudo_izq",    "escudo_upload.png"),
        ("bin_escudo_der",    "escudo2_upload.png"),
        ("bin_fondo",         "fondo.png"),
    ]
    # MP3 de Pausa Activa (10 modelos)
    for i in range(1, 21):  # 20 modelos de pausa activa
        archivos_binarios.append((f"bin_pausa_mp3_{i}", f"pausa_mp3_{i}.mp3"))
    # MP3 de YACHAY QAWAY
    archivos_binarios.append(("bin_qaway_mp3", str(_plk_dir() / "musica_fondo.mp3")))
    for nombre_clave, filepath in archivos_binarios:
        try:
            _restaurar_archivo_binario_gs(nombre_clave, filepath)
        except Exception:
            pass

def _restaurar_datos_desde_gs():
    """Restaura archivos JSON locales desde Google Sheets al iniciar"""
    try:
        gs = _gs()
        if not gs:
            return
        ws = gs._get_hoja('config')
        if not ws:
            return
        data = ws.get_all_values()
        restaurados = 0
        for row in data:
            if not row or len(row) < 2:
                continue
            key, val = row[0], row[1]
            try:
                if key == 'historial_evaluaciones' and not Path('historial_evaluaciones.json').exists():
                    with open('historial_evaluaciones.json', 'w', encoding='utf-8') as f:
                        f.write(val)
                    restaurados += 1
                elif key == 'resultados_json' and not Path('resultados.json').exists():
                    with open('resultados.json', 'w', encoding='utf-8') as f:
                        f.write(val)
                    restaurados += 1
                elif key == 'config_horario' and not Path('config_horario.json').exists():
                    with open('config_horario.json', 'w', encoding='utf-8') as f:
                        f.write(val)
                    restaurados += 1
                elif key == 'diagnostico_data' and not Path('diagnostico_data.json').exists():
                    with open('diagnostico_data.json', 'w', encoding='utf-8') as f:
                        f.write(val)
                    restaurados += 1
            except Exception:
                pass
        return restaurados
    except Exception:
        return 0

def _cargar_historial_evaluaciones():
    """Carga el historial de evaluaciones desde archivo JSON"""
    try:
        if Path('historial_evaluaciones.json').exists():
            with open('historial_evaluaciones.json', 'r', encoding='utf-8') as f:
                return json.load(f)
    except Exception:
        pass
    return {}

def _guardar_historial_evaluaciones(hist_data):
    """Guarda el historial de evaluaciones en archivo JSON + Google Sheets"""
    try:
        with open('historial_evaluaciones.json', 'w', encoding='utf-8') as f:
            json.dump(hist_data, f, ensure_ascii=False, indent=2, default=str)
        # Sync a Google Sheets
        try:
            gs = _gs()
            if gs:
                ws = gs._get_hoja('config')
                if ws:
                    data_str = json.dumps(hist_data, ensure_ascii=False, default=str)
                    all_data = ws.get_all_values()
                    found = False
                    for idx, row in enumerate(all_data):
                        if row and row[0] == 'historial_evaluaciones':
                            ws.update_cell(idx + 1, 2, data_str)
                            found = True
                            break
                    if not found:
                        ws.append_row(['historial_evaluaciones', data_str])
        except Exception:
            pass
        return True
    except Exception:
        return False

def _cargar_diagnostico():
    """Carga diagnósticos guardados desde JSON local"""
    try:
        if Path('diagnostico_data.json').exists():
            with open('diagnostico_data.json', 'r', encoding='utf-8') as f:
                return json.load(f)
    except Exception:
        pass
    return {}

def _guardar_diagnostico(data):
    """Guarda diagnósticos en JSON local + Google Sheets (igual que historial evaluaciones)"""
    try:
        with open('diagnostico_data.json', 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2, default=str)
        try:
            gs = _gs()
            if gs:
                ws = gs._get_hoja('config')
                if ws:
                    data_str = json.dumps(data, ensure_ascii=False, default=str)
                    all_data = ws.get_all_values()
                    found = False
                    for idx, row in enumerate(all_data):
                        if row and row[0] == 'diagnostico_data':
                            ws.update_cell(idx + 1, 2, data_str)
                            found = True
                            break
                    if not found:
                        ws.append_row(['diagnostico_data', data_str])
        except Exception:
            pass
        return True
    except Exception:
        return False

def _restaurar_diagnostico_desde_gs():
    """Restaura diagnósticos desde Google Sheets si el JSON local no existe"""
    try:
        if Path('diagnostico_data.json').exists():
            return
        gs = _gs()
        if not gs:
            return
        ws = gs._get_hoja('config')
        if not ws:
            return
        all_data = ws.get_all_values()
        for row in all_data:
            if row and row[0] == 'diagnostico_data' and len(row) > 1:
                data = json.loads(row[1])
                with open('diagnostico_data.json', 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                break
    except Exception:
        pass

def generar_pdf_diagnostico(grado, anio, estudiantes_df, notas_diag, areas_diag, tipo="entrada", notas_salida=None, config=None):
    """Genera PDF de diagnóstico — una hoja por estudiante con gráficos de barras.
    tipo: 'entrada' o 'salida'. Si es salida, también compara con notas_diag (entrada)."""
    if config is None:
        config = {}
    buf = io.BytesIO()
    from reportlab.graphics.shapes import Drawing, Rect, String, Line
    from reportlab.graphics import renderPDF

    c = canvas.Canvas(buf, pagesize=A4)
    w, h = A4

    colegio = config.get('colegio', 'I.E.P. ALTERNATIVO YACHAY')
    anio_c  = config.get('anio', anio)
    tipo_txt = "DIAGNÓSTICO DE ENTRADA" if tipo == "entrada" else "DIAGNÓSTICO DE SALIDA"

    def draw_bar_chart(c, x, y, chart_w, chart_h, areas, notas, notas_prev=None):
        """Dibuja gráfico de barras horizontales dentro del área dada."""
        n = len(areas)
        if n == 0:
            return
        bar_h   = min(22, (chart_h - 20) / n)
        max_nota = 20

        # Eje Y labels y barras
        for i, (area, nota) in enumerate(zip(areas, notas)):
            bar_y = y + chart_h - 20 - i * (bar_h + 3) - bar_h
            nota_v = float(nota) if nota else 0
            bar_len = (nota_v / max_nota) * chart_w * 0.7

            # Color según escala
            if nota_v >= 18:   col = colors.HexColor('#15803d')
            elif nota_v >= 14: col = colors.HexColor('#2563eb')
            elif nota_v >= 11: col = colors.HexColor('#d97706')
            else:              col = colors.HexColor('#dc2626')

            # Barra principal (entrada o actual)
            c.setFillColor(col)
            c.rect(x + chart_w * 0.3, bar_y, bar_len, bar_h * 0.55, fill=1, stroke=0)

            # Si hay notas previas (comparación entrada vs salida)
            if notas_prev and i < len(notas_prev):
                prev_v = float(notas_prev[i]) if notas_prev[i] else 0
                prev_len = (prev_v / max_nota) * chart_w * 0.7
                c.setFillColor(colors.HexColor('#94a3b8'))
                c.setFillAlpha(0.5)
                c.rect(x + chart_w * 0.3, bar_y + bar_h * 0.55 + 1, prev_len, bar_h * 0.3, fill=1, stroke=0)
                c.setFillAlpha(1.0)

            # Etiqueta área
            c.setFillColor(colors.HexColor('#1e293b'))
            c.setFont('Helvetica', 6)
            area_short = area[:22] + '.' if len(area) > 22 else area
            c.drawRightString(x + chart_w * 0.29, bar_y + bar_h * 0.2, area_short)

            # Nota al final de la barra
            letra = nota_a_letra(nota_v)
            c.setFillColor(col)
            c.setFont('Helvetica-Bold', 7)
            c.drawString(x + chart_w * 0.3 + bar_len + 4, bar_y + bar_h * 0.15,
                         f"{nota_v:.0f} ({letra})")

        # Eje X
        c.setStrokeColor(colors.HexColor('#cbd5e1'))
        c.setLineWidth(0.5)
        c.line(x + chart_w * 0.3, y + chart_h - 20, x + chart_w * 0.3, y + 5)
        for tick in [0, 5, 10, 14, 18, 20]:
            tx = x + chart_w * 0.3 + (tick / max_nota) * chart_w * 0.7
            c.line(tx, y + chart_h - 20, tx, y + 5)
            c.setFillColor(colors.HexColor('#94a3b8'))
            c.setFont('Helvetica', 5)
            c.drawCentredString(tx, y + 2, str(tick))

    if not estudiantes_df.empty:
        est = estudiantes_df.sort_values('Nombre').reset_index(drop=True)
    else:
        est = pd.DataFrame()

    for idx, (_, row) in enumerate(est.iterrows()):
        nombre = str(row.get('Nombre', '')).strip()
        dni    = str(row.get('DNI', '')).strip()
        grade  = str(row.get('Grado', grado)).strip()

        if idx > 0:
            c.showPage()

        # Fondo cabecera
        c.setFillColor(colors.HexColor('#1e3a8a'))
        c.rect(0, h - 70, w, 70, fill=1, stroke=0)

        # Texto cabecera
        c.setFillColor(colors.white)
        c.setFont('Helvetica-Bold', 11)
        c.drawCentredString(w / 2, h - 22, colegio.upper())
        c.setFont('Helvetica-Bold', 13)
        c.drawCentredString(w / 2, h - 40, f"PRUEBA DE {tipo_txt}")
        c.setFont('Helvetica', 9)
        c.drawCentredString(w / 2, h - 55, f"Año Escolar {anio_c}  |  Grado: {grade}")

        # Datos del estudiante
        c.setFillColor(colors.HexColor('#f0f9ff'))
        c.rect(20, h - 110, w - 40, 35, fill=1, stroke=0)
        c.setFillColor(colors.HexColor('#1e293b'))
        c.setFont('Helvetica-Bold', 10)
        c.drawString(30, h - 90, f"Estudiante: {nombre}")
        c.setFont('Helvetica', 9)
        c.drawString(30, h - 104, f"DNI: {dni}   |   N°: {idx + 1}")

        # Notas de este estudiante
        notas_est   = notas_diag.get(nombre, [None] * len(areas_diag))
        notas_sal   = notas_salida.get(nombre, [None] * len(areas_diag)) if notas_salida else None

        # Tabla de notas
        ty = h - 125
        col_widths_t = [190, 60, 55, 55, 55]
        headers_t = ['ÁREA', 'NOTA', 'LITERAL', 'ESCALA', '']
        if tipo == 'salida' and notas_salida:
            headers_t = ['ÁREA', 'ENTRADA', 'SALIDA', 'LITERAL', 'AVANCE']
            col_widths_t = [180, 55, 55, 55, 55]

        # Cabecera tabla
        tx = 20
        c.setFillColor(colors.HexColor('#1e3a8a'))
        c.rect(tx, ty - 16, sum(col_widths_t), 16, fill=1, stroke=0)
        c.setFillColor(colors.white)
        c.setFont('Helvetica-Bold', 7)
        cx = tx
        for hi, hw in zip(headers_t, col_widths_t):
            c.drawCentredString(cx + hw / 2, ty - 11, hi)
            cx += hw

        prom_sum = 0; prom_n = 0
        nota_fila_y = ty - 16
        for ai, (area, nota) in enumerate(zip(areas_diag, notas_est)):
            nota_v = float(nota) if nota is not None and str(nota) not in ('', 'None', 'nan') else 0
            letra  = nota_a_letra(nota_v)
            bg = colors.HexColor('#f8fafc') if ai % 2 == 0 else colors.white
            c.setFillColor(bg)
            c.rect(tx, nota_fila_y - 14, sum(col_widths_t), 14, fill=1, stroke=0)
            c.setFillColor(colors.HexColor('#1e293b'))
            c.setFont('Helvetica', 7)

            if tipo == 'salida' and notas_salida:
                notas_sal_est = notas_salida.get(nombre, [None] * len(areas_diag))
                sal_v = float(notas_sal_est[ai]) if notas_sal_est[ai] is not None and str(notas_sal_est[ai]) not in ('', 'None', 'nan') else 0
                sal_letra = nota_a_letra(sal_v)
                avance = sal_v - nota_v
                avance_txt = f"+{avance:.0f}" if avance > 0 else f"{avance:.0f}"
                avance_color = '#15803d' if avance > 0 else ('#dc2626' if avance < 0 else '#6b7280')
                vals = [area[:28], f"{nota_v:.0f}", f"{sal_v:.0f}", sal_letra, avance_txt]
                val_colors = [None, None, None, color_semaforo(sal_letra), avance_color]
                prom_sum += sal_v
            else:
                vals = [area[:28], f"{nota_v:.0f}", letra, f"({ESCALA_MINEDU[letra]['rango']})", ""]
                val_colors = [None, None, color_semaforo(letra), '#6b7280', None]
                prom_sum += nota_v
            prom_n += 1

            cx = tx
            for vi, (val, cw_t) in enumerate(zip(vals, col_widths_t)):
                vc = val_colors[vi] if val_colors[vi] else '#1e293b'
                c.setFillColor(colors.HexColor(vc))
                if vi == 0:
                    c.drawString(cx + 3, nota_fila_y - 10, str(val))
                else:
                    c.drawCentredString(cx + cw_t / 2, nota_fila_y - 10, str(val))
                cx += cw_t

            nota_fila_y -= 14

        # Promedio
        prom = prom_sum / prom_n if prom_n > 0 else 0
        prom_letra = nota_a_letra(prom)
        c.setFillColor(colors.HexColor('#1e3a8a'))
        c.rect(tx, nota_fila_y - 14, sum(col_widths_t), 14, fill=1, stroke=0)
        c.setFillColor(colors.white)
        c.setFont('Helvetica-Bold', 7)
        c.drawString(tx + 3, nota_fila_y - 10, "PROMEDIO GENERAL")
        c.drawCentredString(tx + sum(col_widths_t) * 0.7, nota_fila_y - 10,
                            f"{prom:.1f}  ({prom_letra} — {ESCALA_MINEDU[prom_letra]['nombre']})")

        # Gráfico de barras
        chart_top = nota_fila_y - 30
        chart_h_px = max(80, min(200, chart_top - 60))
        chart_x = 20
        chart_w_px = w - 40

        c.setFont('Helvetica-Bold', 8)
        c.setFillColor(colors.HexColor('#1e3a8a'))
        if tipo == 'salida' and notas_salida:
            c.drawString(chart_x, chart_top + 5, "📊 COMPARACIÓN ENTRADA vs SALIDA")
        else:
            c.drawString(chart_x, chart_top + 5, "📊 GRÁFICO DE RESULTADOS")

        notas_vals = [float(n) if n is not None and str(n) not in ('', 'None', 'nan') else 0 for n in notas_est]
        notas_sal_vals = None
        if tipo == 'salida' and notas_salida:
            notas_sal_lista = notas_salida.get(nombre, [None] * len(areas_diag))
            notas_sal_vals = [float(n) if n is not None and str(n) not in ('', 'None', 'nan') else 0 for n in notas_sal_lista]

        draw_bar_chart(c, chart_x, chart_top - chart_h_px - 5, chart_w_px, chart_h_px,
                       areas_diag, notas_vals if tipo == 'entrada' else (notas_sal_vals or notas_vals),
                       notas_prev=notas_vals if tipo == 'salida' else None)

        # Leyenda si es salida
        if tipo == 'salida' and notas_salida:
            ly = chart_top - chart_h_px - 15
            c.setFont('Helvetica', 6)
            c.setFillColor(colors.HexColor('#1e3a8a'))
            c.rect(chart_x, ly - 8, 10, 7, fill=1, stroke=0)
            c.setFillColor(colors.HexColor('#1e293b'))
            c.drawString(chart_x + 13, ly - 7, "Resultado Salida")
            c.setFillColor(colors.HexColor('#94a3b8'))
            c.setFillAlpha(0.6)
            c.rect(chart_x + 100, ly - 8, 10, 7, fill=1, stroke=0)
            c.setFillAlpha(1.0)
            c.setFillColor(colors.HexColor('#1e293b'))
            c.drawString(chart_x + 113, ly - 7, "Resultado Entrada")

        # Pie
        c.setFont('Helvetica-Oblique', 6)
        c.setFillColor(colors.HexColor('#94a3b8'))
        c.drawCentredString(w / 2, 20, f"YACHAY PRO — {tipo_txt} — {grade} — Año {anio_c}")

    c.save()
    buf.seek(0)
    return buf


def tab_registrar_notas(config):
    """Módulo para que docentes registren notas — multi-área, sesión limpia, historial"""
    st.header("📝 Registrar Notas")

    usuario = st.session_state.get('usuario_actual', '')
    _di_rn = st.session_state.get('docente_info', {}) or {}
    nombre_completo_doc = _nombre_completo_docente()
    gs = _gs()

    # ─── Determinar grado disponible para el docente ─────────────────────────
    grado_doc = None
    nivel_doc = None
    if st.session_state.docente_info:
        grado_doc = st.session_state.docente_info.get('grado', '')
        nivel_doc = st.session_state.docente_info.get('nivel', '')

    # ─── PESTAÑA: Historial / Nueva Evaluación / Diagnóstico ─────────────────
    vista = st.radio("", ["📋 Nueva Evaluación", "📂 Historial de Evaluaciones",
                          "🔬 Examen Diagnóstico"],
                     horizontal=True, key="rn_vista")

    # ── DIAGNÓSTICO ───────────────────────────────────────────────────────────
    if vista == "🔬 Examen Diagnóstico":
        st.subheader("🔬 Examen Diagnóstico de Entrada / Salida")
        st.info("Registre las notas del diagnóstico por estudiante. Se guardan en Google Sheets automáticamente.")

        tipo_diag = st.radio("Tipo:", ["📥 Diagnóstico de Entrada", "📤 Diagnóstico de Salida"],
                              horizontal=True, key="tipo_diag")
        es_salida = "Salida" in tipo_diag

        # Grado
        if grado_doc and grado_doc not in ('ALL_NIVELES', 'ALL_SEC_PREU', 'ALL_SECUNDARIA', 'N/A', ''):
            grado_diag = grado_doc
            st.info(f"Grado: **{grado_diag}**")
        else:
            grado_diag = st.selectbox("Grado:", GRADOS_OPCIONES, key="grado_diag")

        sec_diag = st.selectbox("Sección:", ["Todas"] + SECCIONES, key="sec_diag")
        df_diag = BaseDatos.obtener_estudiantes_grado(grado_diag, sec_diag if sec_diag != "Todas" else None)

        # Cargar datos guardados permanentes
        _restaurar_diagnostico_desde_gs()
        todos_diag = _cargar_diagnostico()
        clave_ent = f"entrada_{grado_diag}_{sec_diag}"
        clave_sal = f"salida_{grado_diag}_{sec_diag}"
        datos_entrada_guardados = todos_diag.get(clave_ent, {})
        datos_salida_guardados  = todos_diag.get(clave_sal, {})
        areas_guardadas = todos_diag.get(f"areas_{grado_diag}_{sec_diag}", [])

        if df_diag.empty:
            st.warning("No hay estudiantes en este grado.")
        else:
            st.markdown("---")
            st.markdown("**Áreas a evaluar:**")
            na_diag = st.number_input("Número de áreas:", 1, 10,
                                       value=len(areas_guardadas) if areas_guardadas else 4,
                                       key="na_diag")
            areas_diag = []
            defaults = ["Comunicación","Matemática","Ciencia y Tec.","Historia","Personal Social","Arte","Educación Física","Inglés","Religión","Tutoría"]
            for ai in range(int(na_diag)):
                saved_area = areas_guardadas[ai] if ai < len(areas_guardadas) else ""
                placeholder = saved_area if saved_area else (defaults[ai] if ai < len(defaults) else f"Área {ai+1}")
                area_n = st.text_input(f"Área {ai+1}:", value=saved_area, key=f"area_diag_{ai}",
                                        placeholder=placeholder)
                areas_diag.append(area_n.strip() if area_n.strip() else placeholder)

            st.markdown("---")
            st.markdown("**Notas de los estudiantes (0–20):**")
            # Mostrar cabecera de áreas
            hdr_cols = st.columns([3] + [1] * int(na_diag))
            with hdr_cols[0]:
                st.markdown("**Estudiante**")
            for ai, area in enumerate(areas_diag):
                with hdr_cols[ai + 1]:
                    st.markdown(f"**{area[:10]}**")

            # Datos a usar para prellenar según tipo
            datos_prev = datos_salida_guardados if es_salida else datos_entrada_guardados

            notas_ingresadas = {}
            for _, row in df_diag.iterrows():
                nombre_e = str(row.get('Nombre', '')).strip()
                cols_notas = st.columns([3] + [1] * int(na_diag))
                with cols_notas[0]:
                    st.markdown(nombre_e)
                notas_est_diag = []
                prev_notas = datos_prev.get(nombre_e, [0.0] * int(na_diag))
                for ai in range(int(na_diag)):
                    prev_v = float(prev_notas[ai]) if ai < len(prev_notas) and prev_notas[ai] is not None else 0.0
                    with cols_notas[ai + 1]:
                        nota_d = st.number_input(
                            f"n{ai}", min_value=0.0, max_value=20.0,
                            value=prev_v, step=0.5,
                            key=f"nd_{nombre_e}_{ai}_{es_salida}",
                            label_visibility="collapsed"
                        )
                    notas_est_diag.append(nota_d)
                notas_ingresadas[nombre_e] = notas_est_diag

            st.markdown("---")
            tipo_nombre = "Salida" if es_salida else "Entrada"
            cg1, cg2 = st.columns(2)

            with cg1:
                if st.button(f"💾 Guardar {tipo_nombre} en Google Sheets", type="primary",
                             use_container_width=True, key="btn_save_diag"):
                    todos_diag[f"areas_{grado_diag}_{sec_diag}"] = areas_diag
                    todos_diag[clave_sal if es_salida else clave_ent] = notas_ingresadas
                    if _guardar_diagnostico(todos_diag):
                        st.success(f"✅ Diagnóstico de {tipo_nombre} guardado en Google Sheets.")
                    else:
                        st.warning("⚠️ Guardado localmente (sin conexión a Google Sheets).")

            with cg2:
                if st.button(f"🖨️ Generar PDF Diagnóstico {tipo_nombre}", type="primary",
                             use_container_width=True, key="btn_gen_diag"):
                    notas_ent_para_pdf = notas_ingresadas
                    notas_sal_para_pdf = None
                    if es_salida:
                        notas_ent_para_pdf = datos_entrada_guardados if datos_entrada_guardados else notas_ingresadas
                        notas_sal_para_pdf = notas_ingresadas
                    with st.spinner("Generando PDF..."):
                        pdf_diag = generar_pdf_diagnostico(
                            grado_diag, config.get('anio', '2026'),
                            df_diag, notas_ent_para_pdf, areas_diag,
                            tipo="salida" if es_salida else "entrada",
                            notas_salida=notas_sal_para_pdf,
                            config=config
                        )
                    st.session_state['_diag_pdf'] = pdf_diag
                    st.session_state['_diag_tipo'] = tipo_nombre
                    st.session_state['_diag_grado'] = grado_diag

            # Info de comparación
            if es_salida:
                if datos_entrada_guardados:
                    st.caption(f"✅ Se compara con {len(datos_entrada_guardados)} registros de Entrada guardados")
                else:
                    st.caption("⚠️ No hay datos de Entrada guardados aún. Guarde primero el diagnóstico de Entrada.")

            # Botón de descarga — siempre visible si hay PDF generado para este tipo
            if st.session_state.get('_diag_pdf') and st.session_state.get('_diag_tipo') == tipo_nombre:
                st.download_button(
                    f"⬇️ Descargar PDF Diagnóstico {tipo_nombre}",
                    st.session_state['_diag_pdf'],
                    f"Diagnostico_{tipo_nombre}_{st.session_state.get('_diag_grado', grado_diag)}.pdf",
                    mime="application/pdf",
                    type="primary",
                    use_container_width=True,
                    key="dl_diag_pdf"
                )

    if vista == "📂 Historial de Evaluaciones":
        st.markdown("### 📂 Evaluaciones Guardadas")
        hist = _cargar_historial_evaluaciones()
        # Filtrar por rol
        if st.session_state.rol not in ['admin', 'directivo']:
            hist = {k: v for k, v in hist.items() if v.get('docente') == usuario}
        if not hist:
            st.info("📭 No hay evaluaciones guardadas aún.")
            return
        for clave, ev in sorted(hist.items(), reverse=True):
            titulo_h = ev.get('titulo', '') or ''
            label_h = f"📝 {ev['grado']} | {ev['periodo']} | {ev['fecha']}"
            if titulo_h:
                label_h += f" — {titulo_h}"
            with st.expander(label_h):
                areas_h = ev.get('areas', [])
                areas_nombres = [a['nombre'] for a in areas_h] if isinstance(areas_h[0], dict) else areas_h
                st.caption(f"Docente: {ev.get('docente_nombre', ev.get('docente','—'))} | Áreas: {', '.join(areas_nombres)} | Estudiantes: {len(ev.get('ranking',[]))}")
                ranking_h = ev.get('ranking', [])
                if ranking_h:
                    df_h = pd.DataFrame(ranking_h)
                    cols_h = ['Puesto','Medalla','Nombre'] + areas_nombres + ['Promedio']
                    cols_h = [c for c in cols_h if c in df_h.columns]
                    st.dataframe(df_h[cols_h], use_container_width=True, hide_index=True)
                    if st.button("📥 PDF Ranking", key=f"pdf_hist_{clave}", type="primary"):
                        pdf_h = _generar_ranking_pdf(ranking_h, areas_nombres, ev['grado'], ev['periodo'], config)
                        st.download_button("⬇️ Descargar", pdf_h,
                                           f"Ranking_{ev['grado']}_{ev['periodo']}_{ev['fecha']}.pdf",
                                           "application/pdf", key=f"dl_hist_{clave}")
        return

    # ═══════════════════════════════════════════════════════════════════════════
    # NUEVA EVALUACIÓN
    # ═══════════════════════════════════════════════════════════════════════════

    # ─── FASE 1: Configurar evaluación si no hay sesión activa ───────────────
    if 'eval_sesion' not in st.session_state or st.session_state.eval_sesion is None:

        st.markdown("### ⚙️ Configurar Nueva Evaluación")

        # Grado — usa el helper central que filtra por rol
        grado_cfg = _grados_para_selector("rn_cfg")
        if not grado_cfg:
            return

        # Período y título
        c1, c2 = st.columns(2)
        with c1:
            bim_cfg = st.selectbox("📅 Período:", PERIODOS_EVALUACION, key="rn_cfg_bim")
        with c2:
            titulo_cfg = st.text_input("📝 Título (opcional):", placeholder="Ej: Evaluación Semanal 3", key="rn_cfg_titulo")

        # Número de áreas — aplica a TODOS los niveles
        st.markdown("---")
        num_areas = st.radio("📚 Número de áreas a evaluar:", [1, 2, 3, 4, 5, 6], horizontal=True, key="rn_cfg_nareas")

        # Determinar áreas disponibles según el grado seleccionado
        grado_str_cfg = str(grado_cfg)
        if 'Inicial' in grado_str_cfg:
            areas_disp = AREAS_MINEDU.get('INICIAL', AREAS_MINEDU.get('PRIMARIA', []))
        elif any(x in grado_str_cfg for x in ['1° Sec','2° Sec','3° Sec','4° Sec','5° Sec']):
            areas_cepre_all = sorted(set(AREAS_CEPRE_UNSAAC.get('GRUPO AB', []) + AREAS_CEPRE_UNSAAC.get('GRUPO CD', [])))
            areas_disp = AREAS_MINEDU.get('SECUNDARIA', []) + areas_cepre_all
        elif 'GRUPO AB' in grado_str_cfg:
            areas_disp = AREAS_CEPRE_UNSAAC.get('GRUPO AB', [])
        elif 'GRUPO CD' in grado_str_cfg:
            areas_disp = AREAS_CEPRE_UNSAAC.get('GRUPO CD', [])
        elif any(x in grado_str_cfg for x in ['Ciclo','Reforzamiento','Preu','PREU']):
            areas_preu = AREAS_CEPRE_UNSAAC.get('GRUPO AB', []) + AREAS_CEPRE_UNSAAC.get('GRUPO CD', [])
            areas_disp = sorted(set(areas_preu))
        else:
            # Primaria (y cualquier otro)
            areas_disp = AREAS_MINEDU.get('PRIMARIA', [])

        # Si no hay áreas definidas, permitir texto libre
        if not areas_disp:
            areas_disp = ["Matemática", "Comunicación", "Ciencias", "Historia", "Arte", "Educación Física"]

        areas_cfg = []
        cols_a = st.columns(num_areas)
        for i in range(num_areas):
            with cols_a[i]:
                st.markdown(f"**Área/Curso {i+1}**")
                nombre_a = st.selectbox(f"Área:", areas_disp, key=f"rn_cfg_area_{i}")
                npregs_a = st.number_input(f"N° preguntas:", 1, 100, 20, key=f"rn_cfg_npregs_{i}")
                areas_cfg.append({'nombre': nombre_a, 'num_preguntas': int(npregs_a)})

        st.markdown("---")
        if st.button("▶ INICIAR EVALUACIÓN", type="primary", use_container_width=True, key="btn_iniciar_eval"):
            nombres_areas = [a['nombre'] for a in areas_cfg]
            if len(set(nombres_areas)) < len(nombres_areas):
                st.error("⚠️ Las áreas seleccionadas deben ser diferentes entre sí.")
            else:
                # Intentar cargar estudiantes - búsqueda robusta
                dg_cache = BaseDatos.obtener_estudiantes_grado(grado_cfg)
                if dg_cache.empty:
                    # Mostrar info diagnóstico
                    df_all = BaseDatos.cargar_matricula()
                    if df_all.empty:
                        st.error("⚠️ La matrícula está vacía. Registra estudiantes primero.")
                    else:
                        grados_existentes = sorted(df_all['Grado'].dropna().unique().tolist()) if 'Grado' in df_all.columns else []
                        st.error(f"⚠️ No hay estudiantes en **{grado_cfg}**.")
                        st.info(f"💡 Grados con estudiantes: {', '.join(str(g) for g in grados_existentes[:10])}")
                else:
                    import uuid
                    st.session_state.eval_sesion = {
                        'id': str(uuid.uuid4())[:8],
                        'grado': grado_cfg,
                        'periodo': bim_cfg,
                        'titulo': titulo_cfg,
                        'areas': areas_cfg,
                        'fecha': fecha_peru_str(),
                        'docente': usuario,
                        'docente_nombre': nombre_completo_doc,
                    }
                    st.session_state.eval_estudiantes = dg_cache.to_dict('records')
                    st.session_state.notas_sesion = {}
                    st.rerun()
        return

    # ─── FASE 2: Ingresar notas ───────────────────────────────────────────────
    ev = st.session_state.eval_sesion
    areas = ev['areas']  # lista de {nombre, num_preguntas}
    num_areas = len(areas)  # IMPORTANTE: definir aquí para evitar UnboundLocalError
    grado_sel = ev['grado']
    bim_sel = ev['periodo']
    titulo_ev = ev.get('titulo', '')

    # Encabezado de la evaluación activa
    titulo_mostrar = f"{grado_sel} | {bim_sel}"
    if titulo_ev:
        titulo_mostrar += f" — {titulo_ev}"
    st.success(f"✅ Evaluación activa: **{titulo_mostrar}**")
    areas_str = " + ".join([f"{a['nombre']} ({a['num_preguntas']} pregs.)" for a in areas])
    st.caption(f"📚 {areas_str}")

    col_nueva, _ = st.columns([1, 4])
    with col_nueva:
        # Botón NUEVA EVALUACIÓN con color cyan intenso
        st.markdown("""
        <style>
        button[key="btn_nueva_eval"] {
            background: linear-gradient(135deg, #06b6d4 0%, #0891b2 100%) !important;
            color: white !important;
            font-weight: bold !important;
            border: none !important;
            box-shadow: 0 4px 6px rgba(6, 182, 212, 0.4) !important;
        }
        button[key="btn_nueva_eval"]:hover {
            background: linear-gradient(135deg, #0891b2 0%, #0e7490 100%) !important;
            box-shadow: 0 6px 10px rgba(6, 182, 212, 0.6) !important;
        }
        </style>
        """, unsafe_allow_html=True)
        if st.button("🔄 NUEVA EVALUACIÓN", key="btn_nueva_eval", type="primary"):
            st.session_state.eval_sesion = None
            st.session_state.notas_sesion = {}
            st.session_state.eval_estudiantes = []
            st.rerun()

    st.markdown("---")

    # Cargar estudiantes — usar caché de sesión para estabilidad
    if 'eval_estudiantes' in st.session_state and st.session_state.eval_estudiantes:
        dg = pd.DataFrame(st.session_state.eval_estudiantes)
    else:
        # Fallback: intentar cargar de BD
        dg = BaseDatos.obtener_estudiantes_grado(grado_sel)
        if not dg.empty:
            st.session_state.eval_estudiantes = dg.to_dict('records')
    
    if dg.empty:
        st.warning("No hay estudiantes matriculados en este grado.")
        return

    st.markdown(f"### 📋 {len(dg)} estudiantes")

    # Inicializar notas_sesion si no existe
    if 'notas_sesion' not in st.session_state:
        st.session_state.notas_sesion = {}

    # Encabezado tabla — dinámico según número de áreas
    # Layout: [Estudiante, NSP] + [Nota/20 × num_areas] + [Promedio, Lit]
    if num_areas == 1:
        hcols = st.columns([3, 0.7, 1.5, 1, 1])
        headers = ["Estudiante", "NSP", f"{areas[0]['nombre'][:12]}/20", "Lit.", "Estado"]
    else:
        # Para 2-6 áreas: mostrar nota directa por cada área
        col_widths = [2.5, 0.7]  # Estudiante + NSP
        headers_list = ["Estudiante", "NSP"]
        for i in range(num_areas):
            col_widths.append(1.3)  # Solo nota, sin separador
            headers_list.append(f"{areas[i]['nombre'][:8]}/20")
        col_widths.extend([1, 0.8])  # Promedio + Lit
        headers_list.extend(["Prom.", "Lit."])
        hcols = st.columns(col_widths)
        headers = headers_list

    for hc, hdr in zip(hcols, headers):
        with hc:
            if hdr:  # No mostrar header vacío para separadores
                st.markdown(f"**{hdr}**")

    notas_actuales = {}


    for idx, row in dg.iterrows():
        nombre = str(row.get('Nombre', ''))
        dni = str(row.get('DNI', ''))
        sesion_id = ev['id']

        # ── Crear columnas dinámicas ─────────────────────────────────────────
        if num_areas == 1:
            col_widths = [3, 0.7, 1.5, 1, 1]
        else:
            col_widths = [2.5, 0.7]  # Nombre + NSP
            for _ in range(num_areas):
                col_widths.append(1.3)  # Solo nota
            col_widths.extend([1, 0.8])  # Promedio + Lit

        nc = st.columns(col_widths)

        # Columna 0: Nombre
        with nc[0]:
            st.write(f"👤 {nombre}")

        # Columna 1: NSP checkbox
        with nc[1]:
            nsp = st.checkbox("", key=f"nsp_{sesion_id}_{dni}",
                             value=st.session_state.notas_sesion.get(dni, {}).get('nsp', False),
                             label_visibility="collapsed")

        if nsp:
            # Si NSP está marcado, no pedir inputs de notas
            notas_actuales[dni] = {
                'nombre': nombre,
                'nsp': True,
                'areas': {},
                'promedio': 0
            }
            # Mostrar "NSP" en el resto de columnas
            for i in range(2, len(nc)):
                with nc[i]:
                    st.caption("—")
        else:
            # Ingresar correctas y calcular notas por cada área
            correctas_vals = []
            notas_vals = []
            col_idx = 2  # Empieza después de Nombre y NSP

            for i, area in enumerate(areas):
                # NOTA DIRECTA sobre 20 (con decimales)
                with nc[col_idx]:
                    nota_actual = st.session_state.notas_sesion.get(dni, {}).get(f'nota_{i}', 0.0)
                    nota_i = st.number_input("", min_value=0.0, max_value=20.0, value=float(nota_actual),
                                            step=0.5,  # Incrementos de 0.5
                                            key=f"nota_{i}_{sesion_id}_{dni}",
                                            label_visibility="collapsed")
                    notas_vals.append(nota_i)
                col_idx += 1

            # Promedio y literal (solo si hay más de 1 área)
            if num_areas == 1:
                promedio = notas_vals[0]
                lit = nota_a_letra(promedio)
                with nc[col_idx]:
                    st.markdown(f"<span style='color:{color_semaforo(lit)};font-weight:bold;'>{lit}</span>",
                               unsafe_allow_html=True)
                col_idx += 1
                with nc[col_idx]:
                    st.caption(ESCALA_MINEDU.get(lit, {}).get('nombre', '')[:10])
            else:
                promedio = round(sum(notas_vals) / num_areas, 1) if notas_vals else 0
                lit = nota_a_letra(promedio)
                with nc[col_idx]:
                    st.markdown(f"<span style='color:{color_semaforo(lit)};font-weight:bold;'>{promedio}</span>",
                               unsafe_allow_html=True)
                col_idx += 1
                with nc[col_idx]:
                    st.markdown(f"<span style='color:{color_semaforo(lit)};font-weight:bold;'>{lit}</span>",
                               unsafe_allow_html=True)

            # Guardar en diccionario - ahora con notas directas
            notas_actuales[dni] = {
                'nombre': nombre,
                'nsp': False,
                'areas': {areas[i]['nombre']: notas_vals[i] for i in range(num_areas)},
                'promedio': promedio,
                **{f'nota_{i}': notas_vals[i] for i in range(num_areas)}
            }

    # Actualizar sesión con lo ingresado
    st.session_state.notas_sesion = notas_actuales

    # ─── RANKING EN TIEMPO REAL ───────────────────────────────────────────────
    st.markdown("---")
    st.subheader("🏆 Ranking (evaluación actual)")

    areas_nombres = [a['nombre'] for a in areas]
    ranking_filas = []
    sin_nota_filas = []   # alumnos sin nota o NSP
    for dni_r, data_r in notas_actuales.items():
        if data_r.get('nsp', False) or data_r['promedio'] == 0:
            sin_nota_filas.append({'DNI': dni_r, 'Nombre': data_r['nombre']})
            continue
        if data_r['promedio'] > 0:
            fila = {'DNI': dni_r, 'Nombre': data_r['nombre']}
            for a_name in areas_nombres:
                fila[a_name] = data_r['areas'].get(a_name, 0)
            fila['Promedio'] = data_r['promedio']
            ranking_filas.append(fila)
    sin_nota_filas.sort(key=lambda x: x['Nombre'])

    ranking_filas.sort(key=lambda x: x['Promedio'], reverse=True)
    for i, f in enumerate(ranking_filas):
        f['Puesto'] = i + 1
        f['Medalla'] = ['🥇','🥈','🥉'][i] if i < 3 else f'#{i+1}'

    if ranking_filas:
        df_rank = pd.DataFrame(ranking_filas)
        cols_order = ['Puesto', 'Medalla', 'Nombre'] + areas_nombres + ['Promedio']
        cols_exist = [c for c in cols_order if c in df_rank.columns]
        st.dataframe(df_rank[cols_exist], use_container_width=True, hide_index=True, height=350)
        st.caption(f"📊 {len(ranking_filas)} estudiantes con nota > 0")
    else:
        st.info("📭 Ingresa correctas para ver el ranking en tiempo real")

    # ─── GUARDAR Y FINALIZAR ──────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("### 💾 Guardar Evaluación en Historial")
    st.info("💡 Al guardar, la evaluación queda registrada en el historial y podrás iniciar una nueva.")

    # CSS global - Colores SÓLIDOS FIJOS que no cambian
    st.markdown("""
    <style>
    /* TODOS los botones en esta sección */
    div[data-testid="column"] button,
    div.stButton > button {
        color: #000000 !important;
        font-weight: 900 !important;
        font-size: 16px !important;
        border: 2px solid rgba(0,0,0,0.2) !important;
        text-shadow: none !important;
    }
    /* Botón GUARDAR - Verde sólido */
    div[data-testid="column"]:nth-of-type(1) button {
        background: #10b981 !important;
    }
    /* Botón DESCARGAR - Naranja sólido */
    div[data-testid="column"]:nth-of-type(2) button {
        background: #f97316 !important;
    }
    /* Botón WhatsApp - Verde WA sólido */
    button[key="btn_wa_eval"] {
        background: #25D366 !important;
    }
    /* Botón NUEVA EVALUACIÓN - Cyan sólido */
    button[key="btn_nueva_eval"] {
        background: #0891b2 !important;
    }
    /* Hover - solo un poco más oscuro */
    div.stButton > button:hover {
        opacity: 0.9 !important;
        transform: scale(1.02);
    }
    </style>
    """, unsafe_allow_html=True)

    if ranking_filas:
        col_g1, col_g2 = st.columns(2)
        with col_g1:
            if st.button("💾 GUARDAR EN HISTORIAL", type="primary",
                         use_container_width=True, key="btn_guardar_historial"):
                hist = _cargar_historial_evaluaciones()
                clave_hist = f"{grado_sel}_{bim_sel}_{ev['id']}_{fecha_peru_str()}"
                hist[clave_hist] = {
                    'id': ev['id'],
                    'grado': grado_sel,
                    'periodo': bim_sel,
                    'titulo': titulo_ev,
                    'fecha': fecha_peru_str(),
                    'hora': hora_peru_str(),
                    'docente': usuario,
                    'docente_nombre': nombre_completo_doc,
                    'areas': areas,
                    'ranking': ranking_filas,
                }
                if gs:
                    try:
                        ws = gs._get_hoja('config')
                        if ws:
                            ws.append_row([f"histeval_{clave_hist}",
                                           json.dumps(hist[clave_hist], ensure_ascii=False, default=str)])
                    except Exception:
                        pass
                if _guardar_historial_evaluaciones(hist):
                    # También guardar notas individuales para Reporte Integral
                    try:
                        for dni_nota, data_nota in notas_actuales.items():
                            if data_nota.get('nsp', False):
                                continue  # No guardar NSP
                            # Crear registro individual
                            reg = {
                                'dni': dni_nota,
                                'nombre': data_nota['nombre'],
                                'grado': grado_sel,
                                'periodo': bim_sel,
                                'titulo': titulo_ev,
                                'fecha': fecha_peru_str(),
                                'hora': hora_peru_str(),
                                'docente': usuario,
                                'docente_nombre': nombre_completo_doc,
                                'areas': [{'nombre': a_name, 'nota': data_nota['areas'].get(a_name, 0)}
                                         for a_name in areas_nombres],
                                'promedio_general': data_nota['promedio'],
                                '_docente': usuario
                            }
                            # Guardar en resultados.json
                            resultados_actuales = BaseDatos.cargar_todos_resultados()
                            resultados_actuales.append(reg)
                            with open('resultados.json', 'w', encoding='utf-8') as f:
                                json.dump(resultados_actuales, f, ensure_ascii=False, indent=2)
                    except Exception:
                        pass
                    
                    _sync_resultados_a_gs()
                    st.success(f"✅ Evaluación guardada — {len(ranking_filas)} estudiantes")
                    st.balloons()
                    reproducir_beep_exitoso()
                else:
                    st.error("❌ Error al guardar")

        with col_g2:
            if st.button("📥 DESCARGAR RANKING", use_container_width=True, key="btn_pdf_eval", type="primary"):
                pdf_r = _generar_ranking_pdf(ranking_filas, areas_nombres, grado_sel, bim_sel, config, sin_nota=sin_nota_filas)
                st.download_button("⬇️ PDF", pdf_r, f"Ranking_{grado_sel}_{bim_sel}.pdf",
                                   "application/pdf", key="dl_pdf_eval")

        if st.button("📱 ENVIAR POR WHATSAPP", use_container_width=True, key="btn_wa_eval", type="primary"):
            st.session_state['_mostrar_wa_eval'] = True
        if st.session_state.get('_mostrar_wa_eval'):
            st.markdown("### 📱 Enviar Notas por WhatsApp")
            for fila in ranking_filas:
                alumno_wa = BaseDatos.buscar_por_dni(fila.get('DNI', ''))
                cel = str(alumno_wa.get('Celular_Apoderado', '') or '') if alumno_wa else ''
                cel = cel.strip() if cel.lower() not in ('nan', 'none', '') else ''
                if cel and cel.strip():
                    # Mensaje con caracteres seguros
                    msg = f"🏫 *I.E.P. YACHAY - CHINCHERO*\n📊 *REPORTE DE NOTAS*\n\n"
                    msg += f"👤 Alumno: {fila['Nombre']}\n📚 Grado: {grado_sel}\n📅 Periodo: {bim_sel}\n"
                    msg += "━" * 30 + "\n"
                    for a_n in areas_nombres:
                        nota_w = fila.get(a_n, 0)
                        msg += f"📖 {a_n}: *{nota_w}* ({nota_a_letra(nota_w)})\n"
                    msg += "━" * 30 + "\n"
                    msg += f"📊 *PROMEDIO: {fila['Promedio']}*\n🏆 *PUESTO: {fila['Medalla']}*"
                    
                    # Normalizar número
                    cel_c = cel.replace(' ','').replace('+','').replace('-','').strip()
                    if not cel_c.startswith('51'):
                        cel_c = '51' + cel_c
                    
                    # URL que abre DIRECTO en desktop app (no web)
                    # whatsapp:// funciona en desktop, wa.me en móvil
                    msg_encoded = urllib.parse.quote(msg)
                    url_desktop = f"whatsapp://send?phone={cel_c}&text={msg_encoded}"
                    url_movil = f"https://wa.me/{cel_c}?text={msg_encoded}"
                    
                    # Mostrar ambos links con HTML para abrir directo
                    st.markdown(f"""
                    <div style='margin:10px 0; padding:10px; background:#f0fdf4; border-radius:8px;'>
                        <b>📱 {fila['Nombre']}</b> → {cel}<br>
                        <a href="{url_desktop}" style='display:inline-block; margin:5px; padding:8px 15px; background:#25D366; color:white; text-decoration:none; border-radius:5px; font-weight:bold;'>
                            💻 WhatsApp Desktop
                        </a>
                        <a href="{url_movil}" target="_blank" style='display:inline-block; margin:5px; padding:8px 15px; background:#128C7E; color:white; text-decoration:none; border-radius:5px; font-weight:bold;'>
                            📱 WhatsApp Móvil
                        </a>
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.caption(f"⚠️ {fila['Nombre']} — Sin celular registrado")
    else:
        st.warning("⚠️ Ingresa al menos una nota para guardar")


def _generar_ranking_pdf(ranking_filas, areas, grado, periodo, config, sin_nota=None):
    """Genera PDF del ranking — colores por área, nombres completos, separadores visuales"""
    if sin_nota is None:
        sin_nota = []
    buffer = io.BytesIO()
    c_pdf = canvas.Canvas(buffer, pagesize=landscape(A4))
    w, h = landscape(A4)

    # Paleta de colores por área (fondo cabecera y franja de columna)
    PALETA_AREAS = [
        ("#1e3a8a", "#dbeafe"),  # azul
        ("#14532d", "#dcfce7"),  # verde
        ("#7c2d12", "#ffedd5"),  # naranja
        ("#581c87", "#f3e8ff"),  # violeta
        ("#164e63", "#cffafe"),  # cyan
        ("#713f12", "#fef9c3"),  # amarillo oscuro
        ("#1e1b4b", "#e0e7ff"),  # índigo
        ("#831843", "#fce7f3"),  # rosa
    ]

    # ── Marca de agua ────────────────────────────────────────────────────
    if Path("escudo_upload.png").exists():
        try:
            from PIL import Image as PILImage
            img = PILImage.open("escudo_upload.png")
            iw, ih = img.size
            mw = 200; mh = mw / (iw / ih)
            c_pdf.saveState()
            c_pdf.setFillAlpha(0.12)
            c_pdf.drawImage("escudo_upload.png", w/2-mw/2, h/2-mh/2, mw, mh, mask='auto')
            c_pdf.restoreState()
        except Exception:
            pass

    # ── Cabecera superior ────────────────────────────────────────────────
    c_pdf.setFillColor(colors.HexColor("#001e7c"))
    c_pdf.rect(0, h - 55, w, 55, fill=1, stroke=0)

    # Escudos
    ALTO_ESC = 45
    for path_e, lado in [("escudo_upload.png", "izq"),
                          ("escudo2_upload.png" if Path("escudo2_upload.png").exists() else "escudo_upload.png", "der")]:
        try:
            if Path(path_e).exists():
                from PIL import Image as PILImage
                img = PILImage.open(path_e)
                iw, ih = img.size
                aw = ALTO_ESC * (iw / ih)
                xp = 10 if lado == "izq" else w - 10 - aw
                c_pdf.drawImage(path_e, xp, h - 52, aw, ALTO_ESC, mask='auto')
        except Exception:
            pass

    c_pdf.setFillColor(colors.white)
    c_pdf.setFont("Helvetica-Bold", 16)
    c_pdf.drawCentredString(w / 2, h - 22, "RANKING DE ESTUDIANTES")
    c_pdf.setFont("Helvetica-Bold", 10)
    c_pdf.drawCentredString(w / 2, h - 36, f"I.E.P. YACHAY  —  {grado}  —  {periodo}")
    c_pdf.setFont("Helvetica", 8)
    c_pdf.setFillColor(colors.HexColor("#93c5fd"))
    c_pdf.drawCentredString(w / 2, h - 49, hora_peru().strftime('%d/%m/%Y'))

    # ── Layout de columnas ───────────────────────────────────────────────
    x_margin   = 12
    table_w    = w - 2 * x_margin
    col_w_pos  = 38
    col_w_nom  = min(175, table_w * 0.23)
    col_w_prom = 62
    remaining  = table_w - col_w_pos - col_w_nom - col_w_prom
    col_w_area = max(55, remaining / max(len(areas), 1))

    col_widths = [col_w_pos, col_w_nom] + [col_w_area] * len(areas) + [col_w_prom]

    # ── Cabecera de tabla — dos filas: fila área (coloreada) + fila nombre área ──
    y = h - 58
    HEADER_H1 = 14   # franja de color del área
    HEADER_H2 = 22   # nombre del área (puede ser 2 líneas)
    ROW_H     = 18

    # Fila superior: bloques de color por área + encabezados fijos
    # Puesto
    c_pdf.setFillColor(colors.HexColor("#0f172a"))
    c_pdf.rect(x_margin, y - HEADER_H1, col_w_pos, HEADER_H1, fill=1, stroke=0)
    # Nombre
    c_pdf.rect(x_margin + col_w_pos, y - HEADER_H1, col_w_nom, HEADER_H1, fill=1, stroke=0)
    # Promedio
    x_prom = x_margin + col_w_pos + col_w_nom + col_w_area * len(areas)
    c_pdf.setFillColor(colors.HexColor("#0f172a"))
    c_pdf.rect(x_prom, y - HEADER_H1, col_w_prom, HEADER_H1, fill=1, stroke=0)

    # Bloques de color por área en fila superior
    x_area_start = x_margin + col_w_pos + col_w_nom
    for ai, area in enumerate(areas):
        col_hex, _ = PALETA_AREAS[ai % len(PALETA_AREAS)]
        c_pdf.setFillColor(colors.HexColor(col_hex))
        c_pdf.rect(x_area_start + ai * col_w_area, y - HEADER_H1,
                   col_w_area, HEADER_H1, fill=1, stroke=0)
        # Nombre del área centrado en esa franja — font pequeño
        c_pdf.setFillColor(colors.white)
        c_pdf.setFont("Helvetica-Bold", 6)
        area_label = area if len(area) <= 18 else area[:17] + "."
        c_pdf.drawCentredString(x_area_start + ai * col_w_area + col_w_area / 2,
                                y - HEADER_H1 + 3, area_label)

    y -= HEADER_H1

    # Fila inferior de cabecera: "#", "APELLIDOS Y NOMBRES", "NOTA (LIT)", "PROM."
    c_pdf.setFillColor(colors.HexColor("#1e293b"))
    c_pdf.rect(x_margin, y - HEADER_H2, table_w, HEADER_H2, fill=1, stroke=0)
    c_pdf.setFillColor(colors.white)
    c_pdf.setFont("Helvetica-Bold", 8)
    cx = x_margin
    for hi, hw in zip(["#", "APELLIDOS Y NOMBRES"] + ["NOTA  (LIT)"] * len(areas) + ["PROMEDIO"],
                       col_widths):
        c_pdf.drawCentredString(cx + hw / 2, y - HEADER_H2 + 8, hi)
        # Separador vertical entre áreas
        if hi == "NOTA  (LIT)":
            c_pdf.setStrokeColor(colors.HexColor("#475569"))
            c_pdf.setLineWidth(0.4)
            c_pdf.line(cx, y, cx, y - HEADER_H2)
        cx += hw
    y -= HEADER_H2

    PIE_H   = 32   # espacio reservado para el pie de página
    es_primera_pag = True   # para saber si hay que redibujar cabecera

    def _dibujar_cabecera_tabla(c_pdf, y_pos):
        """Redibuja la cabecera de la tabla en páginas adicionales"""
        c_pdf.setFillColor(colors.HexColor("#1e293b"))
        c_pdf.rect(x_margin, y_pos - HEADER_H1, table_w, HEADER_H1, fill=1, stroke=0)
        cx2 = x_margin
        for ai2, area2 in enumerate(areas):
            col_hex2, _ = PALETA_AREAS[ai2 % len(PALETA_AREAS)]
            c_pdf.setFillColor(colors.HexColor(col_hex2))
            c_pdf.rect(x_margin + col_w_pos + col_w_nom + ai2 * col_w_area,
                       y_pos - HEADER_H1, col_w_area, HEADER_H1, fill=1, stroke=0)
            c_pdf.setFillColor(colors.white)
            c_pdf.setFont("Helvetica-Bold", 6)
            area_l2 = area2 if len(area2) <= 18 else area2[:17] + "."
            c_pdf.drawCentredString(
                x_margin + col_w_pos + col_w_nom + ai2 * col_w_area + col_w_area / 2,
                y_pos - HEADER_H1 + 3, area_l2)
        y2 = y_pos - HEADER_H1
        c_pdf.setFillColor(colors.HexColor("#0f172a"))
        c_pdf.rect(x_margin, y2 - HEADER_H2, table_w, HEADER_H2, fill=1, stroke=0)
        c_pdf.setFillColor(colors.white)
        c_pdf.setFont("Helvetica-Bold", 8)
        cx3 = x_margin
        for hi2, hw2 in zip(["#", "APELLIDOS Y NOMBRES"] + ["NOTA  (LIT)"] * len(areas) + ["PROMEDIO"],
                             col_widths):
            c_pdf.drawCentredString(cx3 + hw2 / 2, y2 - HEADER_H2 + 8, hi2)
            cx3 += hw2
        return y2 - HEADER_H2

    def _pie_pagina(c_pdf, grado, periodo, pagina):
        """Dibuja pie de página con línea separadora"""
        c_pdf.setStrokeColor(colors.HexColor("#cbd5e1"))
        c_pdf.setLineWidth(0.5)
        c_pdf.line(x_margin, PIE_H + 12, w - x_margin, PIE_H + 12)
        c_pdf.setFont("Helvetica", 7)
        c_pdf.setFillColor(colors.HexColor("#64748b"))
        c_pdf.drawString(x_margin, PIE_H, f"I.E.P. YACHAY — Ranking {grado} — {periodo}")
        c_pdf.drawString(x_margin, PIE_H - 10, "Doc. referencial. El consolidado oficial lo registra el/la docente.  |  SIGE - Sistema Integral de Gestion Educativa")
        c_pdf.drawRightString(w - x_margin, PIE_H, f"Pág. {pagina}  |  Generado: {hora_peru().strftime('%d/%m/%Y %H:%M')}")

    num_pagina = 1

    # ── Filas de datos ───────────────────────────────────────────────────
    for idx, fila in enumerate(ranking_filas):
        # Salto de página si no queda espacio (respetando pie)
        if y < PIE_H + ROW_H + 15:
            _pie_pagina(c_pdf, grado, periodo, num_pagina)
            c_pdf.showPage()
            num_pagina += 1
            # Cabecera compacta en página nueva
            c_pdf.setFillColor(colors.HexColor("#001e7c"))
            c_pdf.rect(0, h - 28, w, 28, fill=1, stroke=0)
            c_pdf.setFillColor(colors.white)
            c_pdf.setFont("Helvetica-Bold", 10)
            c_pdf.drawCentredString(w / 2, h - 12, f"RANKING — {grado} — {periodo}  (continuacion)")
            c_pdf.setFont("Helvetica", 7)
            c_pdf.setFillColor(colors.HexColor("#93c5fd"))
            c_pdf.drawCentredString(w / 2, h - 24, f"Página {num_pagina}")
            y = h - 30
            y = _dibujar_cabecera_tabla(c_pdf, y)

        # Fondo de fila
        if idx == 0:
            bg = "#fef3c7"   # oro
        elif idx == 1:
            bg = "#e5e7eb"   # plata
        elif idx == 2:
            bg = "#fed7aa"   # bronce
        elif idx % 2 == 0:
            bg = "#f8fafc"
        else:
            bg = "#ffffff"
        c_pdf.setFillColor(colors.HexColor(bg))
        c_pdf.rect(x_margin, y - ROW_H, table_w, ROW_H, fill=1, stroke=0)

        # Línea separadora horizontal sutil
        c_pdf.setStrokeColor(colors.HexColor("#e2e8f0"))
        c_pdf.setLineWidth(0.3)
        c_pdf.line(x_margin, y - ROW_H, x_margin + table_w, y - ROW_H)

        cx = x_margin

        # Puesto / medalla — texto limpio sin emojis
        if idx == 0:
            medalla_txt = "1. ORO"
            c_pdf.setFillColor(colors.HexColor("#92400e"))
        elif idx == 1:
            medalla_txt = "2. PLATA"
            c_pdf.setFillColor(colors.HexColor("#374151"))
        elif idx == 2:
            medalla_txt = "3. BRONCE"
            c_pdf.setFillColor(colors.HexColor("#7c2d12"))
        else:
            medalla_txt = f"#{idx + 1}"
            c_pdf.setFillColor(colors.HexColor("#1e293b"))
        c_pdf.setFont("Helvetica-Bold", 8 if idx < 3 else 9)
        c_pdf.drawCentredString(cx + col_w_pos / 2, y - ROW_H + 5, medalla_txt)
        cx += col_w_pos

        # Nombre
        nombre_full = str(fila.get('Nombre', ''))
        max_ch = int(col_w_nom / 5.8)
        nombre_display = nombre_full[:max_ch] + ("." if len(nombre_full) > max_ch else "")
        c_pdf.setFont("Helvetica-Bold" if idx < 3 else "Helvetica", 9)
        c_pdf.setFillColor(colors.HexColor("#0f172a"))
        c_pdf.drawString(cx + 4, y - ROW_H + 5, nombre_display)
        cx += col_w_nom

        # Notas por área con fondo de color tenue y separador
        for ai, area in enumerate(areas):
            _, light_hex = PALETA_AREAS[ai % len(PALETA_AREAS)]
            # Franja de color muy suave de fondo
            c_pdf.setFillColor(colors.HexColor(light_hex))
            c_pdf.setFillAlpha(0.35)
            c_pdf.rect(cx, y - ROW_H, col_w_area, ROW_H, fill=1, stroke=0)
            c_pdf.setFillAlpha(1.0)

            nota_v = float(fila.get(area, 0) or 0)
            lit_v  = nota_a_letra(nota_v) if nota_v > 0 else "-"
            col_n  = color_semaforo(lit_v) if nota_v > 0 else "#94a3b8"
            c_pdf.setFillColor(colors.HexColor(col_n))
            c_pdf.setFont("Helvetica-Bold", 9)
            nota_txt = f"{nota_v:.0f}  ({lit_v})" if nota_v > 0 else "—"
            c_pdf.drawCentredString(cx + col_w_area / 2, y - ROW_H + 5, nota_txt)

            # Separador vertical entre áreas
            c_pdf.setStrokeColor(colors.HexColor("#94a3b8"))
            c_pdf.setLineWidth(0.5)
            c_pdf.line(cx, y, cx, y - ROW_H)

            cx += col_w_area

        # Separador antes de promedio
        c_pdf.setStrokeColor(colors.HexColor("#475569"))
        c_pdf.setLineWidth(0.8)
        c_pdf.line(cx, y, cx, y - ROW_H)

        # Promedio destacado
        prom = float(fila.get('Promedio', 0) or 0)
        prom_color = "#15803d" if prom >= 14 else "#dc2626" if prom < 11 else "#d97706"
        c_pdf.setFillColor(colors.HexColor("#f0fdf4" if prom >= 14 else "#fff1f2" if prom < 11 else "#fffbeb"))
        c_pdf.rect(cx, y - ROW_H, col_w_prom, ROW_H, fill=1, stroke=0)
        c_pdf.setFillColor(colors.HexColor(prom_color))
        c_pdf.setFont("Helvetica-Bold", 11 if idx < 3 else 10)
        c_pdf.drawCentredString(cx + col_w_prom / 2, y - ROW_H + 5, f"{prom:.1f}")

        y -= ROW_H

    # ── Cuadro: estudiantes sin nota / no presentados ────────────────────
    if sin_nota:
        BOX_PAD   = 8
        FILA_SN   = 14
        BOX_H     = BOX_PAD * 2 + 28 + (len(sin_nota) // (3 if len(sin_nota) > 10 else 2 if len(sin_nota) > 5 else 1) + 1) * FILA_SN + 8
        # ¿Cabe en esta página?
        if y - BOX_H < PIE_H + 15:
            _pie_pagina(c_pdf, grado, periodo, num_pagina)
            c_pdf.showPage()
            num_pagina += 1
            c_pdf.setFillColor(colors.HexColor("#001e7c"))
            c_pdf.rect(0, h - 28, w, 28, fill=1, stroke=0)
            c_pdf.setFillColor(colors.white)
            c_pdf.setFont("Helvetica-Bold", 10)
            c_pdf.drawCentredString(w / 2, h - 18, f"RANKING — {grado} — {periodo}  (continuación)")
            y = h - 38

        # Borde y fondo del cuadro
        box_y = y - 10
        c_pdf.setStrokeColor(colors.HexColor("#dc2626"))
        c_pdf.setLineWidth(1.2)
        c_pdf.setFillColor(colors.HexColor("#fff1f2"))
        c_pdf.roundRect(x_margin, box_y - BOX_H, table_w, BOX_H, 6, fill=1, stroke=1)

        # Título del cuadro
        c_pdf.setFillColor(colors.HexColor("#dc2626"))
        c_pdf.setFont("Helvetica-Bold", 9)
        c_pdf.drawString(x_margin + BOX_PAD,
                         box_y - BOX_PAD - 11,
                         f"SIN NOTA / NO SE PRESENTARON  ({len(sin_nota)} estudiante{'s' if len(sin_nota) != 1 else ''})")

        # Línea separadora dentro del cuadro
        c_pdf.setStrokeColor(colors.HexColor("#fca5a5"))
        c_pdf.setLineWidth(0.6)
        c_pdf.line(x_margin + BOX_PAD,
                   box_y - BOX_PAD - 18,
                   x_margin + table_w - BOX_PAD,
                   box_y - BOX_PAD - 18)

        # Lista de alumnos en columnas (3 columnas si hay muchos)
        ncols_sn = 3 if len(sin_nota) > 10 else 2 if len(sin_nota) > 5 else 1
        col_sn_w = (table_w - BOX_PAD * 2) / ncols_sn
        c_pdf.setFillColor(colors.HexColor("#7f1d1d"))
        c_pdf.setFont("Helvetica", 8)
        for si, alumno_sn in enumerate(sin_nota):
            col_idx = si % ncols_sn
            row_idx = si // ncols_sn
            sx = x_margin + BOX_PAD + col_idx * col_sn_w
            sy = box_y - BOX_PAD - 30 - row_idx * FILA_SN
            nombre_sn = str(alumno_sn.get('Nombre', ''))
            max_ch_sn = int(col_sn_w / 5.5)
            nombre_sn = nombre_sn[:max_ch_sn] + ("." if len(nombre_sn) > max_ch_sn else "")
            c_pdf.drawString(sx, sy, f"• {nombre_sn}")

    # ── Línea de firma — Coordinador Académico PREU ──────────────────────
    firma_y = PIE_H + 42
    firma_w = 180
    firma_x = w / 2 - firma_w / 2
    c_pdf.setStrokeColor(colors.HexColor("#1e293b"))
    c_pdf.setLineWidth(0.8)
    c_pdf.line(firma_x, firma_y, firma_x + firma_w, firma_y)
    c_pdf.setFillColor(colors.HexColor("#1e293b"))
    c_pdf.setFont("Helvetica-Bold", 7)
    c_pdf.drawCentredString(w / 2, firma_y - 9, "COORDINADOR ACADEMICO PREU")
    c_pdf.setFont("Helvetica", 6)
    c_pdf.setFillColor(colors.HexColor("#64748b"))
    c_pdf.drawCentredString(w / 2, firma_y - 18, "Firma y Sello")

    # ── Pie de última página ─────────────────────────────────────────────
    _pie_pagina(c_pdf, grado, periodo, num_pagina)

    c_pdf.save()
    buffer.seek(0)
    return buffer.getvalue()


# ================================================================
# REPORTE INTEGRAL POR ESTUDIANTE — PDF COMPLETO
# ================================================================

def generar_reporte_integral_pdf(nombre, dni, grado, notas, asistencia, config):
    """Genera PDF completo: notas + asistencia + semáforo + recomendaciones"""
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    w, h = A4

    # === PÁGINA 1: Datos + Notas ===
    # Encabezado con colores
    c.setFillColor(colors.HexColor("#1a56db"))
    c.rect(0, h-80, w, 80, fill=True)
    c.setFillColor(colors.white)
    c.setFont("Helvetica-Bold", 22)
    c.drawCentredString(w/2, h-35, "INFORME INTEGRAL DEL ESTUDIANTE")
    c.setFont("Helvetica", 10)
    c.drawCentredString(w/2, h-55, f"I.E.P. ALTERNATIVO YACHAY — Año Escolar {config.get('anio', 2026)}")
    c.drawCentredString(w/2, h-70, f"Chinchero, Cusco — Perú")

    # Datos del estudiante
    c.setFillColor(colors.black)
    y = h - 110
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, f"Estudiante: {nombre}")
    c.drawString(350, y, f"DNI: {dni}")
    y -= 18
    c.drawString(50, y, f"Grado: {grado}")
    c.drawString(350, y, f"Fecha: {fecha_peru_str()}")

    # Línea separadora
    y -= 12
    c.setStrokeColor(colors.HexColor("#1a56db"))
    c.setLineWidth(2)
    c.line(50, y, w-50, y)
    y -= 25

    # === SECCIÓN: NOTAS ===
    c.setFont("Helvetica-Bold", 13)
    c.drawString(50, y, "📊 REGISTRO DE CALIFICACIONES")
    y -= 22

    if notas:
        # Header de tabla
        c.setFillColor(colors.HexColor("#1e293b"))
        c.rect(45, y-2, w-90, 16, fill=True)
        c.setFillColor(colors.white)
        c.setFont("Helvetica-Bold", 8)
        col_x = [50, 160, 300, 370, 420, 480]
        for i, header in enumerate(["Área", "Bimestre", "Nota", "Literal", "Semáforo", "Fecha"]):
            c.drawString(col_x[i], y+2, header)
        y -= 18

        c.setFillColor(colors.black)
        c.setFont("Helvetica", 7)
        promedios = {}
        for n in notas:
            area = str(n.get('area', ''))
            nota_val = float(n.get('nota', 0))
            literal = nota_a_letra(nota_val)
            col = color_semaforo(literal)

            if area not in promedios:
                promedios[area] = []
            promedios[area].append(nota_val)

            c.drawString(col_x[0], y, area[:22])
            c.drawString(col_x[1], y, str(n.get('bimestre', ''))[:15])
            c.drawString(col_x[2], y, f"{nota_val}/20")
            c.drawString(col_x[3], y, literal)
            c.setFillColor(colors.HexColor(col))
            c.circle(col_x[4]+12, y+3, 5, fill=True)
            c.setFillColor(colors.black)
            c.drawString(col_x[5], y, str(n.get('fecha', ''))[:10])
            y -= 13
            if y < 120:
                c.showPage()
                y = h - 60
                c.setFont("Helvetica", 7)

        # Resumen por áreas
        y -= 15
        if y < 200:
            c.showPage()
            y = h - 60

        c.setFont("Helvetica-Bold", 11)
        c.setFillColor(colors.black)
        c.drawString(50, y, "📈 PROMEDIOS POR ÁREA")
        y -= 20

        total_all = []
        for area, notas_area in promedios.items():
            prom = round(sum(notas_area)/len(notas_area), 1)
            total_all.append(prom)
            lit = nota_a_letra(prom)
            col = color_semaforo(lit)

            c.setFont("Helvetica-Bold", 8)
            c.drawString(55, y, f"{area}:")
            c.drawString(220, y, f"{prom}/20 ({lit})")

            # Barra de progreso
            c.setFillColor(colors.HexColor("#e2e8f0"))
            c.rect(320, y-2, 150, 12, fill=True)
            c.setFillColor(colors.HexColor(col))
            c.rect(320, y-2, max(1, (prom/20)*150), 12, fill=True)
            c.setFillColor(colors.black)
            y -= 16
            if y < 100:
                c.showPage()
                y = h - 60

        # Promedio general
        if total_all:
            prom_gen = round(sum(total_all)/len(total_all), 1)
            lit_gen = nota_a_letra(prom_gen)
            col_gen = color_semaforo(lit_gen)
            y -= 10
            c.setFont("Helvetica-Bold", 14)
            c.drawString(55, y, f"PROMEDIO GENERAL: {prom_gen}/20")

            # Círculo semáforo grande
            c.setFillColor(colors.HexColor(col_gen))
            c.circle(350, y+5, 18, fill=True)
            c.setFillColor(colors.white)
            c.setFont("Helvetica-Bold", 14)
            c.drawCentredString(350, y, lit_gen)
            c.setFillColor(colors.black)
    else:
        c.setFont("Helvetica", 10)
        c.drawString(55, y, "Sin calificaciones registradas.")

    # === Asistencia + Recomendaciones (misma página si cabe) ===
    # Solo nueva página si queda poco espacio
    if y < 250:
        c.showPage()
        y = h - 50
    else:
        y -= 25

    c.setFont("Helvetica-Bold", 13)
    c.setFillColor(colors.black)
    c.drawString(50, y, f"REGISTRO DE ASISTENCIA — {nombre}")
    y -= 22

    if asistencia:
        c.setFont("Helvetica", 8)
        total_dias = len(asistencia)
        c.drawString(55, y, f"Total de días registrados: {total_dias}")
        y -= 15
        for fecha_a, datos_a in sorted(asistencia.items())[:60]:
            entrada = datos_a.get('entrada', '—')
            salida = datos_a.get('salida', '—')
            c.drawString(55, y, f"📅 {fecha_a}: Entrada {entrada} | Salida {salida}")
            y -= 12
            if y < 100:
                c.showPage()
                y = h - 50
                c.setFont("Helvetica", 8)
    else:
        c.setFont("Helvetica", 10)
        c.drawString(55, y, "Sin registros de asistencia.")

    # Recomendaciones
    y -= 25
    if y < 200:
        c.showPage()
        y = h - 50

    c.setFont("Helvetica-Bold", 13)
    c.drawString(50, y, "📝 RECOMENDACIONES PEDAGÓGICAS")
    y -= 20
    c.setFont("Helvetica", 8)

    if notas and total_all:
        lit_gen = nota_a_letra(prom_gen)
        info_lit = ESCALA_MINEDU.get(lit_gen, {})
        c.drawString(55, y, f"• Nivel de logro: {lit_gen} — {info_lit.get('nombre', '')}")
        y -= 13
        c.drawString(55, y, f"  {info_lit.get('desc', '')}")
        y -= 18

        recomendaciones = {
            'AD': [
                "Mantener el excelente ritmo académico con retos adicionales.",
                "Participar en concursos académicos para potenciar sus habilidades.",
                "Puede ayudar como tutor de compañeros con dificultades.",
                "Orientación vocacional: explorar carreras afines a sus fortalezas.",
            ],
            'A': [
                "Reforzar las áreas con menor puntaje para alcanzar nivel destacado.",
                "Establecer metas semanales de estudio.",
                "Lectura diaria de 30 minutos para fortalecer comprensión.",
                "Continuar con el buen hábito de estudio.",
            ],
            'B': [
                "Requiere acompañamiento permanente del docente y la familia.",
                "Sesiones de refuerzo en las áreas con menor calificación.",
                "Horario de estudio fijo en casa con supervisión del apoderado.",
                "Reuniones quincenales padres-docente para seguimiento.",
            ],
            'C': [
                "ATENCIÓN PRIORITARIA: Plan de recuperación inmediata.",
                "Evaluación psicopedagógica recomendada.",
                "Sesiones de refuerzo diarias con material adaptado.",
                "Reunión urgente con padres para establecer compromisos.",
                "Considerar factores emocionales o externos que afecten el aprendizaje.",
            ]
        }
        for rec in recomendaciones.get(lit_gen, []):
            c.drawString(55, y, f"• {rec}")
            y -= 12

    # Escala MINEDU
    y -= 20
    c.setFont("Helvetica-Bold", 9)
    c.drawString(50, y, "ESCALA DE CALIFICACIÓN — MINEDU Perú:")
    y -= 14
    c.setFont("Helvetica", 7)
    for sigla, info in ESCALA_MINEDU.items():
        c.setFillColor(colors.HexColor(info['color']))
        c.circle(60, y+3, 4, fill=True)
        c.setFillColor(colors.black)
        c.drawString(70, y, f"{sigla} ({info['min']}-{info['max']}): {info['nombre']}")
        y -= 11

    # Pie de página
    c.setFont("Helvetica-Oblique", 7)
    c.drawCentredString(w/2, 25, f"YACHAY PRO — Sistema de Gestión Educativa © {hora_peru().year}")
    c.drawCentredString(w/2, 15, "Documento generado automáticamente — Válido sin firma ni sello")

    c.save()
    buf.seek(0)
    return buf


# ================================================================
# FUNCIÓN PRINCIPAL
# ================================================================

# ================================================================
# MÓDULO: AULA VIRTUAL — MATERIAL DOCENTE (Estilo Classroom)
# ================================================================
ARCHIVO_MATERIALES = "materiales_docente.json"
ARCHIVO_EXAMENES_SEM = "examenes_semanales.json"
ARCHIVO_FICHAS_REGISTRO = "fichas_registro.json"


def _cargar_fichas_registro():
    """Carga el registro de fichas subidas por docentes."""
    # Intentar desde GS primero
    gs = _gs()
    if gs:
        try:
            ws = gs._get_hoja('config')
            if ws:
                data = ws.get_all_records()
                for row in data:
                    if str(row.get('clave', '')) == 'fichas_registro':
                        contenido = str(row.get('valor', ''))
                        if contenido and contenido.strip() not in ('', '[]'):
                            return json.loads(contenido)
        except Exception:
            pass
    # Fallback: archivo local
    if Path(ARCHIVO_FICHAS_REGISTRO).exists():
        try:
            with open(ARCHIVO_FICHAS_REGISTRO, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            pass
    return []


def _guardar_ficha_registro(ficha):
    """Guarda una ficha en el registro (local + GS)."""
    fichas = _cargar_fichas_registro()
    ficha['id'] = f"FICHA-{int(time.time())}"
    ficha['fecha_subida'] = hora_peru().strftime('%Y-%m-%d %H:%M')
    fichas.append(ficha)
    # Guardar local
    with open(ARCHIVO_FICHAS_REGISTRO, 'w', encoding='utf-8') as f:
        json.dump(fichas, f, indent=2, ensure_ascii=False)
    # Sync a GS
    try:
        gs = _gs()
        if gs:
            ws = gs._get_hoja('config')
            if ws:
                data = ws.get_all_records()
                fila = None
                for i, row in enumerate(data):
                    if str(row.get('clave', '')) == 'fichas_registro':
                        fila = i + 2
                        break
                contenido = json.dumps(fichas, ensure_ascii=False)
                if len(contenido) < 45000:
                    if fila:
                        ws.update_cell(fila, 2, contenido)
                    else:
                        ws.append_row(['fichas_registro', contenido])
    except Exception:
        pass
    return ficha

AREAS_POR_NIVEL = {
    "INICIAL": ["Comunicación", "Matemática", "Personal Social",
                "Ciencia y Tecnología", "Psicomotriz",
                "Castellano como segunda lengua", "Educación Física",
                "Inglés", "Tutoría"],
    "PRIMARIA": ["Comunicación", "Matemática", "Personal Social",
                 "Ciencia y Tecnología", "Educación Religiosa",
                 "Arte y Cultura", "Educación Física", "Inglés",
                 "Castellano como segunda lengua", "Tutoría",
                 "Gramática", "Razonamiento Verbal", "Redacción",
                 "Expresión Oral", "Aritmética", "Geometría",
                 "Razonamiento Matemático", "Álgebra", "Física",
                 "Química", "Biología", "Lenguaje",
                 "Competencia Lingüística", "Historia", "Geografía",
                 "Trigonometría"],
    "SECUNDARIA": ["Comunicación", "Matemática", "Ciencia y Tecnología",
                    "Ciencias Sociales", "Desarrollo Personal, Ciudadanía y Cívica",
                    "Educación para el Trabajo", "Educación Religiosa",
                    "Arte y Cultura", "Educación Física", "Inglés",
                    "Castellano como segunda lengua", "Tutoría"],
    "PREUNIVERSITARIO": ["Razonamiento Matemático", "Aritmética", "Álgebra",
                          "Geometría", "Trigonometría", "Lenguaje", "Literatura",
                          "Razonamiento Verbal", "Historia del Perú",
                          "Historia Universal", "Geografía", "Economía",
                          "Filosofía y Lógica", "Psicología", "Educación Cívica",
                          "Biología", "Química", "Física", "Anatomía",
                          "Educación Física", "Inglés"],
}

TIPOS_EVALUACION = [
    "Evaluación Semanal", "Evaluación Mensual", "Evaluación Bimestral",
    "Examen Parcial", "Examen Final", "Examen de Recuperación",
    "Examen de Nivelación", "Práctica Calificada", "Control de Lectura",
]


def _inicio_escolar(anio=None):
    """Primer día de clases: primer día hábil de marzo"""
    if anio is None:
        anio = hora_peru().year
    d = date(anio, 3, 1)
    while d.weekday() >= 5:  # sáb/dom → avanzar al lunes
        d += timedelta(days=1)
    return d


def _semana_escolar_actual():
    hoy = hora_peru().date()
    inicio = _inicio_escolar(hoy.year)
    if hoy < inicio:
        return 0
    return ((hoy - inicio).days // 7) + 1


def _rango_semana(semana_num, anio=None):
    if anio is None:
        anio = hora_peru().year
    inicio = _inicio_escolar(anio)
    dias_a_lunes = inicio.weekday()
    primer_lunes = inicio - timedelta(days=dias_a_lunes)
    lunes = primer_lunes + timedelta(weeks=semana_num - 1)
    viernes = lunes + timedelta(days=4)
    return lunes, viernes


def _semanas_del_mes(mes, anio=None):
    """Retorna [(semana_num, lunes, viernes)] que caen en el mes dado"""
    if anio is None:
        anio = hora_peru().year
    resultado = []
    for sem in range(1, 45):
        lun, vie = _rango_semana(sem, anio)
        if lun.month == mes or vie.month == mes:
            resultado.append((sem, lun, vie))
        elif lun.month > mes and vie.month > mes:
            break
    return resultado


# Hora límite puntualidad docente: antes de 8:05 = puntual, después = tardanza
# ── HORARIOS DE PUNTUALIDAD ──────────────────────────────────────
HORARIOS = {
    'normal': {'limite': '08:05', 'nombre': '☀️ Normal (8:05am)', 'minutos': 8*60+5},
    'invierno': {'limite': '08:15', 'nombre': '❄️ Invierno (8:15am)', 'minutos': 8*60+15},
}


def _horario_activo():
    """Retorna horario activo - persistente"""
    if "horario_escolar" not in st.session_state:
        try:
            if Path("config_horario.json").exists():
                with open("config_horario.json","r") as fh:
                    st.session_state.horario_escolar = json.load(fh).get("horario","normal")
        except Exception:
            st.session_state.horario_escolar = "normal"
    return st.session_state.get("horario_escolar", "normal")


def _guardar_horario(horario):
    """Guarda horario en archivo persistente"""
    st.session_state.horario_escolar = horario
    try:
        import json as jj
        with open("config_horario.json","w") as fh:
            jj.dump({"horario": horario, "modificado_por": st.session_state.get("usuario_actual",""), "fecha": fecha_peru_str()}, fh)
    except Exception:
        pass



def _limite_minutos():
    return HORARIOS[_horario_activo()]['minutos']


def _es_tardanza(hora_str):
    """Determina si la hora de entrada es tardanza según horario activo"""
    try:
        h, m = hora_str.split(':')[:2]
        minutos = int(h) * 60 + int(m)
        return minutos > _limite_minutos()
    except Exception:
        return False


# Alias para compatibilidad
def _es_tardanza_docente(hora_str):
    return _es_tardanza(hora_str)


def _comprimir_imagen_aula(img_bytes, max_size=400, quality=65):
    try:
        img = Image.open(io.BytesIO(img_bytes))
        if img.mode == 'RGBA':
            img = img.convert('RGB')
        w, h = img.size
        if max(w, h) > max_size:
            ratio = max_size / max(w, h)
            img = img.resize((int(w * ratio), int(h * ratio)), Image.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, format='JPEG', quality=quality, optimize=True)
        return buf.getvalue()
    except Exception:
        return img_bytes


def _img_a_base64(img_bytes):
    return base64.b64encode(img_bytes).decode('utf-8')


def _base64_a_bytes(b64_str):
    return base64.b64decode(b64_str)


def _areas_del_docente():
    info = st.session_state.get('docente_info', {}) or {}
    nivel = str(info.get('nivel', 'PRIMARIA')).upper()
    grado = str(info.get('grado', ''))
    # Secundaria/Preu: incluir todas las áreas de ambos niveles
    es_sec = ('SECUNDARIA' in nivel or 'PREUNIVERSITARIO' in nivel
              or 'GRUPO' in grado or 'Sec' in grado
              or grado in ('ALL_NIVELES', 'ALL_SEC_PREU', 'ALL_SECUNDARIA'))
    if es_sec:
        areas_sec = AREAS_POR_NIVEL.get("SECUNDARIA", [])
        areas_preu = AREAS_POR_NIVEL.get("PREUNIVERSITARIO", [])
        # Combinar sin duplicados, manteniendo orden
        todas = list(areas_sec)
        for a in areas_preu:
            if a not in todas:
                todas.append(a)
        return todas
    for key in AREAS_POR_NIVEL:
        if key in nivel:
            return AREAS_POR_NIVEL[key]
    return AREAS_POR_NIVEL.get("PRIMARIA", ["Comunicación", "Matemática"])


def _grados_del_docente():
    """Retorna la lista de grados disponibles para el docente."""
    info = st.session_state.get('docente_info', {}) or {}
    nivel = str(info.get('nivel', 'PRIMARIA')).upper()
    grado = str(info.get('grado', ''))
    
    # ALL_NIVELES: acceso a TODOS (Ed. Física, Inglés, etc.)
    if grado == 'ALL_NIVELES':
        return TODOS_LOS_GRADOS
    
    es_sec = ('SECUNDARIA' in nivel or 'PREUNIVERSITARIO' in nivel
              or 'GRUPO' in grado or 'Sec' in grado
              or grado in ('ALL_NIVELES', 'ALL_SEC_PREU', 'ALL_SECUNDARIA'))
    if es_sec:
        grados_sec = NIVELES_GRADOS.get('SECUNDARIA', [])
        grados_preu = NIVELES_GRADOS.get('PREUNIVERSITARIO', [])
        return grados_sec + grados_preu
    elif grado and grado != 'N/A':
        return [grado]
    return GRADOS_OPCIONES


def _grados_para_selector(key_prefix="gs"):
    """
    Muestra un selector de grado filtrado según el rol del usuario.
    - Admin/Directivo: todos los grados
    - Docente Secundaria/Preu: solo sus grados (sec + preu)
    - Docente Inicial/Primaria: solo su grado asignado (mostrado como info, sin selector)
    Devuelve el grado seleccionado o None si no aplica.
    """
    rol = st.session_state.get('rol', '')
    info = st.session_state.get('docente_info', {}) or {}
    nivel = str(info.get('nivel', '')).upper()
    grado = str(info.get('grado', ''))

    if rol in ['admin', 'directivo']:
        return st.selectbox("🎓 Grado:", GRADOS_OPCIONES, key=f"{key_prefix}_grado")

    # ALL_NIVELES: todos los grados (Ed. Física, Inglés, etc.)
    if grado == 'ALL_NIVELES':
        return st.selectbox("🎓 Grado:", TODOS_LOS_GRADOS, key=f"{key_prefix}_grado")

    es_sec = ('SECUNDARIA' in nivel or 'PREUNIVERSITARIO' in nivel
              or 'GRUPO' in grado or grado in ('ALL_NIVELES', 'ALL_SEC_PREU', 'ALL_SECUNDARIA'))
    if es_sec:
        grados_disp = _grados_del_docente()
        return st.selectbox("🎓 Grado:", grados_disp, key=f"{key_prefix}_grado")
    else:
        if grado and grado != 'N/A':
            st.info(f"🎓 **Grado asignado: {grado}**")
            return grado
        else:
            st.warning("No tienes grado asignado. Contacta al administrador.")
            return None


# ---- Almacenamiento Materiales ----
def _cargar_materiales():
    gs = _gs()
    if gs:
        try:
            ws = gs._get_hoja('materiales')
            if ws:
                data = ws.get_all_records()
                materiales = []
                for row in data:
                    try:
                        mat = json.loads(str(row.get('data_json', '{}')))
                        mat['id'] = row.get('id', '')
                        materiales.append(mat)
                    except Exception:
                        pass
                return materiales
        except Exception:
            pass
    if Path(ARCHIVO_MATERIALES).exists():
        try:
            with open(ARCHIVO_MATERIALES, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            pass
    return []


def _guardar_material(material):
    materiales = _cargar_materiales()
    material['id'] = f"MAT-{int(time.time())}"
    material['fecha_creacion'] = hora_peru().strftime('%Y-%m-%d %H:%M')
    materiales.append(material)
    gs = _gs()
    if gs:
        try:
            ws = gs._get_hoja('materiales')
            if ws:
                ws.append_row([
                    material['id'], material.get('docente', ''),
                    material.get('grado', ''), material.get('semana', 0),
                    material.get('area', ''), material.get('fecha_creacion', ''),
                    json.dumps(material, ensure_ascii=False)
                ], value_input_option='RAW')
        except Exception:
            pass
    try:
        with open(ARCHIVO_MATERIALES, 'w', encoding='utf-8') as f:
            json.dump(materiales, f, indent=2, ensure_ascii=False)
    except Exception:
        pass
    return material['id']


# ---- Almacenamiento Exámenes ----
def _cargar_examenes_sem():
    gs = _gs()
    if gs:
        try:
            ws = gs._get_hoja('examenes')
            if ws:
                data = ws.get_all_records()
                examenes = []
                for row in data:
                    try:
                        ex = json.loads(str(row.get('data_json', '{}')))
                        ex['id'] = row.get('id', '')
                        examenes.append(ex)
                    except Exception:
                        pass
                return examenes
        except Exception:
            pass
    if Path(ARCHIVO_EXAMENES_SEM).exists():
        try:
            with open(ARCHIVO_EXAMENES_SEM, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            pass
    return []


def _guardar_pregunta_examen(pregunta):
    examenes = _cargar_examenes_sem()
    pregunta['id'] = f"EX-{int(time.time())}-{len(examenes)}"
    pregunta['fecha_creacion'] = hora_peru().strftime('%Y-%m-%d %H:%M')
    examenes.append(pregunta)
    gs = _gs()
    if gs:
        try:
            ws = gs._get_hoja('examenes')
            if ws:
                ws.append_row([
                    pregunta['id'], pregunta.get('docente', ''),
                    pregunta.get('grado', ''), pregunta.get('semana', 0),
                    pregunta.get('area', ''), pregunta.get('fecha_creacion', ''),
                    json.dumps(pregunta, ensure_ascii=False)
                ], value_input_option='RAW')
        except Exception:
            pass
    try:
        with open(ARCHIVO_EXAMENES_SEM, 'w', encoding='utf-8') as f:
            json.dump(examenes, f, indent=2, ensure_ascii=False)
    except Exception:
        pass
    return pregunta['id']


def _dibujar_escudo(c, x, y, alto_deseado):
    """Dibuja escudo manteniendo proporción correcta"""
    if not Path("escudo_upload.png").exists():
        return
    try:
        from PIL import Image as PILImage
        img = PILImage.open("escudo_upload.png")
        iw, ih = img.size
        ratio = iw / ih
        ancho = alto_deseado * ratio
        c.drawImage("escudo_upload.png", x, y, ancho, alto_deseado, mask='auto')
    except Exception:
        pass


# ---- PDF Material Docente ----
def _pdf_encabezado_material(c, w, h, config, semana, area, titulo, grado, docente):
    # ── Barra azul superior ──────────────────────────────────────────────
    c.setFillColor(colors.HexColor("#001e7c"))
    c.rect(0, h - 15, w, 15, fill=1, stroke=0)

    # ── Escudo IZQUIERDA y DERECHA (proporción correcta, alto=65) ───────
    ALTO_ESC = 65
    esc_izq = "escudo_upload.png"
    esc_der = "escudo2_upload.png" if Path("escudo2_upload.png").exists() else "escudo_upload.png"
    if Path(esc_izq).exists():
        try:
            from PIL import Image as PILImage
            img = PILImage.open(esc_izq)
            iw, ih = img.size
            ratio = iw / ih
            ancho_esc = ALTO_ESC * ratio
            c.drawImage(esc_izq, 18, h - 12 - ALTO_ESC, ancho_esc, ALTO_ESC, mask='auto')
        except Exception:
            pass
    if Path(esc_der).exists():
        try:
            from PIL import Image as PILImage
            img2 = PILImage.open(esc_der)
            iw2, ih2 = img2.size
            ancho_esc2 = ALTO_ESC * (iw2 / ih2)
            _alto_der = 80
            _ancho_der = _alto_der * (iw2 / ih2)
            c.drawImage(esc_der, w - 18 - _ancho_der, h - 12 - _alto_der, _ancho_der, _alto_der, mask='auto')
        except Exception:
            pass

    # ── Textos institucionales centrados ────────────────────────────────
    c.setFillColor(colors.HexColor("#001e7c"))
    c.setFont("Helvetica-Bold", 7.5)
    c.drawCentredString(w / 2, h - 28, "MINISTERIO DE EDUCACIÓN — DRE CUSCO — PIONEROS EN LA EDUCACION DE CALIDAD")
    c.setFont("Helvetica-Bold", 11)
    c.drawCentredString(w / 2, h - 43, "I.E.P. YACHAY — CHINCHERO")
    frase = config.get('frase', '')
    if frase:
        c.setFont("Helvetica-Oblique", 7)
        c.drawCentredString(w / 2, h - 56, f'"{frase}"')

    # ── Cuadro de datos con bordes redondeados ───────────────────────────
    c.setStrokeColor(colors.HexColor("#1a56db"))
    c.setLineWidth(1.5)
    c.roundRect(25, h - 148, w - 50, 68, 8, fill=0)
    c.setFillColor(colors.black)
    c.setFont("Helvetica", 9)
    c.drawString(35,    h - 90,  f"GRADO: {grado}")
    c.drawString(220,   h - 90,  f"SEMANA: {semana}")
    c.drawRightString(w - 35, h - 90,  f"FECHA: {hora_peru().strftime('%d/%m/%Y')}")
    c.drawString(35,    h - 107, f"ÁREA: {area}")
    c.drawString(280,   h - 107, f"DOCENTE: {docente}")
    c.drawString(35,    h - 128, "ALUMNO(A): _______________________________________________")
    c.drawRightString(w - 35, h - 128, "N° ______")

    # ── Título del documento en azul + línea ────────────────────────────
    c.setFillColor(colors.HexColor("#1a56db"))
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(w / 2, h - 167, titulo)
    c.setStrokeColor(colors.HexColor("#1a56db"))
    c.setLineWidth(2)
    c.line(60, h - 174, w - 60, h - 174)


def _pdf_pie_material(c, w, grado, area, semana, pagina=None):
    # Marca de agua GRANDE en cada página (sin duplicar con encabezado)
    if Path("escudo_upload.png").exists():
        try:
            from PIL import Image as PILImage
            img = PILImage.open("escudo_upload.png")
            iw, ih = img.size
            ratio = iw / ih
            mw = 420; mh = mw / ratio
            c.saveState()
            c.setFillAlpha(0.35)
            c.drawImage("escudo_upload.png", w/2 - mw/2, A4[1]/2 - mh/2, mw, mh, mask='auto')
            c.restoreState()
        except Exception:
            pass
    c.setStrokeColor(colors.HexColor("#1a56db"))
    c.setLineWidth(0.5)
    c.line(30, 35, w - 30, 35)
    c.setFont("Helvetica", 7)
    c.setFillColor(colors.HexColor("#6b7280"))
    c.drawString(30, 23, f"I.E.P. YACHAY — {grado} — {area} — Semana {semana}")
    if pagina:
        c.drawCentredString(w / 2, 23, f"— {pagina} —")
    c.drawRightString(w - 30, 23, f"Generado: {hora_peru().strftime('%d/%m/%Y %H:%M')}")
    c.setFillColor(colors.black)


def _pdf_encabezado_cont(c, w, h, grado, area, docente, semana):
    """Encabezado compacto para páginas de continuación"""
    c.setFillColor(colors.HexColor("#001e7c"))
    c.rect(0, h - 12, w, 12, fill=1, stroke=0)
    c.setFillColor(colors.HexColor("#374151"))
    c.setFont("Helvetica-Bold", 7)
    c.drawString(30, h - 25, f"I.E.P. YACHAY — {grado} — {area}")
    c.drawRightString(w - 30, h - 25, f"Docente: {docente} — Semana {semana}")
    c.setStrokeColor(colors.HexColor("#d1d5db"))
    c.setLineWidth(0.5)
    c.line(30, h - 30, w - 30, h - 30)


def _generar_pdf_material(material, config):
    buffer = io.BytesIO()
    c_pdf = canvas.Canvas(buffer, pagesize=A4)
    w, h = A4
    semana = material.get('semana', '')
    area = material.get('area', '')
    titulo = material.get('titulo', 'Material de Trabajo')
    grado = material.get('grado', '')
    docente = material.get('docente_nombre', '')
    # Si el nombre almacenado parece un username (sin espacio), resolver nombre completo
    if not docente or (docente and ' ' not in docente and len(docente) < 20):
        docente = _nombre_completo_docente() or docente
    bloques = material.get('bloques', [])
    pagina = [1]  # mutable counter
    LM = 35  # margen izquierdo estrecho
    RM = w - 35  # margen derecho estrecho

    def nueva_pagina():
        _pdf_pie_material(c_pdf, w, grado, area, semana, pagina[0])
        c_pdf.showPage()
        pagina[0] += 1
        _pdf_encabezado_cont(c_pdf, w, h, grado, area, docente, semana)
        return h - 45

    _pdf_encabezado_material(c_pdf, w, h, config, semana, area, titulo, grado, docente)
    y_pos = h - 230
    num_actividad = 1

    for bloque in bloques:
        tipo = bloque.get('tipo', 'texto')
        contenido = bloque.get('contenido', '')
        subtitulo = bloque.get('subtitulo', '')

        if y_pos < 100:
            y_pos = nueva_pagina()

        if subtitulo:
            c_pdf.setFont("Helvetica-Bold", 11)
            c_pdf.setFillColor(colors.HexColor("#1a56db"))
            c_pdf.drawString(LM + 5, y_pos, f"  {num_actividad}. {subtitulo}")
            c_pdf.setFillColor(colors.black)
            y_pos -= 20
            num_actividad += 1

        if tipo == 'texto' and contenido:
            lineas_raw = contenido.split('\n')
            for linea_r in lineas_raw:
                linea_r = linea_r.strip()
                if not linea_r:
                    y_pos -= 6
                    continue
                if y_pos < 70:
                    y_pos = nueva_pagina()
                if linea_r.startswith('## '):
                    c_pdf.setFont("Helvetica-Bold", 12)
                    c_pdf.setFillColor(colors.HexColor("#1a56db"))
                    c_pdf.drawString(LM, y_pos, linea_r[3:].strip())
                    c_pdf.setFillColor(colors.black)
                    y_pos -= 20
                elif linea_r.startswith('### '):
                    c_pdf.setFont("Helvetica-Bold", 10)
                    c_pdf.drawString(LM + 5, y_pos, linea_r[4:].strip())
                    y_pos -= 16
                elif linea_r.startswith('**') and linea_r.endswith('**'):
                    c_pdf.setFont("Helvetica-Bold", 10)
                    for sub_l in textwrap.wrap(linea_r.strip('*').strip(), width=90):
                        if y_pos < 70:
                            y_pos = nueva_pagina()
                        c_pdf.drawString(LM + 5, y_pos, sub_l)
                        y_pos -= 13
                else:
                    c_pdf.setFont("Helvetica", 10)
                    for sub_l in textwrap.wrap(linea_r, width=90):
                        if y_pos < 70:
                            y_pos = nueva_pagina()
                        c_pdf.drawString(LM + 5, y_pos, sub_l.replace('**', ''))
                        y_pos -= 13
            y_pos -= 6

        elif tipo == 'imagen' and bloque.get('imagen_b64'):
            try:
                img_bytes = _base64_a_bytes(bloque['imagen_b64'])
                img = Image.open(io.BytesIO(img_bytes))
                img_w, img_h = img.size
                max_w = w - 80
                max_h = 280
                ratio = min(max_w / img_w, max_h / img_h, 1.0)
                disp_w = img_w * ratio
                disp_h = img_h * ratio
                if y_pos - disp_h < 70:
                    y_pos = nueva_pagina()
                tmp = f"tmp_mat_{int(time.time())}.jpg"
                if img.mode == 'RGBA':
                    img = img.convert('RGB')
                img.save(tmp, 'JPEG', quality=85)
                c_pdf.drawImage(tmp, (w - disp_w) / 2, y_pos - disp_h, disp_w, disp_h)
                y_pos -= disp_h + 12
                try:
                    os.remove(tmp)
                except Exception:
                    pass
            except Exception:
                c_pdf.setFont("Helvetica-Oblique", 9)
                c_pdf.drawString(LM + 5, y_pos, "[Imagen no disponible]")
                y_pos -= 15

        elif tipo == 'instruccion' and contenido:
            c_pdf.setStrokeColor(colors.HexColor("#2563eb"))
            c_pdf.setFillColor(colors.HexColor("#eff6ff"))
            box_h = max(30, len(textwrap.wrap(contenido, width=85)) * 13 + 16)
            if y_pos - box_h < 70:
                y_pos = nueva_pagina()
            c_pdf.roundRect(LM, y_pos - box_h, w - LM * 2, box_h, 5, fill=1)
            c_pdf.setFillColor(colors.HexColor("#1e40af"))
            c_pdf.setFont("Helvetica-Bold", 9)
            c_pdf.drawString(LM + 10, y_pos - 14, "INSTRUCCIONES:")
            c_pdf.setFont("Helvetica", 9)
            c_pdf.setFillColor(colors.black)
            lineas = textwrap.wrap(contenido, width=85)
            ty = y_pos - 28
            for linea in lineas:
                c_pdf.drawString(LM + 10, ty, linea)
                ty -= 13
            y_pos -= box_h + 10

        elif tipo == 'ejercicio' and contenido:
            c_pdf.setFont("Helvetica", 10)
            lineas = contenido.split('\n')
            for linea in lineas:
                if y_pos < 70:
                    y_pos = nueva_pagina()
                linea = linea.strip()
                if linea:
                    c_pdf.drawString(LM + 5, y_pos, linea)
                    y_pos -= 13
                    if bloque.get('espacio_resolver', True):
                        for _ in range(2):
                            if y_pos < 70:
                                break
                            c_pdf.setStrokeColor(colors.HexColor("#d1d5db"))
                            c_pdf.setDash(3, 3)
                            c_pdf.line(LM + 5, y_pos, RM, y_pos)
                            c_pdf.setDash()
                            c_pdf.setStrokeColor(colors.black)
                            y_pos -= 16
                        y_pos -= 4

    _pdf_pie_material(c_pdf, w, grado, area, semana, pagina[0])
    c_pdf.save()
    buffer.seek(0)
    return buffer.getvalue()


# ---- PDF Examen Semanal ----
def _generar_pdf_examen_semanal(preguntas_por_area, config, grado, semana, titulo_examen):
    buffer = io.BytesIO()
    c_pdf = canvas.Canvas(buffer, pagesize=A4)
    w, h = A4
    lunes, viernes = _rango_semana(semana)

    # ── ENCABEZADO OFICIAL ────────────────────────────────────────────────
    # Barra azul superior
    c_pdf.setFillColor(colors.HexColor("#001e7c"))
    c_pdf.rect(0, h - 15, w, 15, fill=1, stroke=0)

    # Logo escudo (izquierda)
    if Path("escudo_upload.png").exists():
        try:
            c_pdf.drawImage("escudo_upload.png", 25, h - 88, 60, 60, mask='auto')
        except Exception:
            pass

    # Textos institucionales centrados
    c_pdf.setFillColor(colors.HexColor("#001e7c"))
    c_pdf.setFont("Helvetica-Bold", 7.5)
    c_pdf.drawCentredString(w / 2, h - 28, "MINISTERIO DE EDUCACIÓN — DRE CUSCO — PIONEROS EN LA EDUCACION DE CALIDAD")
    c_pdf.setFont("Helvetica-Bold", 11)
    c_pdf.drawCentredString(w / 2, h - 43, "I.E.P. YACHAY — CHINCHERO")
    frase = config.get('frase', '')
    if frase:
        c_pdf.setFont("Helvetica-Oblique", 7)
        c_pdf.drawCentredString(w / 2, h - 55, f'"{frase}"')

    # Cuadro de datos (igual que ficha)
    c_pdf.setStrokeColor(colors.HexColor("#1a56db"))
    c_pdf.setLineWidth(1.5)
    c_pdf.roundRect(25, h - 148, w - 50, 68, 8, fill=0)
    c_pdf.setFillColor(colors.black)
    c_pdf.setFont("Helvetica", 9)
    areas_nombres_str = " / ".join(list(preguntas_por_area.keys()))
    c_pdf.drawString(35, h - 90,  f"GRADO: {grado}")
    c_pdf.drawString(220, h - 90, f"SEMANA: {semana}")
    c_pdf.drawRightString(w - 35, h - 90, f"FECHA: {hora_peru().strftime('%d/%m/%Y')}")
    c_pdf.drawString(35, h - 107, f"ÁREA: {areas_nombres_str[:55]}")
    c_pdf.drawRightString(w - 35, h - 107, f"Del {lunes.strftime('%d/%m')} al {viernes.strftime('%d/%m/%Y')}")
    c_pdf.drawString(35, h - 128, "ALUMNO(A): _______________________________________________")
    c_pdf.drawRightString(w - 35, h - 128, "N° ______")

    # Título del examen
    c_pdf.setFillColor(colors.HexColor("#1a56db"))
    c_pdf.setFont("Helvetica-Bold", 15)
    c_pdf.drawCentredString(w / 2, h - 170, (titulo_examen or "EVALUACIÓN SEMANAL").upper())
    c_pdf.setStrokeColor(colors.HexColor("#1a56db"))
    c_pdf.setLineWidth(2)
    c_pdf.line(60, h - 177, w - 60, h - 177)

    # Marca de agua
    if Path("escudo_upload.png").exists():
        try:
            c_pdf.saveState()
            c_pdf.setFillAlpha(0.35)
            c_pdf.drawImage("escudo_upload.png", w/2 - 100, h/2 - 100, 200, 200, mask='auto')
            c_pdf.restoreState()
        except Exception:
            pass

    # Datos del alumno + instrucciones
    y_datos = h - 200
    c_pdf.setFillColor(colors.black)
    c_pdf.setFont("Helvetica", 9)
    c_pdf.drawString(35, y_datos, "SECCIÓN: ________    N° DE ORDEN: ________")
    c_pdf.drawRightString(w - 35, y_datos, f"Año Escolar {config.get('anio', hora_peru().year)}")

    # Instrucciones en cuadro
    y_datos -= 10
    c_pdf.setStrokeColor(colors.HexColor("#d1d5db"))
    c_pdf.roundRect(35, y_datos - 30, w - 70, 26, 5, fill=0)
    c_pdf.setFont("Helvetica-Bold", 8)
    c_pdf.setFillColor(colors.HexColor("#1e40af"))
    c_pdf.drawString(45, y_datos - 15, "INSTRUCCIONES:")
    c_pdf.setFont("Helvetica", 8)
    c_pdf.setFillColor(colors.black)
    c_pdf.drawString(130, y_datos - 15,
                     "Lee cada pregunta. Marca con X la alternativa correcta. No se permiten borrones.")

    y_pos = y_datos - 50
    num_pregunta_global = 1
    total_preguntas = sum(len(ps) for ps in preguntas_por_area.values())
    c_pdf.setFont("Helvetica-Bold", 8)
    c_pdf.setFillColor(colors.HexColor("#1a56db"))
    areas_resumen = "   |   ".join([f"{a}: {len(ps)} preg." for a, ps in preguntas_por_area.items()])
    c_pdf.drawCentredString(w / 2, y_pos, f"TOTAL: {total_preguntas} preguntas — {areas_resumen}")
    c_pdf.setFillColor(colors.black)
    y_pos -= 20
    pagina_num = [1]
    LM_EX = 30  # Margen izquierdo estrecho
    RM_EX = w - 30  # Margen derecho estrecho
    ANCHO_TEXTO = 95  # Ancho de wrap más amplio

    def _nueva_pagina_examen():
        """Página nueva con encabezado compacto y número"""
        # Pie de página actual
        c_pdf.setFont("Helvetica", 7)
        c_pdf.setFillColor(colors.HexColor("#9ca3af"))
        c_pdf.drawString(LM_EX, 18, f"I.E.P. YACHAY — {grado} — Semana {semana}")
        c_pdf.drawCentredString(w/2, 18, f"— Pág. {pagina_num[0]} —")
        c_pdf.drawRightString(RM_EX, 18, hora_peru().strftime('%d/%m/%Y'))
        c_pdf.setFillColor(colors.black)
        c_pdf.showPage()
        pagina_num[0] += 1
        # Encabezado compacto página 2+
        c_pdf.setFillColor(colors.HexColor("#001e7c"))
        c_pdf.rect(0, h - 12, w, 12, fill=1, stroke=0)
        c_pdf.setFillColor(colors.HexColor("#374151"))
        c_pdf.setFont("Helvetica-Bold", 7)
        titulo_corto = (titulo_examen or "EVALUACIÓN SEMANAL").upper()
        c_pdf.drawString(LM_EX, h - 24, f"I.E.P. YACHAY — {grado} — {titulo_corto}")
        c_pdf.drawRightString(RM_EX, h - 24, f"Semana {semana} — Pág. {pagina_num[0]}")
        c_pdf.setStrokeColor(colors.HexColor("#d1d5db"))
        c_pdf.setLineWidth(0.5)
        c_pdf.line(LM_EX, h - 28, RM_EX, h - 28)
        # Marca de agua
        if Path("escudo_upload.png").exists():
            try:
                c_pdf.saveState()
                c_pdf.setFillAlpha(0.04)
                c_pdf.drawImage("escudo_upload.png", w/2-80, h/2-80, 160, 160, mask='auto')
                c_pdf.restoreState()
            except Exception:
                pass
        return h - 40

    # PREGUNTAS POR ÁREA
    for area, preguntas in preguntas_por_area.items():
        if not preguntas:
            continue
        if y_pos < 120:
            y_pos = _nueva_pagina_examen()

        c_pdf.setFillColor(colors.HexColor("#1a56db"))
        c_pdf.roundRect(LM_EX, y_pos - 20, RM_EX - LM_EX, 22, 4, fill=1)
        c_pdf.setFillColor(colors.white)
        c_pdf.setFont("Helvetica-Bold", 11)
        c_pdf.drawCentredString(w / 2, y_pos - 14, f"{area.upper()}")
        c_pdf.setFillColor(colors.black)
        y_pos -= 35

        for pregunta in preguntas:
            texto_p = pregunta.get('texto', '')
            opciones = pregunta.get('opciones', {})
            tiene_imagen = bool(pregunta.get('imagen_b64'))

            lineas_texto = textwrap.wrap(texto_p, width=ANCHO_TEXTO)
            espacio = len(lineas_texto) * 13 + len(opciones) * 15 + 25 + (120 if tiene_imagen else 0)

            if y_pos - espacio < 50:
                y_pos = _nueva_pagina_examen()

            c_pdf.setFont("Helvetica-Bold", 10)
            c_pdf.setFillColor(colors.HexColor("#1a56db"))
            c_pdf.drawString(LM_EX + 5, y_pos, f"{num_pregunta_global}.")
            c_pdf.setFillColor(colors.black)
            c_pdf.setFont("Helvetica", 10)
            x_t = LM_EX + 25
            for linea in lineas_texto:
                c_pdf.drawString(x_t, y_pos, linea)
                y_pos -= 13
            y_pos -= 3

            if tiene_imagen:
                try:
                    img_bytes = _base64_a_bytes(pregunta['imagen_b64'])
                    img = Image.open(io.BytesIO(img_bytes))
                    if img.mode == 'RGBA':
                        img = img.convert('RGB')
                    iw, ih = img.size
                    ratio = min((RM_EX - LM_EX - 40) / iw, 150 / ih, 1.0)
                    dw = iw * ratio
                    dh = ih * ratio
                    if y_pos - dh < 50:
                        y_pos = _nueva_pagina_examen()
                    tmp = f"tmp_ex_{int(time.time())}.jpg"
                    img.save(tmp, 'JPEG', quality=80)
                    c_pdf.drawImage(tmp, (w - dw) / 2, y_pos - dh, dw, dh)
                    y_pos -= dh + 8
                    try:
                        os.remove(tmp)
                    except Exception:
                        pass
                except Exception:
                    pass

            c_pdf.setFont("Helvetica", 10)
            opciones_orden = ['a', 'b', 'c', 'd']
            for letra in opciones_orden:
                txt = opciones.get(letra, '')
                if not txt:
                    continue
                if y_pos < 50:
                    y_pos = _nueva_pagina_examen()
                c_pdf.circle(LM_EX + 35, y_pos + 3, 5, stroke=1, fill=0)
                c_pdf.setFont("Helvetica-Bold", 9)
                c_pdf.drawString(LM_EX + 43, y_pos, f"{letra.upper()})")
                c_pdf.setFont("Helvetica", 9)
                txt_disp = txt[:80] + ('...' if len(txt) > 80 else '')
                c_pdf.drawString(LM_EX + 58, y_pos, txt_disp)
                y_pos -= 15
            y_pos -= 10
            num_pregunta_global += 1

    # Pie de última página
    c_pdf.setFont("Helvetica", 7)
    c_pdf.setFillColor(colors.HexColor("#9ca3af"))
    c_pdf.drawString(LM_EX, 18, f"I.E.P. YACHAY — {grado} — Semana {semana}")
    c_pdf.drawCentredString(w/2, 18, f"— Pág. {pagina_num[0]} —")
    c_pdf.drawRightString(RM_EX, 18, hora_peru().strftime('%d/%m/%Y'))
    c_pdf.setFillColor(colors.black)

    # CLAVE DE RESPUESTAS — Página nueva
    c_pdf.showPage()
    c_pdf.setFont("Helvetica-Bold", 16)
    c_pdf.setFillColor(colors.HexColor("#dc2626"))
    c_pdf.drawCentredString(w / 2, h - 60, "CLAVE DE RESPUESTAS — SOLO DIRECTOR")
    c_pdf.setFillColor(colors.black)
    c_pdf.setFont("Helvetica-Bold", 10)
    c_pdf.drawCentredString(w / 2, h - 80, f"{grado} — Semana {semana}")
    y_c = h - 110
    num = 1
    for area_c, preguntas_c in preguntas_por_area.items():
        c_pdf.setFont("Helvetica-Bold", 11)
        c_pdf.setFillColor(colors.HexColor("#1a56db"))
        c_pdf.drawString(60, y_c, f"{area_c}")
        c_pdf.setFillColor(colors.black)
        y_c -= 18
        c_pdf.setFont("Helvetica", 10)
        for preg in preguntas_c:
            resp = preg.get('respuesta_correcta', '?').upper()
            c_pdf.drawString(80, y_c, f"{num}. {resp}")
            if num % 5 == 0:
                y_c -= 16
            else:
                # Poner en la misma fila
                pass
            num += 1
            y_c -= 16
            if y_c < 60:
                c_pdf.showPage()
                y_c = h - 60
        y_c -= 10

    c_pdf.save()
    buffer.seek(0)
    return buffer.getvalue()


# ================================================================
# TAB: AULA VIRTUAL — MATERIAL DOCENTE
# ================================================================

def tab_material_docente(config):
    """REGISTRAR FICHA — Dos opciones: generar con diseño O cargar tal cual"""
    st.subheader("📄 Registrar Ficha")

    usuario = st.session_state.get('usuario_actual', '')
    info_doc = st.session_state.get('docente_info', {}) or {}
    nombre_doc = _nombre_completo_docente()

    fichas_dir = Path("fichas")
    fichas_dir.mkdir(exist_ok=True)

    tab1, tab2, tab3 = st.tabs([
        "📝 Generar Ficha con Diseño",
        "📤 Cargar Ficha (PDF/Word)",
        "📥 Mis Fichas"
    ])

    # ── TAB 1: GENERAR CON DISEÑO (DOCX → PDF oficial) ─────────────────────
    with tab1:
        st.markdown("### 📝 Generar Ficha con Diseño Oficial")
        st.info("💡 Sube tu Word (.docx) sin diseño. El sistema le agregará "
                "la portada y formato oficial del colegio automáticamente.")

        col1, col2, col3 = st.columns(3)
        with col1:
            titulo_f1 = st.text_input("📝 Título:", placeholder="Ej: Lógica y Formas",
                                       key="ficha_titulo_diseno")
        with col2:
            grado_f1 = _grados_para_selector("ficha_dis")
        with col3:
            semana_f1 = st.number_input("📅 Semana N°:", 1, 52,
                                         int(hora_peru().strftime('%V')), key="ficha_semana_dis")

        col4, col5 = st.columns(2)
        with col4:
            area_f1 = st.text_input("📚 Área/Curso:", placeholder="Ej: Personal Social",
                                     key="ficha_area_diseno")
        with col5:
            tipo_doc = st.selectbox("📄 Tipo:", ["FICHA", "PRÁCTICA", "EXAMEN",
                                                  "ACTIVIDAD", "TALLER"], key="ficha_tipo_dis")

        archivo_docx = st.file_uploader("📎 Subir Word (.docx):",
                                         type=['docx'],
                                         key="upload_ficha_docx_dis")

        if archivo_docx and titulo_f1 and grado_f1 and area_f1:
            st.caption(f"📁 **{archivo_docx.name}** ({archivo_docx.size/1024:.1f} KB)")

            if st.button("🎨 GENERAR PDF CON DISEÑO", type="primary",
                         use_container_width=True, key="btn_gen_diseno"):
                try:
                    with st.spinner("📄 Procesando Word y generando PDF con diseño..."):
                        bloques = _leer_docx(archivo_docx.getvalue())
                        if bloques:
                            pdf_bytes = _generar_pdf_desde_docx(
                                bloques, config, nombre_doc, str(grado_f1),
                                area_f1, str(semana_f1), titulo_f1, tipo_doc)
                            nombre_arch = (f"ficha_{usuario}_{grado_f1}_{fecha_peru_str()}_"
                                           f"{titulo_f1[:25]}.pdf")
                            nombre_arch = nombre_arch.replace(' ', '_').replace('/', '-').replace(':', '-')
                            with open(fichas_dir / nombre_arch, 'wb') as f:
                                f.write(pdf_bytes)
                            _guardar_ficha_registro({
                                'titulo': titulo_f1, 'area': area_f1,
                                'grado': str(grado_f1), 'semana': semana_f1,
                                'mes': list(MESES_ESCOLARES.values())[hora_peru().month - 1],
                                'docente': usuario, 'docente_nombre': nombre_doc,
                                'archivo': nombre_arch, 'tipo': 'pdf (con diseño)',
                            })
                            st.success("✅ ¡Ficha generada con diseño oficial!")
                            st.balloons()
                            st.download_button("📥 Descargar PDF", pdf_bytes,
                                               nombre_arch, "application/pdf",
                                               use_container_width=True, key="dl_ficha_dis")
                        else:
                            st.error("⚠️ No se pudo leer el archivo Word. Verifique el formato.")
                except Exception as e:
                    st.error(f"❌ Error: {str(e)}")
        else:
            st.caption("⚠️ Completa todos los campos y sube el archivo Word (.docx)")

    # ── TAB 2: CARGAR TAL CUAL (PDF o Word ya con diseño) ───────────────────
    with tab2:
        st.markdown("### 📤 Cargar Ficha (ya tiene diseño)")
        st.info("💡 Sube tu ficha en PDF o Word que **ya tiene** los diseños del colegio. "
                "Se guardará tal cual sin modificar, solo para registro y verificación.")

        col1, col2, col3 = st.columns(3)
        with col1:
            titulo_ficha = st.text_input("📝 Título:", placeholder="Ej: Lógica y Formas",
                                          key="ficha_titulo")
        with col2:
            grado_ficha = _grados_para_selector("ficha")
        with col3:
            semana_ficha = st.number_input("📅 Semana N°:", 1, 52,
                                           int(hora_peru().strftime('%V')), key="ficha_semana")

        col4, col5, col6 = st.columns(3)
        with col4:
            area_ficha = st.text_input("📚 Área/Curso:", placeholder="Ej: Personal Social",
                                       key="ficha_area")
        with col5:
            mes_ficha = st.selectbox("📅 Mes:", list(MESES_ESCOLARES.values()),
                                      index=hora_peru().month - 1, key="ficha_mes")
        with col6:
            docente_ficha = st.text_input("👤 Docente:", value=nombre_doc, key="ficha_docente")

        archivo_ficha = st.file_uploader("📎 Subir ficha (.pdf o .docx):",
                                         type=['pdf', 'docx'],
                                         key="upload_ficha_archivo")

        if archivo_ficha and titulo_ficha and grado_ficha and area_ficha:
            ext = archivo_ficha.name.rsplit('.', 1)[-1].lower()
            st.caption(f"📁 **{archivo_ficha.name}** ({archivo_ficha.size/1024:.1f} KB) — .{ext}")

            if st.button("💾 GUARDAR FICHA", type="primary",
                         use_container_width=True, key="btn_guardar_ficha"):
                try:
                    archivo_bytes = archivo_ficha.getvalue()
                    nombre_arch = (f"ficha_{usuario}_{grado_ficha}_{fecha_peru_str()}_"
                                   f"{titulo_ficha[:25]}.{ext}")
                    nombre_arch = nombre_arch.replace(' ', '_').replace('/', '-').replace(':', '-')
                    with open(fichas_dir / nombre_arch, 'wb') as f:
                        f.write(archivo_bytes)
                    _guardar_ficha_registro({
                        'titulo': titulo_ficha, 'area': area_ficha,
                        'grado': str(grado_ficha), 'semana': semana_ficha,
                        'mes': mes_ficha, 'docente': usuario,
                        'docente_nombre': docente_ficha,
                        'archivo': nombre_arch, 'tipo': ext,
                    })
                    st.success(f"✅ Ficha guardada correctamente ({ext.upper()})")
                    st.balloons()
                    mime = "application/pdf" if ext == 'pdf' else \
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    st.download_button(f"📥 Descargar {ext.upper()}", archivo_bytes,
                                       nombre_arch, mime,
                                       use_container_width=True, key="dl_ficha_guardada")
                except Exception as e:
                    st.error(f"❌ Error: {str(e)}")
        else:
            st.caption("⚠️ Completa todos los campos y sube el archivo para continuar.")
    
    # ── TAB 3: MIS FICHAS ──────────────────────────────────────────────────
    with tab3:
        st.markdown("### 📥 Mis Fichas Guardadas")
        fichas_pdf = sorted(fichas_dir.glob(f"ficha_{usuario}_*.pdf"), key=lambda x: x.stat().st_mtime, reverse=True)
        fichas_docx = sorted(fichas_dir.glob(f"ficha_{usuario}_*.docx"), key=lambda x: x.stat().st_mtime, reverse=True)
        fichas_usuario = fichas_pdf + fichas_docx

        if fichas_usuario:
            st.success(f"📚 {len(fichas_usuario)} ficha(s) guardada(s)")
            for ficha in fichas_usuario:
                partes = ficha.stem.split('_')
                grado_f = partes[2] if len(partes) > 2 else 'N/A'
                fecha_f = partes[3] if len(partes) > 3 else 'N/A'
                titulo_f = '_'.join(partes[4:]).replace('_', ' ')[:50] if len(partes) > 4 else ficha.stem
                mime = "application/pdf" if ficha.suffix == '.pdf' else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                with st.expander(f"📄 {titulo_f} — {grado_f} ({fecha_f})"):
                    col_a, col_b, col_c = st.columns([3, 1, 1])
                    with col_a:
                        st.caption(f"Archivo: {ficha.name} | {ficha.stat().st_size/1024:.1f} KB")
                    with col_b:
                        with open(ficha, 'rb') as f:
                            st.download_button("⬇️ Descargar", f.read(), ficha.name, mime,
                                               key=f"dl_{ficha.name}", use_container_width=True)
                    with col_c:
                        if st.button("🗑️ Eliminar", key=f"del_{ficha.name}", type="primary"):
                            ficha.unlink()
                            st.success("Eliminada")
                            time.sleep(0.5)
                            st.rerun()
        else:
            st.info("📭 No has guardado fichas aún. Ve a 'Cargar Ficha' para subir la primera.")


# ---- Funciones para leer Word y convertir a PDF oficial ----

def _leer_docx(file_bytes):
    """Lee un archivo .docx y extrae contenido como lista de bloques."""
    if not HAS_DOCX:
        return []
    doc = DocxDocument(io.BytesIO(file_bytes))
    bloques = []
    contadores_lista = {}  # nivel -> contador para listas numeradas
    for para in doc.paragraphs:
        texto = para.text.strip()
        if not texto:
            bloques.append({'tipo': 'vacio'})
            continue
        style_name = (para.style.name or '').lower()
        is_bold = para.runs and all(r.bold for r in para.runs if r.text.strip())
        is_heading = 'heading' in style_name or 'título' in style_name
        font_size = None
        if para.runs:
            for r in para.runs:
                if r.font.size:
                    font_size = r.font.size.pt
                    break

        # Detectar listas numeradas
        is_lista_num = any(x in style_name for x in ['list number', 'lista con número', 'list paragraph'])
        # Detectar listas con viñetas
        is_lista_bullet = any(x in style_name for x in ['list bullet', 'list paragraph', 'lista con viñeta'])
        # Detectar por formato XML si tiene numeración
        if not is_lista_num and para._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr') is not None:
            numPr = para._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
            ilvl = numPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl')
            nivel = int(ilvl.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 0)) if ilvl is not None else 0
            contadores_lista[nivel] = contadores_lista.get(nivel, 0) + 1
            # Resetear niveles más profundos
            for k in list(contadores_lista.keys()):
                if k > nivel:
                    del contadores_lista[k]
            prefijo = "  " * nivel + f"{contadores_lista[nivel]}. "
            bloques.append({'tipo': 'lista_num', 'contenido': prefijo + texto, 'nivel': nivel})
            continue

        if is_heading or 'heading 1' in style_name:
            contadores_lista.clear()
            bloques.append({'tipo': 'titulo', 'contenido': texto})
        elif 'heading 2' in style_name or (is_bold and font_size and font_size >= 13):
            contadores_lista.clear()
            bloques.append({'tipo': 'subtitulo', 'contenido': texto})
        elif is_bold:
            bloques.append({'tipo': 'negrita', 'contenido': texto})
        else:
            contadores_lista.clear()
            bloques.append({'tipo': 'texto', 'contenido': texto})
    # Extraer imágenes
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_data = rel.target_part.blob
                img_b64 = base64.b64encode(img_data).decode('utf-8')
                bloques.append({'tipo': 'imagen', 'imagen_b64': img_b64})
            except Exception:
                pass
    return bloques


def _generar_pdf_desde_docx(bloques_docx, config, nombre_doc, grado, area, semana, titulo, tipo_doc="FICHA"):
    """Genera PDF con formato oficial del colegio desde contenido de Word — 2 columnas."""
    from reportlab.lib.enums import TA_JUSTIFY
    from reportlab.platypus import Paragraph
    from reportlab.lib.styles import ParagraphStyle
    
    buffer = io.BytesIO()
    c_pdf = canvas.Canvas(buffer, pagesize=A4)
    w, h = A4

    # ENCABEZADO OFICIAL
    _pdf_encabezado_material(c_pdf, w, h, config, semana, area, titulo, grado, nombre_doc)

    # Tipo de documento (pequeño, gris)
    c_pdf.setFont("Helvetica", 8)
    c_pdf.setFillColor(colors.HexColor("#6b7280"))
    c_pdf.drawRightString(w - 35, h - 192, f"{tipo_doc} — Docente: {nombre_doc}")
    c_pdf.setFillColor(colors.black)

    # ── CONFIGURACIÓN 2 COLUMNAS ─────────────────────────────────────────
    MARGEN_IZQ = 30
    MARGEN_DER = 30
    GAP_COLS   = 14          # espacio entre columnas
    CONTENT_W  = w - MARGEN_IZQ - MARGEN_DER
    COL_W      = (CONTENT_W - GAP_COLS) / 2
    COL1_X     = MARGEN_IZQ
    COL2_X     = MARGEN_IZQ + COL_W + GAP_COLS
    Y_TOP      = h - 205     # inicio de contenido
    Y_BOTTOM   = 45          # margen inferior

    # Línea divisoria entre columnas
    def _dibujar_linea_col(c, y_top, y_bot):
        c.setStrokeColor(colors.HexColor("#e5e7eb"))
        c.setLineWidth(0.5)
        c.line(COL2_X - GAP_COLS / 2, y_bot, COL2_X - GAP_COLS / 2, y_top)

    col_actual = 0   # 0 = izquierda, 1 = derecha
    y = Y_TOP
    pagina = [1]

    def x_col():
        return COL1_X if col_actual == 0 else COL2_X

    def max_chars():
        # Ancho de columna en caracteres aproximados (Helvetica 10pt)
        return int(COL_W / 5.5)

    def nueva_columna_o_pagina():
        nonlocal col_actual, y
        if col_actual == 0:
            col_actual = 1
            y = Y_TOP
        else:
            _dibujar_linea_col(c_pdf, Y_TOP, Y_BOTTOM)
            _pdf_pie_material(c_pdf, w, grado, area, semana, pagina[0])
            c_pdf.showPage()
            pagina[0] += 1
            col_actual = 0
            # En páginas de continuación, el contenido empieza MUY arriba
            y = h - 25  # Solo 25 pts del borde superior
            # Dibujar encabezado mínimo
            c_pdf.setFillColor(colors.HexColor("#001e7c"))
            c_pdf.rect(0, h - 12, w, 12, fill=1, stroke=0)
            c_pdf.setFont("Helvetica-Bold", 7)
            c_pdf.setFillColor(colors.HexColor("#6b7280"))
            c_pdf.drawString(30, h - 22, f"I.E.P. YACHAY — {grado} — {area}")
            c_pdf.drawRightString(w - 30, h - 22, f"Docente: {nombre_doc} — Semana {semana}")
            c_pdf.setStrokeColor(colors.HexColor("#d1d5db"))
            c_pdf.setLineWidth(0.5)
            c_pdf.line(30, h - 24, w - 30, h - 24)
            c_pdf.setFillColor(colors.black)
            _dibujar_linea_col(c_pdf, h - 25, Y_BOTTOM)

    _dibujar_linea_col(c_pdf, Y_TOP, Y_BOTTOM)

    for bloque in bloques_docx:
        tipo = bloque.get('tipo', '')
        contenido = bloque.get('contenido', '')

        if tipo == 'vacio':
            y -= 7
            if y < Y_BOTTOM:
                nueva_columna_o_pagina()
            continue

        if tipo == 'titulo':
            if y < Y_BOTTOM + 30:
                nueva_columna_o_pagina()
            c_pdf.setFont("Helvetica-Bold", 13)
            c_pdf.setFillColor(colors.HexColor("#1a56db"))
            from reportlab.platypus import Paragraph
            from reportlab.lib.styles import ParagraphStyle
            style_tit = ParagraphStyle('titulo', fontName='Helvetica-Bold', fontSize=13,
                                      leading=16, alignment=TA_JUSTIFY)
            p_tit = Paragraph(contenido, style_tit)
            w_tit, h_tit = p_tit.wrap(COL_W - 8, 500)
            if y - h_tit < Y_BOTTOM:
                nueva_columna_o_pagina()
            p_tit.drawOn(c_pdf, x_col(), y - h_tit)
            y -= (h_tit + 8)
            c_pdf.setFillColor(colors.black)

        elif tipo == 'subtitulo':
            if y < Y_BOTTOM + 20:
                nueva_columna_o_pagina()
            c_pdf.setFont("Helvetica-Bold", 10)
            c_pdf.setFillColor(colors.HexColor("#1e40af"))
            from reportlab.platypus import Paragraph
            from reportlab.lib.styles import ParagraphStyle
            style_sub = ParagraphStyle('subtitulo', fontName='Helvetica-Bold', fontSize=10,
                                      leading=13, alignment=TA_JUSTIFY)
            p_sub = Paragraph(contenido, style_sub)
            w_sub, h_sub = p_sub.wrap(COL_W - 8, 500)
            if y - h_sub < Y_BOTTOM:
                nueva_columna_o_pagina()
            p_sub.drawOn(c_pdf, x_col(), y - h_sub)
            y -= (h_sub + 6)
            c_pdf.setFillColor(colors.black)

        elif tipo == 'negrita':
            c_pdf.setFont("Helvetica-Bold", 9)
            from reportlab.platypus import Paragraph
            from reportlab.lib.styles import ParagraphStyle
            style_neg = ParagraphStyle('negrita', fontName='Helvetica-Bold', fontSize=9,
                                      leading=12, alignment=TA_JUSTIFY)
            p_neg = Paragraph(contenido, style_neg)
            w_neg, h_neg = p_neg.wrap(COL_W - 8, 500)
            if y - h_neg < Y_BOTTOM:
                nueva_columna_o_pagina()
            p_neg.drawOn(c_pdf, x_col(), y - h_neg)
            y -= (h_neg + 4)

        elif tipo in ('texto', 'lista_num'):
            c_pdf.setFont("Helvetica", 9)
            from reportlab.platypus import Paragraph
            from reportlab.lib.styles import ParagraphStyle
            indent = bloque.get('nivel', 0) * 10 if tipo == 'lista_num' else 0
            style_txt = ParagraphStyle('texto_just', fontName='Helvetica', fontSize=9,
                                      leading=12, alignment=TA_JUSTIFY,
                                      leftIndent=indent)
            p_txt = Paragraph(contenido, style_txt)
            w_txt, h_txt = p_txt.wrap(COL_W - 8 - indent, 500)
            if y - h_txt < Y_BOTTOM:
                nueva_columna_o_pagina()
            p_txt.drawOn(c_pdf, x_col() + indent, y - h_txt)
            y -= h_txt

        elif tipo == 'imagen' and bloque.get('imagen_b64'):
            try:
                img_bytes = base64.b64decode(bloque['imagen_b64'])
                img = Image.open(io.BytesIO(img_bytes))
                if img.mode == 'RGBA':
                    img = img.convert('RGB')
                img_w, img_h = img.size
                max_w_img = COL_W - 4
                max_h_img = 180
                ratio = min(max_w_img / img_w, max_h_img / img_h, 1.0)
                dw, dh = img_w * ratio, img_h * ratio
                if y - dh < Y_BOTTOM:
                    nueva_columna_o_pagina()
                tmp = f"tmp_docx_{int(time.time())}.jpg"
                img.save(tmp, 'JPEG', quality=80)
                c_pdf.drawImage(tmp, x_col(), y - dh, dw, dh)
                y -= dh + 10
                try:
                    os.remove(tmp)
                except Exception:
                    pass
            except Exception:
                pass

    _dibujar_linea_col(c_pdf, Y_TOP, Y_BOTTOM)
    _pdf_pie_material(c_pdf, w, grado, area, semana, pagina[0])
    c_pdf.save()
    buffer.seek(0)
    return buffer.getvalue()


def _vista_docente_material(config, usuario, nombre_doc, grado_doc, semana_actual):
    tab1, tab2, tab3 = st.tabs(["📤 Crear Ficha", "📄 Subir Word", "📋 Mi Material"])
    with tab1:
        st.markdown("### 📝 Crear Ficha de Trabajo")
        st.markdown("""
        <div style="background: #eff6ff; border-radius: 10px; padding: 12px; 
                    border-left: 4px solid #1a56db; margin-bottom: 15px;">
            <strong>📌 IMPORTANTE:</strong> Una vez enviada, la ficha <b>NO se puede eliminar ni editar</b>.
            Revise bien antes de enviar. El material será revisado por la dirección.
        </div>""", unsafe_allow_html=True)

        areas = _areas_del_docente()
        with st.form("form_material", clear_on_submit=True):
            c1, c2 = st.columns(2)
            with c1:
                semana = st.number_input("📅 Semana N°:", min_value=1, max_value=40,
                                         value=semana_actual, key="mat_semana")
            with c2:
                area = st.selectbox("📚 Área:", areas, key="mat_area")
            titulo = st.text_input("📝 Título de la ficha:",
                                   placeholder="Ej: Operaciones con fracciones", key="mat_titulo")
            st.markdown("---")
            st.markdown("### 📄 Contenido de la Ficha")
            st.markdown("""
            <div style="background: #f0fdf4; border-radius: 8px; padding: 10px; font-size: 0.82rem; margin-bottom: 10px;">
                <strong>📝 Formato disponible:</strong><br>
                • <code>**texto**</code> → <b>negrita</b><br>
                • <code>## Subtítulo</code> → subtítulo grande azul<br>
                • <code>### Sub-subtítulo</code> → subtítulo mediano<br>
                • Línea normal → texto regular
            </div>""", unsafe_allow_html=True)

            instrucciones = st.text_area("📌 Instrucciones generales:",
                                         placeholder="Ej: Lee atentamente cada ejercicio...",
                                         key="mat_instrucciones", height=80)
            st.markdown("**📖 Contenido / Teoría / Explicación:**")
            contenido_texto = st.text_area("Texto principal:",
                                           placeholder="Escribe aquí la explicación o contenido...",
                                           key="mat_contenido", height=150)
            img_contenido = st.file_uploader("🖼️ Imagen del contenido (opcional):",
                                             type=["png", "jpg", "jpeg"], key="mat_img_contenido",
                                             help="Suba imágenes de problemas, gráficos, etc.")
            st.markdown("**✏️ Ejercicios / Actividades:**")
            ejercicios = st.text_area("Ejercicios (uno por línea):",
                                      placeholder="1) Resuelve: 2/3 + 1/4 = \n2) Simplifica: 8/12",
                                      key="mat_ejercicios", height=150)
            img_ejercicios = st.file_uploader("🖼️ Imagen de ejercicios (opcional):",
                                              type=["png", "jpg", "jpeg"], key="mat_img_ejercicios",
                                              help="Para ecuaciones, figuras geométricas, tablas...")
            actividad_extra = st.text_area("📝 Actividad complementaria (opcional):",
                                           placeholder="Tarea para casa, investigación...",
                                           key="mat_extra", height=80)
            st.markdown("---")
            espacio_resolver = st.checkbox("Agregar líneas punteadas para resolver", value=True, key="mat_espacio")
            enviado = st.form_submit_button("📤 ENVIAR MATERIAL", type="primary", use_container_width=True)

            if enviado:
                if not titulo or not titulo.strip():
                    st.error("⚠️ Debe ingresar un título para la ficha")
                elif not (contenido_texto or ejercicios or img_contenido):
                    st.error("⚠️ Debe agregar al menos contenido, ejercicios o una imagen")
                else:
                    with st.spinner("📦 Procesando y guardando material..."):
                        bloques = []
                        if instrucciones and instrucciones.strip():
                            bloques.append({'tipo': 'instruccion', 'contenido': instrucciones.strip(), 'subtitulo': ''})
                        if contenido_texto and contenido_texto.strip():
                            bloques.append({'tipo': 'texto', 'contenido': contenido_texto.strip(), 'subtitulo': 'Contenido'})
                        if img_contenido:
                            comp = _comprimir_imagen_aula(img_contenido.getvalue(), max_size=500, quality=70)
                            bloques.append({'tipo': 'imagen', 'imagen_b64': _img_a_base64(comp), 'subtitulo': ''})
                        if ejercicios and ejercicios.strip():
                            bloques.append({'tipo': 'ejercicio', 'contenido': ejercicios.strip(),
                                           'subtitulo': 'Ejercicios', 'espacio_resolver': espacio_resolver})
                        if img_ejercicios:
                            comp = _comprimir_imagen_aula(img_ejercicios.getvalue(), max_size=500, quality=70)
                            bloques.append({'tipo': 'imagen', 'imagen_b64': _img_a_base64(comp), 'subtitulo': ''})
                        if actividad_extra and actividad_extra.strip():
                            bloques.append({'tipo': 'texto', 'contenido': actividad_extra.strip(),
                                           'subtitulo': 'Actividad Complementaria'})
                        material = {
                            'docente': usuario, 'docente_nombre': nombre_doc,
                            'grado': grado_doc, 'semana': semana, 'area': area,
                            'titulo': titulo.strip(), 'bloques': bloques,
                            'anio': config.get('anio', hora_peru().year),
                        }
                        mat_id = _guardar_material(material)
                    st.success(f"✅ Material guardado exitosamente (ID: {mat_id})")
                    st.balloons()
                    try:
                        pdf_bytes = _generar_pdf_material(material, config)
                        st.session_state['_ultimo_pdf_material'] = pdf_bytes
                        st.session_state['_ultimo_pdf_material_nombre'] = f"ficha_{area}_{semana}.pdf"
                    except Exception as e:
                        st.warning(f"⚠️ PDF generado con observaciones: {str(e)[:100]}")

        # Botón de descarga FUERA del form
        if '_ultimo_pdf_material' in st.session_state and st.session_state['_ultimo_pdf_material']:
            st.download_button("📥 Descargar Ficha en PDF",
                               st.session_state['_ultimo_pdf_material'],
                               st.session_state.get('_ultimo_pdf_material_nombre', 'ficha.pdf'),
                               "application/pdf", use_container_width=True, key="dl_material_pdf")

    with tab2:
        st.markdown("### 📄 Subir Archivo Word (.docx)")
        if not HAS_DOCX:
            st.error("⚠️ La librería python-docx no está instalada. Agregue `python-docx` a requirements.txt")
        else:
            st.markdown("""
            <div style="background: #f0fdf4; border-radius: 10px; padding: 12px; 
                        border-left: 4px solid #16a34a; margin-bottom: 15px;">
                <strong>📄 Suba un Word simple</strong> (sin encabezado ni pie de página).<br>
                El sistema le agregará el <b>formato oficial del colegio</b> con logo, datos y pie de página.
                <br>Se reconocen: <b>títulos, subtítulos, negritas</b> e imágenes.
            </div>""", unsafe_allow_html=True)

            areas = _areas_del_docente()
            c1, c2, c3 = st.columns(3)
            with c1:
                w_semana = st.number_input("📅 Semana:", 1, 40, semana_actual, key="w_mat_sem")
            with c2:
                w_area = st.selectbox("📚 Área:", areas, key="w_mat_area")
            with c3:
                w_titulo = st.text_input("📝 Título:", placeholder="Ej: Fracciones", key="w_mat_titulo")

            w_file = st.file_uploader("📎 Subir archivo Word (.docx):",
                                       type=["docx"], key="w_mat_file",
                                       help="Solo archivos .docx (Word 2007+)")
            if w_file and w_titulo:
                with st.spinner("📖 Leyendo documento Word..."):
                    bloques = _leer_docx(w_file.getvalue())
                if bloques:
                    # Vista previa
                    with st.expander("👁️ Vista previa del contenido", expanded=True):
                        for b in bloques:
                            if b['tipo'] == 'titulo':
                                st.markdown(f"## {b['contenido']}")
                            elif b['tipo'] == 'subtitulo':
                                st.markdown(f"### {b['contenido']}")
                            elif b['tipo'] == 'negrita':
                                st.markdown(f"**{b['contenido']}**")
                            elif b['tipo'] == 'texto':
                                st.write(b['contenido'])
                            elif b['tipo'] == 'imagen':
                                try:
                                    img_bytes = base64.b64decode(b['imagen_b64'])
                                    st.image(img_bytes, width=400)
                                except Exception:
                                    st.caption("[Imagen]")
                    st.info(f"📊 {len([b for b in bloques if b['tipo'] != 'vacio'])} bloques de contenido detectados")

                    if st.button("📤 CONVERTIR A PDF OFICIAL", type="primary",
                                 use_container_width=True, key="btn_word_pdf"):
                        with st.spinner("🖨️ Generando PDF con formato oficial..."):
                            pdf_bytes = _generar_pdf_desde_docx(
                                bloques, config, nombre_doc, grado_doc,
                                w_area, w_semana, w_titulo, "FICHA DE TRABAJO"
                            )
                        st.success("🎉 PDF generado con formato oficial del colegio")
                        st.download_button("📥 DESCARGAR PDF OFICIAL",
                                           pdf_bytes,
                                           f"ficha_{w_area}_S{w_semana}.pdf",
                                           "application/pdf",
                                           use_container_width=True,
                                           key="dl_word_pdf")
                        # También guardar como material
                        bloques_mat = []
                        for b in bloques:
                            if b['tipo'] in ('titulo', 'subtitulo', 'negrita'):
                                bloques_mat.append({'tipo': 'texto', 'contenido': b['contenido'],
                                                    'subtitulo': b['contenido'] if b['tipo'] in ('titulo', 'subtitulo') else ''})
                            elif b['tipo'] == 'texto':
                                bloques_mat.append({'tipo': 'texto', 'contenido': b['contenido'], 'subtitulo': ''})
                            elif b['tipo'] == 'imagen':
                                bloques_mat.append({'tipo': 'imagen', 'imagen_b64': b.get('imagen_b64', ''), 'subtitulo': ''})
                        material = {
                            'docente': usuario, 'docente_nombre': nombre_doc,
                            'grado': grado_doc, 'semana': w_semana, 'area': w_area,
                            'titulo': w_titulo.strip(), 'bloques': bloques_mat,
                            'anio': config.get('anio', hora_peru().year),
                            'origen': 'word'
                        }
                        _guardar_material(material)
                else:
                    st.warning("⚠️ No se pudo extraer contenido del archivo Word.")

    with tab3:
        st.markdown("### 📋 Mi Material Subido")
        materiales = _cargar_materiales()
        mis_materiales = [m for m in materiales if m.get('docente') == usuario]
        if not mis_materiales:
            st.info("📭 Aún no has subido material. Ve a la pestaña 'Subir Material'.")
        else:
            por_semana = {}
            for m in mis_materiales:
                s = m.get('semana', 0)
                if s not in por_semana:
                    por_semana[s] = []
                por_semana[s].append(m)
            for sem in sorted(por_semana.keys(), reverse=True):
                lun, vie = _rango_semana(sem)
                with st.expander(f"📅 Semana {sem} ({lun.strftime('%d/%m')} - {vie.strftime('%d/%m')}) — {len(por_semana[sem])} material(es)",
                                 expanded=(sem == semana_actual)):
                    for mat in por_semana[sem]:
                        st.markdown(f"**📚 {mat.get('area', '')}** — *{mat.get('titulo', '')}*")
                        st.caption(f"🕒 Subido: {mat.get('fecha_creacion', '')}")
                        if st.button(f"📥 Descargar PDF", key=f"dl_{mat.get('id', '')}", type="primary"):
                            try:
                                pdf = _generar_pdf_material(mat, config)
                                st.download_button("⬇️ Descargar", pdf,
                                                   f"ficha_{mat.get('area', '')}_{sem}.pdf",
                                                   "application/pdf", key=f"pdf_{mat.get('id', '')}")
                            except Exception:
                                st.error("Error generando PDF")
                        st.markdown("---")


def _vista_directivo_material(config, semana_actual):
    tab1, tab2 = st.tabs(["📊 Vista por Semana", "📈 Panel de Seguimiento"])
    with tab1:
        semana_ver = st.slider("📅 Seleccionar Semana:", 1, 40, semana_actual, key="dir_semana_mat")
        lun, vie = _rango_semana(semana_ver)
        st.markdown(f"**Semana {semana_ver}:** {lun.strftime('%d/%m/%Y')} al {vie.strftime('%d/%m/%Y')}")
        materiales = _cargar_materiales()
        mat_semana = [m for m in materiales if m.get('semana') == semana_ver]
        if not mat_semana:
            st.warning(f"📭 Ningún docente ha subido material para la Semana {semana_ver}")
        else:
            st.success(f"✅ {len(mat_semana)} material(es) subido(s) esta semana")
            por_docente = {}
            for m in mat_semana:
                doc = m.get('docente_nombre', m.get('docente', ''))
                if doc not in por_docente:
                    por_docente[doc] = []
                por_docente[doc].append(m)
            for docente_n, mats in por_docente.items():
                grado_n = mats[0].get('grado', '')
                with st.expander(f"👨‍🏫 {docente_n} — {grado_n} ({len(mats)} material(es))", expanded=True):
                    for mat in mats:
                        c1, c2, c3 = st.columns([3, 1, 1])
                        with c1:
                            st.markdown(f"**📚 {mat.get('area', '')}** — *{mat.get('titulo', '')}*")
                            st.caption(f"Subido: {mat.get('fecha_creacion', '')}")
                        with c2:
                            st.metric("Bloques", len(mat.get('bloques', [])))
                        with c3:
                            try:
                                pdf = _generar_pdf_material(mat, config)
                                st.download_button("📥 PDF", pdf, f"ficha_{mat.get('id', '')}.pdf",
                                                   "application/pdf", key=f"dir_pdf_{mat.get('id', '')}")
                            except Exception:
                                st.caption("Error PDF")

    with tab2:
        st.markdown("### 📈 Seguimiento de Entrega de Materiales")
        materiales = _cargar_materiales()
        if not materiales:
            st.info("📭 Sin datos de materiales aún")
            return
        semanas_rango = range(max(1, semana_actual - 4), semana_actual + 1)
        docentes_activos = set()
        for m in materiales:
            docentes_activos.add(m.get('docente_nombre', m.get('docente', '')))
        datos_tabla = []
        for docente_n in sorted(docentes_activos):
            fila = {'Docente': docente_n}
            for sem in semanas_rango:
                count = len([m for m in materiales
                            if m.get('docente_nombre', m.get('docente', '')) == docente_n
                            and m.get('semana') == sem])
                fila[f'S{sem}'] = f"✅ {count}" if count > 0 else "❌ 0"
            datos_tabla.append(fila)
        if datos_tabla:
            st.dataframe(pd.DataFrame(datos_tabla), use_container_width=True, hide_index=True)


# ================================================================
# TAB: EXÁMENES SEMANALES
# ================================================================

def tab_examenes_semanales(config):
    """GENERAR EXÁMENES - Versión simplificada con solo 2 tabs"""
    st.subheader("📝 Generar Exámenes")
    st.info("💡 Cree exámenes pregunta por pregunta. Se generarán en PDF con formato profesional de 2 columnas.")
    
    usuario = st.session_state.get('usuario_actual', '')
    info = st.session_state.get('docente_info', {}) or {}
    grado_doc = str(info.get('grado', ''))
    
    # Crear directorio de exámenes si no existe
    examenes_dir = Path("examenes")
    examenes_dir.mkdir(exist_ok=True)
    
    tab1, tab2 = st.tabs(["✏️ Crear Examen", "📥 Mis Exámenes"])
    
    # ===== TAB 1: CREAR EXAMEN =====
    with tab1:
        st.markdown("### ✏️ Crear Examen Pregunta por Pregunta")
        
        # Configuración del examen
        col1, col2, col3 = st.columns(3)
        with col1:
            titulo_examen = st.text_input("📝 Título del examen:", 
                                         placeholder="Ej: Evaluación de Matemática - Semana 5",
                                         key="titulo_exam")
        with col2:
            grado_examen = _grados_para_selector("exam")
        with col3:
            num_preguntas = st.number_input("🔢 Número de preguntas:", 
                                           min_value=1, max_value=50, value=10,
                                           key="num_preg")
        
        area_examen = st.text_input("📚 Área/Curso:", 
                                    placeholder="Ej: Matemática, Comunicación, etc.",
                                    key="area_exam")

        # ── Tip para capturas de pantalla ────────────────────────────────────
        st.info("💡 **Capturas de pantalla:** Haz captura (Win+Shift+S), "
                "guárdala como imagen (.png), luego súbela en el campo 🖼️ de cada pregunta.")
        
        st.markdown("---")
        st.markdown("### 📝 Preguntas del Examen")
        st.caption("Las preguntas aparecerán en **negrita** en el PDF. "
                   "Sube imagen por archivo desde tu computadora.")
        
        # Preguntas SIN form (para que file_uploader funcione mejor)
        preguntas = []
        
        for i in range(1, int(num_preguntas) + 1):
            st.markdown(f"#### 📌 Pregunta {i}")
            
            col_texto, col_img = st.columns([3, 1])
            with col_texto:
                texto_pregunta = st.text_area(
                    f"Enunciado de la pregunta {i}:",
                    height=80,
                    key=f"texto_p_{i}",
                    placeholder="Escriba el enunciado de la pregunta..."
                )
            
            with col_img:
                st.caption("🖼️ Imagen (opcional)")
                imagen_pregunta = st.file_uploader(
                    f"Imagen {i}",
                    type=['png', 'jpg', 'jpeg'],
                    key=f"img_p_{i}",
                    label_visibility="collapsed"
                )
            
            # Alternativas en 2 columnas
            col_a, col_b = st.columns(2)
            with col_a:
                alt_a = st.text_input(f"A)", key=f"alt_a_{i}", placeholder="Primera alternativa")
                alt_c = st.text_input(f"C)", key=f"alt_c_{i}", placeholder="Tercera alternativa")
            with col_b:
                alt_b = st.text_input(f"B)", key=f"alt_b_{i}", placeholder="Segunda alternativa")
                alt_d = st.text_input(f"D)", key=f"alt_d_{i}", placeholder="Cuarta alternativa")
            
            # Respuesta correcta
            correcta = st.radio(
                f"✅ Respuesta correcta de la pregunta {i}:",
                ['A', 'B', 'C', 'D'],
                horizontal=True,
                key=f"correcta_{i}"
            )
            
            preguntas.append({
                'numero': i,
                'texto': texto_pregunta,
                'imagen': imagen_pregunta,
                'alternativas': {'A': alt_a, 'B': alt_b, 'C': alt_c, 'D': alt_d},
                'correcta': correcta
            })
            
            if i < num_preguntas:
                st.markdown("---")
        
        # Botón de generar (fuera del form)
        st.markdown("---")
        if st.button("🖨️ GENERAR PDF DEL EXAMEN", type="primary",
                     use_container_width=True, key="btn_generar_exam"):
            if not titulo_examen or not area_examen or not grado_examen:
                st.error("⚠️ Complete: Título, Grado y Área")
            else:
                preguntas_vacias = [p['numero'] for p in preguntas if not p['texto'].strip()]
                if preguntas_vacias:
                    st.warning(f"⚠️ Preguntas vacías: {', '.join(map(str, preguntas_vacias))}")
                else:
                    try:
                        with st.spinner("📄 Generando PDF..."):
                            pdf_bytes = _generar_pdf_examen_2columnas(
                                titulo_examen, area_examen, grado_examen, preguntas, config)
                            fecha_actual = fecha_peru_str()
                            nombre_archivo = f"examen_{usuario}_{grado_examen}_{fecha_actual}_{titulo_examen[:25]}.pdf"
                            nombre_archivo = nombre_archivo.replace(' ','_').replace('/','_').replace(':','_')
                            ruta_archivo = examenes_dir / nombre_archivo
                            with open(ruta_archivo, 'wb') as f:
                                f.write(pdf_bytes)
                            st.session_state['_ultimo_examen_pdf'] = pdf_bytes
                            st.session_state['_ultimo_examen_nombre'] = nombre_archivo
                            st.success("🎉 ¡Examen generado!")
                            st.balloons()
                    except Exception as e:
                        st.error(f"❌ Error: {str(e)}")

        # ── Descarga ────────────────────────────────────────────────────────
        if st.session_state.get('_ultimo_examen_pdf'):
            st.download_button(
                "📥 DESCARGAR EXAMEN PDF",
                st.session_state['_ultimo_examen_pdf'],
                st.session_state.get('_ultimo_examen_nombre', 'examen.pdf'),
                "application/pdf",
                use_container_width=True,
                key="dl_examen_fuera_form"
            )
    
    # ===== TAB 2: MIS EXÁMENES =====
    with tab2:
        st.markdown("### 📥 Mis Exámenes Guardados")
        
        # Buscar exámenes del usuario
        patron = f"examen_{usuario}_*.pdf"
        examenes_usuario = list(examenes_dir.glob(patron))
        
        if examenes_usuario:
            st.success(f"📚 {len(examenes_usuario)} examen(es) guardado(s)")
            
            # Ordenar por fecha (más recientes primero)
            examenes_usuario.sort(key=lambda x: x.stat().st_mtime, reverse=True)
            
            for examen in examenes_usuario:
                # Extraer información del nombre del archivo
                partes = examen.stem.split('_')
                grado = partes[2] if len(partes) > 2 else 'N/A'
                fecha = partes[3] if len(partes) > 3 else 'N/A'
                titulo = '_'.join(partes[4:]) if len(partes) > 4 else examen.stem
                titulo = titulo.replace('_', ' ')[:50]
                
                # Mostrar cada examen
                with st.expander(f"📝 {titulo} - {grado} ({fecha})"):
                    col_a, col_b = st.columns([3, 1])
                    
                    with col_a:
                        st.caption(f"**Archivo:** {examen.name}")
                        st.caption(f"**Tamaño:** {examen.stat().st_size / 1024:.1f} KB")
                    
                    with col_b:
                        with open(examen, 'rb') as f:
                            st.download_button(
                                "⬇️ Descargar",
                                f.read(),
                                examen.name,
                                "application/pdf",
                                key=f"dl_exam_{examen.name}",
                                use_container_width=True
                            )
                    
                    # Botón para eliminar
                    if st.session_state.rol in ['admin', 'docente']:
                        if st.button(f"🗑️ Eliminar", key=f"del_exam_{examen.name}", type="primary"):
                            examen.unlink()
                            st.success("🎉 Examen eliminado")
                            time.sleep(0.5)
                            st.rerun()
        else:
            st.info("📭 No has creado exámenes aún")
            st.caption("💡 Ve a la pestaña 'Crear Examen' para generar tu primer examen")


def _generar_pdf_examen_2columnas(titulo, area, grado, preguntas, config):
    """Genera PDF de examen con encabezado oficial, 2 columnas, imágenes uniformes y hoja de claves"""
    buffer = io.BytesIO()
    c_pdf = canvas.Canvas(buffer, pagesize=A4)
    w, h = A4
    _info_doc = st.session_state.get('docente_info', {}) or {}
    usuario_doc = _nombre_completo_docente()

    # ── ENCABEZADO OFICIAL (igual que ficha) ────────────────────────────────
    c_pdf.setFillColor(colors.HexColor("#001e7c"))
    c_pdf.rect(0, h - 15, w, 15, fill=1, stroke=0)
    
    if Path("escudo_upload.png").exists():
        try:
            c_pdf.drawImage("escudo_upload.png", 25, h - 90, 62, 62, mask='auto')
        except Exception:
            pass
    
    c_pdf.setFillColor(colors.HexColor("#001e7c"))
    c_pdf.setFont("Helvetica-Bold", 7.5)
    c_pdf.drawCentredString(w / 2, h - 28, "MINISTERIO DE EDUCACIÓN — DRE CUSCO — PIONEROS EN LA EDUCACION DE CALIDAD")
    c_pdf.setFont("Helvetica-Bold", 11)
    c_pdf.drawCentredString(w / 2, h - 43, "I.E.P. YACHAY — CHINCHERO")
    frase = config.get('frase', '')
    if frase:
        c_pdf.setFont("Helvetica-Oblique", 7)
        c_pdf.drawCentredString(w / 2, h - 56, f'"{frase}"')
    
    # Cuadro de datos
    c_pdf.setStrokeColor(colors.HexColor("#1a56db"))
    c_pdf.setLineWidth(1.5)
    c_pdf.roundRect(25, h - 148, w - 50, 68, 8, fill=0)
    c_pdf.setFillColor(colors.black)
    c_pdf.setFont("Helvetica", 9)
    c_pdf.drawString(35, h - 90,  f"GRADO: {grado}")
    c_pdf.drawRightString(w - 35, h - 90, f"FECHA: {hora_peru().strftime('%d/%m/%Y')}")
    c_pdf.drawString(35, h - 107, f"ÁREA: {area}")
    c_pdf.drawRightString(w - 35, h - 107, f"DOCENTE: {usuario_doc}")
    c_pdf.drawString(35, h - 128, "ALUMNO(A): _______________________________________________")
    c_pdf.drawRightString(w - 35, h - 128, "N° ______")
    
    # Título del examen
    c_pdf.setFillColor(colors.HexColor("#1a56db"))
    c_pdf.setFont("Helvetica-Bold", 14)
    c_pdf.drawCentredString(w / 2, h - 167, titulo.upper())
    c_pdf.setStrokeColor(colors.HexColor("#1a56db"))
    c_pdf.setLineWidth(2)
    c_pdf.line(60, h - 174, w - 60, h - 174)
    
    # Instrucciones
    c_pdf.setFont("Helvetica-Oblique", 8)
    c_pdf.setFillColor(colors.HexColor("#6b7280"))
    c_pdf.drawCentredString(w / 2, h - 185, "Instrucciones: Marque con X la alternativa correcta para cada pregunta.")
    c_pdf.setFillColor(colors.black)
    
    # Configuración de 2 columnas
    col_width = (w - 80) / 2
    col_gap = 20
    y_start = h - 200
    y = y_start
    x_col1 = 40
    x_col2 = 40 + col_width + col_gap
    columna_actual = 1
    x = x_col1
    y_min = 60
    y_col1 = y_start
    y_col2 = y_start

    # Línea divisoria entre columnas
    c_pdf.setStrokeColor(colors.HexColor("#e5e7eb"))
    c_pdf.setLineWidth(0.5)
    c_pdf.line(w/2, y_start, w/2, y_min)
    
    from reportlab.platypus import Paragraph
    from reportlab.lib.styles import ParagraphStyle
    
    for pregunta in preguntas:
        # Calcular espacio real necesario para la pregunta completa
        n_alternativas = sum(1 for l in ['A','B','C','D'] if pregunta['alternativas'].get(l,''))
        espacio_necesario = (
            (120 if pregunta['imagen'] else 0) +
            max(40, len(pregunta['texto']) // 4) +  # espacio por texto
            n_alternativas * 12 + 25  # alternativas + margen
        )
        espacio_necesario = min(espacio_necesario, 200)  # tope máximo

        if y < y_min + espacio_necesario:
            if columna_actual == 1:
                columna_actual = 2
                x = x_col2
                y = y_start  # col2 misma altura inicio de página
            else:
                # Nueva página
                c_pdf.setFont("Helvetica-Bold", 9)
                c_pdf.setFillColor(colors.HexColor("#6b7280"))
                c_pdf.drawCentredString(w / 2, 20, f"{titulo} — Página {c_pdf.getPageNumber()}")
                c_pdf.setFillColor(colors.black)
                c_pdf.showPage()
                if Path("escudo_upload.png").exists():
                    try:
                        from PIL import Image as PILImage
                        _img = PILImage.open("escudo_upload.png")
                        _iw, _ih = _img.size
                        _mw = 420; _mh = _mw / (_iw / _ih)
                        c_pdf.saveState()
                        c_pdf.setFillAlpha(0.35)
                        c_pdf.drawImage("escudo_upload.png", w/2-_mw/2, h/2-_mh/2, _mw, _mh, mask='auto')
                        c_pdf.restoreState()
                    except Exception:
                        pass
                y_start = h - 35
                columna_actual = 1
                x = x_col1
                y = y_start
                c_pdf.setFillColor(colors.black)
        
        # Número de pregunta
        c_pdf.setFont("Helvetica-Bold", 10)
        c_pdf.setFillColor(colors.HexColor("#1e3a8a"))
        c_pdf.drawString(x, y, f"{pregunta['numero']}.")
        c_pdf.setFillColor(colors.black)
        
        # Texto de la pregunta
        style = ParagraphStyle('pregunta', fontName='Helvetica-Bold', fontSize=9, 
                              leading=11, alignment=TA_JUSTIFY, leftIndent=12)
        p = Paragraph(pregunta['texto'], style)
        w_p, h_p = p.wrap(col_width - 15, 200)
        p.drawOn(c_pdf, x, y - h_p)
        y -= (h_p + 5)
        
        # Imagen UNIFORME si existe
        if pregunta['imagen']:
            try:
                img_bytes = pregunta['imagen'].getvalue()
                img_pil = Image.open(io.BytesIO(img_bytes))
                if img_pil.mode == 'RGBA':
                    img_pil = img_pil.convert('RGB')
                
                # TAMAÑO UNIFORME: 80x80 para todas las imágenes
                IMG_SIZE = 75
                tmp_img = f"tmp_exam_img_{int(time.time())}_{pregunta['numero']}.jpg"
                img_pil.thumbnail((IMG_SIZE, IMG_SIZE), Image.Resampling.LANCZOS)
                img_pil.save(tmp_img, 'JPEG', quality=85)
                
                # Centrar imagen en la columna
                x_img = x + (col_width - IMG_SIZE) / 2
                c_pdf.drawImage(tmp_img, x_img, y - IMG_SIZE, IMG_SIZE, IMG_SIZE)
                y -= (IMG_SIZE + 5)
                try:
                    os.remove(tmp_img)
                except Exception:
                    pass
            except Exception:
                pass
        
        # Alternativas con burbujas
        c_pdf.setFont("Helvetica", 8.5)
        alternativas = pregunta['alternativas']
        for letra in ['A', 'B', 'C', 'D']:
            texto_alt = alternativas.get(letra, '')
            if texto_alt:
                c_pdf.circle(x + 4, y - 2, 3.5, stroke=1, fill=0)
                style_alt = ParagraphStyle('alt', fontName='Helvetica', fontSize=8.5,
                                          leading=10, alignment=TA_JUSTIFY, leftIndent=12)
                p_alt = Paragraph(f"<b>{letra})</b> {texto_alt}", style_alt)
                w_alt, h_alt = p_alt.wrap(col_width - 18, 100)
                p_alt.drawOn(c_pdf, x, y - h_alt)
                y -= (h_alt + 1)
        y -= 8
    
    # ── HOJA DE CLAVES — solo si hay respuestas correctas definidas ────────
    tiene_claves = any(preg.get('correcta', '').strip() for preg in preguntas)
    if not tiene_claves:
        c_pdf.save()
        buffer.seek(0)
        return buffer.getvalue()

    c_pdf.showPage()
    # Marca de agua en clave
    if Path("escudo_upload.png").exists():
        try:
            from PIL import Image as PILImage
            _img = PILImage.open("escudo_upload.png")
            _iw, _ih = _img.size
            _mw = 420; _mh = _mw / (_iw/_ih)
            c_pdf.saveState()
            c_pdf.setFillAlpha(0.35)
            c_pdf.drawImage("escudo_upload.png", w/2-_mw/2, h/2-_mh/2, _mw, _mh, mask='auto')
            c_pdf.restoreState()
        except Exception:
            pass
    c_pdf.setFont("Helvetica-Bold", 16)
    c_pdf.setFillColor(colors.HexColor("#dc2626"))
    c_pdf.drawCentredString(w / 2, h - 50, "CLAVE DE RESPUESTAS — USO EXCLUSIVO DOCENTE")
    c_pdf.setStrokeColor(colors.HexColor("#dc2626"))
    c_pdf.setLineWidth(2)
    c_pdf.line(w / 4, h - 55, 3 * w / 4, h - 55)
    c_pdf.setFillColor(colors.black)
    c_pdf.setFont("Helvetica-Bold", 11)
    c_pdf.drawCentredString(w / 2, h - 75, f"{grado} — {area} — {titulo}")
    
    y_clave = h - 100
    c_pdf.setFont("Helvetica", 10)
    num_cols_clave = 5
    col_w_clave = (w - 80) / num_cols_clave
    
    for i, preg in enumerate(preguntas):
        col_idx = i % num_cols_clave
        row_idx = i // num_cols_clave
        x_c = 40 + col_idx * col_w_clave
        y_c = y_clave - row_idx * 20
        
        if y_c < 60:
            c_pdf.showPage()
            y_clave = h - 60
            row_idx = 0
            y_c = y_clave
        
        c_pdf.setFont("Helvetica-Bold", 10)
        c_pdf.drawString(x_c, y_c, f"{preg['numero']}.")
        c_pdf.setFont("Helvetica", 10)
        resp_correcta = preg.get('correcta', '').upper()
        c_pdf.setFillColor(colors.HexColor("#16a34a"))
        c_pdf.drawString(x_c + 15, y_c, resp_correcta)
        c_pdf.setFillColor(colors.black)
    
    c_pdf.save()
    buffer.seek(0)
    return buffer.getvalue()


def _vista_docente_examenes(config, usuario, nombre_doc, grado_doc, semana_actual):
    tab1, tab2, tab3, tab4 = st.tabs(["📤 Cargar Preguntas", "📄 Subir Word", "📋 Mis Preguntas", "📥 Descargar Examen"])
    with tab1:
        st.markdown("### ✏️ Cargar Preguntas para Evaluación")
        st.markdown("""
        <div style="background: #fef3c7; border-radius: 10px; padding: 12px; 
                    border-left: 4px solid #f59e0b; margin-bottom: 15px;">
            <strong>⚠️ ATENCIÓN:</strong> Las preguntas enviadas <b>NO se pueden borrar</b>.
            Revise bien cada pregunta antes de enviar. El director compilará el examen final.
        </div>""", unsafe_allow_html=True)

        areas = _areas_del_docente()
        lun, vie = _rango_semana(semana_actual)
        st.info(f"📅 Semana {semana_actual}: {lun.strftime('%d/%m')} al {vie.strftime('%d/%m/%Y')} | Grado: **{grado_doc}**")

        # Configuración general
        c1, c2, c3 = st.columns(3)
        with c1:
            sem_input = st.number_input("📅 Semana:", 1, 40, semana_actual, key="ex_semana")
        with c2:
            area = st.selectbox("📚 Área:", areas, key="ex_area")
        with c3:
            tipo_eval = st.selectbox("📋 Tipo de evaluación:", TIPOS_EVALUACION, key="ex_tipo_eval")

        # Cuántas preguntas
        num_preguntas = st.number_input("🔢 ¿Cuántas preguntas desea cargar?",
                                        min_value=1, max_value=30, value=5, key="ex_num_preg")
        st.markdown("---")

        # Mostrar preguntas cargadas previamente
        examenes = _cargar_examenes_sem()
        mis_preg = [e for e in examenes if e.get('docente') == usuario
                    and e.get('semana') == sem_input and e.get('area') == area]
        if mis_preg:
            st.success(f"✅ Ya tienes **{len(mis_preg)}** pregunta(s) de {area} en Semana {sem_input}")

        # Formulario de N preguntas
        with st.form("form_preguntas_multiple", clear_on_submit=True):
            preguntas_data = []
            for i in range(1, num_preguntas + 1):
                st.markdown(f"#### Pregunta {i}")
                texto = st.text_area(f"Enunciado pregunta {i}:",
                                     placeholder=f"Escriba aquí la pregunta {i}...",
                                     key=f"ex_texto_{i}", height=80)
                img_p = st.file_uploader(f"🖼️ Imagen pregunta {i} (opcional):",
                                          type=["png", "jpg", "jpeg"], key=f"ex_img_{i}")
                c1, c2 = st.columns(2)
                with c1:
                    op_a = st.text_input(f"A) Preg {i}:", key=f"ex_a_{i}", placeholder="Opción A")
                    op_c = st.text_input(f"C) Preg {i}:", key=f"ex_c_{i}", placeholder="Opción C")
                with c2:
                    op_b = st.text_input(f"B) Preg {i}:", key=f"ex_b_{i}", placeholder="Opción B")
                    op_d = st.text_input(f"D) Preg {i}:", key=f"ex_d_{i}", placeholder="Opción D")
                resp = st.selectbox(f"✅ Respuesta correcta preg {i}:",
                                    ["a", "b", "c", "d"], key=f"ex_resp_{i}")
                preguntas_data.append({
                    'texto': texto, 'img': img_p,
                    'a': op_a, 'b': op_b, 'c': op_c, 'd': op_d,
                    'resp': resp
                })
                if i < num_preguntas:
                    st.markdown("---")

            enviado = st.form_submit_button(f"📤 GUARDAR {num_preguntas} PREGUNTA(S)",
                                            type="primary", use_container_width=True)
            if enviado:
                guardadas = 0
                errores = 0
                for idx, pd_item in enumerate(preguntas_data, 1):
                    texto = pd_item['texto']
                    if not texto or not texto.strip():
                        continue  # Saltar vacías
                    if not (pd_item['a'] and pd_item['b']):
                        errores += 1
                        st.warning(f"⚠️ Pregunta {idx}: necesita al menos opciones A y B")
                        continue
                    pregunta = {
                        'docente': usuario, 'docente_nombre': nombre_doc,
                        'grado': grado_doc, 'semana': sem_input, 'area': area,
                        'tipo_evaluacion': tipo_eval,
                        'texto': texto.strip(),
                        'opciones': {
                            'a': pd_item['a'].strip(),
                            'b': pd_item['b'].strip(),
                            'c': pd_item['c'].strip() if pd_item['c'] else '',
                            'd': pd_item['d'].strip() if pd_item['d'] else '',
                        },
                        'respuesta_correcta': pd_item['resp'],
                        'imagen_b64': '',
                    }
                    if pd_item['img']:
                        comp = _comprimir_imagen_aula(pd_item['img'].getvalue(), max_size=400, quality=65)
                        pregunta['imagen_b64'] = _img_a_base64(comp)
                    _guardar_pregunta_examen(pregunta)
                    guardadas += 1
                if guardadas > 0:
                    st.success(f"✅ {guardadas} pregunta(s) guardadas correctamente")
                    st.balloons()
                if errores > 0:
                    st.warning(f"⚠️ {errores} pregunta(s) con errores (no guardadas)")

    with tab2:
        st.markdown("### 📄 Subir Examen desde Word (.docx)")
        if not HAS_DOCX:
            st.error("⚠️ La librería python-docx no está instalada.")
        else:
            st.markdown("""
            <div style="background: #fef3c7; border-radius: 10px; padding: 12px; 
                        border-left: 4px solid #f59e0b; margin-bottom: 15px;">
                <strong>📄 Suba su examen en Word</strong> (sin encabezado ni pie).<br>
                El sistema le agrega el <b>formato oficial</b> con logo, datos del colegio,
                nombre del docente y pie de página.
            </div>""", unsafe_allow_html=True)

            areas_ex = _areas_del_docente()
            c1, c2, c3 = st.columns(3)
            with c1:
                we_sem = st.number_input("📅 Semana:", 1, 40, semana_actual, key="we_sem")
            with c2:
                we_area = st.selectbox("📚 Área:", areas_ex, key="we_area")
            with c3:
                we_tipo = st.selectbox("📋 Tipo:", TIPOS_EVALUACION, key="we_tipo")
            we_titulo = st.text_input("📝 Título del examen:", placeholder="Ej: Evaluación Semanal 3",
                                       key="we_titulo")
            we_file = st.file_uploader("📎 Subir examen Word (.docx):",
                                        type=["docx"], key="we_file")
            if we_file and we_titulo:
                with st.spinner("📖 Leyendo examen..."):
                    bloques = _leer_docx(we_file.getvalue())
                if bloques:
                    with st.expander("👁️ Vista previa", expanded=True):
                        for b in bloques:
                            if b['tipo'] == 'titulo':
                                st.markdown(f"## {b['contenido']}")
                            elif b['tipo'] == 'subtitulo':
                                st.markdown(f"### {b['contenido']}")
                            elif b['tipo'] == 'negrita':
                                st.markdown(f"**{b['contenido']}**")
                            elif b['tipo'] == 'texto':
                                st.write(b['contenido'])
                            elif b['tipo'] == 'imagen':
                                try:
                                    st.image(base64.b64decode(b['imagen_b64']), width=400)
                                except Exception:
                                    pass
                    if st.button("📤 CONVERTIR A PDF OFICIAL", type="primary",
                                 use_container_width=True, key="btn_word_ex"):
                        titulo_full = f"{we_tipo} — {we_titulo}"
                        with st.spinner("🖨️ Generando PDF oficial..."):
                            pdf_bytes = _generar_pdf_desde_docx(
                                bloques, config, nombre_doc, grado_doc,
                                we_area, we_sem, titulo_full, "EXAMEN"
                            )
                        st.success("🎉 Examen con formato oficial generado")
                        st.download_button("📥 DESCARGAR EXAMEN PDF",
                                           pdf_bytes,
                                           f"examen_{we_area}_S{we_sem}.pdf",
                                           "application/pdf",
                                           use_container_width=True,
                                           key="dl_word_ex")
                else:
                    st.warning("⚠️ No se pudo leer el archivo Word.")

    with tab3:
        st.markdown("### 📋 Mis Preguntas Cargadas")
        examenes = _cargar_examenes_sem()
        mis_preguntas = [e for e in examenes if e.get('docente') == usuario]
        if not mis_preguntas:
            st.info("📭 Aún no has cargado preguntas.")
        else:
            por_semana = {}
            for p in mis_preguntas:
                s = p.get('semana', 0)
                if s not in por_semana:
                    por_semana[s] = []
                por_semana[s].append(p)
            for sem in sorted(por_semana.keys(), reverse=True):
                with st.expander(f"📅 Semana {sem} — {len(por_semana[sem])} pregunta(s)",
                                 expanded=(sem == semana_actual)):
                    por_area = {}
                    for p in por_semana[sem]:
                        a = p.get('area', 'Sin área')
                        if a not in por_area:
                            por_area[a] = []
                        por_area[a].append(p)
                    for area_n, pregs_area in por_area.items():
                        st.markdown(f"**📚 {area_n}** — {len(pregs_area)} pregunta(s)")
                        for i, p in enumerate(pregs_area, 1):
                            tipo_e = p.get('tipo_evaluacion', 'Semanal')
                            st.caption(f"  {i}. {p.get('texto', '')[:80]}... [Resp: {p.get('respuesta_correcta', '?').upper()}] ({tipo_e})")

    with tab4:
        st.markdown("### 📥 Descargar Mi Examen")
        st.caption("Genera un PDF con tus preguntas cargadas para imprimir.")
        examenes = _cargar_examenes_sem()
        mis_preguntas = [e for e in examenes if e.get('docente') == usuario]
        if not mis_preguntas:
            st.info("📭 Sin preguntas para generar examen.")
        else:
            semanas_disp = sorted(set(p.get('semana', 0) for p in mis_preguntas), reverse=True)
            c1, c2 = st.columns(2)
            with c1:
                sem_dl = st.selectbox("Semana:", semanas_disp, key="ex_dl_sem")
            with c2:
                areas_disp = sorted(set(p.get('area', '') for p in mis_preguntas if p.get('semana') == sem_dl))
                area_dl = st.selectbox("Área:", ["TODAS"] + areas_disp, key="ex_dl_area")

            preg_filtradas = [p for p in mis_preguntas if p.get('semana') == sem_dl]
            if area_dl != "TODAS":
                preg_filtradas = [p for p in preg_filtradas if p.get('area') == area_dl]

            st.info(f"📝 {len(preg_filtradas)} preguntas disponibles")

            if st.button("🖨️ GENERAR MI EXAMEN PDF", type="primary",
                         use_container_width=True, key="btn_gen_mi_examen"):
                if preg_filtradas:
                    areas_agrupadas = {}
                    for p in preg_filtradas:
                        a = p.get('area', 'General')
                        if a not in areas_agrupadas:
                            areas_agrupadas[a] = []
                        areas_agrupadas[a].append(p)
                    titulo = f"{preg_filtradas[0].get('tipo_evaluacion', 'Evaluación')} — Semana {sem_dl}"
                    try:
                        pdf_bytes = _generar_pdf_examen_semanal(areas_agrupadas, config, grado_doc, sem_dl, titulo)
                        st.download_button("📥 DESCARGAR EXAMEN PDF", pdf_bytes,
                                           f"mi_examen_S{sem_dl}.pdf",
                                           "application/pdf", use_container_width=True,
                                           key="dl_mi_examen")
                    except Exception as e:
                        st.error(f"Error: {str(e)[:100]}")
                else:
                    st.warning("Sin preguntas para generar")


def _vista_directivo_examenes(config, semana_actual):
    tab1, tab2, tab3 = st.tabs(["📝 Compilar Examen", "📊 Preguntas Cargadas", "📈 Estado de Entrega"])

    with tab1:
        st.markdown("### 🖨️ Generar Examen Final para Imprimir")
        c1, c2 = st.columns(2)
        with c1:
            semana_ver = st.number_input("📅 Semana:", 1, 40, semana_actual, key="dir_ex_sem")
        with c2:
            examenes = _cargar_examenes_sem()
            grados_disp = sorted(set(e.get('grado', '') for e in examenes
                                     if e.get('semana') == semana_ver and e.get('grado')))
            if grados_disp:
                grado_sel = st.selectbox("🎓 Grado:", grados_disp, key="dir_ex_grado")
            else:
                grado_sel = st.text_input("🎓 Grado:", key="dir_ex_grado_txt")

        titulo_examen = st.text_input("📝 Título del examen:",
                                       value=f"EVALUACIÓN SEMANAL N° {semana_ver}", key="dir_ex_titulo")
        preguntas_filtradas = [e for e in examenes
                               if e.get('semana') == semana_ver and e.get('grado') == grado_sel]

        if not preguntas_filtradas:
            st.warning(f"📭 No hay preguntas cargadas para {grado_sel} en la Semana {semana_ver}")
            st.info("💡 Los docentes deben cargar sus preguntas desde el módulo 'Exámenes Semanales'")
        else:
            por_area = {}
            for p in preguntas_filtradas:
                a = p.get('area', 'Sin área')
                if a not in por_area:
                    por_area[a] = []
                por_area[a].append(p)
            st.success(f"✅ {len(preguntas_filtradas)} preguntas disponibles en {len(por_area)} área(s)")

            st.markdown("**Seleccione áreas a incluir:**")
            areas_incluir = {}
            for area_s, pregs in por_area.items():
                incluir = st.checkbox(f"📚 {area_s} ({len(pregs)} preg.)", value=True, key=f"inc_{area_s}")
                if incluir:
                    areas_incluir[area_s] = pregs

            for area_s, pregs in areas_incluir.items():
                with st.expander(f"📚 {area_s} — {len(pregs)} preguntas"):
                    for i, p in enumerate(pregs):
                        st.markdown(f"**{i + 1}.** {p.get('texto', '')[:100]}")
                        st.caption(f"   Resp: {p.get('respuesta_correcta', '?').upper()} | Docente: {p.get('docente_nombre', '')}")

            st.markdown("---")
            if st.button("🖨️ GENERAR EXAMEN PDF", type="primary", use_container_width=True, key="btn_gen_examen"):
                if not areas_incluir:
                    st.error("⚠️ Seleccione al menos un área")
                else:
                    with st.spinner("📄 Generando examen profesional..."):
                        try:
                            pdf_bytes = _generar_pdf_examen_semanal(areas_incluir, config, grado_sel,
                                                                     semana_ver, titulo_examen)
                            st.download_button("📥 DESCARGAR EXAMEN PDF", pdf_bytes,
                                               f"examen_{grado_sel}_semana{semana_ver}.pdf",
                                               "application/pdf", use_container_width=True,
                                               key="dl_examen_final")
                            st.success(f"✅ Examen generado: {len(preguntas_filtradas)} preguntas. ¡Listo para imprimir!")
                        except Exception as e:
                            st.error(f"❌ Error generando examen: {str(e)[:200]}")

    with tab2:
        st.markdown("### 📊 Todas las Preguntas Cargadas")
        examenes = _cargar_examenes_sem()
        semana_filtro = st.slider("Semana:", 1, 40, semana_actual, key="dir_filtro_sem")
        preg_sem = [e for e in examenes if e.get('semana') == semana_filtro]
        if not preg_sem:
            st.info(f"Sin preguntas para Semana {semana_filtro}")
        else:
            datos = []
            for p in preg_sem:
                datos.append({
                    'Docente': p.get('docente_nombre', ''), 'Grado': p.get('grado', ''),
                    'Área': p.get('area', ''), 'Pregunta': p.get('texto', '')[:60] + '...',
                    'Resp': p.get('respuesta_correcta', '?').upper(), 'Fecha': p.get('fecha_creacion', ''),
                })
            st.dataframe(pd.DataFrame(datos), use_container_width=True, hide_index=True)

    with tab3:
        st.markdown("### 📈 Estado de Entrega de Preguntas")
        examenes = _cargar_examenes_sem()
        if not examenes:
            st.info("Sin datos aún")
            return
        sem_ver = st.number_input("Semana:", 1, 40, semana_actual, key="estado_sem_ex")
        preg_sem = [e for e in examenes if e.get('semana') == sem_ver]
        por_doc = {}
        for p in preg_sem:
            doc = p.get('docente_nombre', p.get('docente', ''))
            if doc not in por_doc:
                por_doc[doc] = {'total': 0, 'areas': set(), 'grado': ''}
            por_doc[doc]['total'] += 1
            por_doc[doc]['areas'].add(p.get('area', ''))
            por_doc[doc]['grado'] = p.get('grado', '')
        if por_doc:
            datos = []
            for doc, info_d in sorted(por_doc.items()):
                datos.append({
                    'Docente': doc, 'Grado': info_d['grado'],
                    'Preguntas': info_d['total'],
                    'Áreas': ', '.join(sorted(info_d['areas'])),
                    'Estado': '✅ Entregado' if info_d['total'] >= 3 else '⚠️ Pocas'
                })
            st.dataframe(pd.DataFrame(datos), use_container_width=True, hide_index=True)
        else:
            st.warning(f"Ningún docente ha cargado preguntas para la Semana {sem_ver}")



# ================================================================
# MÓDULO: GENERADOR DE EXÁMENES
# ================================================================

def generar_examen_pdf(titulo, preguntas_data, num_columnas=2, buffer=None):
    """
    Genera un examen en PDF con preguntas y alternativas
    preguntas_data: lista de dict con 'pregunta', 'alternativas' (lista), 'respuesta_correcta'
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    
    if buffer is None:
        buffer = io.BytesIO()
    
    c = canvas.Canvas(buffer, pagesize=A4)
    ancho, alto = A4
    
    # Encabezado
    c.setFont("Helvetica-Bold", 16)
    c.drawCentredString(ancho/2, alto - 30*mm, "I.E.P. YACHAY - CHINCHERO")
    
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(ancho/2, alto - 40*mm, titulo)
    
    # Datos del estudiante
    c.setFont("Helvetica", 10)
    y_pos = alto - 50*mm
    c.drawString(30*mm, y_pos, "Nombre: _" + "_" * 50)
    c.drawString(30*mm, y_pos - 10, "Grado: ________  Fecha: ________")
    
    y_pos -= 25
    
    # Preguntas en columnas
    margen_izq = 20*mm
    margen_der = ancho - 20*mm
    ancho_columna = (margen_der - margen_izq) / num_columnas
    
    col_actual = 0
    x_base = margen_izq
    
    for idx, p in enumerate(preguntas_data, 1):
        # Verificar espacio
        if y_pos < 40*mm:
            c.showPage()
            y_pos = alto - 30*mm
            col_actual = 0
            x_base = margen_izq
        
        # Calcular posición X según columna
        x_pos = x_base + (col_actual * ancho_columna)
        
        # Pregunta (negrita)
        c.setFont("Helvetica-Bold", 10)
        pregunta_texto = f"{idx}. {p['pregunta']}"
        
        # Wrap texto de pregunta
        from textwrap import wrap
        lineas_pregunta = wrap(pregunta_texto, width=40 if num_columnas == 2 else 80)
        for linea in lineas_pregunta:
            c.drawString(x_pos, y_pos, linea)
            y_pos -= 12
        
        # Alternativas (burbujas A B C D)
        c.setFont("Helvetica", 9)
        alternativas = p.get('alternativas', [])
        letras = ['A', 'B', 'C', 'D', 'E']
        
        for i, alt in enumerate(alternativas[:5]):
            letra = letras[i]
            # Dibujar círculo para marcar
            c.circle(x_pos + 5, y_pos - 3, 3, fill=0)
            c.drawString(x_pos + 12, y_pos - 5, f"{letra}) {alt}")
            y_pos -= 12
        
        y_pos -= 8  # Espacio entre preguntas
        
        # Cambiar de columna
        col_actual += 1
        if col_actual >= num_columnas:
            col_actual = 0
            y_pos -= 10  # Espacio extra entre filas
    
    # Hoja de respuestas al final
    c.showPage()
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(ancho/2, alto - 30*mm, "HOJA DE RESPUESTAS")
    
    y_pos = alto - 50*mm
    c.setFont("Helvetica", 10)
    
    for idx, p in enumerate(preguntas_data, 1):
        if y_pos < 40*mm:
            c.showPage()
            y_pos = alto - 30*mm
        
        respuesta = p.get('respuesta_correcta', '-')
        c.drawString(30*mm, y_pos, f"{idx}. Respuesta correcta: {respuesta}")
        y_pos -= 15
    
    c.save()
    buffer.seek(0)
    return buffer

def tab_generar_examen():
    """Pestaña para generar exámenes"""
    st.markdown("## 📝 GENERADOR DE EXÁMENES")
    st.markdown("---")
    
    opcion = st.radio(
        "Seleccione método de generación:",
        ["✍️ Crear examen manual", "📄 Cargar examen desde Word"],
        horizontal=True
    )
    
    if opcion == "✍️ Crear examen manual":
        st.markdown("### Crear Examen Manual")
        
        titulo_examen = st.text_input("Título del Examen", "EXAMEN BIMESTRAL - MATEMÁTICA")
        num_preguntas = st.number_input("Número de preguntas", min_value=1, max_value=50, value=20)
        num_columnas = st.selectbox("Columnas", [1, 2], index=1)
        
        if 'preguntas_examen' not in st.session_state:
            st.session_state['preguntas_examen'] = []
        
        st.markdown("---")
        st.markdown("### ✏️ Ingrese las Preguntas")
        
        preguntas_data = []
        
        for i in range(int(num_preguntas)):
            with st.expander(f"📌 Pregunta {i+1}", expanded=(i==0)):
                pregunta_texto = st.text_area(
                    f"Pregunta {i+1}", 
                    key=f"preg_{i}",
                    height=80,
                    placeholder="Escriba la pregunta aquí..."
                )
                
                col1, col2 = st.columns(2)
                alternativas = []
                
                with col1:
                    alt_a = st.text_input(f"A)", key=f"alt_a_{i}", placeholder="Alternativa A")
                    alt_b = st.text_input(f"B)", key=f"alt_b_{i}", placeholder="Alternativa B")
                    alternativas.extend([alt_a, alt_b])
                
                with col2:
                    alt_c = st.text_input(f"C)", key=f"alt_c_{i}", placeholder="Alternativa C")
                    alt_d = st.text_input(f"D)", key=f"alt_d_{i}", placeholder="Alternativa D")
                    alternativas.extend([alt_c, alt_d])
                
                respuesta_correcta = st.selectbox(
                    "Respuesta correcta:", 
                    ["A", "B", "C", "D"],
                    key=f"resp_{i}"
                )
                
                if pregunta_texto and all(alternativas):
                    preguntas_data.append({
                        'pregunta': pregunta_texto,
                        'alternativas': alternativas,
                        'respuesta_correcta': respuesta_correcta
                    })
        
        st.markdown("---")
        
        if st.button("🎓 GENERAR EXAMEN PDF", type="primary", use_container_width=True):
            if len(preguntas_data) >= num_preguntas:
                buffer = generar_examen_pdf(titulo_examen, preguntas_data, num_columnas)
                st.download_button(
                    "⬇️ DESCARGAR EXAMEN",
                    buffer.getvalue(),
                    file_name=f"Examen_{titulo_examen.replace(' ', '_')}.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
                st.success("🎉 Examen generado exitosamente!")
            else:
                st.error("⚠️ Complete todas las preguntas antes de generar")
    
    else:  # Cargar desde Word
        st.markdown("### 📄 Cargar Examen desde Word")
        st.info("📝 Suba un documento Word con el examen. El formato se respetará: negritas, numeración, guiones, etc.")
        
        archivo_word = st.file_uploader(
            "Seleccione archivo Word (.docx)",
            type=['docx'],
            key="upload_examen_word"
        )
        
        if archivo_word:
            try:
                if HAS_DOCX:
                    doc = DocxDocument(archivo_word)
                    
                    # Leer contenido del Word
                    contenido_completo = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            estilo = {
                                'texto': para.text,
                                'negrita': para.runs[0].bold if para.runs else False,
                                'alineacion': para.alignment
                            }
                            contenido_completo.append(estilo)
                    
                    st.success(f"✅ Documento cargado: {len(contenido_completo)} párrafos")
                    
                    # Generar PDF respetando formato
                    if st.button("🎓 CONVERTIR A PDF", type="primary", use_container_width=True):
                        buffer = io.BytesIO()
                        c = canvas.Canvas(buffer, pagesize=A4)
                        ancho, alto = A4
                        
                        # Encabezado
                        c.setFont("Helvetica-Bold", 16)
                        c.drawCentredString(ancho/2, alto - 30*mm, "I.E.P. YACHAY - CHINCHERO")
                        
                        y_pos = alto - 50*mm
                        margen = 30*mm
                        
                        for item in contenido_completo:
                            texto = item['texto']
                            negrita = item['negrita']
                            
                            # Seleccionar fuente
                            if negrita:
                                c.setFont("Helvetica-Bold", 11)
                            else:
                                c.setFont("Helvetica", 10)
                            
                            # Wrap texto
                            from textwrap import wrap
                            lineas = wrap(texto, width=90)
                            
                            for linea in lineas:
                                if y_pos < 40*mm:
                                    c.showPage()
                                    y_pos = alto - 30*mm
                                
                                c.drawString(margen, y_pos, linea)
                                y_pos -= 14
                        
                        c.save()
                        buffer.seek(0)
                        
                        st.download_button(
                            "⬇️ DESCARGAR EXAMEN PDF",
                            buffer.getvalue(),
                            file_name="Examen_desde_Word.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                        st.success("🎉 Examen convertido exitosamente!")
                else:
                    st.error("❌ Módulo python-docx no disponible")
            except Exception as e:
                st.error(f"Error al procesar documento: {str(e)}")



def generar_registro_mensual_pdf(docente, mes, grado, area, notas_data, buffer=None):
    """
    Genera PDF con todas las notas del mes para un área
    notas_data: {'Alumno1': [nota1, nota2, nota3, ...], 'Alumno2': [...]}
    """
    if buffer is None:
        buffer = io.BytesIO()
    
    c = canvas.Canvas(buffer, pagesize=landscape(A4))
    ancho, alto = landscape(A4)
    
    # Encabezado
    c.setFont("Helvetica-Bold", 18)
    c.drawCentredString(ancho/2, alto - 20*mm, "I.E.P. YACHAY - CHINCHERO")
    
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(ancho/2, alto - 30*mm, f"REGISTRO DE NOTAS - {mes.upper()}")
    
    c.setFont("Helvetica", 11)
    c.drawString(20*mm, alto - 40*mm, f"Docente: {docente}")
    c.drawString(20*mm, alto - 47*mm, f"Grado: {grado}  |  Área: {area}")
    
    # Crear tabla con las notas
    if not notas_data:
        c.drawString(20*mm, alto - 60*mm, "No hay notas registradas para este mes")
        c.save()
        buffer.seek(0)
        return buffer
    
    # Preparar datos para tabla
    num_evaluaciones = max(len(notas) for notas in notas_data.values())
    
    # Headers
    headers = ["N°", "APELLIDOS Y NOMBRES"]
    for i in range(num_evaluaciones):
        headers.append(f"Eval {i+1}")
    headers.append("PROMEDIO")
    
    # Datos
    tabla_data = [headers]
    for idx, (alumno, notas) in enumerate(sorted(notas_data.items()), 1):
        fila = [str(idx), alumno]
        fila.extend([str(n) if n else "-" for n in notas])
        # Calcular promedio
        notas_validas = [n for n in notas if n]
        if notas_validas:
            promedio = round(sum(notas_validas) / len(notas_validas), 1)
            fila.append(str(promedio))
        else:
            fila.append("-")
        tabla_data.append(fila)
    
    # Crear tabla con ReportLab
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib import colors
    
    col_widths = [15*mm, 60*mm] + [15*mm] * (num_evaluaciones + 1)
    
    tabla = Table(tabla_data, colWidths=col_widths, repeatRows=1)
    tabla.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
    ]))
    
    # Posicionar tabla
    tabla.wrapOn(c, ancho, alto)
    tabla.drawOn(c, 15*mm, alto - 140*mm)
    
    # Pie de página
    c.setFont("Helvetica-Oblique", 8)
    c.drawString(20*mm, 15*mm, f"Generado el {fecha_peru_str()} a las {hora_peru_str()}")
    
    c.save()
    buffer.seek(0)
    return buffer

def mostrar_registro_mensual_notas():
    """Interfaz para generar registro mensual de notas"""
    st.markdown("### 📅 REGISTRO MENSUAL DE NOTAS")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        mes = st.selectbox(
            "Mes",
            ["Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", 
             "Septiembre", "Octubre", "Noviembre", "Diciembre"],
            key="mes_registro"
        )
    
    with col2:
        grado = st.selectbox("Grado", BaseDatos.grados_unicos(), key="grado_reg_mes")
    
    with col3:
        # Áreas según el grado
        if "INICIAL" in grado.upper():
            areas = AREAS_INICIAL
        elif "PRIMARIA" in grado.upper() or any(g in grado for g in ["1°", "2°", "3°", "4°", "5°", "6°"]):
            areas = AREAS_PRIMARIA
        else:
            areas = AREAS_SECUNDARIA + AREAS_PREUNIVERSITARIO
        
        area = st.selectbox("Área", areas, key="area_reg_mes")
    
    if st.button("📊 GENERAR REGISTRO MENSUAL", type="primary", use_container_width=True):
        # Buscar notas del mes en session_state
        if 'historial_evaluaciones' in st.session_state:
            historial = st.session_state['historial_evaluaciones']
            
            # Filtrar notas del mes, grado y área
            notas_mes = {}
            
            for clave, datos in historial.items():
                if (mes.lower() in clave.lower() and 
                    datos.get('grado') == grado):
                    # Buscar notas del área en los datos
                    if 'ranking' in datos:
                        for alumno_data in datos['ranking']:
                            alumno = alumno_data.get('Nombre', '')
                            nota = alumno_data.get(area, 0)
                            
                            if alumno not in notas_mes:
                                notas_mes[alumno] = []
                            
                            if nota and nota > 0:
                                notas_mes[alumno].append(nota)
            
            if notas_mes:
                docente = _nombre_completo_docente()
                buffer = generar_registro_mensual_pdf(
                    docente, mes, grado, area, notas_mes
                )
                
                st.download_button(
                    f"⬇️ DESCARGAR REGISTRO {mes.upper()}",
                    buffer.getvalue(),
                    file_name=f"Registro_{mes}_{grado}_{area}.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
                st.success(f"✅ Registro de {mes} generado con {len(notas_mes)} estudiantes")
            else:
                st.warning(f"⚠️ No se encontraron notas de {area} para {grado} en {mes}")
        else:
            st.info("ℹ️ No hay evaluaciones guardadas en el historial")


# ================================================================
# FIN MÓDULOS AULA VIRTUAL + EXÁMENES SEMANALES
# ================================================================

# ================================================================
# MODULO: YACHAY QAWAY - Evaluacion con QR
# ================================================================


# ================================================================
# MODULO: YACHAY QAWAY v2 — Sync PC + Celular
# ================================================================
# PC: crear preguntas, proyectar, ver respuestas en vivo
# Celular: escanear QR caminando
# Datos compartidos via archivo JSON

def _plk_dir():
    d = Path("plickers_data")
    d.mkdir(exist_ok=True)
    return d

def _qaway_guardar_musica(audio_bytes):
    """Guarda MP3 subido por admin para musica de fondo + sincroniza a GSheets"""
    p = _plk_dir() / "musica_fondo.mp3"
    with open(p, 'wb') as f:
        f.write(audio_bytes)
    # Persistir en Google Sheets como base64
    try:
        _guardar_archivo_binario_gs("bin_qaway_mp3", str(p))
    except Exception:
        pass
    return p

def _qaway_cargar_musica():
    """Carga MP3 guardado como base64"""
    p = _plk_dir() / "musica_fondo.mp3"
    if p.exists():
        with open(p, 'rb') as f:
            return base64.b64encode(f.read()).decode('utf-8')
    return None

def _plk_guardar_sesion(sesion_id, data):
    """Guarda sesion activa en archivo + Google Sheets"""
    p = _plk_dir() / f"sesion_{sesion_id}.json"
    with open(p, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

def _plk_sync_quiz_a_gs(sesion_id, quiz_data):
    """Sincroniza quiz completo a Google Sheets para persistencia"""
    try:
        gs = _gs()
        if gs:
            ws = gs._get_hoja('config')
            if ws:
                key = f"qaway_quiz_{sesion_id}"
                data_str = json.dumps(quiz_data, ensure_ascii=False, default=str)
                # Buscar si ya existe para actualizar
                all_data = ws.get_all_values()
                found = False
                for idx, row in enumerate(all_data):
                    if row and row[0] == key:
                        ws.update_cell(idx + 1, 2, data_str)
                        found = True
                        break
                if not found:
                    ws.append_row([key, data_str])
    except Exception:
        pass

def _plk_sync_resp_a_gs(sesion_id):
    """Sincroniza respuestas a Google Sheets"""
    try:
        gs = _gs()
        if gs:
            resp = _plk_cargar_respuestas(sesion_id)
            if resp:
                ws = gs._get_hoja('config')
                if ws:
                    key = f"qaway_resp_{sesion_id}"
                    data_str = json.dumps(resp, ensure_ascii=False, default=str)
                    all_data = ws.get_all_values()
                    found = False
                    for idx, row in enumerate(all_data):
                        if row and row[0] == key:
                            ws.update_cell(idx + 1, 2, data_str)
                            found = True
                            break
                    if not found:
                        ws.append_row([key, data_str])
    except Exception:
        pass

def _plk_restaurar_desde_gs():
    """Restaura quizzes y respuestas desde Google Sheets al iniciar"""
    try:
        gs = _gs()
        if not gs:
            return
        ws = gs._get_hoja('config')
        if not ws:
            return
        data = ws.get_all_values()
        plk_dir = _plk_dir()
        restaurados = 0
        for row in data:
            if not row or len(row) < 2:
                continue
            key, val = row[0], row[1]
            if key.startswith('qaway_quiz_'):
                sesion_id = key.replace('qaway_quiz_', '')
                p = plk_dir / f"quiz_{sesion_id}.json"
                if not p.exists():
                    try:
                        qdata = json.loads(val)
                        with open(p, 'w', encoding='utf-8') as f:
                            json.dump(qdata, f, indent=2, ensure_ascii=False)
                        restaurados += 1
                    except Exception:
                        pass
            elif key.startswith('qaway_resp_'):
                sesion_id = key.replace('qaway_resp_', '')
                p = plk_dir / f"resp_{sesion_id}.json"
                if not p.exists():
                    try:
                        rdata = json.loads(val)
                        with open(p, 'w', encoding='utf-8') as f:
                            json.dump(rdata, f, indent=2, ensure_ascii=False)
                        restaurados += 1
                    except Exception:
                        pass
        return restaurados
    except Exception:
        return 0

def _plk_cargar_sesion(sesion_id):
    """Carga sesion desde archivo"""
    p = _plk_dir() / f"sesion_{sesion_id}.json"
    if p.exists():
        with open(p, 'r', encoding='utf-8') as f:
            return json.load(f)
    return None

def _plk_guardar_respuesta(sesion_id, dni, nombre, pregunta_idx, respuesta, correcta_es):
    """Guarda respuesta individual en archivo compartido"""
    p = _plk_dir() / f"resp_{sesion_id}.json"
    resp = {}
    if p.exists():
        with open(p, 'r', encoding='utf-8') as f:
            resp = json.load(f)
    if dni not in resp:
        resp[dni] = {'nombre': nombre, 'respuestas': {}}
    resp[dni]['respuestas'][str(pregunta_idx)] = {
        'resp': respuesta, 'correcta': correcta_es,
        'ok': respuesta == correcta_es
    }
    with open(p, 'w', encoding='utf-8') as f:
        json.dump(resp, f, indent=2, ensure_ascii=False)

def _plk_guardar_en_reportes(quiz_data, sesion_id):
    """Guarda resultados de QAWAY en historial_evaluaciones y resultados.json"""
    try:
        resp = _plk_cargar_respuestas(sesion_id)
        if not resp:
            return False
        preguntas = quiz_data.get('preguntas', [])
        total_p = len(preguntas)
        titulo = quiz_data.get('titulo', 'YACHAY QAWAY')
        area = quiz_data.get('area', '')
        grado = quiz_data.get('grado', '')
        docente = quiz_data.get('docente', '')
        usuario = quiz_data.get('usuario', '')
        fecha = quiz_data.get('fecha', fecha_peru_str())

        # 1. Guardar en historial_evaluaciones.json
        hist = _cargar_historial_evaluaciones()
        clave = f"qaway_{sesion_id}"
        ranking_filas = []
        for dni_r, pr in resp.items():
            nm = pr.get('nombre', dni_r)
            resps = pr.get('respuestas', {})
            cor = sum(1 for r in resps.values() if r.get('ok'))
            nota = round(cor / max(total_p, 1) * 20, 1)
            ranking_filas.append({
                'nombre': nm, 'dni': dni_r,
                'promedio': nota, 'correctas': cor, 'total': total_p
            })
        ranking_filas.sort(key=lambda x: x['promedio'], reverse=True)

        hist[clave] = {
            'id': sesion_id,
            'grado': grado,
            'periodo': 'QAWAY',
            'titulo': f"QAWAY: {titulo}",
            'fecha': fecha,
            'hora': hora_peru_str(),
            'docente': usuario,
            'docente_nombre': docente,
            'areas': [area] if area else [],
            'ranking': ranking_filas,
            'tipo': 'qaway'
        }
        _guardar_historial_evaluaciones(hist)

        # 2. Guardar en resultados.json (por alumno)
        resultados = BaseDatos.cargar_todos_resultados()
        for r_item in ranking_filas:
            reg = {
                'dni': r_item['dni'],
                'nombre': r_item['nombre'],
                'grado': grado,
                'periodo': 'QAWAY',
                'titulo': f"QAWAY: {titulo}",
                'fecha': fecha,
                'hora': hora_peru_str(),
                'docente': usuario,
                'docente_nombre': docente,
                'areas': [{'nombre': area or 'General', 'nota': r_item['promedio']}],
                'promedio_general': r_item['promedio'],
                '_docente': usuario,
                'tipo': 'qaway'
            }
            resultados.append(reg)
        with open('resultados.json', 'w', encoding='utf-8') as f:
            json.dump(resultados, f, ensure_ascii=False, indent=2, default=str)
        return True
    except Exception:
        return False

def _plk_format_quiz(path):
    """Muestra titulo del quiz en selectbox"""
    try:
        with open(str(path), 'r', encoding='utf-8') as f:
            q = json.load(f)
        parts = []
        for k in ['titulo', 'area', 'grado', 'docente', 'fecha']:
            v = str(q.get(k, '')).strip()
            if v:
                parts.append(v)
        return ' | '.join(parts) if parts else str(path.stem).replace('quiz_','')
    except Exception:
        return str(path.stem).replace('quiz_','')

def _plk_cargar_respuestas(sesion_id):
    """Carga todas las respuestas"""
    p = _plk_dir() / f"resp_{sesion_id}.json"
    if p.exists():
        with open(p, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def _generar_tarjeta_plickers(nombre, dni, numero):
    """Tarjeta PLEGABLE: doblar en 4, cada cara = 1 QR"""
    import qrcode
    from PIL import Image, ImageDraw, ImageFont
    W, H = 600, 600
    img = Image.new('RGB', (W, H), 'white')
    draw = ImageDraw.Draw(img)
    try:
        fs = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 12)
        fx = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 48)
        fm = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 14)
    except Exception:
        fs = fx = fm = ImageFont.load_default()
    # Lineas de doblez
    for x in range(0, W, 8):
        draw.line([(x, H//2), (x+4, H//2)], fill='#999', width=1)
    for y in range(0, H, 8):
        draw.line([(W//2, y), (W//2, y+4)], fill='#999', width=1)
    ops = ['A','B','C','D']
    cls = ['#16a34a','#2563eb','#d97706','#db2777']
    bgs = ['#f0fdf4','#eff6ff','#fffbeb','#fdf2f8']
    pnl = [(0,0),(W//2,0),(0,H//2),(W//2,H//2)]
    pw, ph = W//2, H//2
    for i in range(4):
        op, cl, bg = ops[i], cls[i], bgs[i]
        px, py = pnl[i]
        draw.rectangle([px+2,py+2,px+pw-2,py+ph-2], fill=bg, outline=cl, width=3)
        draw.text((px+pw//2, py+25), op, fill=cl, anchor='mm', font=fx)
        qr_data = f"YP|{dni}|{op}|{numero}"
        qr = qrcode.QRCode(version=2, box_size=5, border=1, error_correction=qrcode.constants.ERROR_CORRECT_H)
        qr.add_data(qr_data)
        qr.make(fit=True)
        qi = qr.make_image(fill_color='black', back_color='white').resize((150,150))
        img.paste(qi, (px+pw//2-75, py+55))
        draw.text((px+pw//2, py+ph-25), f"#{numero:03d} {nombre[:18]}", fill='#666', anchor='mm', font=fs)
        draw.text((px+pw//2, py+ph-10), f"Doblar y mostrar {op}", fill=cl, anchor='mm', font=fm)
    return img


def _generar_pdf_tarjetas_plickers(estudiantes_df, grado):
    """PDF tarjetas: 2 por hoja A4 (arriba y abajo), grandes con nombre visible"""
    from reportlab.lib.utils import ImageReader
    buffer = io.BytesIO()
    c_pdf = canvas.Canvas(buffer, pagesize=A4)
    w, h = A4
    est = estudiantes_df.sort_values('Nombre').reset_index(drop=True)
    tam = 340  # Tarjeta grande
    mitad = h / 2  # Linea de corte central
    for idx in range(len(est)):
        pos_en_pagina = idx % 2  # 0=arriba, 1=abajo
        if pos_en_pagina == 0:
            if idx > 0:
                c_pdf.showPage()
            # Header morado
            c_pdf.setFillColor(colors.HexColor("#7c3aed"))
            c_pdf.rect(0, h - 22, w, 22, fill=1, stroke=0)
            c_pdf.setFillColor(colors.white)
            c_pdf.setFont('Helvetica-Bold', 9)
            c_pdf.drawCentredString(w/2, h - 16, f'YACHAY QAWAY - TARJETAS QR | {grado}')
            c_pdf.setFillColor(colors.HexColor("#666"))
            c_pdf.setFont('Helvetica', 6)
            c_pdf.drawCentredString(w/2, h - 28, 'Recortar por linea punteada. Doblar en cruz. Mostrar solo la cara de su respuesta.')
            # ── LINEA DE CORTE HORIZONTAL (centro de hoja) ──
            c_pdf.setStrokeColor(colors.HexColor("#aaa"))
            c_pdf.setDash(6, 4)
            c_pdf.setLineWidth(0.8)
            c_pdf.line(15, mitad, w - 15, mitad)
            c_pdf.setDash()
            c_pdf.setLineWidth(1)
            # Tijera en linea de corte
            c_pdf.setFillColor(colors.HexColor("#999"))
            c_pdf.setFont('Helvetica', 8)
            c_pdf.drawString(5, mitad - 3, '- - -')
        row = est.iloc[idx]
        nombre = str(row.get('Nombre', ''))
        dni = str(row.get('DNI', ''))
        try:
            img = _generar_tarjeta_plickers(nombre, dni, idx + 1)
            img_buf = io.BytesIO()
            img.save(img_buf, format='PNG')
            img_buf.seek(0)
            xp = (w - tam) / 2  # Centrado horizontal
            if pos_en_pagina == 0:
                yp = mitad + (mitad - 30 - tam) / 2  # Centrado en mitad superior
            else:
                yp = (mitad - tam - 15) / 2  # Centrado en mitad inferior
            # Nombre del estudiante grande encima
            c_pdf.setFillColor(colors.HexColor("#1e1b4b"))
            c_pdf.setFont('Helvetica-Bold', 11)
            c_pdf.drawCentredString(w/2, yp + tam + 8, f'#{idx+1:03d}  {nombre}')
            # Tarjeta QR
            c_pdf.drawImage(ImageReader(img_buf), xp, yp, width=tam, height=tam,
                            preserveAspectRatio=True, mask='auto')
            # Marco punteado de recorte
            c_pdf.setStrokeColor(colors.HexColor("#bbb"))
            c_pdf.setDash(4, 3)
            c_pdf.rect(xp - 4, yp - 4, tam + 8, tam + 24)
            c_pdf.setDash()
            # ── NOMBRE DEL SISTEMA LATERAL (rotado vertical) ──
            c_pdf.saveState()
            c_pdf.setFillColor(colors.HexColor("#c4b5fd"))
            c_pdf.setFont('Helvetica-Bold', 7)
            c_pdf.translate(xp - 12, yp + tam / 2)
            c_pdf.rotate(90)
            c_pdf.drawCentredString(0, 0, 'YACHAY QAWAY - I.E.P. ALTERNATIVO YACHAY')
            c_pdf.restoreState()
            c_pdf.setFillColor(colors.black)
        except Exception as e:
            c_pdf.drawString(50, h/2, f'Error: {str(e)[:60]}')
    c_pdf.save()
    buffer.seek(0)
    return buffer



def _generar_pdf_cuestionario_qaway(quiz_data):
    """Genera PDF bonito del cuestionario con preguntas y alternativas"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm, cm
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            topMargin=15*mm, bottomMargin=15*mm,
                            leftMargin=20*mm, rightMargin=20*mm)
    styles = getSampleStyleSheet()
    elements = []
    w_page = A4[0] - 40*mm

    # === ENCABEZADO INSTITUCIONAL ===
    style_inst = ParagraphStyle('inst', parent=styles['Normal'], fontSize=7,
                                 textColor=colors.HexColor('#555555'), alignment=TA_CENTER)
    style_title = ParagraphStyle('title_q', parent=styles['Title'], fontSize=16,
                                  textColor=colors.HexColor('#7c3aed'), alignment=TA_CENTER,
                                  spaceAfter=4*mm)
    style_sub = ParagraphStyle('sub_q', parent=styles['Normal'], fontSize=10,
                                textColor=colors.HexColor('#374151'), alignment=TA_CENTER,
                                spaceAfter=2*mm)
    style_preg = ParagraphStyle('preg', parent=styles['Normal'], fontSize=11,
                                 textColor=colors.HexColor('#1e1b4b'), leading=14,
                                 fontName='Helvetica-Bold', spaceBefore=5*mm, spaceAfter=3*mm)
    style_opt = ParagraphStyle('opt', parent=styles['Normal'], fontSize=10,
                                textColor=colors.HexColor('#374151'), leading=13,
                                leftIndent=10*mm)
    style_opt_ok = ParagraphStyle('opt_ok', parent=styles['Normal'], fontSize=10,
                                   textColor=colors.HexColor('#16a34a'), leading=13,
                                   fontName='Helvetica-Bold', leftIndent=10*mm)
    style_footer = ParagraphStyle('foot', parent=styles['Normal'], fontSize=7,
                                   textColor=colors.HexColor('#999'), alignment=TA_CENTER)

    # Header con barra morada + escudos
    esc_izq = None
    esc_der = None
    try:
        if Path("escudo_upload.png").exists():
            esc_izq = RLImage("escudo_upload.png", width=12*mm, height=12*mm)
        if Path("escudo2_upload.png").exists():
            esc_der = RLImage("escudo2_upload.png", width=12*mm, height=12*mm)
    except Exception:
        pass
    h_center = Paragraph("I.E.P. ALTERNATIVO YACHAY", ParagraphStyle('h1', parent=styles['Normal'],
                   fontSize=12, fontName='Helvetica-Bold', textColor=colors.white, alignment=TA_CENTER))
    if esc_izq and esc_der:
        header_data = [[esc_izq, h_center, esc_der]]
        header_table = Table(header_data, colWidths=[16*mm, w_page - 32*mm, 16*mm])
    elif esc_izq:
        header_data = [[esc_izq, h_center]]
        header_table = Table(header_data, colWidths=[16*mm, w_page - 16*mm])
    else:
        header_data = [[h_center]]
        header_table = Table(header_data, colWidths=[w_page])
    header_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#7c3aed')),
        ('TOPPADDING', (0,0), (-1,-1), 8),
        ('BOTTOMPADDING', (0,0), (-1,-1), 8),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('ROUNDEDCORNERS', [8, 8, 0, 0]),
    ]))
    elements.append(header_table)

    # Sub-header
    sub_data = [[
        Paragraph(f"PIONEROS EN LA EDUCACION DE CALIDAD", style_inst),
    ]]
    sub_table = Table(sub_data, colWidths=[w_page])
    sub_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#ede9fe')),
        ('TOPPADDING', (0,0), (-1,-1), 4),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('ROUNDEDCORNERS', [0, 0, 8, 8]),
    ]))
    elements.append(sub_table)
    elements.append(Spacer(1, 5*mm))

    # Titulo del cuestionario
    titulo = quiz_data.get('titulo', 'Cuestionario')
    elements.append(Paragraph(f"YACHAY QAWAY", style_title))
    elements.append(Paragraph(f"<b>{titulo}</b>", ParagraphStyle('tit2', parent=styles['Normal'],
                    fontSize=13, textColor=colors.HexColor('#1e1b4b'), alignment=TA_CENTER, spaceAfter=3*mm)))

    # Info del cuestionario
    area = quiz_data.get('area', '')
    grado = quiz_data.get('grado', '')
    fecha = quiz_data.get('fecha', '')
    docente = quiz_data.get('docente', '')
    total_pregs = len(quiz_data.get('preguntas', []))

    info_data = [[
        Paragraph(f"<b>Area:</b> {area}", styles['Normal']),
        Paragraph(f"<b>Grado:</b> {grado}", styles['Normal']),
        Paragraph(f"<b>Fecha:</b> {fecha}", styles['Normal']),
    ], [
        Paragraph(f"<b>Docente:</b> {docente}", styles['Normal']),
        Paragraph(f"<b>Preguntas:</b> {total_pregs}", styles['Normal']),
        Paragraph("", styles['Normal']),
    ]]
    info_table = Table(info_data, colWidths=[w_page/3]*3)
    info_table.setStyle(TableStyle([
        ('FONTSIZE', (0,0), (-1,-1), 9),
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f5f3ff')),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#c4b5fd')),
        ('TOPPADDING', (0,0), (-1,-1), 5),
        ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ('LEFTPADDING', (0,0), (-1,-1), 8),
    ]))
    elements.append(info_table)

    # Linea nombre y seccion
    elements.append(Spacer(1, 4*mm))
    elements.append(Paragraph("Nombre: ________________________________________ Seccion: _______", 
                    ParagraphStyle('nombre', parent=styles['Normal'], fontSize=10, spaceAfter=3*mm)))
    elements.append(Spacer(1, 3*mm))

    # Linea separadora
    sep_data = [[Paragraph(f"PREGUNTAS ({total_pregs})", ParagraphStyle('sep', parent=styles['Normal'],
                 fontSize=9, fontName='Helvetica-Bold', textColor=colors.white, alignment=TA_CENTER))]]
    sep_table = Table(sep_data, colWidths=[w_page])
    sep_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#4c1d95')),
        ('TOPPADDING', (0,0), (-1,-1), 4),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
    ]))
    elements.append(sep_table)
    elements.append(Spacer(1, 3*mm))

    # === PREGUNTAS ===
    letras_color = {
        'A': '#16a34a', 'B': '#2563eb', 'C': '#d97706', 'D': '#db2777'
    }
    letras_bg = {
        'A': '#f0fdf4', 'B': '#eff6ff', 'C': '#fffbeb', 'D': '#fdf2f8'
    }

    for idx, preg in enumerate(quiz_data.get('preguntas', [])):
        # Numero y pregunta
        num_preg = f"<font color='#7c3aed'><b>{idx+1}.</b></font> {preg['pregunta']}"
        elements.append(Paragraph(num_preg, style_preg))

        # Imagen si existe
        if preg.get('imagen'):
            try:
                import base64 as b64mod
                img_data = b64mod.b64decode(preg['imagen'])
                img_buf = io.BytesIO(img_data)
                from PIL import Image as PILImg
                pil_img = PILImg.open(img_buf)
                iw, ih = pil_img.size
                max_w = 120*mm
                ratio = min(max_w / iw, 60*mm / ih)
                img_buf2 = io.BytesIO(img_data)
                rl_img = RLImage(img_buf2, width=iw*ratio, height=ih*ratio)
                elements.append(rl_img)
                elements.append(Spacer(1, 2*mm))
            except Exception:
                pass

        # Opciones en tabla 2x2
        opciones = preg.get('opciones', {})
        opt_cells = []
        row1 = []
        row2 = []
        for letra in ['A', 'B', 'C', 'D']:
            texto = opciones.get(letra, '')
            color = letras_color.get(letra, '#333')
            circulo = f"<font color='{color}'><b>  {letra})</b></font> {texto}"
            p_opt = Paragraph(circulo, style_opt)
            if letra in ('A', 'B'):
                row1.append(p_opt)
            else:
                row2.append(p_opt)
        opt_data = [row1, row2]
        opt_table = Table(opt_data, colWidths=[w_page/2]*2)
        opt_table.setStyle(TableStyle([
            ('FONTSIZE', (0,0), (-1,-1), 10),
            ('TOPPADDING', (0,0), (-1,-1), 4),
            ('BOTTOMPADDING', (0,0), (-1,-1), 4),
            ('LEFTPADDING', (0,0), (-1,-1), 6),
            ('BACKGROUND', (0,0), (0,0), colors.HexColor(letras_bg['A'])),
            ('BACKGROUND', (1,0), (1,0), colors.HexColor(letras_bg['B'])),
            ('BACKGROUND', (0,1), (0,1), colors.HexColor(letras_bg['C'])),
            ('BACKGROUND', (1,1), (1,1), colors.HexColor(letras_bg['D'])),
            ('BOX', (0,0), (-1,-1), 0.3, colors.HexColor('#e5e7eb')),
            ('INNERGRID', (0,0), (-1,-1), 0.3, colors.HexColor('#e5e7eb')),
        ]))
        elements.append(opt_table)
        elements.append(Spacer(1, 2*mm))

        # Separador sutil entre preguntas
        if idx < total_pregs - 1:
            elements.append(Spacer(1, 1*mm))
            sep_line = Table([['']], colWidths=[w_page])
            sep_line.setStyle(TableStyle([
                ('LINEBELOW', (0,0), (-1,-1), 0.5, colors.HexColor('#e5e7eb')),
            ]))
            elements.append(sep_line)

    # Footer
    elements.append(Spacer(1, 8*mm))
    elements.append(Paragraph(f"I.E.P. ALTERNATIVO YACHAY | Sistema YACHAY PRO | {fecha}", style_footer))

    def _qaway_page_footer(canv, doc_obj):
        canv.saveState()
        canv.setFont('Helvetica', 7)
        canv.setFillColor(colors.HexColor('#999'))
        canv.drawCentredString(A4[0]/2, 10*mm,
            f"I.E.P. ALTERNATIVO YACHAY | YACHAY QAWAY | Pag. {doc_obj.page}")
        canv.restoreState()
    doc.build(elements, onFirstPage=_qaway_page_footer, onLaterPages=_qaway_page_footer)
    buf.seek(0)
    return buf


# ================================================================
# PAUSA ACTIVA — 10 MODELOS — Admin sube música, Docentes presentan
# ================================================================

PAUSA_MODELOS = [
    {
        "id": 1,
        "nombre": "🌟 Estiramiento Solar",
        "nivel": ["TODOS"],
        "color_fondo": "#1e3a8a",
        "color_acento": "#fbbf24",
        "emoji_principal": "☀️",
        "descripcion": "Activa tu energia con estiramientos de cuerpo completo",
        "pasos": [
            ("🙆", "Brazos arriba — respira hondo", 15),
            ("🤸", "Inclinate hacia la derecha — 5 seg cada lado", 10),
            ("🙆‍♂️", "Abre el pecho — brazos atras", 10),
            ("🧘", "Gira el cuello suavemente — 3 veces", 12),
            ("💃", "Sacude los brazos y relaja", 8),
        ]
    },
    {
        "id": 2,
        "nombre": "🔥 Energia Rapida",
        "nivel": ["PRIMARIA","SECUNDARIA"],
        "color_fondo": "#dc2626",
        "color_acento": "#fef08a",
        "emoji_principal": "⚡",
        "descripcion": "Movimientos rapidos para despertar la mente",
        "pasos": [
            ("👐", "Aplaude 10 veces rapido", 8),
            ("🦵", "Levanta rodillas alternando — 15 veces", 12),
            ("💪", "Puno arriba izquierda — derecha — 10 veces", 10),
            ("🤜", "Boxeo en el aire — 20 golpes", 15),
            ("🌬️", "Respira profundo 3 veces", 10),
        ]
    },
    {
        "id": 3,
        "nombre": "🌊 Relajacion Cosmica",
        "nivel": ["SECUNDARIA"],
        "color_fondo": "#0e7490",
        "color_acento": "#a5f3fc",
        "emoji_principal": "🧘",
        "descripcion": "Tecnicas de respiracion y mindfulness",
        "pasos": [
            ("😮‍💨", "Inhala por la nariz — 4 segundos", 8),
            ("🫁", "Retiene el aire — 4 segundos", 6),
            ("💨", "Exhala por la boca — 8 segundos", 10),
            ("🧠", "Visualiza un lugar tranquilo", 15),
            ("✨", "Abre los ojos — sonrie", 5),
        ]
    },
    {
        "id": 4,
        "nombre": "🦁 Poder Animal",
        "nivel": ["INICIAL","PRIMARIA"],
        "color_fondo": "#7c2d12",
        "color_acento": "#fed7aa",
        "emoji_principal": "🐾",
        "descripcion": "Imita animales para mover todo el cuerpo",
        "pasos": [
            ("🦒", "El JIRAFA — estira el cuello al maximo", 10),
            ("🦅", "El AGUILA — abre los brazos y vuela", 12),
            ("🐊", "El COCODRILO — abre y cierra los brazos", 10),
            ("🐻", "El OSO — balancea de lado a lado", 12),
            ("🐱", "El GATO — estira como al despertar", 10),
        ]
    },
    {
        "id": 5,
        "nombre": "🎯 Concentracion Max",
        "nivel": ["PRIMARIA","SECUNDARIA"],
        "color_fondo": "#4c1d95",
        "color_acento": "#c4b5fd",
        "emoji_principal": "🧩",
        "descripcion": "Ejercicios para el cerebro y la concentracion",
        "pasos": [
            ("👁️", "Cierra los ojos — cuenta hasta 10", 12),
            ("🤔", "Toca nariz con mano derecha — oreja con izquierda", 10),
            ("🔄", "Ahora cambia — oreja con derecha — nariz con izquierda", 10),
            ("✍️", "Escribe tu nombre en el aire con el codo", 12),
            ("🧠", "Cuenta desde 20 hasta 1 rapidamente", 15),
        ]
    },
    {
        "id": 6,
        "nombre": "🌈 Danza Libre",
        "nivel": ["TODOS"],
        "color_fondo": "#065f46",
        "color_acento": "#6ee7b7",
        "emoji_principal": "🕺",
        "descripcion": "Muevete con musica y libertad total",
        "pasos": [
            ("💃", "Mueve los hombros al ritmo", 12),
            ("🦶", "Pasos de baile — izquierda — derecha", 10),
            ("🙌", "Manos arriba y muevelas", 10),
            ("🫶", "Gira sobre tu lugar dos veces", 8),
            ("🎉", "Celebra con un salto y aplauso", 5),
        ]
    },
    {
        "id": 7,
        "nombre": "💧 Hidratacion Activa",
        "nivel": ["TODOS"],
        "color_fondo": "#1e40af",
        "color_acento": "#bfdbfe",
        "emoji_principal": "🫗",
        "descripcion": "Pausas para hidratarse y moverse",
        "pasos": [
            ("💧", "Toma agua — minimo 3 sorbos", 8),
            ("🙆", "Estira mientras tomas agua", 10),
            ("👀", "Mira lejos — descansa los ojos 20 seg", 20),
            ("🤲", "Masajea tus manos y dedos", 12),
            ("😌", "Relaja los hombros — baja la tension", 10),
        ]
    },
    {
        "id": 8,
        "nombre": "🏆 Desafio Fisico",
        "nivel": ["PRIMARIA","SECUNDARIA"],
        "color_fondo": "#713f12",
        "color_acento": "#fef3c7",
        "emoji_principal": "🏅",
        "descripcion": "Pequeños retos fisicos divertidos",
        "pasos": [
            ("🦷", "Equilibrio — parate en un pie 10 seg cada uno", 20),
            ("🤼", "Sentadilla — baja despacio 5 veces", 15),
            ("🤸", "Salta en el lugar 10 veces", 12),
            ("🧗", "Escala imaginaria — trepa con brazos y piernas", 15),
            ("🎖️", "Pose de victoria — has lo mejor!", 8),
        ]
    },
    {
        "id": 9,
        "nombre": "🌺 Yoga Express",
        "nivel": ["SECUNDARIA"],
        "color_fondo": "#831843",
        "color_acento": "#fce7f3",
        "emoji_principal": "🪷",
        "descripcion": "Posturas de yoga adaptadas al aula",
        "pasos": [
            ("🌲", "Postura del ARBOL — un pie apoyado en tobillo", 15),
            ("⚡", "Postura del GUERRERO — un paso al frente", 12),
            ("🌉", "Postura del PUENTE — en silla inclina el torso", 12),
            ("🦅", "Postura del AGUILA — brazos entrelazados arriba", 10),
            ("🧘", "Postura del LOTO — sienta y cierra ojos", 15),
        ]
    },
    {
        "id": 10,
        "nombre": "🎵 Ritmo y Musica",
        "nivel": ["TODOS"],
        "color_fondo": "#1e1b4b",
        "color_acento": "#e0e7ff",
        "emoji_principal": "🎶",
        "descripcion": "Sigue el ritmo de la musica con todo el cuerpo",
        "pasos": [
            ("🥁", "Toca la bateria imaginaria al ritmo", 12),
            ("🎸", "Toca la guitarra imaginaria", 10),
            ("🎹", "Toca el piano en tu escritorio", 10),
            ("🎤", "Canta la melodia — aunque sea mentalmente", 12),
            ("🎊", "Gran final — mueve todo junto!", 10),
        ]
    },
    # ── NUEVOS: por nivel ──────────────────────────────────────────────
    {
        "id": 11,
        "nombre": "🐸 Animales Locos",
        "nivel": ["INICIAL"],
        "color_fondo": "#065f46",
        "color_acento": "#6ee7b7",
        "emoji_principal": "🐸",
        "descripcion": "Imita animales con todo tu cuerpo — para los mas pequeños",
        "pasos": [
            ("🐸", "Salta como RANA — 5 saltos con las dos piernas juntas!", 12),
            ("🐘", "Camina como ELEFANTE — brazos colgando y balanceate", 10),
            ("🦋", "Vuela como MARIPOSA — abre y cierra los brazos", 10),
            ("🐍", "Muevete como SERPIENTE — ondula todo el cuerpo sentado", 12),
            ("🦁", "RUGE como Leon — abre la boca y estira los brazos!", 8),
        ]
    },
    {
        "id": 12,
        "nombre": "🌈 Colores y Cuerpo",
        "nivel": ["INICIAL"],
        "color_fondo": "#7c2d12",
        "color_acento": "#fde68a",
        "emoji_principal": "🎨",
        "descripcion": "Aprender colores moviendose — especial para inicial",
        "pasos": [
            ("🔴", "ROJO — toca algo rojo en el salon sin moverte del lugar!", 12),
            ("💛", "AMARILLO — levanta los brazos como el sol brillante", 8),
            ("🔵", "AZUL — da 3 pasos lentos como el agua del mar", 10),
            ("💚", "VERDE — agachate como un arbol en el viento", 10),
            ("🌈", "ARCO IRIS — pinta el arco iris con tus dos brazos!", 10),
        ]
    },
    {
        "id": 13,
        "nombre": "🤸 Simon Dice",
        "nivel": ["INICIAL","PRIMARIA"],
        "color_fondo": "#1e1b4b",
        "color_acento": "#a5b4fc",
        "emoji_principal": "🎭",
        "descripcion": "Simon dice — el docente guia los movimientos",
        "pasos": [
            ("👆", "Simon dice: toca tu NARIZ con el dedo indice!", 8),
            ("🦷", "Simon dice: toca tu OREJA IZQUIERDA con la mano derecha!", 8),
            ("🦵", "Simon dice: levanta la RODILLA DERECHA 5 veces!", 10),
            ("🙌", "Simon dice: aplaude 10 veces lo mas rapido que puedas!", 8),
            ("🤐", "Simon dice: cierra los ojos y cuenta hasta 5 en silencio...", 8),
        ]
    },
    {
        "id": 14,
        "nombre": "🥊 Karate Kids",
        "nivel": ["PRIMARIA"],
        "color_fondo": "#450a0a",
        "color_acento": "#fca5a5",
        "emoji_principal": "🥋",
        "descripcion": "Movimientos de karate para descargar energia positiva",
        "pasos": [
            ("🥋", "Posicion de KARATE — pies separados, manos al frente!", 8),
            ("🤜", "GOLPE DERECHO — extiende el brazo derecho lento x5", 12),
            ("🤛", "GOLPE IZQUIERDO — alterna con el brazo izquierdo x5", 12),
            ("🦵", "PATADA LATERAL — levanta la rodilla y extiende x3 cada lado", 15),
            ("🙏", "REVERENCIA — inclinate hacia adelante en senal de respeto", 8),
        ]
    },
    {
        "id": 15,
        "nombre": "🦸 Superheroes",
        "nivel": ["PRIMARIA"],
        "color_fondo": "#1e3a8a",
        "color_acento": "#fbbf24",
        "emoji_principal": "🦸",
        "descripcion": "Poses de superheroes para activar la confianza",
        "pasos": [
            ("🦸", "Pose de SUPERMAN — un brazo arriba, corre en el lugar!", 10),
            ("🕷️", "Pose de SPIDERMAN — trepa la pared imaginaria!", 12),
            ("🦇", "Pose de BATMAN — abre el cape con los brazos bien abiertos", 8),
            ("💪", "Pose de HULK — hincha los musculos y ruge fuerte!", 8),
            ("⭐", "Eres un SUPERHEROE — tu pose especial inventada!", 12),
        ]
    },
    {
        "id": 16,
        "nombre": "🎪 Circo Express",
        "nivel": ["PRIMARIA","SECUNDARIA"],
        "color_fondo": "#4c1d95",
        "color_acento": "#f0abfc",
        "emoji_principal": "🎡",
        "descripcion": "Actos de circo imaginarios que activan coordinacion",
        "pasos": [
            ("🤹", "Malabarista — lanza y atrapa pelotas imaginarias x10", 12),
            ("🎪", "Equilibrista — camina sobre la cuerda floja imaginaria", 10),
            ("🦁", "Domador — da ordenes a los leones con la mano", 8),
            ("🎠", "Acrobata — rueda imaginaria con los brazos extendidos", 12),
            ("🎉", "Gran final de circo — todos aplauden al artista!", 8),
        ]
    },
    {
        "id": 17,
        "nombre": "🧠 Brain Gym",
        "nivel": ["PRIMARIA","SECUNDARIA"],
        "color_fondo": "#134e4a",
        "color_acento": "#99f6e4",
        "emoji_principal": "🧩",
        "descripcion": "Ejercicios de gimnasia cerebral para activar ambos hemisferios",
        "pasos": [
            ("✋", "Toca tu nariz DERECHA con mano IZQUIERDA — alterna 10x", 15),
            ("🔄", "Dibuja un INFINITO en el aire con tu dedo — grande y lento", 12),
            ("👣", "Marcha cruzada — rodilla izquierda con codo derecho x10", 15),
            ("✍️", "Escribe tu nombre EN EL AIRE con el codo izquierdo", 12),
            ("🧘", "Bosteza exagerado 3 veces — activa el cerebro!", 8),
        ]
    },
    {
        "id": 18,
        "nombre": "🌬️ Respiracion Magica",
        "nivel": ["TODOS"],
        "color_fondo": "#0c4a6e",
        "color_acento": "#bae6fd",
        "emoji_principal": "🫧",
        "descripcion": "Tecnicas de respiracion para calmarse y concentrarse",
        "pasos": [
            ("🫧", "Inhala por la nariz 4 segundos — hincha la barriga", 6),
            ("✋", "Cuenta en tus 4 dedos — retiene el aire 4 segundos", 5),
            ("💨", "Exhala soplando como apagar velas — 8 segundos", 10),
            ("🌊", "Respira como las OLAS del mar — sube y baja los brazos", 15),
            ("😌", "Sonrie y abre los ojos — listos para aprender!", 8),
        ]
    },
    {
        "id": 19,
        "nombre": "💃 Macarena 2.0",
        "nivel": ["TODOS"],
        "color_fondo": "#7c2d12",
        "color_acento": "#fed7aa",
        "emoji_principal": "🕺",
        "descripcion": "Secuencia de baile coordinado al ritmo — todos juntos",
        "pasos": [
            ("🤲", "Extiende el BRAZO DERECHO hacia adelante — palma abajo", 8),
            ("🤲", "Extiende el BRAZO IZQUIERDO — palma abajo — al ritmo", 8),
            ("🔄", "Gira las dos palmas — ahora miran hacia arriba!", 8),
            ("🙌", "Brazos al pecho — derecho e izquierdo cruzados", 8),
            ("💃", "Manos a la cintura y mueve las caderas — 8 tiempos!", 12),
        ]
    },
    {
        "id": 20,
        "nombre": "🏕️ Naturaleza Viva",
        "nivel": ["INICIAL","PRIMARIA"],
        "color_fondo": "#14532d",
        "color_acento": "#bbf7d0",
        "emoji_principal": "🌿",
        "descripcion": "Conectar con la naturaleza a traves del movimiento",
        "pasos": [
            ("🌱", "Eres una SEMILLA — enrroscate pequeno en tu silla", 8),
            ("🌿", "CRECE lentamente — estira brazos y piernas al cielo", 12),
            ("🌳", "Eres un ARBOL — balancea tus ramas con el viento", 12),
            ("🌻", "El SOL te da energia — abre bien los brazos y sonrie", 10),
            ("🍃", "Una BRISA te mueve — ondula todo el cuerpo suavemente", 10),
        ]
    },
]

ARCHIVO_PAUSA_MUSICA = "pausa_activa_musica.json"

def _cargar_pausa_config():
    try:
        if Path(ARCHIVO_PAUSA_MUSICA).exists():
            with open(ARCHIVO_PAUSA_MUSICA, 'r', encoding='utf-8') as f:
                return json.load(f)
    except Exception:
        pass
    return {}

def _guardar_pausa_config(data):
    try:
        with open(ARCHIVO_PAUSA_MUSICA, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False)
        return True
    except Exception:
        return False

def tab_pausa_activa(config):
    st.header("🏃 PAUSA ACTIVA")
    es_admin = st.session_state.get('rol', '') in ['admin', 'directivo']

    # ── Admin: subir MP3 ─────────────────────────────────────────────────
    if es_admin:
        with st.expander("🎵 Cargar Música MP3 por Modelo (solo Admin/Directivo)", expanded=False):
            pausa_cfg = _cargar_pausa_config()
            st.caption("Sube un archivo MP3 para cada modelo. Se guardará en el servidor.")
            for m in PAUSA_MODELOS:
                c1, c2, c3 = st.columns([2, 2, 1])
                with c1:
                    st.markdown(f"**{m['nombre']}**")
                    mp3_path = f"pausa_mp3_{m['id']}.mp3"
                    if Path(mp3_path).exists():
                        st.caption("🎵 MP3 cargado ✅")
                    else:
                        st.caption("🔇 Sin música")
                with c2:
                    mp3_file = st.file_uploader(
                        f"MP3 para modelo {m['id']}",
                        type=["mp3", "ogg", "wav"],
                        key=f"pausa_mp3_{m['id']}",
                        label_visibility="collapsed"
                    )
                    if mp3_file is not None:
                        ext_up = mp3_file.name.split(".")[-1].lower()
                        with st.spinner("☁️ Guardando y sincronizando..."):
                            saved_path = _pausa_guardar_mp3(m['id'], mp3_file.read(), ext_up)
                        pausa_cfg[str(m['id'])] = {'mp3': saved_path}
                        _guardar_pausa_config(pausa_cfg)
                        st.success("✅ Audio guardado y sincronizado en Google Sheets")
                        st.rerun()
                with c3:
                    if Path(f"pausa_mp3_{m['id']}.mp3").exists():
                        if st.button("🗑️", key=f"del_mp3_{m['id']}", help="Eliminar MP3"):
                            try:
                                Path(f"pausa_mp3_{m['id']}.mp3").unlink()
                                pausa_cfg[str(m['id'])] = {'mp3': ''}
                                _guardar_pausa_config(pausa_cfg)
                                # Limpiar en GSheets
                                try:
                                    gs = _gs()
                                    if gs:
                                        ws = gs._get_hoja('config')
                                        if ws:
                                            all_v = ws.get_all_values()
                                            for ri, row in enumerate(all_v):
                                                if row and row[0] == f"bin_pausa_mp3_{m['id']}":
                                                    ws.update_cell(ri + 1, 2, '')
                                                    break
                                except Exception:
                                    pass
                                st.rerun()
                            except Exception:
                                pass

    pausa_cfg = _cargar_pausa_config()

    # ── Selección de modelo con filtro por nivel ─────────────────────────
    st.markdown("### 🎯 Elige tu Pausa Activa")

    NIVELES_PAUSA = ["TODOS", "INICIAL", "PRIMARIA", "SECUNDARIA"]
    CN = {
        "TODOS":      ("#0f172a", "#e2e8f0", "⭐"),
        "INICIAL":    ("#065f46", "#6ee7b7", "🌱"),
        "PRIMARIA":   ("#1e3a8a", "#93c5fd", "📚"),
        "SECUNDARIA": ("#7c2d12", "#fca5a5", "🎓"),
    }
    nivel_sel = st.session_state.get('_pausa_nivel_filtro', 'TODOS')

    # ── CSS global único — estilos por key de botón ──────────────────────
    css_btns = "<style>"
    # Filtros de nivel
    for nv, (bg, fg, _ic) in CN.items():
        activo = nivel_sel == nv
        bg_btn = fg if activo else bg
        fg_btn = bg if activo else fg
        brd    = f"3px solid {fg}" if activo else f"2px solid {fg}"
        css_btns += f"""
        button[data-testid="baseButton-secondary"][key="nfilt_{nv}"],
        div[data-testid="stButton"] > button[kind="secondary"]:has(div > p:contains("{nv}")) {{
            background: {bg_btn} !important; color: {fg_btn} !important;
            border: {brd} !important; border-radius: 12px !important;
            font-weight: 800 !important; font-size: 0.9rem !important;
            min-height: 50px !important;
        }}"""
    css_btns += "</style>"
    st.markdown(css_btns, unsafe_allow_html=True)

    # Filtros — radio visual con st.button
    nf_cols = st.columns(4)
    for ni, nv in enumerate(NIVELES_PAUSA):
        bg, fg, ico = CN[nv]
        activo = nivel_sel == nv
        with nf_cols[ni]:
            bg_btn = fg if activo else bg
            fg_btn = bg if activo else fg
            brd    = f"3px solid {fg}" if activo else f"2px solid {fg}"
            st.markdown(f"""<style>
            div[data-testid="stButton"]:has(button[key="nfilt_{nv}"]) button {{
                background: {bg_btn} !important;
                color: {fg_btn} !important;
                border: {brd} !important;
                border-radius: 12px !important;
                font-weight: 800 !important;
                font-size: 0.88rem !important;
                min-height: 50px !important;
                width: 100% !important;
            }}
            div[data-testid="stButton"]:has(button[key="nfilt_{nv}"]) button:hover {{
                filter: brightness(1.15) !important;
            }}
            </style>""", unsafe_allow_html=True)
            label = f"{'✅ ' if activo else ico+' '}{nv}"
            if st.button(label, key=f"nfilt_{nv}", use_container_width=True):
                st.session_state['_pausa_nivel_filtro'] = nv
                st.rerun()

    nivel_sel = st.session_state.get('_pausa_nivel_filtro', 'TODOS')
    modelos_filtrados = [m for m in PAUSA_MODELOS
                         if nivel_sel == 'TODOS' or nivel_sel in m.get('nivel', ['TODOS'])]
    st.caption(f"📋 {len(modelos_filtrados)} pausas activas disponibles")

    modelo_seleccionado = st.session_state.get('_pausa_modelo_id', None)

    # CSS para botones Iniciar — uno por modelo usando su key único
    css_iniciar = "<style>"
    for m in PAUSA_MODELOS:
        ca = m['color_acento']
        cf = m['color_fondo']
        css_iniciar += f"""
        div[data-testid="stButton"]:has(button[key="sel_pausa_{m['id']}"]) button {{
            background: {ca} !important;
            color: {cf} !important;
            border: none !important;
            border-radius: 12px !important;
            font-weight: 800 !important;
            font-size: 1rem !important;
            min-height: 48px !important;
        }}
        div[data-testid="stButton"]:has(button[key="sel_pausa_{m['id']}"]) button:hover {{
            filter: brightness(1.12) !important;
            transform: scale(1.02) !important;
        }}"""
    css_iniciar += "</style>"
    st.markdown(css_iniciar, unsafe_allow_html=True)

    cols = st.columns(2)
    for i, m in enumerate(modelos_filtrados):
        with cols[i % 2]:
            tiene_musica = any(Path(f"pausa_mp3_{m['id']}.{ext}").exists()
                               for ext in ["mp3","ogg","wav"])
            musica_badge = "🎵" if tiene_musica else "🔇"
            seleccionado = modelo_seleccionado == m['id']
            niveles_m = m.get('nivel', ['TODOS'])
            nivel_badges = " ".join([
                f"<span style='background:{CN.get(n,('',''))[0]};color:{CN.get(n,('','#fff'))[1]};"
                f"padding:3px 9px;border-radius:10px;font-size:0.7rem;font-weight:800;'>{n}</span>"
                for n in niveles_m
            ])
            borde = f"3px solid {m['color_acento']}" if seleccionado else "2px solid rgba(255,255,255,0.15)"
            st.markdown(f"""
            <div style='background:{m["color_fondo"]}; border:{borde};
                        border-radius:16px; padding:20px 14px 12px; margin-bottom:4px;
                        text-align:center; box-shadow:0 4px 16px rgba(0,0,0,0.25);'>
                <div style='font-size:3rem; margin-bottom:8px;'>{m["emoji_principal"]}</div>
                <div style='color:white; font-size:1rem; font-weight:bold; margin:4px 0 6px;'>{m["nombre"]}</div>
                <div style='color:{m["color_acento"]}; font-size:0.78rem; margin-bottom:8px;'>{m["descripcion"]}</div>
                <div style='margin-bottom:8px; display:flex; gap:5px; justify-content:center; flex-wrap:wrap;'>{nivel_badges}</div>
                <div style='color:rgba(255,255,255,0.65); font-size:0.72rem;'>{musica_badge} {"Con musica" if tiene_musica else "Sin musica"}</div>
            </div>
            """, unsafe_allow_html=True)
            if st.button(f"▶  Iniciar  {m['emoji_principal']}", key=f"sel_pausa_{m['id']}", use_container_width=True):
                st.session_state['_pausa_modelo_id'] = m['id']
                st.session_state['_pausa_paso_actual'] = 0
                st.session_state['_pausa_activa'] = True
                st.rerun()

    # ── Presentación en pantalla grande ──────────────────────────────────
    if st.session_state.get('_pausa_activa') and modelo_seleccionado:
        modelo = next((m for m in PAUSA_MODELOS if m['id'] == modelo_seleccionado), None)
        if not modelo:
            return

        paso_idx = st.session_state.get('_pausa_paso_actual', 0)
        pasos    = modelo['pasos']
        total_pasos = len(pasos)

        # Botón cerrar arriba
        if st.button("✖ Cerrar Pausa Activa", key="pausa_cerrar_top"):
            st.session_state['_pausa_activa']    = False
            st.session_state['_pausa_modelo_id'] = None
            st.session_state['_pausa_paso_actual'] = 0
            st.rerun()

        # Preparar datos para el iframe completo
        _b64_audio, _ext_audio = _pausa_cargar_mp3_b64(modelo['id'])
        _mime = "audio/mpeg"
        _audio_tag = ""
        if _b64_audio:
            _mime = "audio/mpeg" if _ext_audio == "mp3" else f"audio/{_ext_audio}"
            _audio_src = f"data:{_mime};base64,{_b64_audio}"
            _audio_tag = f'''<audio id="bgm" loop autoplay style="display:none">
                <source src="{_audio_src}" type="{_mime}"></audio>'''

        # Construir lista JSON de pasos para JS
        import json as _json_pa
        pasos_js = _json_pa.dumps([
            {"emoji": p[0], "texto": p[1], "seg": p[2]} for p in pasos
        ])
        
        _col_fondo  = modelo['color_fondo']
        _col_acento = modelo['color_acento']

        import streamlit.components.v1 as _comp_pa
        _comp_pa.html(f"""
<!DOCTYPE html><html><head>
<meta charset="utf-8">
<style>
  * {{ box-sizing:border-box; margin:0; padding:0; }}
  body {{
    background: {_col_fondo};
    font-family: 'Segoe UI', sans-serif;
    min-height: 96vh;
    display: flex; flex-direction: column;
    overflow: hidden;
  }}

  /* Fondo animado con partículas */
  .bg-particles {{
    position:fixed; top:0; left:0; width:100%; height:100%;
    pointer-events:none; z-index:0; overflow:hidden;
  }}
  .particle {{
    position:absolute; border-radius:50%;
    background: {_col_acento};
    opacity:0.12;
    animation: float linear infinite;
  }}
  @keyframes float {{
    0%   {{ transform: translateY(110vh) scale(0.5); opacity:0.15; }}
    100% {{ transform: translateY(-10vh)  scale(1.2); opacity:0;    }}
  }}

  /* Barra de progreso */
  .prog-bar-wrap {{
    position:relative; z-index:2;
    background:rgba(255,255,255,0.15);
    height:8px; width:100%; border-radius:4px;
  }}
  .prog-bar {{
    height:8px; border-radius:4px;
    background:{_col_acento};
    transition: width 0.5s ease;
  }}

  /* Header */
  .header {{
    position:relative; z-index:2;
    padding:14px 24px 6px;
    display:flex; justify-content:space-between; align-items:center;
  }}
  .header-title {{
    color:{_col_acento}; font-size:0.95rem;
    font-weight:700; letter-spacing:2px; text-transform:uppercase;
  }}
  .step-badge {{
    background:rgba(255,255,255,0.18);
    color:white; padding:4px 14px;
    border-radius:20px; font-size:0.85rem; font-weight:600;
  }}

  /* Contenido central */
  .main-card {{
    position:relative; z-index:2;
    flex:1; display:flex; flex-direction:column;
    align-items:center; justify-content:center;
    padding:10px 30px 0;
    text-align:center;
  }}
  .emoji-big {{
    font-size:9rem; line-height:1;
    animation: pop 0.5s cubic-bezier(0.34,1.56,0.64,1) both,
               bounce 2.5s ease-in-out 0.5s infinite;
    filter: drop-shadow(0 0 30px {_col_acento});
  }}
  @keyframes pop {{
    0%   {{ transform: scale(0.2) rotate(-15deg); opacity:0; }}
    100% {{ transform: scale(1)   rotate(0deg);   opacity:1; }}
  }}
  @keyframes bounce {{
    0%,100% {{ transform: translateY(0);    }}
    50%      {{ transform: translateY(-18px); }}
  }}
  .instruccion {{
    color:white; font-size:2.4rem; font-weight:800;
    line-height:1.25; margin:22px 0 18px;
    text-shadow: 0 3px 12px rgba(0,0,0,0.4);
    animation: slideUp 0.4s ease both;
  }}
  @keyframes slideUp {{
    from {{ opacity:0; transform:translateY(30px); }}
    to   {{ opacity:1; transform:translateY(0);    }}
  }}

  /* Timer ring */
  .timer-wrap {{
    position:relative; width:120px; height:120px;
    margin:0 auto 16px;
    animation: slideUp 0.4s 0.15s ease both;
  }}
  .timer-svg {{ transform:rotate(-90deg); }}
  .timer-ring {{ transition: stroke-dashoffset 1s linear; }}
  .timer-text {{
    position:absolute; top:50%; left:50%;
    transform:translate(-50%,-50%);
    font-size:2rem; font-weight:900; color:white;
    text-shadow: 0 2px 8px rgba(0,0,0,0.5);
  }}
  .timer-label {{
    font-size:0.7rem; color:{_col_acento};
    font-weight:700; letter-spacing:2px;
    display:block; margin-top:2px;
  }}

  /* Dots */
  .dots {{
    display:flex; gap:8px; justify-content:center;
    padding:8px 0; position:relative; z-index:2;
  }}
  .dot {{
    width:10px; height:10px; border-radius:50%;
    background:rgba(255,255,255,0.3);
    transition: all 0.3s ease;
  }}
  .dot.active {{
    background:{_col_acento};
    width:28px; border-radius:5px;
    box-shadow: 0 0 10px {_col_acento};
  }}
  .dot.done {{
    background:rgba(255,255,255,0.55);
  }}

  /* Botones de navegación */
  .nav-bar {{
    position:relative; z-index:10;
    display:flex; justify-content:space-between; align-items:center;
    padding:10px 20px 16px;
    gap:12px;
  }}
  .btn {{
    border:none; cursor:pointer; border-radius:50px;
    font-size:1rem; font-weight:800; padding:13px 28px;
    transition: transform 0.15s, box-shadow 0.15s;
    display:flex; align-items:center; gap:8px;
  }}
  .btn:hover {{ transform:scale(1.06); box-shadow:0 6px 20px rgba(0,0,0,0.4); }}
  .btn:active {{ transform:scale(0.96); }}

  .btn-prev {{
    background:#1e293b; color:white;
    min-width:120px;
  }}
  .btn-prev:disabled {{
    background:#334155; color:#64748b; cursor:default;
  }}
  .btn-next {{
    background:{_col_acento}; color:{_col_fondo};
    min-width:140px; font-size:1.1rem;
  }}
  .btn-fin {{
    background:#16a34a; color:white;
    min-width:140px; font-size:1.1rem;
    animation: pulse-green 1.5s ease-in-out infinite;
  }}
  @keyframes pulse-green {{
    0%,100% {{ box-shadow:0 0 0 0 rgba(22,163,74,0.5); }}
    50%      {{ box-shadow:0 0 0 12px rgba(22,163,74,0); }}
  }}
  .btn-music {{
    background:rgba(255,255,255,0.15);
    color:white; border:2px solid rgba(255,255,255,0.3);
    padding:10px 18px; font-size:0.85rem;
  }}
  .btn-music.playing {{
    background:rgba(255,255,255,0.25);
    border-color:{_col_acento};
    color:{_col_acento};
  }}

  /* Flecha parpadeante */
  .arrow-hint {{
    color:rgba(255,255,255,0.35);
    font-size:2rem;
    animation: blink 1.8s ease-in-out infinite;
  }}
  @keyframes blink {{
    0%,100% {{ opacity:0.3; }} 50% {{ opacity:1; }}
  }}

  /* Pantalla completa hint */
  .fs-btn {{
    background:rgba(255,255,255,0.1);
    color:white; border:1px solid rgba(255,255,255,0.2);
    padding:7px 14px; font-size:0.78rem;
    border-radius:20px; cursor:pointer;
    transition:all 0.2s;
  }}
  .fs-btn:hover {{ background:rgba(255,255,255,0.2); }}
</style>
</head>
<body>
{_audio_tag}

<!-- Partículas de fondo -->
<div class="bg-particles" id="particles"></div>

<!-- Barra de progreso superior -->
<div class="prog-bar-wrap">
  <div class="prog-bar" id="progbar" style="width:0%"></div>
</div>

<!-- Header -->
<div class="header">
  <span class="header-title">{modelo['nombre']}</span>
  <div style="display:flex;gap:10px;align-items:center;">
    <span class="step-badge" id="step-badge">Paso 1 / {total_pasos}</span>
    <button class="fs-btn" onclick="toggleFS()">⛶ Pantalla completa</button>
  </div>
</div>

<!-- Contenido principal -->
<div class="main-card">
  <div class="emoji-big" id="emoji">⏳</div>
  <div class="instruccion" id="instruccion">Cargando...</div>
  <div class="timer-wrap">
    <svg class="timer-svg" width="120" height="120" viewBox="0 0 120 120">
      <circle cx="60" cy="60" r="52" fill="none"
        stroke="rgba(255,255,255,0.15)" stroke-width="8"/>
      <circle id="timer-ring" class="timer-ring" cx="60" cy="60" r="52"
        fill="none" stroke="{_col_acento}" stroke-width="8"
        stroke-dasharray="326.7" stroke-dashoffset="326.7"
        stroke-linecap="round"/>
    </svg>
    <div class="timer-text">
      <span id="timer-num">0</span>
      <span class="timer-label">SEG</span>
    </div>
  </div>
</div>

<!-- Dots de pasos -->
<div class="dots" id="dots"></div>

<!-- Barra de navegación -->
<div class="nav-bar">
  <button class="btn btn-prev" id="btn-prev" onclick="navegarPrev()">&#8592; Anterior</button>
  <button class="btn btn-music" id="btn-music" onclick="toggleMusica()">&#9654; Musica</button>
  <button class="btn btn-next" id="btn-next" onclick="navegarNext()">Siguiente &#8594;</button>
</div>

<script>
const PASOS   = {pasos_js};
const TOTAL   = PASOS.length;
let   pasoActual = {paso_idx};
let   timerInterval = null;
let   tiempoRestante = 0;
const CIRCUM = 326.7;

function crearParticulas() {{
  const cont = document.getElementById('particles');
  for(let i=0;i<22;i++) {{
    const p = document.createElement('div');
    p.className = 'particle';
    const sz = Math.random()*40+10;
    p.style.cssText = `width:${{sz}}px;height:${{sz}}px;
      left:${{Math.random()*100}}%;
      animation-duration:${{Math.random()*8+6}}s;
      animation-delay:${{Math.random()*8}}s;`;
    cont.appendChild(p);
  }}
}}

function construirDots() {{
  const d = document.getElementById('dots');
  d.innerHTML = '';
  for(let i=0;i<TOTAL;i++) {{
    const dot = document.createElement('div');
    dot.className = 'dot' + (i===pasoActual?' active':i<pasoActual?' done':'');
    d.appendChild(dot);
  }}
}}

function actualizarPaso() {{
  if(pasoActual >= TOTAL) return;
  const p = PASOS[pasoActual];

  // Animar salida y entrada
  const emoji = document.getElementById('emoji');
  const inst  = document.getElementById('instruccion');
  emoji.style.animation = 'none'; inst.style.animation = 'none';
  void emoji.offsetWidth;
  emoji.style.animation = 'pop 0.5s cubic-bezier(0.34,1.56,0.64,1) both, bounce 2.5s ease-in-out 0.5s infinite';
  inst.style.animation  = 'slideUp 0.4s ease both';

  emoji.textContent = p.emoji;
  inst.textContent  = p.texto;

  // Badge
  document.getElementById('step-badge').textContent = `Paso ${{pasoActual+1}} / ${{TOTAL}}`;

  // Progreso
  const pct = ((pasoActual+1)/TOTAL)*100;
  document.getElementById('progbar').style.width = pct+'%';

  // Dots
  construirDots();

  // Botones
  document.getElementById('btn-prev').disabled = (pasoActual===0);
  const btnNext = document.getElementById('btn-next');
  if(pasoActual === TOTAL-1) {{
    btnNext.textContent = '🎉 FINALIZAR';
    btnNext.className   = 'btn btn-fin';
  }} else {{
    btnNext.textContent = 'Siguiente →';
    btnNext.className   = 'btn btn-next';
  }}

  // Timer
  iniciarTimer(p.seg);
}}

function iniciarTimer(seg) {{
  if(timerInterval) clearInterval(timerInterval);
  tiempoRestante = seg;
  const ring = document.getElementById('timer-ring');
  const num  = document.getElementById('timer-num');

  function tick() {{
    num.textContent = tiempoRestante;
    const offset = CIRCUM * (1 - tiempoRestante/seg);
    ring.style.strokeDashoffset = CIRCUM - offset;
    if(tiempoRestante > 0) tiempoRestante--;
    else {{ clearInterval(timerInterval); flashTimer(); }}
  }}
  tick();
  timerInterval = setInterval(tick, 1000);
}}

function flashTimer() {{
  const wrap = document.querySelector('.timer-wrap');
  let v = 0;
  const f = setInterval(()=>{{
    wrap.style.opacity = (v%2===0)?'0.3':'1'; v++;
    if(v>8) {{ clearInterval(f); wrap.style.opacity='1'; }}
  }},300);
}}

function navegarNext() {{
  if(timerInterval) clearInterval(timerInterval);
  if(pasoActual < TOTAL-1) {{
    pasoActual++;
    actualizarPaso();
    notificarStreamlit('next');
  }} else {{
    notificarStreamlit('fin');
  }}
}}

function navegarPrev() {{
  if(timerInterval) clearInterval(timerInterval);
  if(pasoActual > 0) {{
    pasoActual--;
    actualizarPaso();
    notificarStreamlit('prev');
  }}
}}

function notificarStreamlit(accion) {{
  // Enviar mensaje al padre (Streamlit iframe)
  try {{ window.parent.postMessage({{type:'pausa_nav', accion, paso:pasoActual}}, '*'); }} catch(e){{}}
}}

// Teclado — flechas izquierda/derecha
document.addEventListener('keydown', e => {{
  if(e.key==='ArrowRight'||e.key===' ') navegarNext();
  if(e.key==='ArrowLeft')               navegarPrev();
}});

// Música
let musicaActiva = true;
function toggleMusica() {{
  const bgm = document.getElementById('bgm');
  const btn = document.getElementById('btn-music');
  if(!bgm) return;
  if(musicaActiva) {{ bgm.pause(); btn.textContent='▶ Musica'; btn.classList.remove('playing'); }}
  else             {{ bgm.play();  btn.textContent='⏸ Musica'; btn.classList.add('playing');    }}
  musicaActiva = !musicaActiva;
}}

// Pantalla completa
function toggleFS() {{
  if(!document.fullscreenElement) {{
    document.documentElement.requestFullscreen().catch(()=>{{}});
  }} else {{
    document.exitFullscreen();
  }}
}}

// Escuchar mensajes desde Streamlit para sincronizar paso
window.addEventListener('message', e => {{
  if(e.data && e.data.type==='pausa_set_paso') {{
    pasoActual = e.data.paso;
    actualizarPaso();
  }}
}});

// Init
crearParticulas();
actualizarPaso();
if(document.getElementById('bgm')) {{
  document.getElementById('btn-music').textContent = '⏸ Musica';
  document.getElementById('btn-music').classList.add('playing');
}}
</script>
</body></html>
""", height=680, scrolling=False)

        # Botones de Streamlit para sincronizar estado real (flechas del iframe no cambian session_state)
        _cf = modelo['color_fondo']
        _ca = modelo['color_acento']
        st.markdown(f"""<style>
        /* Anterior */
        div[data-testid="column"]:nth-child(1) button {
            background: #1e293b !important; color: #f1f5f9 !important;
            border: 2px solid #475569 !important; border-radius: 50px !important;
            font-weight: 800 !important; font-size: 1rem !important;
        }
        div[data-testid="column"]:nth-child(1) button:hover {
            background: #334155 !important;
        }
        /* Cerrar */
        div[data-testid="column"]:nth-child(2) button {
            background: #dc2626 !important; color: white !important;
            border: none !important; border-radius: 50px !important;
            font-weight: 800 !important;
        }
        div[data-testid="column"]:nth-child(2) button:hover {
            background: #b91c1c !important;
        }
        /* Siguiente / Finalizar */
        div[data-testid="column"]:nth-child(3) button {
            background: {_ca} !important; color: {_cf} !important;
            border: none !important; border-radius: 50px !important;
            font-weight: 800 !important; font-size: 1rem !important;
        }
        div[data-testid="column"]:nth-child(3) button:hover {
            opacity: 0.88 !important;
        }
        </style>""", unsafe_allow_html=True)
        nc1, nc2, nc3 = st.columns([1, 2, 1])
        with nc1:
            if paso_idx > 0:
                if st.button("◀  Anterior", use_container_width=True, key="pausa_prev"):
                    st.session_state['_pausa_paso_actual'] -= 1
                    st.rerun()
            else:
                st.markdown("&nbsp;")
        with nc2:
            if st.button("✖  Cerrar Pausa Activa", use_container_width=True, key="pausa_cerrar"):
                st.session_state['_pausa_activa']    = False
                st.session_state['_pausa_modelo_id'] = None
                st.session_state['_pausa_paso_actual'] = 0
                st.rerun()
        with nc3:
            if paso_idx < total_pasos - 1:
                if st.button("Siguiente  ▶", use_container_width=True, key="pausa_next"):
                    st.session_state['_pausa_paso_actual'] += 1
                    st.rerun()
            else:
                if st.button("🎉  FINALIZAR", use_container_width=True, key="pausa_fin"):
                    st.balloons()
                    st.session_state['_pausa_activa']    = False
                    st.session_state['_pausa_modelo_id'] = None
                    st.session_state['_pausa_paso_actual'] = 0
                    st.rerun()


def tab_yachay_plickers(config):
    """YACHAY QAWAY v2 — sync PC + celular"""
    # Restaurar quizzes desde Google Sheets si no hay locales
    if 'qaway_restored' not in st.session_state:
        try:
            _plk_restaurar_desde_gs()
            st.session_state.qaway_restored = True
        except Exception:
            st.session_state.qaway_restored = True
    st.markdown("""<div style='background:linear-gradient(135deg,#7c3aed,#2563eb);color:white;
        padding:15px 20px;border-radius:12px;text-align:center;margin-bottom:15px;'>
        <h2 style='margin:0;color:white;'>YACHAY QAWAY</h2>
        <p style='margin:4px 0 0;color:#c4b5fd;'>Qaway rikuy yachayta | PC: proyectar | Celular: escanear</p>
    </div>""", unsafe_allow_html=True)

    tab_cards, tab_quiz, tab_proj, tab_scan, tab_results = st.tabs([
        "Tarjetas QR", "Crear Preguntas", "Proyectar (PC)", "Escanear (Celular)", "Resultados"])

    usuario = st.session_state.get('usuario_actual', '')
    plickers_dir = _plk_dir()

    # ================================================================
    # TAB 1: TARJETAS
    # ================================================================
    with tab_cards:
        st.markdown("### Generar Tarjetas QR Plegables")
        st.markdown("""**Instrucciones:** Imprima 1 hoja por alumno. Recorte el cuadrado.
        Doble por las lineas punteadas (en cruz). El alumno muestra solo la cara de su respuesta.""")
        grado_p = st.selectbox("Grado:", GRADOS_OPCIONES, key="plik_grado")
        sec_p = st.selectbox("Seccion:", ["Todas"] + SECCIONES, key="plik_sec")
        df_est = BaseDatos.obtener_estudiantes_grado(grado_p, sec_p)
        st.caption(f"{len(df_est)} estudiantes en {grado_p}")
        if not df_est.empty:
            if st.button("GENERAR TARJETAS PDF", type="primary", use_container_width=True, key="btn_gen_t"):
                with st.spinner("Generando tarjetas plegables..."):
                    try:
                        pdf = _generar_pdf_tarjetas_plickers(df_est, grado_p)
                        st.success(f"{len(df_est)} tarjetas generadas")
                        st.download_button("Descargar PDF Tarjetas", pdf,
                                           f"Tarjetas_Qaway_{grado_p}.pdf", "application/pdf",
                                           use_container_width=True, type="primary", key="dl_tarj")
                    except Exception as e:
                        st.error(f"Error: {e}")
            with st.expander("Vista previa"):
                if len(df_est) > 0:
                    est_sel = st.selectbox("Estudiante:", df_est['Nombre'].tolist(), key="plik_prev_est")
                    fila = df_est[df_est['Nombre'] == est_sel].iloc[0]
                    try:
                        img = _generar_tarjeta_plickers(str(fila['Nombre']), str(fila['DNI']), 1)
                        st.image(img, width=400)
                    except Exception as e:
                        st.error(f"Error: {e}")
        st.markdown("---")
        with st.expander("🎵 Configurar musica de fondo"):
            st.caption("Suba un archivo MP3 para que suene durante las evaluaciones.")
            musica_actual = _qaway_cargar_musica()
            if musica_actual:
                st.success("Ya hay musica cargada.")
                if st.button("Eliminar musica", key="plik_del_mus"):
                    p_mus = _plk_dir() / "musica_fondo.mp3"
                    if p_mus.exists():
                        p_mus.unlink()
                    # Limpiar en GSheets
                    try:
                        gs = _gs()
                        if gs:
                            ws = gs._get_hoja('config')
                            if ws:
                                all_v = ws.get_all_values()
                                for ri, row in enumerate(all_v):
                                    if row and row[0] == "bin_qaway_mp3":
                                        ws.update_cell(ri + 1, 2, '')
                                        break
                    except Exception:
                        pass
                    st.success("Musica eliminada")
                    st.rerun()
            mp3_file = st.file_uploader("Subir MP3:", type=['mp3'], key="plik_mp3_up")
            if mp3_file:
                _qaway_guardar_musica(mp3_file.read())
                st.success(f"Musica guardada: {mp3_file.name}")
                st.rerun()

    # ================================================================
    # TAB 2: CREAR PREGUNTAS
    # ================================================================
    with tab_quiz:
        st.markdown("### Crear Cuestionario")
        if 'plickers_preguntas' not in st.session_state:
            st.session_state.plickers_preguntas = []
        if 'plickers_titulo' not in st.session_state:
            st.session_state.plickers_titulo = ""
        st.session_state.plickers_titulo = st.text_input("Titulo:",
            value=st.session_state.plickers_titulo,
            placeholder="Ej: Evaluacion Matematica Semana 1", key="plik_tit")
        ca, cg = st.columns(2)
        with ca:
            area_q = st.text_input("Area:", key="plik_aq", placeholder="Matematica")
        with cg:
            grado_q = st.selectbox("Grado:", GRADOS_OPCIONES, key="plik_gq")
        st.markdown("---")
        st.markdown("#### Agregar Pregunta")
        preg_txt = st.text_area("Pregunta:", key="plik_pt", height=80)
        img_file = st.file_uploader("Imagen opcional (ecuaciones, figuras, geometria):",
                                     type=['png','jpg','jpeg','gif','webp'], key="plik_img")
        img_b64 = ""
        img_mime = "image/png"
        if img_file:
            img_bytes = img_file.read()
            img_b64 = base64.b64encode(img_bytes).decode('utf-8')
            img_mime = img_file.type or 'image/png'
            st.image(img_bytes, caption="Vista previa", width=300)
        ca2, cb2 = st.columns(2)
        with ca2:
            opt_a = st.text_input("A:", key="plik_oa")
            opt_c = st.text_input("C:", key="plik_oc")
        with cb2:
            opt_b = st.text_input("B:", key="plik_ob")
            opt_d = st.text_input("D:", key="plik_od")
        correcta = st.radio("Respuesta correcta:", ['A','B','C','D'], horizontal=True, key="plik_cor")
        if st.button("AGREGAR PREGUNTA", type="primary", key="btn_add_p"):
            if preg_txt and opt_a and opt_b:
                q_item = {'pregunta': preg_txt,
                    'opciones': {'A': opt_a, 'B': opt_b, 'C': opt_c, 'D': opt_d},
                    'correcta': correcta}
                if img_b64:
                    q_item['imagen'] = img_b64
                    q_item['imagen_mime'] = img_mime
                st.session_state.plickers_preguntas.append(q_item)
                st.success(f"Pregunta {len(st.session_state.plickers_preguntas)} agregada")
                st.rerun()
        if st.session_state.plickers_preguntas:
            st.markdown("---")
            st.markdown(f"#### Preguntas ({len(st.session_state.plickers_preguntas)})")
            for i, p in enumerate(st.session_state.plickers_preguntas):
                with st.expander(f"P{i+1}: {p['pregunta'][:60]}"):
                    if p.get('imagen'):
                        try:
                            st.image(base64.b64decode(p['imagen']), width=250)
                        except Exception:
                            pass
                    for letra, texto in p['opciones'].items():
                        marca = ">>>" if letra == p['correcta'] else "   "
                        st.write(f"{marca} **{letra}:** {texto}")
                    if st.button(f"Eliminar P{i+1}", key=f"del_p_{i}"):
                        st.session_state.plickers_preguntas.pop(i)
                        st.rerun()
            cs, cc2 = st.columns(2)
            with cs:
                if st.button("GUARDAR CUESTIONARIO", type="primary", use_container_width=True, key="btn_sv_q"):
                    sesion_id = f"{usuario}_{fecha_peru_str()}"
                    nombre_doc = _nombre_completo_docente()
                    qd = {'titulo': st.session_state.plickers_titulo, 'area': area_q,
                           'grado': grado_q, 'fecha': fecha_peru_str(), 'docente': nombre_doc, 'usuario': usuario,
                           'preguntas': st.session_state.plickers_preguntas,
                           'pregunta_actual': 0, 'sesion_id': sesion_id}
                    _plk_guardar_sesion(sesion_id, qd)
                    # Tambien guardar como quiz
                    aq = plickers_dir / f"quiz_{sesion_id}.json"
                    with open(aq, 'w', encoding='utf-8') as fq:
                        json.dump(qd, fq, indent=2, ensure_ascii=False)
                    # Sync a Google Sheets para persistencia
                    _plk_sync_quiz_a_gs(sesion_id, qd)
                    st.success(f"Cuestionario guardado. ID: {sesion_id}")
            with cc2:
                if st.button("Limpiar todo", use_container_width=True, type="primary", key="btn_cl_q"):
                    st.session_state.plickers_preguntas = []
                    st.session_state.plickers_titulo = ""
                    st.rerun()

            # Boton PDF del cuestionario
            if st.session_state.plickers_preguntas:
                st.markdown("---")
                if st.button("📄 DESCARGAR PDF DEL CUESTIONARIO", type="primary", use_container_width=True, key="btn_pdf_quiz"):
                    qd_pdf = {"titulo": st.session_state.plickers_titulo, "area": area_q,
                              "grado": grado_q, "fecha": fecha_peru_str(), "docente": usuario,
                              "preguntas": st.session_state.plickers_preguntas}
                    try:
                        pdf_quiz = _generar_pdf_cuestionario_qaway(qd_pdf)
                        st.download_button("Descargar PDF", pdf_quiz,
                                           f"Cuestionario_{grado_q}_{area_q}.pdf",
                                           "application/pdf", type="primary", key="dl_quiz_pdf")
                    except Exception as e:
                        st.error(f"Error generando PDF: {e}")

    # ================================================================
    # TAB 3: PROYECTAR (PC) — Pantalla grande + nombres en vivo
    # ================================================================
    with tab_proj:
        st.markdown("### Proyectar Preguntas (abrir en PC/proyector)")
        st.caption("Los estudiantes responden con sus tarjetas QR. El profesor escanea desde el celular en otra pestana.")

        qf = sorted([q for q in plickers_dir.glob("quiz_*.json")
                              if st.session_state.get('rol','') in ('directivo','admin') or str(usuario) in str(q.stem)], reverse=True)
        if not qf:
            st.warning("Primero cree un cuestionario.")
        else:
            qs = st.selectbox("Cuestionario:", qf,
                              format_func=_plk_format_quiz, key="plik_proj_qs")
            with open(qs, 'r', encoding='utf-8') as fq:
                quiz = json.load(fq)
            sesion_id = quiz.get('sesion_id', qs.stem.replace('quiz_',''))
            total_p = len(quiz['preguntas'])

            # Controles de navegacion
            if 'plik_pidx' not in st.session_state:
                st.session_state.plik_pidx = quiz.get('pregunta_actual', 0)
            pidx = st.session_state.plik_pidx

            if pidx < total_p:
                preg = quiz['preguntas'][pidx]
                op = preg['opciones']

                # Actualizar pregunta actual en archivo para sync
                quiz['pregunta_actual'] = pidx
                _plk_guardar_sesion(sesion_id, quiz)

                # Modo pantalla completa + musica de fondo
                col_fs, col_mus = st.columns([1, 1])
                with col_fs:
                    fs_on = st.checkbox("🖥️ PANTALLA GIGANTE", key="plik_fs_chk")
                with col_mus:
                    musica_on = st.checkbox("🎵 Musica de fondo", key="plik_musica", value=st.session_state.get("plik_musica_v", False))
                    st.session_state.plik_musica_v = musica_on
                if musica_on:
                    import streamlit.components.v1 as comp_m
                    musica_b64 = _qaway_cargar_musica()
                    if musica_b64:
                        audio_src = 'data:audio/mpeg;base64,' + musica_b64
                    else:
                        audio_src = 'https://cdn.pixabay.com/audio/2022/02/22/audio_d1718ab41b.mp3'
                    comp_m.html('<div style="text-align:center;padding:4px;">'
                        '<audio id="bgm" loop autoplay><source src="' + audio_src + '" type="audio/mpeg"></audio>'
                        '<button onclick="var a=document.getElementById(\'bgm\');if(a.paused){a.play();this.textContent=\'🔊 Pausar\'}else{a.pause();this.textContent=\'🔈 Reproducir\'}"'
                        ' style="background:#7c3aed;color:white;border:none;padding:6px 16px;border-radius:8px;cursor:pointer">🔊 Pausar</button>'
                        '</div>', height=50)

                # PANTALLA DE PROYECCION (grande, para proyector)
                st.markdown(f"""<div style='background:#1e1b4b;color:white;padding:30px;
                    border-radius:16px;text-align:center;margin:10px 0;'>
                    <p style='color:#a5b4fc;margin:0;font-size:1.2rem;'>Pregunta {pidx+1} de {total_p}</p>
                    <h1 style='color:white;margin:15px 0;font-size:2.2rem;'>{preg['pregunta']}</h1>
                    <div style='display:grid;grid-template-columns:1fr 1fr;gap:15px;margin-top:20px;max-width:700px;margin-left:auto;margin-right:auto;'>
                        <div style='background:#16a34a;padding:15px 20px;border-radius:12px;font-size:1.2rem;'>
                            <strong>A)</strong> {op.get('A','')}</div>
                        <div style='background:#2563eb;padding:15px 20px;border-radius:12px;font-size:1.2rem;'>
                            <strong>B)</strong> {op.get('B','')}</div>
                        <div style='background:#d97706;padding:15px 20px;border-radius:12px;font-size:1.2rem;'>
                            <strong>C)</strong> {op.get('C','')}</div>
                        <div style='background:#db2777;padding:15px 20px;border-radius:12px;font-size:1.2rem;'>
                            <strong>D)</strong> {op.get('D','')}</div>
                    </div></div>""", unsafe_allow_html=True)
                # Imagen de la pregunta (ecuaciones, geometria)
                if preg.get('imagen'):
                    try:
                        st.image(base64.b64decode(preg['imagen']), width=500)
                    except Exception:
                        pass

                # PANTALLA GIGANTE (overlay completo)
                if fs_on:
                    import streamlit.components.v1 as comp_giant
                    img_html = ''
                    if preg.get('imagen'):
                        img_html = '<img src="data:image/png;base64,' + preg['imagen'] + '" style="max-width:500px;max-height:250px;border-radius:12px;margin:10px auto;display:block;">'
                    opts = preg['opciones']
                    giant_html = '''
                    <style>
                    * { margin:0; padding:0; box-sizing:border-box; }
                    body { background:#0f0a2e; color:white; font-family:sans-serif; display:flex; flex-direction:column; align-items:center; justify-content:center; min-height:100vh; padding:20px; }
                    .q-num { color:#a5b4fc; font-size:1.5rem; margin-bottom:10px; }
                    .q-text { font-size:2.5rem; font-weight:bold; text-align:center; margin-bottom:20px; line-height:1.3; }
                    .opts { display:grid; grid-template-columns:1fr 1fr; gap:18px; max-width:900px; width:100%; }
                    .opt { padding:22px 30px; border-radius:14px; font-size:1.5rem; font-weight:bold; text-align:center; }
                    .opt-a { background:#16a34a; } .opt-b { background:#2563eb; } .opt-c { background:#d97706; } .opt-d { background:#db2777; }
                    </style>
                    '''
                    giant_body = f'''
                    <div class='q-num'>Pregunta {pidx+1} de {total_p}</div>
                    <div class='q-text'>{preg['pregunta']}</div>
                    ''' + img_html + f'''
                    <div class='opts'>
                        <div class='opt opt-a'>A) {opts.get('A','')}</div>
                        <div class='opt opt-b'>B) {opts.get('B','')}</div>
                        <div class='opt opt-c'>C) {opts.get('C','')}</div>
                        <div class='opt opt-d'>D) {opts.get('D','')}</div>
                    </div>
                    '''
                    comp_giant.html(giant_html + giant_body + '<div style="margin-top:15px;text-align:center;"><button onclick="document.documentElement.requestFullscreen().catch(e=>{})" style="background:#7c3aed;color:white;border:none;padding:10px 25px;border-radius:10px;font-size:1.1rem;cursor:pointer;">⛶ Click aqui para pantalla completa</button></div>', height=750, scrolling=False)

                # PANEL LATERAL: nombres que respondieron (en vivo)
                resp_all = _plk_cargar_respuestas(sesion_id)
                resp_preg = {}
                for dni_r, datos_r in resp_all.items():
                    r_item = datos_r.get('respuestas', {}).get(str(pidx))
                    if r_item:
                        resp_preg[dni_r] = {'nombre': datos_r['nombre'], **r_item}

                st.markdown("---")
                n_resp = len(resp_preg)
                correctos = sum(1 for r in resp_preg.values() if r.get('ok'))
                incorrectos = n_resp - correctos

                cr1, cr2, cr3 = st.columns(3)
                cr1.metric("Respondieron", n_resp)
                cr2.metric("Correctas", correctos)
                cr3.metric("Incorrectas", incorrectos)

                if resp_preg:
                    st.progress(correctos / max(n_resp, 1))
                    # Mostrar nombres con color
                    cols_nombres = st.columns(4)
                    for idx_r, (dni_r, r_data) in enumerate(resp_preg.items()):
                        col_i = idx_r % 4
                        color = '#16a34a' if r_data.get('ok') else '#dc2626'
                        emoji = 'OK' if r_data.get('ok') else 'X'
                        nombre_corto = r_data['nombre'].split()[-1] if ' ' in r_data['nombre'] else r_data['nombre']
                        cols_nombres[col_i].markdown(
                            f"<div style='background:{color};color:white;padding:5px 8px;border-radius:6px;margin:2px;font-size:0.85rem;text-align:center;'>"
                            f"{emoji} {nombre_corto} = {r_data['resp']}</div>", unsafe_allow_html=True)

                # Boton refrescar (para ver nuevas respuestas)
                st.markdown("---")
                col_ref, col_prev, col_next = st.columns([1, 1, 1])
                with col_ref:
                    if st.button("🔄 REFRESCAR", use_container_width=True, key="plik_refresh", type="primary"):
                        st.rerun()
                    # Auto-refresh cada 5 seg via time
                    auto_ref = st.checkbox("Auto-refrescar (5s)", key="plik_auto_ref", value=False)
                    if auto_ref:
                        import time as _t_ref
                        if 'plik_last_ref' not in st.session_state:
                            st.session_state.plik_last_ref = _t_ref.time()
                        if _t_ref.time() - st.session_state.plik_last_ref > 5:
                            st.session_state.plik_last_ref = _t_ref.time()
                            st.rerun()

                with col_prev:
                    if pidx > 0 and st.button("ANTERIOR", use_container_width=True, type="primary", key="plik_prev"):
                        st.session_state.plik_pidx -= 1
                        st.rerun()
                with col_next:
                    if pidx < total_p - 1:
                        if st.button("SIGUIENTE", type="primary", use_container_width=True, key="plik_next"):
                            st.session_state.plik_pidx += 1
                            st.rerun()
                    else:
                        if st.button("FINALIZAR", type="primary", use_container_width=True, key="plik_fin"):
                            st.session_state.plik_pidx = total_p
                            st.rerun()
            else:
                # ============ PODIUM ESTILO KAHOOT ============
                resp_final = _plk_cargar_respuestas(sesion_id)
                if resp_final:
                    tpr_f = len(quiz["preguntas"])
                    ranking = []
                    for dni_f, pr_f in resp_final.items():
                        nm_f = pr_f.get("nombre", dni_f)
                        resps_f = pr_f.get("respuestas", {})
                        cor_f = sum(1 for r in resps_f.values() if r.get("ok"))
                        nota_f = round(cor_f / max(tpr_f, 1) * 20, 1)
                        ranking.append({"nombre": nm_f, "correctas": cor_f, "total": tpr_f, "nota": nota_f, "dni": dni_f})
                    ranking.sort(key=lambda x: x["nota"], reverse=True)

                    # Podium animado Top 3
                    top3 = ranking[:3]
                    medallas = ["🥇", "🥈", "🥉"]
                    colores_pod = ["#FFD700", "#C0C0C0", "#CD7F32"]
                    alturas = [220, 170, 130]

                    # Fanfarria
                    import streamlit.components.v1 as comp_pod
                    comp_pod.html("""
                    <audio autoplay><source src="https://cdn.pixabay.com/audio/2021/08/04/audio_0625c1539c.mp3" type="audio/mpeg"></audio>
                    <style>
                    @keyframes slideUp { from { transform: translateY(100px); opacity:0; } to { transform: translateY(0); opacity:1; } }
                    @keyframes glow { 0%,100% { text-shadow: 0 0 10px gold; } 50% { text-shadow: 0 0 30px gold, 0 0 60px orange; } }
                    @keyframes confetti { 0% { transform: translateY(-10px) rotate(0deg); opacity:1; } 100% { transform: translateY(400px) rotate(720deg); opacity:0; } }
                    .podium-title { font-size:2.5rem; text-align:center; animation: glow 2s infinite; color: #FFD700; font-weight:bold; margin:15px 0; }
                    .confetti { position:absolute; width:10px; height:10px; border-radius:2px; animation: confetti 3s ease-out forwards; }
                    </style>
                    <div style="position:relative;overflow:hidden;padding:20px;">
                    <div class="podium-title">🏆 RANKING FINAL 🏆</div>
                    <div id="confetti-box" style="position:absolute;top:0;left:0;width:100%;height:100%;pointer-events:none;"></div>
                    </div>
                    <script>
                    const box=document.getElementById("confetti-box");
                    const colors=["#FFD700","#FF6B6B","#4ECDC4","#45B7D1","#96CEB4","#FFEAA7","#DDA0DD","#98D8C8"];
                    for(let i=0;i<60;i++){const c=document.createElement("div");c.className="confetti";c.style.left=Math.random()*100+"%";c.style.background=colors[Math.floor(Math.random()*colors.length)];c.style.animationDelay=Math.random()*2+"s";c.style.animationDuration=(2+Math.random()*2)+"s";box.appendChild(c)}
                    </script>
                    """, height=120)

                    # Podium visual
                    if len(top3) >= 1:
                        # Orden visual: 2do - 1ro - 3ro
                        orden = []
                        if len(top3) >= 2: orden.append((top3[1], medallas[1], colores_pod[1], alturas[1]))
                        orden.append((top3[0], medallas[0], colores_pod[0], alturas[0]))
                        if len(top3) >= 3: orden.append((top3[2], medallas[2], colores_pod[2], alturas[2]))

                        cols_pod = st.columns(len(orden))
                        for ci, (alumno, medalla, color_p, altura) in enumerate(orden):
                            with cols_pod[ci]:
                                nm_pod = alumno["nombre"].split()[-1] if " " in alumno["nombre"] else alumno["nombre"]
                                st.markdown(f"""
                                <div style="text-align:center;animation:slideUp 1s ease-out;">
                                    <div style="font-size:3rem;margin-bottom:5px;">{medalla}</div>
                                    <div style="font-size:1.3rem;font-weight:bold;color:#1e1b4b;">{nm_pod}</div>
                                    <div style="font-size:1.1rem;color:#7c3aed;font-weight:bold;">{alumno["nota"]}/20</div>
                                    <div style="font-size:0.9rem;color:#666;">{alumno["correctas"]}/{alumno["total"]} correctas</div>
                                    <div style="background:{color_p};width:80%;margin:10px auto 0;height:{altura}px;
                                        border-radius:12px 12px 0 0;display:flex;align-items:center;justify-content:center;
                                        font-size:2rem;font-weight:bold;color:white;text-shadow:1px 1px 3px rgba(0,0,0,0.3);
                                        animation:slideUp 1.5s ease-out;">{medalla}</div>
                                </div>
                                """, unsafe_allow_html=True)

                    st.markdown("---")
                    # Lista completa
                    st.markdown("### 📋 Ranking Completo")
                    for idx_rk, al in enumerate(ranking):
                        puesto = medallas[idx_rk] if idx_rk < 3 else f"{idx_rk+1}to"
                        color_rk = "#16a34a" if al["nota"] >= 14 else ("#2563eb" if al["nota"] >= 11 else "#dc2626")
                        st.markdown(f"""
                        <div style="display:flex;align-items:center;padding:8px 12px;margin:4px 0;border-radius:8px;
                            background:linear-gradient(90deg,{color_rk}15,transparent);border-left:4px solid {color_rk};">
                            <span style="font-size:1.3rem;margin-right:10px;min-width:35px;">{puesto}</span>
                            <span style="flex:1;font-weight:bold;">{al["nombre"]}</span>
                            <span style="color:{color_rk};font-weight:bold;font-size:1.1rem;">{al["nota"]}/20</span>
                            <span style="color:#666;margin-left:10px;font-size:0.85rem;">{al["correctas"]}/{al["total"]}</span>
                        </div>
                        """, unsafe_allow_html=True)
                else:
                    st.info("No hay respuestas registradas.")

                if st.button("🔄 Reiniciar", key="plik_reset", type="primary"):
                    st.session_state.plik_pidx = 0
                    st.rerun()


    # ================================================================
    # TAB 4: ESCANEAR (CELULAR) — Camara foto o manual
    # ================================================================
    with tab_scan:
        st.markdown("### Escanear QR (abrir en celular)")
        st.caption("Abra esta pestana desde su celular. Escanee las tarjetas de los alumnos.")

        qf2 = sorted([q for q in plickers_dir.glob("quiz_*.json")
                              if st.session_state.get('rol','') in ('directivo','admin') or str(usuario) in str(q.stem)], reverse=True)
        if not qf2:
            st.warning("No hay cuestionarios creados.")
        else:
            qs2 = st.selectbox("Cuestionario:", qf2,
                               format_func=_plk_format_quiz, key="plik_scan_qs")
            with open(qs2, 'r', encoding='utf-8') as fq:
                quiz2 = json.load(fq)
            sesion_id2 = quiz2.get('sesion_id', qs2.stem.replace('quiz_',''))
            total_p2 = len(quiz2['preguntas'])

            # Leer pregunta actual desde archivo (sync con PC)
            sesion_data = _plk_cargar_sesion(sesion_id2)
            pidx2 = sesion_data.get('pregunta_actual', 0) if sesion_data else 0

            if pidx2 < total_p2:
                preg2 = quiz2['preguntas'][pidx2]
                op2 = preg2['opciones']

                # Mini display de la pregunta
                st.markdown(f"""<div style='background:#1e1b4b;color:white;padding:12px;
                    border-radius:10px;text-align:center;margin:8px 0;'>
                    <small style='color:#a5b4fc;'>P{pidx2+1}/{total_p2}</small><br>
                    <strong style='font-size:1.1rem;'>{preg2['pregunta']}</strong><br>
                    <small>A:{op2.get('A','')} | B:{op2.get('B','')} | C:{op2.get('C','')} | D:{op2.get('D','')}</small>
                </div>""", unsafe_allow_html=True)
                # Imagen pregunta
                if preg2.get('imagen'):
                    try:
                        st.image(base64.b64decode(preg2['imagen']), width=300)
                    except Exception:
                        pass

                # CAMARA PARA ESCANEAR
                scan_modo = st.radio("Modo de escaneo:", ["Manual", "Camara QR"], horizontal=True, key="plik_scan_m")

                if scan_modo == "Camara QR":
                    foto = st.camera_input("Apunte al QR del alumno:", key="plik_foto2")
                    if foto:
                        d = decodificar_qr_imagen(foto.getvalue())
                        if d and (d.startswith("YP|") or d.startswith("YQ_")):
                            if d.startswith("YP|"):
                                partes = d.split("|")
                            else:
                                partes = d.replace("YQ_", "").split("_")
                                partes = ["", partes[0], partes[1]] if len(partes) >= 2 else partes
                            if len(partes) >= 3:
                                dni_qr = partes[1]
                                resp_qr = partes[2].upper()
                                if resp_qr in ('A','B','C','D'):
                                    pqr = BaseDatos.buscar_por_dni(dni_qr)
                                    nqr = pqr.get('Nombre', dni_qr) if pqr else dni_qr
                                    esc = resp_qr == preg2['correcta']
                                    # GUARDAR EN ARCHIVO (sync con PC)
                                    _plk_guardar_respuesta(sesion_id2, dni_qr, nqr, pidx2, resp_qr, preg2['correcta'])
                                    col_r = "#16a34a" if esc else "#dc2626"
                                    em = "CORRECTO" if esc else "INCORRECTO"
                                    st.markdown(f"<div style='background:{col_r};color:white;padding:12px;border-radius:8px;text-align:center;font-size:1.3rem;'>{em}: <strong>{nqr}</strong> = <strong>{resp_qr}</strong></div>", unsafe_allow_html=True)
                                    reproducir_beep_exitoso()
                        elif d:
                            st.warning(f"QR no reconocido: {d[:30]}")
                        else:
                            st.warning("No se detecto QR.")
                elif scan_modo == "Manual":
                    # MANUAL
                    gsc2 = quiz2.get('grado', '')
                    dfs2 = BaseDatos.obtener_estudiantes_grado(gsc2, "Todas")
                    if not dfs2.empty:
                        em2 = st.selectbox("Alumno:", dfs2['Nombre'].tolist(), key="plik_me2")
                        rm2 = st.radio("Respuesta:", ['A','B','C','D'], horizontal=True, key="plik_mr2")
                        if st.button("Registrar", type="primary", key="btn_mr2"):
                            fm2 = dfs2[dfs2['Nombre']==em2].iloc[0]
                            dm2 = str(fm2['DNI'])
                            ec2 = rm2 == preg2['correcta']
                            _plk_guardar_respuesta(sesion_id2, dm2, em2, pidx2, rm2, preg2['correcta'])
                            st.success(f"{'CORRECTO' if ec2 else 'INCORRECTO'}: {em2} = {rm2}")
                            st.rerun()

                # Resumen de esta pregunta
                resp_scan = _plk_cargar_respuestas(sesion_id2)
                n_scan = sum(1 for d in resp_scan.values() if str(pidx2) in d.get('respuestas', {}))
                st.caption(f"{n_scan} respuestas registradas para esta pregunta")

                # Sync: refrescar para ver siguiente pregunta
                if st.button("REFRESCAR pregunta", use_container_width=True, type="primary", key="plik_scan_ref"):
                    st.rerun()
            else:
                st.success("Cuestionario finalizado!")

    # ================================================================
    # TAB 5: RESULTADOS (funciona en PC y celular)
    # ================================================================
    with tab_results:
        st.markdown("### Resultados y Ranking")
        qfr = sorted([q for q in plickers_dir.glob("quiz_*.json")
                              if st.session_state.get('rol','') in ('directivo','admin') or str(usuario) in str(q.stem)], reverse=True)
        if not qfr:
            st.info("No hay cuestionarios.")
        else:
            qsr = st.selectbox("Cuestionario:", qfr,
                               format_func=_plk_format_quiz, key="plik_qr")
            with open(qsr, 'r', encoding='utf-8') as fq:
                qr2 = json.load(fq)
            sesion_id_r = qr2.get('sesion_id', qsr.stem.replace('quiz_',''))
            # Guardar en reportes de alumnos
            if st.button("💾 GUARDAR EN REPORTES DE ALUMNOS", use_container_width=True, type="primary", key="btn_qaway_save_rep"):
                try:
                    if _plk_guardar_en_reportes(qr2, sesion_id_r):
                        # Sync respuestas a GS
                        _plk_sync_resp_a_gs(sesion_id_r)
                        st.success("Resultados guardados en el historial y Google Sheets.")
                        reproducir_beep_exitoso()
                    else:
                        st.warning("No hay respuestas para guardar.")
                except Exception as e:
                    st.error(f"Error al guardar: {e}")
            col_pdf_q, col_pdf_r2 = st.columns(2)
            with col_pdf_q:
                if st.button("📄 PDF Preguntas", use_container_width=True, type="primary", key="btn_dl_quiz_r"):
                    try:
                        pdf_q = _generar_pdf_cuestionario_qaway(qr2)
                        st.download_button("Descargar PDF", pdf_q,
                                           'Cuestionario_' + qr2.get('grado','') + '.pdf',
                                           "application/pdf", type="primary", key="dl_qr_pdf")
                    except Exception as e:
                        st.error(f"Error: {e}")
            with col_pdf_r2:
                if st.button("🏆 PDF Ranking", use_container_width=True, type="primary", key="btn_dl_rank_r"):
                    resp_rk = _plk_cargar_respuestas(sesion_id_r)
                    if not resp_rk:
                        st.warning("No hay respuestas aun.")
                    else:
                        tpr_rk = len(qr2['preguntas'])
                        res_rk = []
                        for dni_rk, pr_rk in resp_rk.items():
                            nm_rk = pr_rk.get('nombre', dni_rk)
                            resps_rk = pr_rk.get('respuestas', {})
                            cor_rk = sum(1 for r in resps_rk.values() if r.get('ok'))
                            nota_rk = round(cor_rk / max(tpr_rk, 1) * 20, 1)
                            res_rk.append({'Puesto': '', 'Nombre': nm_rk, 'Correctas': f"{cor_rk}/{tpr_rk}", 'Nota': str(nota_rk)})
                        res_rk.sort(key=lambda x: float(x['Nota']), reverse=True)
                        med_rk = ['1ro','2do','3ro']
                        for idx_rk, r in enumerate(res_rk):
                            r['Puesto'] = med_rk[idx_rk] if idx_rk < 3 else f"{idx_rk+1}to"
                        buf_rk = io.BytesIO()
                        cp_rk = canvas.Canvas(buf_rk, pagesize=A4)
                        wp_rk, hp_rk = A4
                        cp_rk.setFillColor(colors.HexColor("#7c3aed"))
                        cp_rk.rect(0, hp_rk-55, wp_rk, 55, fill=1, stroke=0)
                        cp_rk.setFillColor(colors.white)
                        cp_rk.setFont("Helvetica-Bold", 16)
                        cp_rk.drawCentredString(wp_rk/2, hp_rk-22, "YACHAY QAWAY - RANKING")
                        cp_rk.setFont("Helvetica", 10)
                        cp_rk.drawCentredString(wp_rk/2, hp_rk-38, f"{qr2.get('titulo','')} | {qr2.get('area','')} | {qr2.get('grado','')}")
                        cp_rk.setFont("Helvetica", 8)
                        pm_rk = sum(float(r['Nota']) for r in res_rk) / max(len(res_rk), 1)
                        ap_rk = sum(1 for r in res_rk if float(r['Nota']) >= 11)
                        cp_rk.drawCentredString(wp_rk/2, hp_rk-50, f"Docente: {qr2.get('docente','')} | Promedio: {pm_rk:.1f}/20 | Aprobados: {ap_rk}/{len(res_rk)} | {qr2.get('fecha','')}")
                        cp_rk.setFillColor(colors.black)
                        hdr_rk = ["PUESTO", "ALUMNO", "CORRECTAS", "NOTA /20"]
                        rows_rk = [hdr_rk] + [[r['Puesto'], r['Nombre'], r['Correctas'], r['Nota']] for r in res_rk]
                        tb_rk = Table(rows_rk, colWidths=[50, 250, 80, 70])
                        tb_rk.setStyle(TableStyle([
                            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                            ('FONTSIZE', (0,0), (-1,-1), 9),
                            ('GRID', (0,0), (-1,-1), 0.5, colors.black),
                            ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#7c3aed")),
                            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
                            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                            ('ALIGN', (1,1), (1,-1), 'LEFT'),
                            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.Color(0.95,0.93,1.0)])]))
                        tw_rk, th_rk = tb_rk.wrap(wp_rk-40, hp_rk-80)
                        tb_rk.drawOn(cp_rk, 20, hp_rk-70-th_rk)
                        cp_rk.setFont("Helvetica", 7)
                        cp_rk.drawString(20, 12, f"I.E.P. ALTERNATIVO YACHAY | YACHAY QAWAY | {qr2.get('fecha','')}")
                        cp_rk.save()
                        buf_rk.seek(0)
                        st.download_button("Descargar Ranking PDF", buf_rk, "Ranking_Qaway.pdf", "application/pdf", type="primary", key="dl_rank_top")
            resp = _plk_cargar_respuestas(sesion_id_r)
            if not resp:
                st.info("No hay respuestas registradas.")
            else:
                tpr = len(qr2['preguntas'])
                res = []
                for dni_r, pr in resp.items():
                    nm_r = pr.get('nombre', dni_r)
                    resps = pr.get('respuestas', {})
                    cor_r = sum(1 for r in resps.values() if r.get('ok'))
                    det_r = []
                    for pi_r in range(tpr):
                        r_item = resps.get(str(pi_r))
                        if r_item:
                            det_r.append('OK' if r_item.get('ok') else 'X')
                        else:
                            det_r.append("-")
                    nota_r = round(cor_r / max(tpr, 1) * 20, 1)
                    res.append({'DNI': dni_r, 'Nombre': nm_r, 'Correctas': cor_r,
                                'Total': tpr, 'Nota': nota_r, 'Detalle': ' '.join(det_r)})
                res.sort(key=lambda x: x['Nota'], reverse=True)
                puestos = {0: '1ro', 1: '2do', 2: '3ro'}
                for i_r, r_item in enumerate(res):
                    r_item['Puesto'] = puestos.get(i_r, f'{i_r+1}to')
                dfr = pd.DataFrame(res)
                cols_r = ['Puesto', 'Nombre', 'Correctas', 'Total', 'Nota', 'Detalle']
                cols_r = [c for c in cols_r if c in dfr.columns]
                st.dataframe(dfr[cols_r], use_container_width=True, hide_index=True)
                # METRICAS
                cm1, cm2, cm3 = st.columns(3)
                cm1.metric("Participantes", len(res))
                pm_r = sum(r['Nota'] for r in res) / max(len(res), 1)
                cm2.metric("Promedio", f"{pm_r:.1f}/20")
                ap_r = sum(1 for r in res if r['Nota'] >= 11)
                cm3.metric("Aprobados", f"{ap_r}/{len(res)}")
                # GRAFICO
                try:
                    import altair as alt
                    cd_r = [{'Alumno': r['Nombre'].split()[-1][:12] if ' ' in r['Nombre'] else r['Nombre'][:12], 'Nota': r['Nota']} for r in res]
                    dfc_r = pd.DataFrame(cd_r)
                    bar_r = alt.Chart(dfc_r).mark_bar().encode(
                        x=alt.X('Alumno:N', sort='-y'),
                        y=alt.Y('Nota:Q', scale=alt.Scale(domain=[0, 20])),
                        color=alt.condition(alt.datum.Nota >= 11, alt.value('#22c55e'), alt.value('#ef4444'))
                    ).properties(height=300, title=qr2.get('titulo', 'Resultados'))
                    st.altair_chart(bar_r, use_container_width=True)
                except Exception:
                    pass
                # DESCARGAS Y WHATSAPP
                st.markdown("---")
                col_pdf_r, col_wa_r = st.columns(2)
                with col_pdf_r:
                    if st.button("PDF Resultados", type="primary", use_container_width=True, key="plik_pdf"):
                        buf_r = io.BytesIO()
                        cp_r = canvas.Canvas(buf_r, pagesize=A4)
                        wp_r, hp_r = A4
                        cp_r.setFillColor(colors.HexColor("#7c3aed"))
                        cp_r.rect(0, hp_r-55, wp_r, 55, fill=1, stroke=0)
                        cp_r.setFillColor(colors.white)
                        cp_r.setFont("Helvetica-Bold", 16)
                        cp_r.drawCentredString(wp_r/2, hp_r-22, "YACHAY QAWAY - RESULTADOS")
                        cp_r.setFont("Helvetica", 10)
                        cp_r.drawCentredString(wp_r/2, hp_r-38, f"{qr2.get('titulo','')} | {qr2.get('area','')} | {qr2.get('grado','')}")
                        cp_r.setFont("Helvetica", 8)
                        cp_r.drawCentredString(wp_r/2, hp_r-50, f"Promedio: {pm_r:.1f}/20 | Aprobados: {ap_r}/{len(res)} | {fecha_peru_str()}")
                        cp_r.setFillColor(colors.black)
                        hdr_r = ["PUESTO", "ALUMNO", "CORRECTAS", "NOTA /20"]
                        rows_r = [hdr_r]
                        for r2 in res:
                            rows_r.append([r2['Puesto'], r2['Nombre'], f"{r2['Correctas']}/{r2['Total']}", str(r2['Nota'])])
                        tb_r = Table(rows_r, colWidths=[50, 250, 80, 70])
                        tb_r.setStyle(TableStyle([
                            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                            ('FONTSIZE', (0,0), (-1,-1), 9),
                            ('GRID', (0,0), (-1,-1), 0.5, colors.black),
                            ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#7c3aed")),
                            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
                            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                            ('ALIGN', (1,1), (1,-1), 'LEFT'),
                            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.Color(0.95,0.93,1.0)])]))
                        tw_r, th_r = tb_r.wrap(wp_r-40, hp_r-80)
                        tb_r.drawOn(cp_r, 20, hp_r-70-th_r)
                        cp_r.setFont("Helvetica", 7)
                        cp_r.drawString(20, 12, f"I.E.P. ALTERNATIVO YACHAY | {fecha_peru_str()}")
                        cp_r.save()
                        buf_r.seek(0)
                        st.download_button("Descargar PDF", buf_r,
                                           "Qaway_Resultados.pdf", "application/pdf", type="primary", key="dl_plik_pdf")
                with col_wa_r:
                    if st.button("WhatsApp a Padres", type="primary", use_container_width=True, key="plik_wa"):
                        st.session_state.plik_wa_show = True
                if st.session_state.get('plik_wa_show'):
                    st.markdown("### Enviar Resultados por WhatsApp")
                    dfm_wa = BaseDatos.cargar_matricula()
                    envios_wa = 0
                    for r_wa in res:
                        cel_wa = ''
                        if not dfm_wa.empty and 'DNI' in dfm_wa.columns:
                            fi_wa = dfm_wa[dfm_wa['DNI'].astype(str).str.strip() == r_wa['DNI'].strip()]
                            if not fi_wa.empty:
                                cel_wa = str(fi_wa.iloc[0].get('Celular_Apoderado',
                                             fi_wa.iloc[0].get('Celular', ''))).strip()
                                if cel_wa and cel_wa not in ('nan', 'None', ''):
                                    if '.' in cel_wa:
                                        cel_wa = cel_wa.split('.')[0]
                                    cel_wa = ''.join(c for c in cel_wa if c.isdigit())
                                    cel_wa = '' if len(cel_wa) < 7 else cel_wa
                        if cel_wa:
                            elogio = 'Excelente!' if r_wa['Nota'] >= 14 else ('Buen trabajo.' if r_wa['Nota'] >= 11 else 'Puede mejorar.')
                            msg_wa = (f"Estimado apoderado, {r_wa['Nombre']} obtuvo "
                                      f"{r_wa['Nota']}/20 ({r_wa['Correctas']}/{r_wa['Total']} correctas) "
                                      f"en: {qr2.get('titulo', '')}. {elogio} - I.E.P. YACHAY")
                            link_wa = generar_link_whatsapp(cel_wa, msg_wa)
                            st.markdown(f'<a href="{link_wa}" target="_blank" class="wa-btn">{r_wa["Nombre"]} - {r_wa["Nota"]}/20</a>', unsafe_allow_html=True)
                            envios_wa += 1
                        else:
                            st.caption(f"Sin celular: {r_wa['Nombre']}")


def main():
    if st.session_state.rol is None:
        pantalla_login()
        st.stop()

    # Restaurar datos desde Google Sheets si archivos locales no existen
    if 'datos_restaurados' not in st.session_state:
        try:
            _restaurar_datos_desde_gs()
            _restaurar_todos_archivos_binarios()  # escudos, fondo, mp3 pausas y qaway
            st.session_state.datos_restaurados = True
        except Exception:
            st.session_state.datos_restaurados = True

    config = configurar_sidebar()

    # Saludo personalizado
    usuario = st.session_state.get('usuario_actual', '')
    usuarios = cargar_usuarios()
    # Nombre completo: busca en Docentes, docente_info, usuarios
    if st.session_state.rol == 'docente':
        nombre_usuario = _nombre_completo_docente()
    else:
        _di = st.session_state.get('docente_info') or {}
        nombre_usuario = (_di.get('label') or _di.get('nombre') or
                          usuarios.get(usuario, {}).get('label', '') or
                          usuario.replace('.', ' ').title())
    hora_actual = hora_peru().hour
    if hora_actual < 12:
        saludo = "☀️ Buenos días"
    elif hora_actual < 18:
        saludo = "🌤️ Buenas tardes"
    else:
        saludo = "🌙 Buenas noches"

    # ========================================
    # AUXILIAR — Asistencia + Reportes + Incidencias
    # ========================================
    if st.session_state.rol == "auxiliar":
        st.markdown(f"### {saludo}, **{nombre_usuario}** 👋")
        st.markdown("*¿Qué vamos a hacer hoy?*")
        ca1, ca2, ca3, ca4 = st.columns(4)
        with ca1:
            if st.button("📋\n\n**Asistencia**", use_container_width=True, key="aux_asist", type="primary"):
                st.session_state.modulo_activo = "asistencia"
        with ca2:
            if st.button("📈\n\n**Reportes**", use_container_width=True, key="aux_rep", type="primary"):
                st.session_state.modulo_activo = "reportes"
        with ca3:
            if st.button("📝\n\n**Incidencias**", use_container_width=True, key="aux_inc", type="primary"):
                st.session_state.modulo_activo = "incidencias"
        with ca4:
            if st.button("📋\n\n**Registros PDF**", use_container_width=True, key="aux_regpdf", type="primary"):
                st.session_state.modulo_activo = "registros_pdf"

        mod = st.session_state.get('modulo_activo', 'asistencia')
        st.markdown("---")
        if mod == "asistencia":
            tab_asistencias()
        elif mod == "reportes":
            tab_reportes(config)
        elif mod == "incidencias":
            tab_incidencias(config)
        elif mod == "registros_pdf":
            st.subheader("📋 Registros PDF — Asistencia y Auxiliar")
            _seccion_registros_pdf(config)

    # ========================================
    # DOCENTE — Su grado solamente
    # ========================================
    elif st.session_state.rol == "docente":
        # Si no hay módulo seleccionado, mostrar dashboard
        if 'modulo_activo' not in st.session_state:
            st.session_state.modulo_activo = None

        if st.session_state.modulo_activo is None:
            # === DASHBOARD PRINCIPAL ===
            st.markdown(f"""
            <div class='main-header'>
                <h2 style='color:white;margin:0;'>{saludo}, {nombre_usuario} 👋</h2>
                <p style='color:#ccc;'>¿Qué vamos a hacer hoy?</p>
            </div>
            """, unsafe_allow_html=True)

            # Grid de módulos para docentes — TODOS los permitidos
            modulos = [
                ("📝", "Registrar Notas", "reg_notas", "#059669"),
                ("📝", "Registro Auxiliar", "reg_auxiliar", "#2563eb"),
                ("📋", "Registro de Asistencia", "reg_pdf", "#0891b2"),
                ("📄", "Registrar Ficha", "aula_virtual", "#7c3aed"),
                ("📝", "Exámenes Sem.", "examenes_sem", "#b91c1c"),
                ("📝", "YACHAY QAWAY", "plickers", "#7c3aed"),
                ("📊", "Calificación YACHAY", "calificacion", "#dc2626"),
                ("🏃", "Pausa Activa", "pausa_activa", "#059669"),
            ]

            # Grid de módulos
            for i in range(0, len(modulos), 3):
                cols = st.columns(3)
                for j, col in enumerate(cols):
                    idx = i + j
                    if idx < len(modulos):
                        icono, nombre, key, color = modulos[idx]
                        with col:
                            st.markdown(f"""
                            <div style='background: {color}; 
                                        color: white; 
                                        padding: 40px 20px; 
                                        border-radius: 12px; 
                                        text-align: center;
                                        margin-bottom: 10px;
                                        box-shadow: 0 4px 15px rgba(0,0,0,0.2);'>
                                <div style='font-size: 3rem; margin-bottom: 15px;'>{icono}</div>
                                <div style='font-size: 1.3rem; font-weight: bold;'>{nombre}</div>
                            </div>
                            """, unsafe_allow_html=True)
                            
                            if st.button(f"▶ Abrir {nombre}", key=f"dash_doc_{key}", 
                                        type="primary", use_container_width=True):
                                st.session_state.modulo_activo = key
                                st.rerun()

            # Estadísticas del grado
            st.markdown("---")
            info_doc = st.session_state.get('docente_info', {}) or {}
            grado_doc = info_doc.get('grado', '')
            if grado_doc:
                s1 = st.columns(1)[0]
                with s1:
                    st.markdown(f"""<div class="stat-card">
                        <h3>🎓 {grado_doc}</h3>
                        <p>Tu Grado Asignado</p>
                    </div>""", unsafe_allow_html=True)

        else:
            # === MÓDULO SELECCIONADO ===
            col_back, col_space = st.columns([1, 4])
            with col_back:
                if st.button("⬅️ REGRESAR", key="btn_volver_doc", use_container_width=True, type="primary"):
                    st.session_state.modulo_activo = None
                    st.rerun()
            
            st.markdown("---")
            st.markdown(f"### {saludo}, **{nombre_usuario}** 👋")

            mod = st.session_state.modulo_activo
            if mod == "calificacion":
                tab_calificacion_yachay(config)
            elif mod == "reg_notas":
                tab_registrar_notas(config)
            elif mod == "aula_virtual":
                tab_material_docente(config)
            elif mod == "examenes_sem":
                tab_examenes_semanales(config)
            elif mod == "plickers":
                tab_yachay_plickers(config)
            elif mod == "pausa_activa":
                tab_pausa_activa(config)
            elif mod == "registros_pdf":
                st.subheader("📋 Registros PDF — Asistencia")
                _seccion_registros_pdf(config)
            elif mod == "reg_auxiliar":
                info_d = st.session_state.get('docente_info', {}) or {}
                grado_d = info_d.get('grado', '')
                _tab_registro_auxiliar_docente(grado_d, config)
            elif mod == "reg_pdf":
                info_d = st.session_state.get('docente_info', {}) or {}
                grado_d = info_d.get('grado', '')
                _tab_registro_pdf_docente(grado_d, config)

    # ========================================
    # ADMIN / DIRECTIVO — Dashboard con íconos
    # ========================================
    elif st.session_state.rol in ["directivo", "admin"]:
        # Si no hay módulo seleccionado, mostrar dashboard
        if 'modulo_activo' not in st.session_state:
            st.session_state.modulo_activo = None

        if st.session_state.modulo_activo is None:
            # === DASHBOARD PRINCIPAL ===
            st.markdown(f"""
            <div class='main-header'>
                <h2 style='color:white;margin:0;'>{saludo}, {nombre_usuario} 👋</h2>
                <p style='color:#ccc;'>¿Qué vamos a hacer hoy?</p>
            </div>
            """, unsafe_allow_html=True)

            # Grid de módulos
            # Grid de módulos
            modulos = [
                ("📝", "Matrícula", "matricula", "#2563eb"),
                ("📋", "Asistencia", "asistencia", "#16a34a"),
                ("📄", "Documentos", "documentos", "#7c3aed"),
                ("🪪", "Carnets", "carnets", "#0891b2"),
                ("📊", "Calificación", "calificacion", "#dc2626"),
                ("📝", "Registrar Notas", "reg_notas", "#059669"),
                ("📈", "Reportes", "reportes", "#ea580c"),
                ("📝", "Incidencias", "incidencias", "#be185d"),
                ("💾", "Base Datos", "base_datos", "#4f46e5"),
                ("📄", "Registrar Ficha", "aula_virtual", "#7c3aed"),
                ("📝", "Exámenes Sem.", "examenes_sem", "#b91c1c"),
                ("📝", "YACHAY QAWAY", "plickers", "#7c3aed"),
                ("📋", "Registros PDF", "registros_pdf", "#0d9488"),
                ("🏃", "Pausa Activa", "pausa_activa", "#059669"),
            ]
            if st.session_state.rol == "admin":
                modulos.append(("📕", "Reclamaciones", "reclamaciones", "#92400e"))

            # Grid de módulos - SOLUCIÓN SIMPLE Y VISIBLE
            for i in range(0, len(modulos), 3):
                cols = st.columns(3)
                for j, col in enumerate(cols):
                    idx = i + j
                    if idx < len(modulos):
                        icono, nombre, key, color = modulos[idx]
                        with col:
                            # Cuadrado HTML de color sólido - MUY VISIBLE
                            st.markdown(f"""
                            <div style='background: {color}; 
                                        color: white; 
                                        padding: 40px 20px; 
                                        border-radius: 12px; 
                                        text-align: center;
                                        margin-bottom: 10px;
                                        box-shadow: 0 4px 15px rgba(0,0,0,0.2);'>
                                <div style='font-size: 3rem; margin-bottom: 15px;'>{icono}</div>
                                <div style='font-size: 1.3rem; font-weight: bold;'>{nombre}</div>
                            </div>
                            """, unsafe_allow_html=True)
                            
                            # Botón visible debajo para hacer click
                            if st.button(f"▶ Abrir {nombre}", key=f"dash_{key}", 
                                        type="primary", use_container_width=True):
                                st.session_state.modulo_activo = key
                                st.rerun()

            # Estadísticas rápidas
            st.markdown("---")
            stats = BaseDatos.obtener_estadisticas()
            s1, s2, s3 = st.columns(3)
            with s1:
                st.markdown(f"""<div class="stat-card">
                    <h3>📚 {stats['total_alumnos']}</h3>
                    <p>Alumnos Matriculados</p>
                </div>""", unsafe_allow_html=True)
            with s2:
                st.markdown(f"""<div class="stat-card">
                    <h3>👨‍🏫 {stats['total_docentes']}</h3>
                    <p>Docentes Registrados</p>
                </div>""", unsafe_allow_html=True)
            with s3:
                asis_hoy = BaseDatos.obtener_asistencias_hoy()
                st.markdown(f"""<div class="stat-card">
                    <h3>📋 {len(asis_hoy)}</h3>
                    <p>Asistencias Hoy</p>
                </div>""", unsafe_allow_html=True)

        else:
            # === MÓDULO SELECCIONADO ===
            col_back, col_space = st.columns([1, 4])
            
            with col_back:
                if st.button("⬅️ REGRESAR", key="btn_volver", use_container_width=True, type="primary"):
                    st.session_state.modulo_activo = None
                    st.rerun()
            
            st.markdown("---")
            st.markdown(f"### {saludo}, **{nombre_usuario}** 👋")

            mod = st.session_state.modulo_activo
            if mod == "matricula":
                tab_matricula(config)
            elif mod == "asistencia":
                tab_asistencias()
            elif mod == "documentos":
                tab_documentos(config)
            elif mod == "carnets":
                tab_carnets(config)
            elif mod == "calificacion":
                tab_calificacion_yachay(config)
            elif mod == "reg_notas":
                tab_registrar_notas(config)
            elif mod == "reportes":
                tab_reportes(config)
            elif mod == "incidencias":
                tab_incidencias(config)
            elif mod == "base_datos":
                tab_base_datos()
            elif mod == "reclamaciones":
                tab_libro_reclamaciones(config)
            elif mod == "aula_virtual":
                tab_material_docente(config)
            elif mod == "examenes_sem":
                tab_examenes_semanales(config)
            elif mod == "plickers":
                tab_yachay_plickers(config)
            elif mod == "pausa_activa":
                tab_pausa_activa(config)
            elif mod == "registros_pdf":
                st.subheader("📋 Registros PDF — Asistencia")
                _seccion_registros_pdf(config)


# ================================================================
# LIBRO DE RECLAMACIONES VIRTUAL
# ================================================================

def tab_libro_reclamaciones(config):
    """Libro de Reclamaciones Virtual según normativa MINEDU"""
    st.subheader("📕 Libro de Reclamaciones Virtual")
    st.markdown("*Según normativa del Ministerio de Educación*")

    gs = _gs()

    col1, col2 = st.columns([2, 1])
    with col1:
        with st.form("form_reclamo", clear_on_submit=True):
            st.markdown("### 📋 Registrar Reclamo")
            r_nombre = st.text_input("Nombre completo del reclamante:", key="r_nombre")
            r_dni = st.text_input("DNI:", key="r_dni")
            r_celular = st.text_input("Celular:", key="r_cel")
            r_tipo = st.selectbox("Tipo:", [
                "Queja", "Reclamo", "Sugerencia", "Denuncia"
            ], key="r_tipo")
            r_detalle = st.text_area("Detalle del reclamo:", key="r_detalle")
            r_submit = st.form_submit_button("📩 ENVIAR RECLAMO",
                                              type="primary",
                                              use_container_width=True)
            if r_submit:
                if r_nombre and r_dni and r_detalle:
                    codigo_rec = f"REC-{hora_peru().year}-{int(time.time()) % 10000:04d}"
                    if gs:
                        try:
                            ws = gs._get_hoja('config')
                            if ws:
                                ws.append_row([
                                    f"reclamo_{codigo_rec}",
                                    json.dumps({
                                        'codigo': codigo_rec,
                                        'nombre': r_nombre,
                                        'dni': r_dni,
                                        'celular': r_celular,
                                        'tipo': r_tipo,
                                        'detalle': r_detalle,
                                        'fecha': fecha_peru_str(),
                                        'hora': hora_peru_str(),
                                        'estado': 'Pendiente',
                                    }, ensure_ascii=False)
                                ])
                        except Exception:
                            pass
                    st.success(f"✅ Reclamo registrado exitosamente. Código: **{codigo_rec}**")
                    st.info("📌 Su reclamo será revisado por la dirección en un plazo de 72 horas.")
                else:
                    st.error("⚠️ Complete todos los campos obligatorios")

    with col2:
        st.markdown("### 📋 Reclamos Recibidos")
        if gs:
            try:
                ws = gs._get_hoja('config')
                if ws:
                    data = ws.get_all_records()
                    reclamos = [json.loads(d['valor']) for d in data
                               if str(d.get('clave', '')).startswith('reclamo_')]
                    if reclamos:
                        for rec in reversed(reclamos[-15:]):
                            estado = rec.get('estado', 'Pendiente')
                            emoji = "🟡" if estado == "Pendiente" else "🟢"
                            with st.expander(
                                f"{emoji} {rec.get('codigo', '')} — {rec.get('nombre', '')}"):
                                st.write(f"**Tipo:** {rec.get('tipo', '')}")
                                st.write(f"**Fecha:** {rec.get('fecha', '')}")
                                st.write(f"**Detalle:** {rec.get('detalle', '')}")
                                st.write(f"**Estado:** {estado}")
                    else:
                        st.info("📭 Sin reclamos registrados")
            except Exception:
                st.info("📭 Sin reclamos aún")
        else:
            st.warning("⚠️ Conecta Google Sheets")


if __name__ == "__main__":
    main()
